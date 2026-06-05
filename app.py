import json
import os
import re
import time
import unicodedata
import uuid
import csv
import io
import subprocess
import shutil
from datetime import datetime, date, timedelta
from types import SimpleNamespace
from decimal import Decimal
from pathlib import Path
from urllib.request import urlopen
from urllib.parse import urlparse, urlunparse, parse_qsl, urlencode
from werkzeug.security import generate_password_hash, check_password_hash
from flask import Flask, render_template, request, redirect, url_for, flash, send_file, Response, session, abort, has_request_context, jsonify
from werkzeug.utils import secure_filename
from sqlalchemy import event, func, or_
from sqlalchemy.engine import Engine
from flask_sqlalchemy import SQLAlchemy
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = Path(__file__).resolve().parent
DB_PATH = BASE_DIR / 'instance' / 'business.db'
IMPORT_PREVIEW_DIR = BASE_DIR / 'instance' / 'import_previews'
IMPORT_PREVIEW_SESSION_KEY = 'product_import_preview_id'
CATEGORY_IMPORT_SESSION_KEY = 'category_import_preview_id'
SUPPLIER_INTAKE_IMPORT_SESSION_KEY = 'supplier_intake_import_preview_id'
OUTPUT_DIR = BASE_DIR / 'output'
TEMPLATE_DOCX = BASE_DIR / 'templates_docx' / 'hop_dong_template.docx'
TEMPLATE_QUOTE_SALE_CONTRACT = BASE_DIR / 'templates_docx' / 'hop_dong_mua_ban_bao_gia.docx'
TEMPLATE_QUOTE_HANDOVER = BASE_DIR / 'templates_docx' / 'bien_ban_nhan_hang.docx'
QUOTE_DOCS_DIR = OUTPUT_DIR / 'quotes'
PRODUCT_UPLOAD_DIR = BASE_DIR / 'static' / 'uploads' / 'products'
COMPANY_LOGO_DIR = BASE_DIR / 'static' / 'uploads' / 'company'
CONTRACT_SIGNED_UPLOAD_DIR = BASE_DIR / 'static' / 'uploads' / 'contracts' / 'signed'
CONTRACT_DRAFT_SESSION_KEY = 'contract_draft_doc_path'
ACTIVE_COMPANY_SESSION_KEY = 'active_company_id'
TEMP_DOC_DIR = BASE_DIR / 'instance' / 'temp_docs'
DOCX_MIMETYPE = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
PDF_MIMETYPE = 'application/pdf'
DOCX_TO_PDF_TIMEOUT = 180
ORDER_HANDOVER_UPLOAD_DIR = BASE_DIR / 'static' / 'uploads' / 'orders' / 'handover'
PAYMENT_RECEIPT_UPLOAD_DIR = BASE_DIR / 'static' / 'uploads' / 'payments' / 'receipts'
SUPPLIER_PAYMENT_RECEIPT_UPLOAD_DIR = BASE_DIR / 'static' / 'uploads' / 'suppliers' / 'receipts'
ALLOWED_ORDER_HANDOVER_EXTENSIONS = {'.jpg', '.jpeg', '.png', '.webp'}
ALLOWED_IMAGE_EXTENSIONS = {'.jpg', '.jpeg', '.png', '.webp', '.gif'}
ALLOWED_CONTRACT_SCAN_EXTENSIONS = {'.jpg', '.jpeg', '.png', '.webp', '.pdf'}
MAX_PRODUCT_IMAGE_BYTES = 5 * 1024 * 1024
MAX_PRODUCT_IMAGES = 5
MAX_CONTRACT_SCAN_BYTES = 10 * 1024 * 1024

app = Flask(__name__)
app.config['SECRET_KEY'] = 'change-this-secret-key'
app.config['SQLALCHEMY_DATABASE_URI'] = f'sqlite:///{DB_PATH}'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

db = SQLAlchemy(app)

def vietnamese_search_fold(value):
    """Bỏ dấu tiếng Việt để so khớp tìm kiếm gần đúng."""
    if value is None:
        return ''
    text = str(value).strip().lower()
    if not text:
        return ''
    text = unicodedata.normalize('NFD', text)
    text = ''.join(ch for ch in text if unicodedata.category(ch) != 'Mn')
    return text.replace('đ', 'd')


def _sqlite_vn_fold(value):
    return vietnamese_search_fold(value)


@event.listens_for(Engine, 'connect')
def _register_sqlite_search_functions(dbapi_connection, connection_record):
    try:
        dbapi_connection.create_function('vn_fold', 1, _sqlite_vn_fold)
    except Exception:
        pass


class Customer(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    customer_type = db.Column(db.String(20), default='company')
    name = db.Column(db.String(255), nullable=False)
    tax_code = db.Column(db.String(64), default='')
    id_card = db.Column(db.String(32), default='')
    address = db.Column(db.Text, default='')
    phone = db.Column(db.String(64), default='')
    email = db.Column(db.String(120), default='')
    representative = db.Column(db.String(120), default='')
    position = db.Column(db.String(120), default='')
    bank_account = db.Column(db.String(120), default='')
    bank_name = db.Column(db.String(255), default='')
    note = db.Column(db.Text, default='')
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

class Category(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(120), nullable=False)
    parent_id = db.Column(db.Integer, db.ForeignKey('category.id'), nullable=True)
    sort_order = db.Column(db.Integer, default=0)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    parent = db.relationship('Category', remote_side=[id], backref=db.backref('children', lazy='dynamic'))

class Brand(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(120), unique=True, nullable=False)
    category_id = db.Column(db.Integer, db.ForeignKey('category.id'), nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    category = db.relationship('Category')

class Product(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    sku = db.Column(db.String(80), unique=True, nullable=False)
    name = db.Column(db.String(255), nullable=False)
    category = db.Column(db.String(120), default='')
    brand = db.Column(db.String(120), default='')
    category_id = db.Column(db.Integer, db.ForeignKey('category.id'), nullable=True)
    brand_id = db.Column(db.Integer, db.ForeignKey('brand.id'), nullable=True)
    category_ref = db.relationship('Category', foreign_keys=[category_id])
    brand_ref = db.relationship('Brand', foreign_keys=[brand_id])
    model = db.Column(db.String(120), default='')
    variant = db.Column(db.String(120), default='')
    variant_prices = db.Column(db.Text, default='')
    variant_stocks = db.Column(db.Text, default='')
    variant_cost_prices = db.Column(db.Text, default='')
    variant_dealer_prices = db.Column(db.Text, default='')
    variant_project_prices = db.Column(db.Text, default='')
    variant_image_urls = db.Column(db.Text, default='')
    variant_brand_ids = db.Column(db.Text, default='')
    unit = db.Column(db.String(40), default='cái')
    base_unit = db.Column(db.String(40), default='')
    purchase_unit = db.Column(db.String(40), default='')
    unit_conversion_enabled = db.Column(db.Boolean, default=False)
    conversion_factor = db.Column(db.Float, default=1.0)
    lot_unit = db.Column(db.String(40), default='')
    lot_factor = db.Column(db.Float, default=0)
    sale_unit_mode = db.Column(db.String(20), default='')
    stock_qty = db.Column('stock', db.Float, default=0)
    warranty = db.Column(db.String(120), default='')
    cost_price = db.Column(db.Integer, default=0)
    retail_price = db.Column(db.Integer, default=0)
    dealer_price = db.Column(db.Integer, default=0)
    project_price = db.Column(db.Integer, default=0)
    low_stock = db.Column(db.Float, default=5)
    image_path = db.Column(db.String(255), default='')
    image_paths = db.Column(db.Text, default='')
    image_url = db.Column(db.String(512), default='')
    is_active = db.Column(db.Boolean, default=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

    @property
    def stock(self):
        return float(self.stock_qty or 0)

    @stock.setter
    def stock(self, value):
        self.stock_qty = float(value or 0)

class PriceHistory(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    product_id = db.Column(db.Integer, db.ForeignKey('product.id'), nullable=False)
    price_type = db.Column(db.String(50), nullable=False)
    old_price = db.Column(db.Integer, default=0)
    new_price = db.Column(db.Integer, default=0)
    note = db.Column(db.String(255), default='')
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    product = db.relationship('Product')

class Quote(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    quote_code = db.Column(db.String(80), unique=True, nullable=False)
    customer_id = db.Column(db.Integer, db.ForeignKey('customer.id'), nullable=False)
    status = db.Column(db.String(40), default='Mới tạo')
    vat_rate = db.Column(db.Integer, default=10)
    discount = db.Column(db.Integer, default=0)
    subtotal = db.Column(db.Integer, default=0)
    vat_amount = db.Column(db.Integer, default=0)
    total = db.Column(db.Integer, default=0)
    note = db.Column(db.Text, default='')
    walkin_display_name = db.Column(db.String(255), default='')
    valid_until = db.Column(db.Date, nullable=True)
    sale_contract_path = db.Column(db.String(512), default='')
    handover_doc_path = db.Column(db.String(512), default='')
    doc_company_id = db.Column(db.Integer, db.ForeignKey('company_profile.id'), nullable=True)
    sale_contract_signed_date = db.Column(db.Date, nullable=True)
    sale_contract_payment_note = db.Column(db.Text, default='')
    sale_contract_delivery_note = db.Column(db.Text, default='')
    handover_date = db.Column(db.Date, nullable=True)
    handover_place = db.Column(db.Text, default='')
    handover_condition_note = db.Column(db.Text, default='')
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    customer = db.relationship('Customer')
    doc_company = db.relationship('CompanyProfile', foreign_keys=[doc_company_id])
    items = db.relationship('QuoteItem', cascade='all, delete-orphan', backref='quote')

class QuoteItem(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    quote_id = db.Column(db.Integer, db.ForeignKey('quote.id'), nullable=False)
    product_id = db.Column(db.Integer, db.ForeignKey('product.id'), nullable=False)
    variant_label = db.Column(db.String(120), default='')
    product_name = db.Column(db.String(255), nullable=False)
    sku = db.Column(db.String(80), default='')
    unit = db.Column(db.String(40), default='cái')
    qty_unit_mode = db.Column(db.String(20), default='')
    qty = db.Column(db.Float, default=1)
    price = db.Column(db.Integer, default=0)
    amount = db.Column(db.Integer, default=0)
    product = db.relationship('Product')

CONTRACT_TYPE_FRAMEWORK = 'framework'

class Contract(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    contract_code = db.Column(db.String(80), unique=True, nullable=False)
    customer_id = db.Column(db.Integer, db.ForeignKey('customer.id'), nullable=False)
    company_id = db.Column(db.Integer, db.ForeignKey('company_profile.id'), nullable=True)
    contract_type = db.Column(db.String(20), default='framework')
    quote_id = db.Column(db.Integer, db.ForeignKey('quote.id'), nullable=True)
    signed_date = db.Column(db.Date, default=date.today)
    expired_date = db.Column(db.Date, nullable=True)
    file_path = db.Column(db.String(255), default='')
    signed_scan_path = db.Column(db.String(255), default='')
    signed_scan_uploaded_at = db.Column(db.DateTime, nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    customer = db.relationship('Customer')
    quote = db.relationship('Quote')
    company = db.relationship('CompanyProfile')

class Order(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    order_code = db.Column(db.String(80), unique=True, nullable=False)
    customer_id = db.Column(db.Integer, db.ForeignKey('customer.id'), nullable=False)
    quote_id = db.Column(db.Integer, db.ForeignKey('quote.id'), nullable=True)
    status = db.Column(db.String(40), default='Mới tạo')
    total = db.Column(db.Integer, default=0)
    paid_amount = db.Column(db.Integer, default=0)
    handover_scan_path = db.Column(db.String(512), default='')
    handover_scan_uploaded_at = db.Column(db.DateTime, nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    customer = db.relationship('Customer')
    quote = db.relationship('Quote')

class Payment(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    order_id = db.Column(db.Integer, db.ForeignKey('order.id'), nullable=False)
    amount = db.Column(db.Integer, default=0)
    method = db.Column(db.String(80), default='Chuyển khoản')
    note = db.Column(db.String(255), default='')
    payment_date = db.Column(db.Date, default=date.today)
    receipt_path = db.Column(db.String(512), default='')
    receipt_uploaded_at = db.Column(db.DateTime, nullable=True)
    batch_id = db.Column(db.String(64), default='')
    order = db.relationship('Order')

class CompanyProfile(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(255), nullable=False, default='')
    short_name = db.Column(db.String(80), default='')
    tax_code = db.Column(db.String(64), default='')
    address = db.Column(db.Text, default='')
    phone = db.Column(db.String(120), default='')
    email = db.Column(db.String(120), default='')
    bank_account = db.Column(db.String(120), default='')
    bank_name = db.Column(db.String(255), default='')
    director_name = db.Column(db.String(120), default='')
    director_title = db.Column(db.String(80), default='Giám đốc')
    logo_path = db.Column(db.String(255), default='')
    is_default = db.Column(db.Boolean, default=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)

class Supplier(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    code = db.Column(db.String(80), unique=True, nullable=False)
    name = db.Column(db.String(255), nullable=False)
    tax_code = db.Column(db.String(64), default='')
    contact_person = db.Column(db.String(120), default='')
    phone = db.Column(db.String(64), default='')
    email = db.Column(db.String(120), default='')
    address = db.Column(db.Text, default='')
    bank_account = db.Column(db.String(120), default='')
    bank_name = db.Column(db.String(255), default='')
    note = db.Column(db.Text, default='')
    is_active = db.Column(db.Boolean, default=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

class ProductSupplier(db.Model):
    __tablename__ = 'product_supplier'
    id = db.Column(db.Integer, primary_key=True)
    product_id = db.Column(db.Integer, db.ForeignKey('product.id'), nullable=False)
    supplier_id = db.Column(db.Integer, db.ForeignKey('supplier.id'), nullable=False)
    cost_price = db.Column(db.Integer, default=0)
    supplier_sku = db.Column(db.String(120), default='')
    note = db.Column(db.String(255), default='')
    is_primary = db.Column(db.Boolean, default=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    product = db.relationship('Product', backref=db.backref('supplier_links', lazy='dynamic', cascade='all, delete-orphan'))
    supplier = db.relationship('Supplier')
    __table_args__ = (db.UniqueConstraint('product_id', 'supplier_id', name='uq_product_supplier'),)

class StockMovement(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    product_id = db.Column(db.Integer, db.ForeignKey('product.id'), nullable=False)
    supplier_id = db.Column(db.Integer, db.ForeignKey('supplier.id'), nullable=True)
    quote_id = db.Column(db.Integer, db.ForeignKey('quote.id'), nullable=True)
    movement_type = db.Column(db.String(20), nullable=False)  # IN/OUT
    qty = db.Column(db.Float, default=0)
    purchase_qty = db.Column(db.Float, default=0)
    ref_code = db.Column(db.String(80), default='')
    method = db.Column(db.String(80), default='Nhập tay')
    warehouse = db.Column(db.String(80), default='Kho chính')
    note = db.Column(db.String(255), default='')
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    product = db.relationship('Product')
    supplier = db.relationship('Supplier')
    quote = db.relationship('Quote')

SUPPLIER_INTAKE_STATUS_PENDING = 'Chờ nhập kho'
SUPPLIER_INTAKE_STATUS_RECEIVED = 'Đã nhập kho'

class SupplierIntake(db.Model):
    __tablename__ = 'supplier_intake'
    id = db.Column(db.Integer, primary_key=True)
    supplier_id = db.Column(db.Integer, db.ForeignKey('supplier.id'), nullable=False)
    intake_code = db.Column(db.String(80), default='')
    ref_code = db.Column(db.String(80), default='')
    note = db.Column(db.String(255), default='')
    status = db.Column(db.String(40), default=SUPPLIER_INTAKE_STATUS_PENDING)
    total_amount = db.Column(db.Integer, default=0)
    paid_amount = db.Column(db.Integer, default=0)
    stock_received_at = db.Column(db.DateTime, nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    supplier = db.relationship('Supplier', backref=db.backref('intakes', lazy='dynamic'))
    lines = db.relationship(
        'SupplierIntakeLine',
        backref='intake',
        lazy='dynamic',
        cascade='all, delete-orphan',
    )
    payments = db.relationship('SupplierPayment', backref='intake', lazy='dynamic')

class SupplierIntakeLine(db.Model):
    __tablename__ = 'supplier_intake_line'
    id = db.Column(db.Integer, primary_key=True)
    intake_id = db.Column(db.Integer, db.ForeignKey('supplier_intake.id'), nullable=False)
    product_id = db.Column(db.Integer, db.ForeignKey('product.id'), nullable=False)
    qty = db.Column(db.Float, default=0)
    purchase_qty = db.Column(db.Float, default=0)
    cost_price = db.Column(db.Integer, default=0)
    line_amount = db.Column(db.Integer, default=0)
    supplier_sku = db.Column(db.String(120), default='')
    note = db.Column(db.String(255), default='')
    stock_movement_id = db.Column(db.Integer, db.ForeignKey('stock_movement.id'), nullable=True)
    product = db.relationship('Product')
    stock_movement = db.relationship('StockMovement')

    @property
    def qty_base(self):
        return float(self.qty or 0)

class SupplierIntakeCustomerAllocation(db.Model):
    """Phân bổ số lượng nhập kho cho khách hàng (đơn/báo giá đã chốt)."""
    __tablename__ = 'supplier_intake_customer_allocation'
    id = db.Column(db.Integer, primary_key=True)
    intake_line_id = db.Column(db.Integer, db.ForeignKey('supplier_intake_line.id'), nullable=False)
    customer_id = db.Column(db.Integer, db.ForeignKey('customer.id'), nullable=False)
    order_id = db.Column(db.Integer, db.ForeignKey('order.id'), nullable=True)
    quote_id = db.Column(db.Integer, db.ForeignKey('quote.id'), nullable=True)
    qty = db.Column(db.Float, default=0)
    customer = db.relationship('Customer')
    order = db.relationship('Order')
    quote = db.relationship('Quote')

class PurchaseOrder(db.Model):
    __tablename__ = 'purchase_order'
    id = db.Column(db.Integer, primary_key=True)
    po_code = db.Column(db.String(80), unique=True, nullable=False)
    supplier_id = db.Column(db.Integer, db.ForeignKey('supplier.id'), nullable=True)
    receipt_mode = db.Column(db.String(40), default='stock_in')
    customer_id = db.Column(db.Integer, db.ForeignKey('customer.id'), nullable=True)
    order_id = db.Column(db.Integer, db.ForeignKey('order.id'), nullable=True)
    quote_id = db.Column(db.Integer, db.ForeignKey('quote.id'), nullable=True)
    status = db.Column(db.String(40), default='Nháp')
    ref_code = db.Column(db.String(80), default='')
    note = db.Column(db.String(255), default='')
    total_amount = db.Column(db.Integer, default=0)
    paid_amount = db.Column(db.Integer, default=0)
    received_at = db.Column(db.DateTime, nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    supplier = db.relationship('Supplier', backref=db.backref('purchase_orders', lazy='dynamic'))
    customer = db.relationship('Customer')
    order = db.relationship('Order')
    quote = db.relationship('Quote')
    lines = db.relationship(
        'PurchaseOrderLine',
        backref='purchase_order',
        lazy='dynamic',
        cascade='all, delete-orphan',
    )
    payments = db.relationship('SupplierPayment', backref='purchase_order', lazy='dynamic')

class PurchaseOrderLine(db.Model):
    __tablename__ = 'purchase_order_line'
    id = db.Column(db.Integer, primary_key=True)
    purchase_order_id = db.Column(db.Integer, db.ForeignKey('purchase_order.id'), nullable=False)
    supplier_id = db.Column(db.Integer, db.ForeignKey('supplier.id'), nullable=True)
    product_id = db.Column(db.Integer, db.ForeignKey('product.id'), nullable=False)
    qty = db.Column(db.Float, default=0)
    purchase_qty = db.Column(db.Float, default=0)
    cost_price = db.Column(db.Integer, default=0)
    line_amount = db.Column(db.Integer, default=0)
    supplier_sku = db.Column(db.String(120), default='')
    note = db.Column(db.String(255), default='')
    stock_movement_id = db.Column(db.Integer, db.ForeignKey('stock_movement.id'), nullable=True)
    product = db.relationship('Product')
    supplier = db.relationship('Supplier')
    stock_movement = db.relationship('StockMovement')

class SupplierPayment(db.Model):
    __tablename__ = 'supplier_payment'
    id = db.Column(db.Integer, primary_key=True)
    supplier_id = db.Column(db.Integer, db.ForeignKey('supplier.id'), nullable=False)
    intake_id = db.Column(db.Integer, db.ForeignKey('supplier_intake.id'), nullable=True)
    purchase_order_id = db.Column(db.Integer, db.ForeignKey('purchase_order.id'), nullable=True)
    amount = db.Column(db.Integer, default=0)
    method = db.Column(db.String(80), default='Chuyển khoản')
    note = db.Column(db.String(255), default='')
    payment_date = db.Column(db.Date, default=date.today)
    receipt_path = db.Column(db.String(512), default='')
    receipt_uploaded_at = db.Column(db.DateTime, nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    supplier = db.relationship('Supplier', backref=db.backref('supplier_payments', lazy='dynamic'))

class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), unique=True, nullable=False)
    password_hash = db.Column(db.String(255), nullable=False)
    full_name = db.Column(db.String(120), default='')
    role = db.Column(db.String(20), default='sales', nullable=False)
    is_active = db.Column(db.Boolean, default=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

USER_ROLES = {
    'admin': 'Quản trị viên',
    'sales': 'Bán hàng',
    'warehouse': 'Quản lý kho',
}

ADMIN_ONLY_ENDPOINTS = frozenset({
    'company_settings',
    'company_create',
    'company_delete',
    'users',
    'update_user',
    'toggle_user_active',
    'delete_user',
    'products_delete_all',
    'delete_customer',
    'dashboard_clear_data',
})

SALES_ENDPOINTS = frozenset({
    'dashboard',
    'customers',
    'customer_options_json',
    'update_customer',
    'upload_customer_signed_contract',
    'delete_customer_signed_contract',
    'preview_customer_framework_contract',
    'preview_customer_signed_contract',
    'download_customer_framework_contract',
    'download_customer_framework_contract_pdf',
    'download_customer_signed_contract',
    'view_customer_framework_contract',
    'view_customer_framework_contract_pdf',
    'view_customer_signed_contract',
    'quotes',
    'quote_create_preview',
    'quotes_export',
    'quote_preview',
    'quote_print',
    'update_quote',
    'update_quote_status',
    'delete_quote',
    'create_order_from_quote',
    'preview_quote_sale_contract',
    'preview_quote_handover',
    'view_quote_sale_contract',
    'view_quote_sale_contract_pdf',
    'view_quote_handover',
    'view_quote_handover_pdf',
    'download_quote_sale_contract',
    'download_quote_sale_contract_pdf',
    'download_quote_handover',
    'download_quote_handover_pdf',
    'update_quote_sale_contract',
    'update_quote_handover',
    'preview_quote_sale_contract_draft',
    'preview_quote_handover_draft',
    'view_quote_sale_contract_draft',
    'view_quote_sale_contract_draft_pdf',
    'download_quote_sale_contract_draft',
    'download_quote_sale_contract_draft_pdf',
    'view_quote_handover_draft',
    'view_quote_handover_draft_pdf',
    'download_quote_handover_draft',
    'download_quote_handover_draft_pdf',
    'orders',
    'orders_export',
    'order_preview',
    'order_print',
    'update_order_status',
    'complete_order',
    'delete_order',
    'view_order_handover_scan',
    'add_payment',
    'debt',
    'debt_allocate_payment',
    'debt_export',
    'delete_debt_payment',
    'view_payment_receipt',
    'contracts',
    'preview_contract',
    'preview_contract_draft',
    'view_contract',
    'view_contract_pdf',
    'view_contract_draft',
    'view_contract_draft_pdf',
    'download_contract',
    'download_contract_pdf',
    'download_contract_draft',
    'download_contract_draft_pdf',
    'update_contract',
})

WAREHOUSE_ENDPOINTS = frozenset({
    'dashboard',
    'products',
    'product_detail',
    'product_edit',
    'update_product',
    'update_price',
    'toggle_product_active',
    'delete_product',
    'adjust_product_stock',
    'products_export',
    'products_import_template',
    'products_import_preview',
    'products_import_confirm',
    'products_import_cancel',
    'stock',
    'stock_export',
    'stock_intake_requests',
    'stock_intake_request_detail',
    'approve_supplier_intake_stock',
    'reject_supplier_intake_stock',
    'create_category',
    'delete_category',
    'categories_export',
    'categories_import_template',
    'categories_import_preview',
    'categories_import_confirm',
    'categories_import_cancel',
    'create_brand',
    'delete_brand',
    'suppliers',
    'update_supplier',
    'delete_supplier',
    'supplier_options_json',
    'supplier_products_intake',
    'supplier_catalog_add',
    'supplier_catalog_new',
    'supplier_catalog_remove',
    'supplier_quick_create_product',
    'supplier_intake_import_template',
    'supplier_intake_import_preview',
    'supplier_intake_import_confirm',
    'supplier_intake_import_cancel',
    'supplier_intake_payment',
    'delete_supplier_intake',
    'view_supplier_payment_receipt',
    'product_options_json',
    'customer_orders_json',
    'product_suppliers_json',
    'supplier_debt',
})

SIDEBAR_SECTIONS = [
    {
        'title': None,
        'items': [
            {'key': 'dashboard', 'endpoint': 'dashboard', 'label': 'Dashboard', 'icon': 'bi-house-door'},
        ],
    },
    {
        'title': 'DANH MỤC',
        'items': [
            {'key': 'customers', 'endpoint': 'customers', 'label': 'Khách hàng', 'icon': 'bi-people'},
            {'key': 'suppliers', 'endpoint': 'suppliers', 'label': 'Nhà cung cấp', 'icon': 'bi-truck'},
            {'key': 'products', 'endpoint': 'products', 'label': 'Sản phẩm', 'icon': 'bi-box-seam'},
        ],
    },
    {
        'title': 'KINH DOANH',
        'items': [
            {'key': 'quotes', 'endpoint': 'quotes', 'label': 'Báo giá', 'icon': 'bi-file-earmark-text'},
            {'key': 'orders', 'endpoint': 'orders', 'label': 'Đơn hàng', 'icon': 'bi-cart3'},
            {'key': 'contracts', 'endpoint': 'contracts', 'label': 'Hợp đồng', 'icon': 'bi-file-earmark-ruled'},
        ],
    },
    {
        'title': 'KHO VẬN',
        'items': [
            {'key': 'stock', 'endpoint': 'stock', 'label': 'Tồn kho', 'icon': 'bi-boxes'},
            {'key': 'stock_intake_requests', 'endpoint': 'stock_intake_requests', 'label': 'Yêu cầu nhập kho', 'icon': 'bi-inbox'},
        ],
    },
    {
        'title': 'TÀI CHÍNH',
        'items': [
            {'key': 'debt', 'endpoint': 'debt', 'label': 'Công nợ khách hàng', 'icon': 'bi-wallet2'},
            {'key': 'supplier_debt', 'endpoint': 'supplier_debt', 'label': 'Công nợ nhà cung cấp', 'icon': 'bi-wallet2'},
        ],
    },
    {
        'title': 'HỆ THỐNG',
        'items': [
            {'key': 'users', 'endpoint': 'users', 'label': 'Người dùng', 'icon': 'bi-people'},
            {'key': 'company_settings', 'endpoint': 'company_settings', 'label': 'Thông tin công ty', 'icon': 'bi-gear'},
        ],
    },
]

ROLE_NAV_KEYS = {
    'admin': {
        'dashboard', 'customers', 'suppliers', 'products', 'quotes', 'orders', 'contracts',
        'stock', 'stock_intake_requests', 'debt', 'supplier_debt', 'users', 'company_settings',
    },
    'sales': {'dashboard', 'customers', 'quotes', 'orders', 'contracts', 'debt'},
    'warehouse': {'dashboard', 'suppliers', 'products', 'stock', 'stock_intake_requests'},
}

def get_current_user():
    uid = session.get('user_id')
    if not uid:
        return None
    user = db.session.get(User, uid)
    if not user or not user.is_active:
        return None
    return user

def role_can_access(role, endpoint):
    if endpoint in ('login', 'logout', 'static', 'switch_company', None):
        return True
    if role == 'admin':
        return True
    if endpoint in ADMIN_ONLY_ENDPOINTS:
        return False
    if role == 'sales':
        return endpoint in SALES_ENDPOINTS
    if role == 'warehouse':
        return endpoint in WAREHOUSE_ENDPOINTS
    return False

def default_home_for_role(role):
    if role == 'warehouse':
        return url_for('stock')
    if role == 'sales':
        return url_for('quotes')
    return url_for('dashboard')

def safe_next_url(target):
    if not target:
        return None
    parsed = urlparse(target)
    if parsed.scheme or parsed.netloc:
        return None
    return target

def set_user_password(user, raw_password):
    user.password_hash = generate_password_hash(raw_password)

def verify_user_password(user, raw_password):
    return check_password_hash(user.password_hash, raw_password)

def money(v):
    try:
        return f"{int(v):,}".replace(',', '.') + ' đ'
    except Exception:
        return '0 đ'

def money_plain(v):
    try:
        return f"{int(v):,}".replace(',', '.')
    except Exception:
        return '0'

_DIGITS = ('không', 'một', 'hai', 'ba', 'bốn', 'năm', 'sáu', 'bảy', 'tám', 'chín')

def _read_hundreds(n, full=False):
    n = int(n)
    if n == 0:
        return 'không' if full else ''
    if n < 10:
        return _DIGITS[n]
    if n < 20:
        if n == 10:
            return 'mười'
        if n == 15:
            return 'mười lăm'
        return 'mười ' + _DIGITS[n % 10]
    if n < 100:
        tens, ones = divmod(n, 10)
        text = _DIGITS[tens] + ' mươi'
        if ones == 0:
            return text
        if ones == 1:
            return text + ' mốt'
        if ones == 4:
            return text + ' tư'
        if ones == 5:
            return text + ' lăm'
        return text + ' ' + _DIGITS[ones]
    hundreds, rem = divmod(n, 100)
    text = _DIGITS[hundreds] + ' trăm'
    if rem == 0:
        return text
    if rem < 10:
        return text + ' lẻ ' + _DIGITS[rem]
    return text + ' ' + _read_hundreds(rem)

def money_in_words(amount):
    try:
        n = int(amount)
    except (TypeError, ValueError):
        return 'Không đồng'
    if n == 0:
        return 'Không đồng'
    if n < 0:
        return 'Âm ' + money_in_words(-n)

    parts = []
    scales = ('', ' nghìn', ' triệu', ' tỷ')
    chunks = []
    temp = n
    while temp > 0:
        chunks.append(temp % 1000)
        temp //= 1000
    for idx in range(len(chunks) - 1, -1, -1):
        block = chunks[idx]
        if block == 0:
            continue
        block_text = _read_hundreds(block, full=(idx < len(chunks) - 1))
        if block_text:
            parts.append(block_text + scales[idx])
    text = ' '.join(parts).strip()
    if not text:
        return 'Không đồng'
    return text[0].upper() + text[1:] + ' đồng'

app.jinja_env.filters['money'] = money
app.jinja_env.filters['money_plain'] = money_plain
app.jinja_env.filters['money_in_words'] = money_in_words

STATUS_FLASH_CATEGORIES = {
    'Đã chốt': 'success',
    'Hoàn thành': 'success',
    'Chờ xác nhận': 'info',
    'Đang xử lý': 'info',
    'Đang giao': 'info',
    'Mới tạo': 'info',
    'Nháp': 'secondary',
    'Đã hủy': 'danger',
}

def flash_status_updated(status):
    category = STATUS_FLASH_CATEGORIES.get(status, 'info')
    flash(f'Đã cập nhật trạng thái: {status}', category)

CATEGORY_TAG_CLASSES = ['tag-blue', 'tag-green', 'tag-purple', 'tag-orange', 'tag-pink', 'tag-cyan']

def category_tag_class(name):
    if not name:
        return 'tag-gray'
    return CATEGORY_TAG_CLASSES[sum(ord(c) for c in name) % len(CATEGORY_TAG_CLASSES)]

app.jinja_env.globals['category_tag'] = category_tag_class

def ensure_category_brand_schema():
    from sqlalchemy import inspect, text
    ensure_product_columns()
    insp = inspect(db.engine)
    tables = set(insp.get_table_names())
    if 'brand' not in tables:
        Brand.__table__.create(db.engine)
    if 'category' in tables:
        cols = {c['name'] for c in insp.get_columns('category')}
        with db.engine.begin() as conn:
            if 'parent_id' not in cols:
                conn.execute(text('ALTER TABLE category ADD COLUMN parent_id INTEGER'))
            if 'sort_order' not in cols:
                conn.execute(text('ALTER TABLE category ADD COLUMN sort_order INTEGER DEFAULT 0'))
    if 'product' in tables:
        cols = {c['name'] for c in insp.get_columns('product')}
        with db.engine.begin() as conn:
            if 'category_id' not in cols:
                conn.execute(text('ALTER TABLE product ADD COLUMN category_id INTEGER'))
            if 'brand_id' not in cols:
                conn.execute(text('ALTER TABLE product ADD COLUMN brand_id INTEGER'))
    migrate_product_taxonomy_ids()

def migrate_product_taxonomy_ids():
    for product in Product.query.all():
        changed = False
        if product.category and not product.category_id:
            cid = find_category_id_by_path(product.category)
            if cid:
                product.category_id = cid
                changed = True
        if product.brand and not product.brand_id:
            brand = Brand.query.filter_by(name=product.brand).first()
            if not brand:
                brand = Brand(name=product.brand)
                db.session.add(brand)
                db.session.flush()
            product.brand_id = brand.id
            changed = True
        if changed:
            sync_product_taxonomy_strings(product)
    try:
        db.session.commit()
    except Exception:
        db.session.rollback()

def category_display_path(category_id):
    if not category_id:
        return ''
    cat = Category.query.get(category_id)
    if not cat:
        return ''
    if cat.parent_id:
        parent = Category.query.get(cat.parent_id)
        if parent:
            return f'{parent.name} > {cat.name}'
    return cat.name

def product_category_path(product):
    if getattr(product, 'category_id', None):
        return category_display_path(product.category_id)
    return product.category or ''

def product_brand_name(product):
    if getattr(product, 'brand_id', None):
        brand = Brand.query.get(product.brand_id)
        if brand:
            return brand.name
    return product.brand or ''

def parent_categories():
    return Category.query.filter(Category.parent_id.is_(None)).order_by(Category.sort_order, Category.name).all()

def child_categories(parent_id):
    if not parent_id:
        return []
    return Category.query.filter_by(parent_id=parent_id).order_by(Category.sort_order, Category.name).all()

def build_category_tree():
    all_cats = Category.query.order_by(Category.sort_order, Category.name).all()
    by_parent = {}
    for cat in all_cats:
        by_parent.setdefault(cat.parent_id, []).append(cat)

    def walk(parent_id=None):
        nodes = []
        for cat in by_parent.get(parent_id, []):
            nodes.append({'category': cat, 'children': walk(cat.id)})
        return nodes

    return walk(None)

def category_name_exists(name, parent_id=None):
    q = Category.query.filter(Category.name == name)
    if parent_id:
        q = q.filter_by(parent_id=parent_id)
    else:
        q = q.filter(Category.parent_id.is_(None))
    return q.first() is not None

def find_category_id_by_path(path):
    path = (path or '').strip()
    if not path:
        return None
    if '>' in path:
        parts = [p.strip() for p in path.split('>') if p.strip()]
        if len(parts) >= 2:
            parent = Category.query.filter(Category.parent_id.is_(None), Category.name == parts[0]).first()
            if not parent:
                return None
            child = Category.query.filter_by(parent_id=parent.id, name=parts[-1]).first()
            return child.id if child else None
    cat = Category.query.filter(Category.name == path).first()
    return cat.id if cat else None

def sync_product_taxonomy_strings(product):
    product.category = category_display_path(product.category_id) if product.category_id else ''
    if product.brand_id:
        brand = Brand.query.get(product.brand_id)
        product.brand = brand.name if brand else ''
    elif not product.brand:
        product.brand = ''

def assign_product_taxonomy_from_form(product):
    child_id = request.form.get('category_id', type=int)
    parent_id = request.form.get('category_parent_id', type=int)
    if child_id:
        product.category_id = child_id
    elif parent_id:
        product.category_id = parent_id
    else:
        legacy = request.form.get('category', '').strip()
        product.category_id = find_category_id_by_path(legacy) if legacy else None
    variant_raw = request.form.get('variant', product.variant or '')
    variant_labels = split_csv_field(variant_raw)
    if len(variant_labels) <= 1:
        brand_id = request.form.get('brand_id', type=int)
        if brand_id:
            product.brand_id = brand_id
        else:
            legacy_brand = request.form.get('brand', '').strip()
            if legacy_brand:
                brand = Brand.query.filter_by(name=legacy_brand).first()
                if not brand:
                    brand = Brand(name=legacy_brand)
                    db.session.add(brand)
                    db.session.flush()
                product.brand_id = brand.id
            else:
                product.brand_id = None
    sync_product_taxonomy_strings(product)

def ensure_category_from_path(path):
    path = (path or '').strip()
    if not path:
        return None
    existing = find_category_id_by_path(path)
    if existing:
        return existing
    if '>' in path:
        parts = [p.strip() for p in path.split('>') if p.strip()]
        parent_name, child_name = parts[0], parts[-1]
        parent = category_name_exists(parent_name, None) and Category.query.filter(
            Category.parent_id.is_(None), Category.name == parent_name
        ).first()
        if not parent:
            parent = Category(name=parent_name)
            db.session.add(parent)
            db.session.flush()
        child = Category.query.filter_by(parent_id=parent.id, name=child_name).first()
        if not child:
            child = Category(name=child_name, parent_id=parent.id)
            db.session.add(child)
            db.session.flush()
        return child.id
    if not category_name_exists(path, None):
        cat = Category(name=path)
        db.session.add(cat)
        db.session.flush()
        return cat.id
    return Category.query.filter(Category.parent_id.is_(None), Category.name == path).first().id

def ensure_brand_name(name, category_id=None):
    name = (name or '').strip()
    if not name:
        return None
    brand = Brand.query.filter_by(name=name).first()
    if not brand:
        brand = Brand(name=name, category_id=category_id)
        db.session.add(brand)
        db.session.flush()
    elif category_id and not brand.category_id:
        brand.category_id = category_id
    return brand.id

def brands_for_category(category_id):
    q = Brand.query.order_by(Brand.name)
    if not category_id:
        return q.all()
    cat = Category.query.get(category_id)
    if not cat:
        return q.all()
    scope_ids = {category_id}
    if cat.parent_id:
        scope_ids.add(cat.parent_id)
    else:
        scope_ids.update(c.id for c in Category.query.filter_by(parent_id=cat.id))
    return q.filter(or_(Brand.category_id.is_(None), Brand.category_id.in_(scope_ids))).all()

def brands_catalog_json():
    return [
        {'id': b.id, 'name': b.name, 'category_id': b.category_id}
        for b in Brand.query.order_by(Brand.name).all()
    ]

def categories_catalog_json():
    parents = parent_categories()
    children = {}
    for p in parents:
        children[str(p.id)] = [
            {'id': c.id, 'name': c.name, 'path': category_display_path(c.id)}
            for c in child_categories(p.id)
        ]
    return {
        'parents': [{'id': p.id, 'name': p.name} for p in parents],
        'children': children,
    }

app.jinja_env.globals['product_category_path'] = product_category_path
app.jinja_env.globals['product_brand_name'] = product_brand_name
app.jinja_env.globals['category_display_path'] = category_display_path

def ensure_product_columns():
    from sqlalchemy import inspect, text
    try:
        cols = {c['name'] for c in inspect(db.engine).get_columns('product')}
    except Exception:
        return
    with db.engine.begin() as conn:
        if 'image_path' not in cols:
            conn.execute(text('ALTER TABLE product ADD COLUMN image_path VARCHAR(255) DEFAULT ""'))
        if 'image_paths' not in cols:
            conn.execute(text('ALTER TABLE product ADD COLUMN image_paths TEXT DEFAULT ""'))
        if 'image_url' not in cols:
            conn.execute(text('ALTER TABLE product ADD COLUMN image_url VARCHAR(512) DEFAULT ""'))
        if 'is_active' not in cols:
            conn.execute(text('ALTER TABLE product ADD COLUMN is_active BOOLEAN DEFAULT 1'))
        if 'base_unit' not in cols:
            conn.execute(text("ALTER TABLE product ADD COLUMN base_unit VARCHAR(40) DEFAULT ''"))
        if 'purchase_unit' not in cols:
            conn.execute(text("ALTER TABLE product ADD COLUMN purchase_unit VARCHAR(40) DEFAULT ''"))
        if 'unit_conversion_enabled' not in cols:
            conn.execute(text('ALTER TABLE product ADD COLUMN unit_conversion_enabled BOOLEAN DEFAULT 0'))
        if 'conversion_factor' not in cols:
            conn.execute(text('ALTER TABLE product ADD COLUMN conversion_factor REAL DEFAULT 1'))
        if 'lot_unit' not in cols:
            conn.execute(text("ALTER TABLE product ADD COLUMN lot_unit VARCHAR(40) DEFAULT ''"))
        if 'lot_factor' not in cols:
            conn.execute(text('ALTER TABLE product ADD COLUMN lot_factor REAL DEFAULT 0'))
        if 'sale_unit_mode' not in cols:
            conn.execute(text("ALTER TABLE product ADD COLUMN sale_unit_mode VARCHAR(20) DEFAULT ''"))
        if 'variant_prices' not in cols:
            conn.execute(text("ALTER TABLE product ADD COLUMN variant_prices TEXT DEFAULT ''"))
        if 'variant_stocks' not in cols:
            conn.execute(text("ALTER TABLE product ADD COLUMN variant_stocks TEXT DEFAULT ''"))
        if 'variant_cost_prices' not in cols:
            conn.execute(text("ALTER TABLE product ADD COLUMN variant_cost_prices TEXT DEFAULT ''"))
        if 'variant_image_urls' not in cols:
            conn.execute(text("ALTER TABLE product ADD COLUMN variant_image_urls TEXT DEFAULT ''"))
        if 'variant_brand_ids' not in cols:
            conn.execute(text("ALTER TABLE product ADD COLUMN variant_brand_ids TEXT DEFAULT ''"))
        if 'variant_dealer_prices' not in cols:
            conn.execute(text("ALTER TABLE product ADD COLUMN variant_dealer_prices TEXT DEFAULT ''"))
        if 'variant_project_prices' not in cols:
            conn.execute(text("ALTER TABLE product ADD COLUMN variant_project_prices TEXT DEFAULT ''"))
    _migrate_product_image_paths()
    _migrate_product_unit_fields()

def _migrate_product_unit_fields():
    rows = Product.query.filter(
        or_(Product.base_unit == '', Product.base_unit.is_(None)),
    ).all()
    if not rows:
        return
    for p in rows:
        base = (p.unit or 'cái').strip() or 'cái'
        p.base_unit = base
        if not p.purchase_unit:
            p.purchase_unit = base
        if not p.conversion_factor or p.conversion_factor <= 0:
            p.conversion_factor = 1.0
    db.session.commit()

def product_base_unit(product):
    if not product:
        return 'cái'
    base = (getattr(product, 'base_unit', None) or '').strip()
    if base:
        return base
    return (product.unit or 'cái').strip() or 'cái'

def product_purchase_unit(product):
    if not product:
        return 'cái'
    if not product_has_unit_conversion(product):
        return product_base_unit(product)
    purchase = (getattr(product, 'purchase_unit', None) or '').strip()
    return purchase or product_base_unit(product)

def product_conversion_factor(product):
    if not product or not product_has_unit_conversion(product):
        return 1.0
    try:
        factor = float(getattr(product, 'conversion_factor', None) or 1.0)
    except (TypeError, ValueError):
        factor = 1.0
    return factor if factor > 0 else 1.0

def product_has_unit_conversion(product):
    """SP có quy đổi khi bật cờ hoặc (đơn vị nhập ≠ tồn và hệ số ≠ 1)."""
    if not product:
        return False
    try:
        factor = float(getattr(product, 'conversion_factor', None) or 1.0)
    except (TypeError, ValueError):
        factor = 1.0
    if factor <= 0:
        return False
    bu = product_base_unit(product)
    pu_raw = (getattr(product, 'purchase_unit', None) or '').strip() or bu
    units_differ = bool(pu_raw and bu and pu_raw.lower() != bu.lower())
    if not units_differ and factor == 1.0:
        return False
    if getattr(product, 'unit_conversion_enabled', False):
        return units_differ or factor != 1.0
    return units_differ and factor != 1.0

def purchase_qty_to_base(product, purchase_qty):
    return float(purchase_qty or 0) * product_conversion_factor(product)

def base_unit_cost_from_purchase_cost(product, purchase_unit_cost):
    """Giá vốn theo đơn vị tồn = giá theo đơn vị nhập ÷ hệ số quy đổi."""
    cost = int(purchase_unit_cost or 0)
    if cost <= 0 or not product_has_unit_conversion(product):
        return cost
    factor = product_conversion_factor(product)
    if factor <= 0:
        return cost
    return int(round(cost / factor))

def purchase_unit_cost_from_base_cost(product, base_unit_cost):
    """Giá theo đơn vị nhập (cuộn/hộp) từ giá vốn đơn vị tồn."""
    cost = int(base_unit_cost or 0)
    if cost <= 0 or not product_has_unit_conversion(product):
        return cost
    return int(round(cost * product_conversion_factor(product)))

def product_lot_unit(product):
    if not product:
        return ''
    return (getattr(product, 'lot_unit', None) or '').strip()

def product_lot_factor(product):
    if not product:
        return 0.0
    try:
        factor = float(getattr(product, 'lot_factor', None) or 0)
    except (TypeError, ValueError):
        factor = 0.0
    return factor if factor > 0 else 0.0

def product_has_lot_unit(product):
    return (
        product_has_unit_conversion(product)
        and bool(product_lot_unit(product))
        and product_lot_factor(product) > 1
    )

def lot_qty_to_purchase(product, lot_qty):
    return float(lot_qty or 0) * product_lot_factor(product)

def lot_qty_to_base(product, lot_qty):
    return purchase_qty_to_base(product, lot_qty_to_purchase(product, lot_qty))

def product_default_sale_unit_mode(product):
    """Đơn vị mặc định khi bán / hiển thị báo giá: base | purchase | lot."""
    if not product:
        return 'base'
    mode = (getattr(product, 'sale_unit_mode', None) or '').strip().lower()
    if mode == 'lot' and product_has_lot_unit(product):
        return 'lot'
    if mode == 'purchase' and product_has_unit_conversion(product):
        return 'purchase'
    if mode == 'base':
        return 'base'
    if product_has_unit_conversion(product):
        return 'purchase'
    return 'base'

def product_unit_label(product, unit_mode=None):
    if not product:
        return 'cái'
    mode = (unit_mode or product_default_sale_unit_mode(product)).strip().lower()
    if mode == 'lot' and product_has_lot_unit(product):
        return product_lot_unit(product)
    if mode == 'purchase' and product_has_unit_conversion(product):
        return product_purchase_unit(product)
    return product_base_unit(product)

def product_unit_mode_options(product):
    """Các chế độ đơn vị có thể chọn (lô / thùng / kg…)."""
    if not product:
        return [{'mode': 'base', 'label': 'cái'}]
    bu = product_base_unit(product)
    pu = (getattr(product, 'purchase_unit', None) or '').strip() or bu
    opts = [{'mode': 'base', 'label': bu}]
    conv_on = bool(getattr(product, 'unit_conversion_enabled', False)) or product_has_unit_conversion(product)
    if conv_on and pu and pu.lower() != bu.lower():
        opts.append({'mode': 'purchase', 'label': pu})
    if product_has_lot_unit(product):
        opts.append({'mode': 'lot', 'label': product_lot_unit(product)})
    return opts

def qty_to_base_unit(product, qty, unit_mode=None):
    _, base_qty, _ = resolve_supplier_intake_quantities(
        product, qty, unit_mode or product_default_sale_unit_mode(product),
    )
    return float(base_qty or 0)

def qty_from_base_unit(product, base_qty, unit_mode=None):
    base = float(base_qty or 0)
    if not product:
        return base
    mode = (unit_mode or product_default_sale_unit_mode(product)).strip().lower()
    if not product_has_unit_conversion(product):
        return base
    factor = product_conversion_factor(product)
    if mode == 'base':
        return base
    if mode == 'lot' and product_has_lot_unit(product):
        lf = product_lot_factor(product)
        if lf <= 0:
            return base
        purchase_qty = base / factor if factor > 0 else base
        return purchase_qty / lf
    if factor > 0:
        return base / factor
    return base

def quote_item_unit_mode(product, item):
    mode = (getattr(item, 'qty_unit_mode', None) or '').strip().lower()
    if mode in ('base', 'purchase', 'lot'):
        if mode == 'lot' and not product_has_lot_unit(product):
            mode = 'purchase' if product_has_unit_conversion(product) else 'base'
        elif mode == 'purchase' and not product_has_unit_conversion(product):
            mode = 'base'
        return mode
    unit = (item.unit or '').strip()
    if product_has_lot_unit(product) and unit == product_lot_unit(product):
        return 'lot'
    if product_has_unit_conversion(product) and unit == product_purchase_unit(product):
        return 'purchase'
    return 'base'

def quote_item_base_qty(product, item):
    return qty_to_base_unit(product, item.qty, quote_item_unit_mode(product, item))

def resolve_supplier_intake_quantities(product, entered_qty, unit_mode=''):
    """Trả về (purchase_qty, base_qty, mode) — nhập theo lô / đơn vị nhập / đơn vị tồn."""
    entered = float(entered_qty or 0)
    if not product_has_unit_conversion(product):
        return entered, entered, 'base'
    mode = (unit_mode or '').strip().lower()
    if product_has_lot_unit(product) and mode != 'base' and mode != 'purchase':
        mode = 'lot'
    if mode not in ('base', 'purchase', 'lot'):
        mode = 'purchase'
    factor = product_conversion_factor(product)
    if mode == 'base':
        base_qty = entered
        purchase_qty = (entered / factor) if factor > 0 else entered
        return purchase_qty, base_qty, 'base'
    if mode == 'lot' and product_has_lot_unit(product):
        purchase_qty = lot_qty_to_purchase(product, entered)
        base_qty = purchase_qty_to_base(product, purchase_qty)
        return purchase_qty, base_qty, 'lot'
    purchase_qty = entered
    base_qty = purchase_qty_to_base(product, purchase_qty)
    return purchase_qty, base_qty, 'purchase'

def product_cost_sale_unit_mode(product):
    """Đơn vị dùng cho giá nhập trên form / danh sách (theo sale_unit_mode)."""
    if not product:
        return 'base'
    mode = product_default_sale_unit_mode(product)
    if mode == 'purchase' and product_has_unit_conversion(product):
        return 'purchase'
    if mode == 'lot' and product_has_lot_unit(product):
        return 'lot'
    return 'base'

def product_cost_display_amount(product):
    """Giá nhập hiển thị: theo thùng nếu bán theo thùng, theo kg nếu bán theo kg."""
    base = int(product.cost_price or 0) if product else 0
    if not product or not product_has_unit_conversion(product):
        return base
    if product_cost_sale_unit_mode(product) == 'purchase':
        return purchase_unit_cost_from_base_cost(product, base)
    return base

def product_cost_unit_label(product):
    return product_unit_label(product, product_cost_sale_unit_mode(product))

def parse_product_cost_from_form(product):
    """Giá nhập từ form → lưu theo đơn vị tồn (kg)."""
    entered = parse_int(request.form.get('cost_price'), 0)
    if not product:
        return entered
    if not product_has_unit_conversion(product):
        return entered
    mode = (request.form.get('sale_unit_mode') or '').strip().lower()
    if mode not in ('base', 'purchase', 'lot'):
        mode = product_cost_sale_unit_mode(product)
    if mode == 'purchase':
        return base_unit_cost_from_purchase_cost(product, entered)
    return entered

def product_list_unit_display(product):
    """Nhãn đơn vị chính trên danh sách SP theo sale_unit_mode."""
    if not product:
        return {'primary': 'cái', 'secondary': '', 'conv': '', 'mode': 'base', 'mode_label': ''}
    mode = product_default_sale_unit_mode(product)
    primary = product_unit_label(product, mode)
    mode_labels = {'base': 'Theo ' + product_base_unit(product), 'purchase': 'Theo ' + product_purchase_unit(product), 'lot': 'Theo lô'}
    mode_label = mode_labels.get(mode, '')
    secondary = ''
    if product_has_unit_conversion(product):
        bu = product_base_unit(product)
        pu = product_purchase_unit(product)
        if mode == 'purchase':
            secondary = f'Tồn kho: {bu}'
        elif mode == 'base':
            secondary = f'Đóng gói: {pu}'
        elif mode == 'lot' and product_has_lot_unit(product):
            secondary = f'{bu} · {pu}'
    conv = product_unit_conversion_label(product) if product_has_unit_conversion(product) else ''
    return {
        'primary': primary,
        'secondary': secondary,
        'conv': conv,
        'mode': mode,
        'mode_label': mode_label,
    }

def product_qty_equivalent_text(product, base_qty):
    """Quy đổi số lượng tồn (đơn vị gốc) sang text đơn vị nhập/lô."""
    if not product_has_unit_conversion(product):
        return None
    stock = float(base_qty or 0)
    factor = product_conversion_factor(product)
    if factor <= 0:
        return None
    full_units = int(stock // factor)
    remainder = stock - (full_units * factor)
    pu = product_purchase_unit(product)
    bu = product_base_unit(product)
    parts = []
    if full_units > 0:
        parts.append(f'{format_qty_display(full_units)} {pu}')
    if remainder > 1e-6:
        parts.append(f'{format_qty_display(remainder)} {bu}')
    if not parts:
        return f'0 {bu}'
    return ' + '.join(parts)

def product_stock_equivalent_text(product):
    return product_qty_equivalent_text(product, product.stock)

def product_unit_conversion_label(product):
    if not product_has_unit_conversion(product):
        return '—'
    pu = product_purchase_unit(product)
    bu = product_base_unit(product)
    factor = product_conversion_factor(product)
    label = f'1 {pu} = {format_qty_display(factor)} {bu}'
    if product_has_lot_unit(product):
        lu = product_lot_unit(product)
        lf = product_lot_factor(product)
        label += f' · 1 {lu} = {format_qty_display(lf)} {pu} = {format_qty_display(lf * factor)} {bu}'
    return label

def format_qty_display(value, max_decimals=3):
    v = float(value or 0)
    if abs(v - round(v)) < 1e-6:
        return f'{int(round(v)):,}'.replace(',', '.')
    text = f'{v:.{max_decimals}f}'.rstrip('0').rstrip('.')
    if '.' in text:
        whole, frac = text.split('.', 1)
        whole_fmt = f'{int(whole):,}'.replace(',', '.') if whole else '0'
        return f'{whole_fmt},{frac}'
    return f'{int(v):,}'.replace(',', '.')

app.jinja_env.globals['product_base_unit'] = product_base_unit
app.jinja_env.globals['product_purchase_unit'] = product_purchase_unit
app.jinja_env.globals['product_conversion_factor'] = product_conversion_factor
app.jinja_env.globals['product_has_unit_conversion'] = product_has_unit_conversion
app.jinja_env.globals['product_has_lot_unit'] = product_has_lot_unit
app.jinja_env.globals['product_lot_unit'] = product_lot_unit
app.jinja_env.globals['product_unit_conversion_label'] = product_unit_conversion_label
app.jinja_env.globals['product_stock_equivalent_text'] = product_stock_equivalent_text
app.jinja_env.globals['product_qty_equivalent_text'] = product_qty_equivalent_text
app.jinja_env.globals['format_qty_display'] = format_qty_display
app.jinja_env.globals['base_unit_cost_from_purchase_cost'] = base_unit_cost_from_purchase_cost
app.jinja_env.globals['purchase_unit_cost_from_base_cost'] = purchase_unit_cost_from_base_cost
app.jinja_env.globals['product_default_sale_unit_mode'] = product_default_sale_unit_mode
app.jinja_env.globals['product_list_unit_display'] = product_list_unit_display
app.jinja_env.globals['product_cost_display_amount'] = product_cost_display_amount
app.jinja_env.globals['product_cost_unit_label'] = product_cost_unit_label
app.jinja_env.globals['product_cost_sale_unit_mode'] = product_cost_sale_unit_mode
app.jinja_env.globals['product_unit_label'] = product_unit_label
app.jinja_env.globals['product_unit_mode_options'] = product_unit_mode_options
app.jinja_env.globals['qty_from_base_unit'] = qty_from_base_unit
app.jinja_env.globals['qty_to_base_unit'] = qty_to_base_unit
app.jinja_env.globals['quote_item_unit_mode'] = quote_item_unit_mode

def ensure_quote_item_columns():
    from sqlalchemy import inspect, text
    db.create_all()
    try:
        cols = {c['name'] for c in inspect(db.engine).get_columns('quote_item')}
    except Exception:
        return
    with db.engine.begin() as conn:
        if 'qty_unit_mode' not in cols:
            conn.execute(text("ALTER TABLE quote_item ADD COLUMN qty_unit_mode VARCHAR(20) DEFAULT ''"))
        if 'variant_label' not in cols:
            conn.execute(text("ALTER TABLE quote_item ADD COLUMN variant_label VARCHAR(120) DEFAULT ''"))

def apply_product_unit_from_form(product):
    ensure_product_columns()
    enabled = request.form.get('unit_conversion_enabled') == '1'
    base = request.form.get('base_unit', '').strip() or request.form.get('unit', '').strip() or 'cái'
    purchase = request.form.get('purchase_unit', '').strip() or base
    factor = parse_qty(request.form.get('conversion_factor'), 1.0)
    if factor <= 0:
        factor = 1.0
    lot_enabled = request.form.get('lot_unit_enabled') == '1'
    lot_name = request.form.get('lot_unit', '').strip()
    lot_factor = parse_qty(request.form.get('lot_factor'), 0)
    product.unit_conversion_enabled = enabled
    product.base_unit = base
    product.unit = base
    if enabled:
        product.purchase_unit = purchase or base
        product.conversion_factor = factor
        if lot_enabled and lot_name and lot_factor > 1:
            product.lot_unit = lot_name
            product.lot_factor = lot_factor
        else:
            product.lot_unit = ''
            product.lot_factor = 0
    else:
        product.purchase_unit = base
        product.conversion_factor = 1.0
        product.lot_unit = ''
        product.lot_factor = 0
    sale_mode = (request.form.get('sale_unit_mode') or '').strip().lower()
    if sale_mode not in ('base', 'purchase', 'lot'):
        sale_mode = 'purchase' if enabled else 'base'
    if sale_mode == 'lot' and not (enabled and lot_enabled and lot_name and lot_factor > 1):
        sale_mode = 'purchase' if enabled else 'base'
    if sale_mode == 'purchase' and not enabled:
        sale_mode = 'base'
    product.sale_unit_mode = sale_mode

def _migrate_product_image_paths():
    rows = Product.query.filter(
        Product.image_path != '',
        or_(Product.image_paths == '', Product.image_paths.is_(None)),
    ).all()
    if not rows:
        return
    for p in rows:
        sync_product_images(p, [p.image_path])
    db.session.commit()

def ensure_product_image_column():
    ensure_product_columns()

def product_image_paths(product):
    if not product:
        return []
    paths = []
    raw = getattr(product, 'image_paths', '') or ''
    if raw:
        try:
            parsed = json.loads(raw)
            if isinstance(parsed, list):
                paths = [p for p in parsed if p]
        except json.JSONDecodeError:
            paths = []
    if not paths and getattr(product, 'image_path', ''):
        paths = [product.image_path]
    return paths[:MAX_PRODUCT_IMAGES]

def is_external_image_url(url):
    u = (url or '').strip()
    return u.startswith('http://') or u.startswith('https://')

def product_external_image_url(product):
    if not product:
        return ''
    return (getattr(product, 'image_url', None) or '').strip() if is_external_image_url(
        getattr(product, 'image_url', None) or ''
    ) else ''

def parse_variant_image_urls_raw(value):
    if value is None or value == '':
        return []
    s = str(value).strip()
    if s.startswith('['):
        try:
            data = json.loads(s)
            if isinstance(data, list):
                return [str(u).strip() for u in data]
        except Exception:
            pass
    if '||' in s:
        return [u.strip() for u in s.split('||')]
    if s:
        return [s]
    return []

def parse_variant_image_urls_list(value, expected_len=None):
    """Parse URL ảnh biến thể — giữ đủ vị trí, cho phép ô trống."""
    urls = parse_variant_image_urls_raw(value)
    if expected_len is not None:
        while len(urls) < expected_len:
            urls.append('')
        urls = urls[:expected_len]
    return urls

def product_stored_variant_image_urls(product):
    raw = (getattr(product, 'variant_image_urls', None) or '').strip()
    if not raw:
        return []
    return parse_variant_image_urls_raw(raw)

def variant_image_urls_to_storage(urls):
    normalized = []
    for u in urls or []:
        u = str(u or '').strip()
        if u and is_external_image_url(u):
            normalized.append(u)
        else:
            normalized.append('')
    if not any(normalized):
        return ''
    return json.dumps(normalized, ensure_ascii=False)

def product_image_urls(product):
    variant_urls = [u for u in product_stored_variant_image_urls(product) if u]
    if variant_urls:
        return variant_urls
    ext = product_external_image_url(product)
    if ext:
        return [ext]
    return [url_for('static', filename=path) for path in product_image_paths(product)]

def sync_product_images(product, paths):
    clean = [p for p in paths if p][:MAX_PRODUCT_IMAGES]
    product.image_paths = json.dumps(clean, ensure_ascii=False)
    product.image_path = clean[0] if clean else ''

def product_image_url(product):
    ext = product_external_image_url(product)
    if ext:
        return ext
    paths = product_image_paths(product)
    if paths:
        return url_for('static', filename=paths[0])
    return None

def product_is_active(product):
    return bool(getattr(product, 'is_active', True))

def product_spec_short(product):
    parts = [x for x in (product.model or '', product.variant or '') if x]
    return ' - '.join(parts) if parts else ''

def split_csv_field(value):
    if value is None:
        return []
    return [p.strip() for p in str(value).split(',') if p.strip()]

def split_variant_labels(value):
    """Tách tên biến thể — form ghép bằng «Tên1, Tên2» (dấu phẩy + space)."""
    if value is None or not str(value).strip():
        return []
    s = str(value).strip()
    if ', ' in s:
        parts = [p.strip() for p in s.split(', ') if p.strip()]
        if parts:
            return parts
    return split_csv_field(s)

def parse_price_csv(value):
    """Parse một hoặc nhiều giá VNĐ. Nhiều biến thể: «1000000, 500000»; một giá: số thuần hoặc 1.000.000."""
    if value is None or value == '':
        return []
    s = str(value).strip()
    if ', ' in s:
        return [parse_int(part, 0) for part in s.split(', ')]
    if ',' not in s:
        return [parse_int(s, 0)]
    parts = split_csv_field(s)
    if len(parts) > 1 and all(p.isdigit() for p in parts) and all(len(p) == 3 for p in parts[1:]):
        return [parse_int(''.join(parts), 0)]
    return [parse_int(part, 0) for part in parts]

def parse_qty_csv(value):
    """Parse một hoặc nhiều số lượng, phân tách bằng dấu phẩy (theo biến thể)."""
    if value is None or value == '':
        return []
    s = str(value).strip()
    if ',' not in s:
        return [parse_qty(s, 0)]
    return [parse_qty(part, 0) for part in split_csv_field(s)]

def _quote_variant_prices_for_labels(product, labels):
    """Giá lẻ theo từng biến thể — đủ N giá, hoặc một giá áp cho tất cả."""
    if not labels:
        return []
    stored = (getattr(product, 'variant_prices', None) or '').strip()
    prices = parse_price_csv(stored) if stored else []
    if len(prices) == len(labels):
        return prices
    if len(prices) == 1:
        return prices * len(labels)
    if product.retail_price:
        return [product.retail_price] * len(labels)
    if prices:
        last = prices[-1]
        return prices + [last] * (len(labels) - len(prices))
    return [0] * len(labels)

def _product_variant_price_pairs(product):
    labels = split_csv_field(getattr(product, 'variant', '') or '')
    if len(labels) <= 1:
        return []
    prices = _quote_variant_prices_for_labels(product, labels)
    if len(prices) != len(labels):
        return []
    return [{'label': label, 'price': price} for label, price in zip(labels, prices)]

def product_variant_price_rows(product):
    return _product_variant_price_pairs(product)

def product_variant_rows(product):
    """Ghép biến thể với giá, tồn, giá nhập, ảnh — khi đủ dữ liệu khớp."""
    price_pairs = _product_variant_price_pairs(product)
    if not price_pairs:
        return []
    labels = [r['label'] for r in price_pairs]
    stocks_stored = (getattr(product, 'variant_stocks', None) or '').strip()
    if not stocks_stored:
        return []
    stocks = parse_qty_csv(stocks_stored)
    if len(stocks) != len(labels):
        return []
    costs = parse_price_csv((getattr(product, 'variant_cost_prices', None) or '').strip())
    dealers = parse_price_csv((getattr(product, 'variant_dealer_prices', None) or '').strip())
    projects = parse_price_csv((getattr(product, 'variant_project_prices', None) or '').strip())
    images = product_stored_variant_image_urls(product)
    brand_ids = product_stored_variant_brand_ids(product)
    rows = []
    for i, r in enumerate(price_pairs):
        row = {'label': r['label'], 'price': r['price'], 'qty': stocks[i]}
        if len(costs) == len(labels):
            row['cost'] = costs[i]
        if len(dealers) == len(labels):
            row['dealer_price'] = dealers[i]
        if len(projects) == len(labels):
            row['project_price'] = projects[i]
        if len(images) == len(labels):
            row['image_url'] = images[i]
        if len(brand_ids) == len(labels) and brand_ids[i]:
            row['brand_id'] = brand_ids[i]
            row['brand_name'] = brand_name_by_id(brand_ids[i])
        rows.append(row)
    return rows

def product_list_preview_variants(product):
    """Dữ liệu biến thể cho hover preview trên danh sách sản phẩm."""
    rows = product_variant_rows(product) or product_variant_price_rows(product)
    items = []
    for r in rows:
        item = {
            'label': r['label'],
            'price': r.get('price', 0),
            'qty': r.get('qty'),
            'brand_name': r.get('brand_name') or '',
            'image_url': r.get('image_url') or '',
        }
        if r.get('cost') is not None:
            item['cost'] = r['cost']
        items.append(item)
    return items

def quote_catalog_variant_rows(product):
    """Biến thể cho catalog báo giá — tách theo cột biến thể, giá lẻ nếu có."""
    labels = split_csv_field(getattr(product, 'variant', '') or '')
    if len(labels) <= 1:
        return []
    prices = _quote_variant_prices_for_labels(product, labels)
    price_pairs = [{'label': label, 'price': price} for label, price in zip(labels, prices)]
    full_by_label = {r['label']: r for r in product_variant_rows(product)}
    stocks_raw = (getattr(product, 'variant_stocks', None) or '').strip()
    stocks = parse_qty_csv(stocks_raw) if stocks_raw else []
    has_stocks = len(stocks) == len(labels)
    images = product_stored_variant_image_urls(product)
    brand_ids = product_stored_variant_brand_ids(product)
    rows = []
    for i, pp in enumerate(price_pairs):
        label = pp['label']
        if label in full_by_label:
            rows.append(full_by_label[label])
            continue
        row = {'label': label, 'price': pp['price'], 'qty': stocks[i] if has_stocks else 0}
        if len(images) == len(labels):
            row['image_url'] = images[i]
        if len(brand_ids) == len(labels) and brand_ids[i]:
            row['brand_id'] = brand_ids[i]
            row['brand_name'] = brand_name_by_id(brand_ids[i])
        rows.append(row)
    return rows

def apply_product_variant_retail_prices(product, retail_raw):
    """Lưu giá lẻ — mỗi biến thể một giá (có thể trùng nhau)."""
    prices = parse_price_csv(retail_raw)
    labels = split_variant_labels(product.variant or '')
    if len(labels) > 1:
        if len(prices) != len(labels):
            return False, (
                f'Cần nhập đủ {len(labels)} giá lẻ cho {len(labels)} biến thể, '
                f'hiện có {len(prices)} giá.'
            )
        product.variant_prices = ', '.join(str(p) for p in prices)
        product.retail_price = prices[0]
        return True, ''
    if len(prices) > 1:
        return False, (
            'Bạn đang nhập nhiều giá lẻ nhưng sản phẩm chỉ có một biến thể. '
            'Bấm «Thêm biến thể» để thêm dòng, hoặc chỉ nhập một giá lẻ.'
        )
    product.retail_price = prices[0] if prices else 0
    product.variant_prices = ''
    return True, ''

def product_retail_price_input_value(product):
    if not product:
        return ''
    stored = (getattr(product, 'variant_prices', None) or '').strip()
    if stored:
        return ', '.join(money_plain(p) for p in parse_price_csv(stored))
    if product.retail_price:
        return money_plain(product.retail_price)
    return ''

def apply_product_variant_stocks(product, stock_raw):
    """Lưu tồn kho — mỗi biến thể một số lượng riêng; tổng tồn = tổng các biến thể."""
    qtys = parse_qty_csv(stock_raw)
    labels = split_variant_labels(product.variant or '')
    if len(labels) > 1:
        if len(qtys) != len(labels):
            return False, (
                f'Cần nhập đủ {len(labels)} tồn kho cho {len(labels)} biến thể, '
                f'hiện có {len(qtys)} số lượng.'
            )
        product.variant_stocks = ', '.join(
            str(int(q)) if float(q) == int(q) else str(q) for q in qtys
        )
        product.stock = sum(qtys)
        return True, ''
    if len(qtys) > 1:
        return False, (
            'Bạn đang nhập nhiều tồn kho nhưng sản phẩm chỉ có một biến thể. '
            'Bấm «Thêm biến thể» để thêm dòng, hoặc chỉ nhập một số lượng.'
        )
    product.stock = qtys[0] if qtys else 0
    product.variant_stocks = ''
    return True, ''

def product_variant_stock_input_value(product):
    if not product:
        return '0'
    stored = (getattr(product, 'variant_stocks', None) or '').strip()
    if stored:
        return stored
    if product.stock is not None:
        val = float(product.stock or 0)
        return str(int(val)) if val == int(val) else str(val)
    return '0'

def _apply_product_variant_price_csv(product, raw, *, variant_attr, product_attr, label):
    prices = parse_price_csv(raw)
    labels = split_variant_labels(product.variant or '')
    if len(labels) > 1:
        if len(prices) != len(labels):
            return False, (
                f'Cần nhập đủ {len(labels)} {label} cho {len(labels)} biến thể, '
                f'hiện có {len(prices)} giá.'
            )
        setattr(product, variant_attr, ', '.join(str(p) for p in prices))
        setattr(product, product_attr, prices[0])
        return True, ''
    if len(prices) > 1:
        return False, (
            f'{label} không hợp lệ: hệ thống đọc được nhiều giá trong một ô. '
            'Nếu sản phẩm chỉ có một biến thể thì chỉ nhập một giá; '
            'nếu cần nhiều giá theo biến thể thì bấm «Thêm biến thể» để thêm từng dòng.'
        )
    setattr(product, product_attr, prices[0] if prices else 0)
    setattr(product, variant_attr, '')
    return True, ''

def _product_variant_price_input_value(product, variant_attr, product_attr):
    if not product:
        return ''
    stored = (getattr(product, variant_attr, None) or '').strip()
    if stored:
        return ', '.join(money_plain(p) for p in parse_price_csv(stored))
    val = getattr(product, product_attr, None)
    if val:
        return money_plain(val)
    return ''

def apply_product_variant_cost_prices(product, cost_raw):
    return _apply_product_variant_price_csv(
        product, cost_raw,
        variant_attr='variant_cost_prices', product_attr='cost_price', label='giá nhập',
    )

def product_variant_cost_input_value(product):
    return _product_variant_price_input_value(product, 'variant_cost_prices', 'cost_price')

def apply_product_variant_dealer_prices(product, raw):
    return _apply_product_variant_price_csv(
        product, raw,
        variant_attr='variant_dealer_prices', product_attr='dealer_price', label='giá sỉ',
    )

def product_variant_dealer_input_value(product):
    return _product_variant_price_input_value(product, 'variant_dealer_prices', 'dealer_price')

def apply_product_variant_project_prices(product, raw):
    return _apply_product_variant_price_csv(
        product, raw,
        variant_attr='variant_project_prices', product_attr='project_price', label='giá công trình',
    )

def product_variant_project_input_value(product):
    return _product_variant_price_input_value(product, 'variant_project_prices', 'project_price')

def product_edit_form_snapshot(product):
    """Ảnh chụp form chỉnh sửa — dùng xem trước thay đổi trước khi lưu."""
    parent_id = ''
    if product.category_id:
        cat = Category.query.get(product.category_id)
        if cat and cat.parent_id:
            parent_id = str(cat.parent_id)
    return {
        'name': (product.name or '').strip(),
        'sku': (product.sku or '').strip(),
        'model': (product.model or '').strip(),
        'warranty': (product.warranty or '').strip(),
        'low_stock': str(product.low_stock or 0),
        'category_parent_id': parent_id,
        'category_id': str(product.category_id or ''),
        'category_label': product_category_path(product) or '—',
        'brand_id': str(product.brand_id or ''),
        'brand_label': product_brand_name(product) or '—',
        'variant': (product.variant or '').strip(),
        'image_url': (product_external_image_url(product) or '').strip(),
        'base_unit': product_base_unit(product),
        'purchase_unit': product_purchase_unit(product),
        'unit_conversion_enabled': '1' if product_has_unit_conversion(product) else '0',
        'conversion_factor': str(product_conversion_factor(product) or 1),
        'lot_unit_enabled': '1' if product_has_lot_unit(product) else '0',
        'lot_unit': (product_lot_unit(product) or '').strip(),
        'lot_factor': str(product_lot_factor(product) or 0),
        'sale_unit_mode': product_default_sale_unit_mode(product) or 'base',
        'stock': str(product.stock or 0),
        'cost_price': product_variant_cost_input_value(product),
        'retail_price': product_retail_price_input_value(product),
        'dealer_price': product_variant_dealer_input_value(product),
        'project_price': product_variant_project_input_value(product),
        'variant_image_urls': product_variant_image_urls_input_value(product),
        'variant_brand_ids': product_variant_brand_ids_input_value(product),
    }

def apply_product_variant_image_urls(product, raw):
    labels = split_csv_field(product.variant or '')
    if len(labels) > 1:
        urls = parse_variant_image_urls_list(raw, len(labels))
        valid = []
        for u in urls:
            u = u.strip()
            if not u:
                valid.append('')
                continue
            if not is_external_image_url(u):
                return False, f'URL ảnh biến thể không hợp lệ: {u[:80]}'
            valid.append(u)
        product.variant_image_urls = variant_image_urls_to_storage(valid)
        product.image_url = next((u for u in valid if u), '')
        return True, ''
    product.variant_image_urls = ''
    return True, ''

def product_variant_image_urls_input_value(product):
    labels = split_csv_field(getattr(product, 'variant', '') or '')
    if len(labels) > 1:
        raw = (getattr(product, 'variant_image_urls', None) or '').strip()
        urls = parse_variant_image_urls_list(raw if raw else '[]', len(labels))
        return json.dumps(urls, ensure_ascii=False)
    urls = product_stored_variant_image_urls(product)
    return json.dumps(urls, ensure_ascii=False) if urls else ''

def parse_variant_brand_ids_raw(value):
    if value is None or not str(value).strip():
        return []
    ids = []
    for piece in str(value).split(','):
        piece = piece.strip()
        if not piece:
            ids.append(0)
            continue
        try:
            ids.append(int(piece))
        except (TypeError, ValueError):
            ids.append(0)
    return ids

def product_stored_variant_brand_ids(product):
    raw = (getattr(product, 'variant_brand_ids', None) or '').strip()
    if not raw:
        return []
    return parse_variant_brand_ids_raw(raw)

def variant_brand_ids_to_storage(ids):
    return ', '.join(str(int(i)) for i in ids)

def brand_name_by_id(brand_id):
    if not brand_id:
        return ''
    brand = Brand.query.get(int(brand_id))
    return brand.name if brand else ''

def apply_product_variant_brand_ids(product, raw):
    labels = split_csv_field(product.variant or '')
    if len(labels) > 1:
        ids = parse_variant_brand_ids_raw(raw)
        while len(ids) < len(labels):
            ids.append(0)
        ids = ids[:len(labels)]
        for i, bid in enumerate(ids):
            if bid and not Brand.query.get(bid):
                return False, f'Thương hiệu không hợp lệ cho biến thể "{labels[i]}".'
        product.variant_brand_ids = variant_brand_ids_to_storage(ids)
        first_brand = next((bid for bid in ids if bid), None)
        product.brand_id = first_brand
        sync_product_taxonomy_strings(product)
        return True, ''
    product.variant_brand_ids = ''
    return True, ''

def product_variant_brand_ids_input_value(product):
    ids = product_stored_variant_brand_ids(product)
    return ', '.join(str(i) for i in ids) if ids else ''

def product_brands_display(product):
    rows = product_variant_rows(product)
    if rows:
        names = []
        for row in rows:
            name = row.get('brand_name') or ''
            if name and name not in names:
                names.append(name)
        if names:
            return ', '.join(names)
    return product_brand_name(product) or ''

def product_variant_brands_export_value(product):
    ids = product_stored_variant_brand_ids(product)
    labels = split_csv_field(product.variant or '')
    if len(labels) > 1 and len(ids) == len(labels):
        names = [brand_name_by_id(bid) for bid in ids if bid]
        if names:
            return ', '.join(names)
    return product_brand_name(product)

def _import_variant_brand_ids_from_data(data):
    labels = split_csv_field(data.get('variant') or '')
    brand_raw = (data.get('brand') or '').strip()
    if len(labels) <= 1 or not brand_raw:
        return ''
    if ',' in brand_raw:
        brand_names = [x.strip() for x in brand_raw.split(',')]
    else:
        brand_names = [brand_raw] * len(labels)
    if len(brand_names) != len(labels):
        return ''
    category_path = resolve_import_category_path(data)
    category_id = ensure_category_from_path(category_path) if category_path else None
    ids = []
    for name in brand_names:
        if not name:
            return ''
        ids.append(ensure_brand_name(name, category_id))
    return variant_brand_ids_to_storage(ids)

def _import_variant_prices_from_data(data):
    labels = split_csv_field(data.get('variant') or '')
    retail_prices = parse_price_csv(data.get('retail_price_raw', data.get('retail_price')))
    if len(labels) > 1:
        if len(retail_prices) == len(labels):
            return ', '.join(str(p) for p in retail_prices)
        return ''
    if len(retail_prices) > 1:
        return ', '.join(str(p) for p in retail_prices)
    return ''

def _import_variant_stocks_from_data(data):
    labels = split_csv_field(data.get('variant') or '')
    stocks = parse_qty_csv(data.get('stock_raw', data.get('stock')))
    if len(labels) > 1:
        if len(stocks) == len(labels):
            return ', '.join(
                str(int(q)) if float(q) == int(q) else str(q) for q in stocks
            )
        return ''
    return ''

def _import_variant_cost_prices_from_data(data):
    labels = split_csv_field(data.get('variant') or '')
    costs = parse_price_csv(data.get('cost_price_raw', data.get('cost_price')))
    if len(labels) > 1:
        if len(costs) == len(labels):
            return ', '.join(str(p) for p in costs)
        return ''
    return ''

def _import_variant_image_urls_from_data(data):
    labels = split_csv_field(data.get('variant') or '')
    urls = parse_variant_image_urls_raw(data.get('image_url_raw', data.get('image_url')))
    if len(labels) > 1:
        if len(urls) == len(labels) and all(is_external_image_url(u) for u in urls):
            return variant_image_urls_to_storage(urls)
        return ''
    return ''

def _quote_variant_inventory(qty, low_stock):
    qty = float(qty or 0)
    low = float(low_stock or 5)
    if qty <= 0:
        return 'out', 'Hết hàng'
    if qty <= low:
        return 'low', 'Sắp hết'
    return 'ok', 'Còn nhiều'

def _quote_catalog_base_entry(p):
    status_key, status_label = product_inventory_level(p)
    return {
        'id': p.id,
        'sku': p.sku,
        'name': p.name,
        'unit': product_unit_label(p, product_default_sale_unit_mode(p)),
        'base_unit': product_base_unit(p),
        'purchase_unit': product_purchase_unit(p),
        'unit_conversion_enabled': bool(getattr(p, 'unit_conversion_enabled', False))
            or product_has_unit_conversion(p),
        'conversion_factor': product_conversion_factor(p),
        'lot_unit': product_lot_unit(p),
        'lot_factor': product_lot_factor(p),
        'has_lot_unit': product_has_lot_unit(p),
        'sale_unit_mode': product_default_sale_unit_mode(p),
        'stock': p.stock or 0,
        'low_stock': p.low_stock or 5,
        'stock_status': status_key,
        'stock_label': status_label,
        'retail_price': p.retail_price or 0,
        'dealer_price': p.dealer_price or 0,
        'project_price': p.project_price or 0,
        'image_url': product_image_url(p) or '',
    }

def quote_product_catalog(products):
    items = []
    for p in products:
        if not product_is_active(p):
            continue
        base = _quote_catalog_base_entry(p)
        variant_rows = quote_catalog_variant_rows(p)
        if variant_rows:
            for vr in variant_rows:
                vl = vr['label']
                st_key, st_label = _quote_variant_inventory(vr['qty'], p.low_stock)
                brand_name = vr.get('brand_name') or product_brand_name(p)
                spec_parts = [x for x in (brand_name, p.model or '', vl) if x]
                spec = ' · '.join(spec_parts)
                image = vr.get('image_url') or base['image_url']
                items.append({
                    **base,
                    'variant_label': vl,
                    'catalog_key': f'{p.id}:{vl}',
                    'display_name': vl,
                    'brand_name': brand_name,
                    'spec': spec,
                    'retail_price': vr['price'],
                    'dealer_price': vr.get('dealer_price', base['dealer_price']),
                    'project_price': vr.get('project_price', base['project_price']),
                    'stock': vr['qty'],
                    'stock_status': st_key,
                    'stock_label': st_label,
                    'image_url': image,
                    'label': f'{p.sku} - {vl}',
                })
        else:
            items.append({
                **base,
                'variant_label': '',
                'catalog_key': str(p.id),
                'display_name': p.name,
                'spec': product_spec_short(p),
                'label': f'{p.sku} - {p.name}',
            })
    return items

app.jinja_env.globals['product_variant_rows'] = product_variant_rows
app.jinja_env.globals['product_variant_label_list'] = lambda p: split_variant_labels(getattr(p, 'variant', '') or '')
app.jinja_env.globals['product_list_preview_variants'] = product_list_preview_variants
app.jinja_env.globals['product_variant_price_rows'] = product_variant_price_rows
app.jinja_env.globals['product_retail_price_input_value'] = product_retail_price_input_value
app.jinja_env.globals['product_variant_stock_input_value'] = product_variant_stock_input_value
app.jinja_env.globals['product_variant_cost_input_value'] = product_variant_cost_input_value
app.jinja_env.globals['product_variant_dealer_input_value'] = product_variant_dealer_input_value
app.jinja_env.globals['product_variant_project_input_value'] = product_variant_project_input_value
app.jinja_env.globals['product_variant_image_urls_input_value'] = product_variant_image_urls_input_value
app.jinja_env.globals['product_variant_brand_ids_input_value'] = product_variant_brand_ids_input_value
app.jinja_env.globals['product_brands_display'] = product_brands_display
app.jinja_env.globals['product_image_url'] = product_image_url
app.jinja_env.globals['product_image_urls'] = product_image_urls
app.jinja_env.globals['product_external_image_url'] = product_external_image_url
app.jinja_env.globals['product_image_paths'] = product_image_paths
app.jinja_env.globals['max_product_images'] = MAX_PRODUCT_IMAGES
app.jinja_env.globals['child_categories'] = child_categories

def delete_product_image_file(image_path):
    if not image_path:
        return
    file_path = BASE_DIR / 'static' / image_path
    if file_path.is_file():
        file_path.unlink(missing_ok=True)

def delete_all_product_image_files(product):
    for path in product_image_paths(product):
        delete_product_image_file(path)
    sync_product_images(product, [])

def _save_product_image_file(product, file_storage):
    if not file_storage or not file_storage.filename:
        return None
    ext = Path(secure_filename(file_storage.filename)).suffix.lower()
    if ext not in ALLOWED_IMAGE_EXTENSIONS:
        flash('Ảnh phải là JPG, PNG, WEBP hoặc GIF', 'warning')
        return False
    data = file_storage.read()
    file_storage.seek(0)
    if len(data) > MAX_PRODUCT_IMAGE_BYTES:
        flash('Ảnh tối đa 5MB', 'warning')
        return False
    PRODUCT_UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
    pid = product.id or 0
    filename = f'{pid}_{uuid.uuid4().hex[:10]}{ext}'
    rel_path = f'uploads/products/{filename}'
    file_storage.save(PRODUCT_UPLOAD_DIR / filename)
    return rel_path

def save_product_images(product, file_storages):
    uploads = [f for f in (file_storages or []) if f and f.filename]
    if not uploads:
        return True
    current = product_image_paths(product)
    slots = MAX_PRODUCT_IMAGES - len(current)
    if slots <= 0:
        flash(f'Mỗi sản phẩm tối đa {MAX_PRODUCT_IMAGES} ảnh', 'warning')
        return False
    if len(uploads) > slots:
        flash(f'Chỉ thêm được {slots} ảnh nữa (tối đa {MAX_PRODUCT_IMAGES} ảnh)', 'warning')
        uploads = uploads[:slots]
    for file_storage in uploads:
        rel_path = _save_product_image_file(product, file_storage)
        if rel_path is False:
            return False
        if rel_path:
            current.append(rel_path)
    sync_product_images(product, current)
    return True

def save_product_image(product, file_storage):
    return save_product_images(product, [file_storage] if file_storage and file_storage.filename else [])

def remove_product_images(product, indices):
    paths = product_image_paths(product)
    if not paths or not indices:
        return
    for idx in sorted({int(i) for i in indices if str(i).isdigit()}, reverse=True):
        if 0 <= idx < len(paths):
            delete_product_image_file(paths.pop(idx))
    sync_product_images(product, paths)

QUOTE_STATUSES = ['Nháp', 'Mới tạo', 'Đã chốt', 'Đã hủy']
QUOTE_FILTER_STATUSES = ['Mới tạo', 'Đã chốt', 'Đã hủy']
QUOTE_STATUS_KEYS = {
    'Nháp': 'draft',
    'Mới tạo': 'new',
    'Đã chốt': 'won',
    'Đã hủy': 'cancelled',
}
LEGACY_QUOTE_STATUS_MAP = {
    'Đã gửi': 'Mới tạo',
    'Chờ duyệt': 'Mới tạo',
    'Hết hiệu lực': 'Mới tạo',
}
DEFAULT_COMPANY_PROFILE = {
    'name': 'CÔNG TY TNHH ĐẠI PHÁT',
    'short_name': 'ĐẠI PHÁT',
    'address': 'Tòa nhà IC, 82 Duy Tân, Cầu Giấy, Hà Nội',
    'phone': '024 407 781 999 - 0912 335 694',
    'email': 'contact@daiphat.vn',
    'tax_code': '0123456789',
    'bank_account': '18900217',
    'bank_name': 'Ngân hàng ACB - CN Hà Thành',
    'director_name': '',
    'director_title': 'Giám đốc',
}

def company_bank_display(row):
    parts = []
    if row.bank_account:
        parts.append(f'Số TK: {row.bank_account}')
    if row.bank_name:
        parts.append(row.bank_name)
    return ' - '.join(parts)

def company_profile_to_dict(row):
    return {
        'id': row.id,
        'name': row.name or '',
        'short_name': row.short_name or '',
        'address': row.address or '',
        'phone': row.phone or '',
        'email': row.email or '',
        'tax_code': row.tax_code or '',
        'bank': company_bank_display(row),
        'bank_account': row.bank_account or '',
        'bank_name': row.bank_name or '',
        'director': row.director_title or 'Giám đốc',
        'director_name': row.director_name or '',
        'logo_url': url_for('static', filename=row.logo_path) if row.logo_path else None,
        'is_default': bool(row.is_default),
    }

def ensure_company_columns():
    from sqlalchemy import inspect, text
    db.create_all()
    try:
        cols = {c['name'] for c in inspect(db.engine).get_columns('company_profile')}
    except Exception:
        return
    with db.engine.begin() as conn:
        if 'is_default' not in cols:
            conn.execute(text('ALTER TABLE company_profile ADD COLUMN is_default BOOLEAN DEFAULT 0'))
        if 'created_at' not in cols:
            conn.execute(text('ALTER TABLE company_profile ADD COLUMN created_at DATETIME'))
    if CompanyProfile.query.count() and not CompanyProfile.query.filter_by(is_default=True).first():
        first = CompanyProfile.query.order_by(CompanyProfile.id.asc()).first()
        if first:
            first.is_default = True
            db.session.commit()

def _company_from_defaults(**overrides):
    d = {**DEFAULT_COMPANY_PROFILE, **overrides}
    return CompanyProfile(
        name=d['name'],
        short_name=d.get('short_name', ''),
        tax_code=d.get('tax_code', ''),
        address=d.get('address', ''),
        phone=d.get('phone', ''),
        email=d.get('email', ''),
        bank_account=d.get('bank_account', ''),
        bank_name=d.get('bank_name', ''),
        director_name=d.get('director_name', ''),
        director_title=d.get('director_title', 'Giám đốc'),
    )

def ensure_company_profiles():
    ensure_company_columns()
    if CompanyProfile.query.count() == 0:
        row = _company_from_defaults(is_default=True)
        db.session.add(row)
        db.session.commit()
    return CompanyProfile.query.order_by(CompanyProfile.is_default.desc(), CompanyProfile.name).all()

def list_company_profiles():
    ensure_company_profiles()
    return CompanyProfile.query.order_by(CompanyProfile.is_default.desc(), CompanyProfile.name).all()

def get_company_row(company_id=None):
    ensure_company_profiles()
    if company_id:
        return CompanyProfile.query.get(company_id)
    if has_request_context():
        cid = session.get(ACTIVE_COMPANY_SESSION_KEY)
        if cid:
            row = CompanyProfile.query.get(cid)
            if row:
                return row
    row = CompanyProfile.query.filter_by(is_default=True).first()
    if row:
        return row
    return CompanyProfile.query.order_by(CompanyProfile.id.asc()).first()

def get_company_profile(company_id=None):
    row = get_company_row(company_id)
    if not row:
        ensure_company_profiles()
        row = get_company_row()
    return company_profile_to_dict(row) if row else {}

def resolve_company_id_from_form():
    cid = request.form.get('company_id', type=int)
    if cid and CompanyProfile.query.get(cid):
        return cid
    row = get_company_row()
    return row.id if row else None

def apply_company_form(row, form, files):
    row.name = form.get('name', '').strip()
    if not row.name:
        raise ValueError('Vui lòng nhập tên công ty')
    row.short_name = form.get('short_name', '').strip()
    row.tax_code = form.get('tax_code', '').strip()
    row.address = form.get('address', '').strip()
    row.phone = form.get('phone', '').strip()
    row.email = form.get('email', '').strip()
    row.bank_account = form.get('bank_account', '').strip()
    row.bank_name = form.get('bank_name', '').strip()
    row.director_name = form.get('director_name', '').strip()
    row.director_title = form.get('director_title', '').strip() or 'Giám đốc'
    if form.get('remove_logo') == '1' and row.logo_path:
        old = BASE_DIR / 'static' / row.logo_path
        if old.exists():
            old.unlink(missing_ok=True)
        row.logo_path = ''
    if not save_company_logo(row, files.get('logo') if files else None):
        raise ValueError('logo')
    row.updated_at = datetime.utcnow()

def delete_company_logo_files(row):
    if row.logo_path:
        old = BASE_DIR / 'static' / row.logo_path
        if old.exists():
            old.unlink(missing_ok=True)

def ensure_company_profile():
    """Giữ tương thích code cũ — trả về công ty đang chọn."""
    return get_company_row()

def save_company_logo(row, file_storage):
    if not file_storage or not file_storage.filename:
        return True
    ext = Path(file_storage.filename).suffix.lower()
    if ext not in ALLOWED_IMAGE_EXTENSIONS:
        flash('Logo phải là JPG, PNG, WEBP hoặc GIF', 'warning')
        return False
    data = file_storage.read()
    if len(data) > MAX_PRODUCT_IMAGE_BYTES:
        flash('Logo tối đa 5MB', 'warning')
        return False
    file_storage.seek(0)
    COMPANY_LOGO_DIR.mkdir(parents=True, exist_ok=True)
    if row.logo_path:
        old = BASE_DIR / 'static' / row.logo_path
        if old.exists():
            old.unlink(missing_ok=True)
    filename = f'logo_{uuid.uuid4().hex[:10]}{ext}'
    rel_path = f'uploads/company/{filename}'
    file_storage.save(COMPANY_LOGO_DIR / filename)
    row.logo_path = rel_path
    return True

def resolve_active_nav_key():
    ep = request.endpoint
    if ep == 'stock':
        return 'stock'
    if ep in (
        'stock_intake_requests',
        'stock_intake_request_detail',
        'approve_supplier_intake_stock',
        'reject_supplier_intake_stock',
    ):
        return 'stock_intake_requests'
    endpoint_keys = {
        'dashboard': 'dashboard',
        'customers': 'customers',
        'suppliers': 'suppliers',
        'products': 'products',
        'quotes': 'quotes',
        'orders': 'orders',
        'contracts': 'contracts',
        'debt': 'debt',
        'supplier_debt': 'supplier_debt',
        'users': 'users',
        'company_settings': 'company_settings',
    }
    return endpoint_keys.get(ep, ep)

def build_sidebar_sections(user):
    if not user:
        return []
    allowed = ROLE_NAV_KEYS.get(user.role, set())
    active_key = resolve_active_nav_key()
    sections = []
    for block in SIDEBAR_SECTIONS:
        items = []
        for item in block['items']:
            if user.role != 'admin' and item['key'] not in allowed:
                continue
            url_args = dict(item.get('url_args') or {})
            items.append({
                'key': item['key'],
                'href': url_for(item['endpoint'], **url_args),
                'label': item['label'],
                'icon': item['icon'],
                'active': item['key'] == active_key,
            })
        if items:
            sections.append({'title': block['title'], 'links': items})
    return sections

@app.context_processor
def inject_company():
    try:
        row = get_company_row()
        company = company_profile_to_dict(row) if row else {}
        company_list = list_company_profiles()
    except Exception:
        d = DEFAULT_COMPANY_PROFILE
        bank = f"Số TK: {d['bank_account']} - {d['bank_name']}" if d.get('bank_account') else d.get('bank_name', '')
        company = {**d, 'bank': bank, 'director': d.get('director_title', 'Giám đốc'), 'logo_url': None}
        company_list = []
        row = None
    user = get_current_user()
    return {
        'company': company,
        'company_list': company_list,
        'active_company_id': row.id if row else None,
        'nav_endpoint': request.endpoint,
        'active_nav_key': resolve_active_nav_key(),
        'current_user': user,
        'user_role_label': USER_ROLES.get(user.role, '') if user else '',
        'sidebar_sections': build_sidebar_sections(user),
    }

@app.before_request
def enforce_auth():
    endpoint = request.endpoint
    if endpoint in ('login', 'static') or endpoint is None:
        return
    user = get_current_user()
    if not user:
        session.pop('user_id', None)
        return redirect(url_for('login', next=request.full_path if request.query_string else request.path))
    if not role_can_access(user.role, endpoint):
        flash('Bạn không có quyền truy cập chức năng này', 'danger')
        return redirect(default_home_for_role(user.role))

def ensure_quote_columns():
    from sqlalchemy import inspect, text
    try:
        cols = {c['name'] for c in inspect(db.engine).get_columns('quote')}
    except Exception:
        return
    with db.engine.begin() as conn:
        if 'valid_until' not in cols:
            conn.execute(text('ALTER TABLE quote ADD COLUMN valid_until DATE'))
        if 'sale_contract_path' not in cols:
            conn.execute(text("ALTER TABLE quote ADD COLUMN sale_contract_path VARCHAR(512) DEFAULT ''"))
        if 'handover_doc_path' not in cols:
            conn.execute(text("ALTER TABLE quote ADD COLUMN handover_doc_path VARCHAR(512) DEFAULT ''"))
        if 'doc_company_id' not in cols:
            conn.execute(text('ALTER TABLE quote ADD COLUMN doc_company_id INTEGER REFERENCES company_profile(id)'))
        if 'sale_contract_signed_date' not in cols:
            conn.execute(text('ALTER TABLE quote ADD COLUMN sale_contract_signed_date DATE'))
        if 'sale_contract_payment_note' not in cols:
            conn.execute(text("ALTER TABLE quote ADD COLUMN sale_contract_payment_note TEXT DEFAULT ''"))
        if 'sale_contract_delivery_note' not in cols:
            conn.execute(text("ALTER TABLE quote ADD COLUMN sale_contract_delivery_note TEXT DEFAULT ''"))
        if 'handover_date' not in cols:
            conn.execute(text('ALTER TABLE quote ADD COLUMN handover_date DATE'))
        if 'handover_place' not in cols:
            conn.execute(text("ALTER TABLE quote ADD COLUMN handover_place TEXT DEFAULT ''"))
        if 'handover_condition_note' not in cols:
            conn.execute(text("ALTER TABLE quote ADD COLUMN handover_condition_note TEXT DEFAULT ''"))
        if 'walkin_display_name' not in cols:
            conn.execute(text("ALTER TABLE quote ADD COLUMN walkin_display_name VARCHAR(255) DEFAULT ''"))

def quote_status_key(status):
    return QUOTE_STATUS_KEYS.get(status or '', 'new')

app.jinja_env.globals['quote_status_key'] = quote_status_key

def quote_valid_until(quote):
    if quote.valid_until:
        return quote.valid_until
    return (quote.created_at.date() if quote.created_at else date.today()) + timedelta(days=30)

def quote_display_status(quote):
    status = (quote.status or 'Mới tạo').strip()
    if status in LEGACY_QUOTE_STATUS_MAP:
        return LEGACY_QUOTE_STATUS_MAP[status]
    if status in QUOTE_STATUSES:
        return status
    return 'Mới tạo'

def quote_is_expired(quote):
    if quote_display_status(quote) != 'Mới tạo':
        return False
    return quote_valid_until(quote) < date.today()

def get_quote_order(quote):
    qid = quote if isinstance(quote, int) else quote.id
    return Order.query.filter_by(quote_id=qid).first()

def orders_by_quote_ids(quote_ids):
    if not quote_ids:
        return {}
    return {o.quote_id: o for o in Order.query.filter(Order.quote_id.in_(quote_ids)).all()}

def create_order_from_quote_record(quote):
    existing = get_quote_order(quote)
    if existing:
        return existing, False
    order = Order(
        order_code=next_code('DH', Order, 'order_code'),
        customer_id=quote.customer_id,
        quote_id=quote.id,
        total=quote.total,
        paid_amount=0,
        status='Đang xử lý',
    )
    db.session.add(order)
    db.session.flush()
    return order, True

def quote_chot_stock_in_qty(product_id, quote_id):
    ensure_stock_movement_columns()
    return float(
        db.session.query(func.coalesce(func.sum(StockMovement.qty), 0))
        .filter(
            StockMovement.product_id == product_id,
            StockMovement.quote_id == quote_id,
            StockMovement.movement_type == 'IN',
        )
        .scalar() or 0,
    )

def quote_shortage_lines(quote):
    """Dòng báo giá thiếu tồn — hiển thị cảnh báo, gợi nhập qua NCC."""
    lines = []
    for item in quote.items:
        product = item.product if getattr(item, 'product', None) else Product.query.get(item.product_id)
        if not product:
            continue
        qty = quote_item_base_qty(product, item)
        if qty <= 0:
            continue
        stock = float(product.stock or 0)
        shortfall = max(qty - stock, 0)
        if shortfall <= 1e-9:
            continue
        link = (
            ProductSupplier.query.filter_by(product_id=product.id, is_primary=True).first()
            or ProductSupplier.query.filter_by(product_id=product.id)
            .order_by(ProductSupplier.cost_price.asc())
            .first()
        )
        purchase_cost = (link.cost_price if link else 0) or product.cost_price or 0
        if product_has_unit_conversion(product) and purchase_cost:
            purchase_cost = purchase_unit_cost_from_base_cost(product, purchase_cost)
        lines.append({
            'product': product,
            'item': item,
            'shortfall': shortfall,
            'supplier': link.supplier if link else None,
            'supplier_id': link.supplier_id if link else None,
            'cost': int(purchase_cost or 0),
            'supplier_sku': (link.supplier_sku if link else '') or '',
        })
    return lines

def quote_stock_shortages(quote):
    """Sản phẩm trong báo giá không đủ tồn kho (hiển thị thông báo)."""
    shortages = []
    for row in quote_shortage_lines(quote):
        item = row['item']
        product = row['product']
        shortages.append({
            'name': item.product_name or product.name,
            'sku': item.sku or product.sku,
            'stock': float(product.stock or 0),
            'need': float(item.qty or 0),
        })
    return shortages

def quote_stock_already_deducted(quote_id):
    ensure_stock_movement_columns()
    return StockMovement.query.filter_by(
        quote_id=quote_id,
        movement_type='OUT',
    ).first() is not None

def apply_quote_stock_deduction(quote, order=None):
    """Trừ tồn theo báo giá (mỗi BG một lần)."""
    ensure_stock_movement_columns()
    if quote_stock_already_deducted(quote.id):
        return False
    order_code = order.order_code if order else ''
    for item in quote.items:
        product = item.product if getattr(item, 'product', None) else Product.query.get(item.product_id)
        if not product:
            continue
        qty = float(item.qty or 0)
        if qty <= 0:
            continue
        product.stock = max(float(product.stock or 0) - qty, 0)
        out_note = f'Xuất bán — báo giá {quote.quote_code}'
        if order_code:
            out_note += f' → đơn {order_code}'
        db.session.add(StockMovement(
            product_id=product.id,
            movement_type='OUT',
            qty=qty,
            purchase_qty=0,
            ref_code=quote.quote_code,
            method='Bán hàng',
            warehouse=DEFAULT_WAREHOUSE,
            note=out_note,
            quote_id=quote.id,
        ))
    return True

def quote_stock_shortage_flash_message(quote):
    """Thông báo sản phẩm không đủ tồn khi chốt báo giá."""
    rows = quote_stock_shortages(quote)
    if not rows:
        return None
    parts = []
    for row in rows[:4]:
        label = row['sku'] or row['name']
        parts.append(f'{label} (cần {int(row["need"])}, tồn {int(row["stock"])})')
    extra = f' và {len(rows) - 4} sản phẩm khác' if len(rows) > 4 else ''
    return (
        'Kho không đủ: ' + ', '.join(parts) + extra
        + '. Vui lòng nhập hàng từ NCC (gửi yêu cầu kho) nếu cần bổ sung tồn.'
    )

def quote_chot_flash_detail(stock_deducted, had_shortage=False):
    parts = []
    if stock_deducted:
        parts.append('trừ kho')
    if had_shortage:
        parts.append('có dòng thiếu tồn')
    return ' — '.join(parts)

def chot_quote(quote):
    """Chốt báo giá: tạo đơn hàng và trừ kho (không tự tạo phiếu mua khi thiếu tồn)."""
    had_shortage = bool(quote_shortage_lines(quote))
    if quote_display_status(quote) != 'Đã chốt':
        quote.status = 'Đã chốt'
    order, created = create_order_from_quote_record(quote)
    stock_deducted = apply_quote_stock_deduction(quote, order)
    return order, created, stock_deducted, had_shortage

def effective_quote_status(quote):
    return quote_display_status(quote)

def normalize_quote_statuses():
    changed = False
    for quote in Quote.query.filter(Quote.status.in_(list(LEGACY_QUOTE_STATUS_MAP.keys()))).all():
        quote.status = LEGACY_QUOTE_STATUS_MAP[quote.status]
        changed = True
    if changed:
        db.session.commit()

app.jinja_env.globals['quote_display_status'] = quote_display_status
app.jinja_env.globals['effective_quote_status'] = effective_quote_status
app.jinja_env.globals['quote_is_expired'] = quote_is_expired
app.jinja_env.globals['quote_valid_until'] = quote_valid_until
app.jinja_env.globals['get_quote_order'] = get_quote_order

def quote_note_preview(note, limit=40):
    text = (note or '').strip()
    if not text:
        return '—'
    return text if len(text) <= limit else text[:limit] + '…'

app.jinja_env.globals['quote_note_preview'] = quote_note_preview

def quote_stats(query=None):
    base = query if query is not None else Quote.query
    total = base.count()
    if total == 0:
        return {'total': 0, 'new': 0, 'won': 0, 'cancelled': 0,
                'new_pct': 0, 'won_pct': 0, 'cancelled_pct': 0}
    counts = {'new': 0, 'won': 0, 'cancelled': 0}
    for q in base.all():
        st = quote_display_status(q)
        if st == 'Mới tạo':
            counts['new'] += 1
        elif st == 'Đã chốt':
            counts['won'] += 1
        elif st == 'Đã hủy':
            counts['cancelled'] += 1
    def pct(n):
        return round(n * 1000 / total) / 10 if total else 0
    return {
        'total': total,
        **counts,
        'new_pct': pct(counts['new']),
        'won_pct': pct(counts['won']),
        'cancelled_pct': pct(counts['cancelled']),
    }

QUOTE_SORT_FIELDS = {'quote_code', 'created_at', 'valid_until', 'total', 'status'}

def quote_filters_from_request():
    sort = request.args.get('sort', 'created_at')
    sort = sort if sort in QUOTE_SORT_FIELDS else 'created_at'
    return {
        'customer_id': request.args.get('customer_id', ''),
        'date_from': request.args.get('date_from', ''),
        'date_to': request.args.get('date_to', ''),
        'status': request.args.get('status', ''),
        'q': request.args.get('q', '').strip(),
        'tab': request.args.get('tab', 'list'),
        'sort': sort,
        'order': 'asc' if request.args.get('order') == 'asc' else 'desc',
    }

def quote_sort_url(field, list_args, sort, order):
    next_order = 'desc' if sort == field and order == 'asc' else 'asc'
    if sort != field:
        next_order = 'asc'
    args = {k: v for k, v in list_args.items() if k not in ('sort', 'order', 'page')}
    return url_for('quotes', sort=field, order=next_order, page=1, **args)

app.jinja_env.globals['quote_sort_url'] = quote_sort_url

ORDER_STATUSES = ['Đang xử lý', 'Hoàn thành', 'Đã hủy']
ORDER_FILTER_STATUSES = ORDER_STATUSES
ORDER_STATUS_KEYS = {
    'Đang xử lý': 'processing',
    'Hoàn thành': 'completed',
    'Đã hủy': 'cancelled',
}
ORDER_LEGACY_STATUS_MAP = {
    'Mới tạo': 'Đang xử lý',
    'Chờ xác nhận': 'Đang xử lý',
    'Đang giao': 'Đang xử lý',
    'Đã chốt': 'Hoàn thành',
    'Đã thanh toán': 'Hoàn thành',
    'Thanh toán 1 phần': 'Hoàn thành',
}
ORDER_LEGACY_STATUS = ORDER_LEGACY_STATUS_MAP
ORDER_STATUS_FILTER_VALUES = {
    'Đang xử lý': ['Đang xử lý', 'Chờ xác nhận', 'Mới tạo', 'Đang giao'],
    'Hoàn thành': ['Hoàn thành', 'Đã chốt', 'Đã thanh toán', 'Thanh toán 1 phần'],
    'Đã hủy': ['Đã hủy'],
}
ORDER_ALLOWED_TRANSITIONS = {
    'Đang xử lý': {'Hoàn thành', 'Đã hủy'},
}
PAYMENT_STATUS_KEYS = {
    'Đã thanh toán': 'paid',
    'Thanh toán 1 phần': 'partial',
    'Chưa thanh toán': 'unpaid',
    '—': 'none',
}

def effective_order_status(order):
    raw = (order.status or 'Đang xử lý').strip()
    if raw in ORDER_LEGACY_STATUS_MAP:
        return ORDER_LEGACY_STATUS_MAP[raw]
    return raw if raw in ORDER_STATUSES else 'Đang xử lý'

def normalize_order_statuses():
    changed = False
    for order in Order.query.all():
        raw = (order.status or '').strip()
        if raw in ORDER_STATUSES:
            continue
        mapped = ORDER_LEGACY_STATUS_MAP.get(raw, 'Đang xử lý')
        if order.status != mapped:
            order.status = mapped
            changed = True
    if changed:
        db.session.commit()

def order_next_statuses(status):
    return sorted(ORDER_ALLOWED_TRANSITIONS.get(status, set()))

def order_payment_status(order):
    if effective_order_status(order) == 'Đã hủy':
        return '—'
    if (order.total or 0) <= 0:
        return 'Chưa thanh toán'
    if (order.paid_amount or 0) >= order.total:
        return 'Đã thanh toán'
    if (order.paid_amount or 0) > 0:
        return 'Thanh toán 1 phần'
    return 'Chưa thanh toán'

def order_status_key(status):
    display = ORDER_LEGACY_STATUS_MAP.get(status, status)
    return ORDER_STATUS_KEYS.get(display or '', 'processing')

def order_payment_key(status):
    return PAYMENT_STATUS_KEYS.get(status or '', 'unpaid')

def order_counts_toward_debt(order):
    return effective_order_status(order) != 'Đã hủy'

def order_balance(order):
    if not order_counts_toward_debt(order):
        return 0
    return max((order.total or 0) - (order.paid_amount or 0), 0)

def total_outstanding_debt(query=None):
    base = query if query is not None else Order.query
    return sum(order_balance(o) for o in base.all())

def debt_orders(query=None):
    base = query if query is not None else Order.query
    return [o for o in base.order_by(Order.created_at.desc()).all() if order_balance(o) > 0]

def customer_prior_debt(customer_id, exclude_order_id=None):
    """Công nợ còn lại của khách từ các đơn khác (không tính đơn đang xem)."""
    q = Order.query.filter_by(customer_id=customer_id)
    if exclude_order_id:
        q = q.filter(Order.id != exclude_order_id)
    return sum(order_balance(o) for o in q.all())

DEBT_OVERDUE_DAYS = 30

def customer_initials(name):
    if not name:
        return '?'
    parts = [p for p in name.strip().split() if p]
    if len(parts) >= 2:
        return (parts[0][0] + parts[-1][0]).upper()
    return name.strip()[:2].upper()

def customer_debt_aggregates():
    agg = {}
    for o in Order.query.order_by(Order.created_at.asc()).all():
        if not order_counts_toward_debt(o):
            continue
        cid = o.customer_id
        if cid not in agg:
            agg[cid] = {
                'total': 0, 'paid': 0, 'remaining': 0,
                'order_count': 0, 'unpaid_count': 0,
                'oldest_unpaid_at': None,
            }
        a = agg[cid]
        a['total'] += o.total or 0
        a['paid'] += o.paid_amount or 0
        bal = order_balance(o)
        a['remaining'] += bal
        a['order_count'] += 1
        if bal > 0:
            a['unpaid_count'] += 1
            created = o.created_at or datetime.utcnow()
            if a['oldest_unpaid_at'] is None or created < a['oldest_unpaid_at']:
                a['oldest_unpaid_at'] = created
    return agg

def customer_debt_status_from_agg(agg_row):
    if not agg_row or agg_row.get('remaining', 0) <= 0:
        return 'Đã thanh toán'
    cutoff = datetime.utcnow() - timedelta(days=DEBT_OVERDUE_DAYS)
    oldest = agg_row.get('oldest_unpaid_at')
    if oldest and oldest < cutoff:
        return 'Quá hạn'
    return 'Còn nợ'

def customer_debt_status_key(status):
    return {'Còn nợ': 'owing', 'Đã thanh toán': 'paid', 'Quá hạn': 'overdue'}.get(status or '', 'owing')

def customer_unpaid_orders_oldest_first(customer_id):
    return [
        o for o in Order.query.filter_by(customer_id=customer_id)
        .order_by(Order.created_at.asc(), Order.id.asc()).all()
        if order_counts_toward_debt(o) and order_balance(o) > 0
    ]

def customer_outstanding_balance(customer_id):
    return sum(order_balance(o) for o in customer_unpaid_orders_oldest_first(customer_id))

def apply_order_payment(order, amount, method='Chuyển khoản', note='',
                        receipt_path='', receipt_uploaded_at=None, batch_id=''):
    pay = min(parse_int(amount), order_balance(order))
    if pay <= 0:
        return None
    p = Payment(
        order_id=order.id,
        amount=pay,
        method=method or 'Chuyển khoản',
        note=(note or '').strip(),
        receipt_path=receipt_path or '',
        receipt_uploaded_at=receipt_uploaded_at,
        batch_id=batch_id or '',
    )
    order.paid_amount = min((order.paid_amount or 0) + pay, order.total or 0)
    db.session.add(p)
    return {'payment': p, 'amount': pay, 'order': order}

def allocate_customer_payment(customer_id, amount, method='Chuyển khoản', note='',
                              receipt_path='', receipt_uploaded_at=None):
    """Phân bổ thanh toán vào các đơn còn nợ, đơn cũ nhất trước."""
    pay_total = parse_int(amount)
    if pay_total <= 0:
        return []
    batch_id = uuid.uuid4().hex
    allocations = []
    remaining = pay_total
    user_note = (note or '').strip()
    for o in customer_unpaid_orders_oldest_first(customer_id):
        if remaining <= 0:
            break
        pay = min(remaining, order_balance(o))
        if pay <= 0:
            continue
        p_note = user_note or 'Thanh toán công nợ tổng'
        result = apply_order_payment(
            o, pay, method, p_note,
            receipt_path=receipt_path,
            receipt_uploaded_at=receipt_uploaded_at,
            batch_id=batch_id,
        )
        if result:
            allocations.append(result)
            remaining -= pay
    if allocations:
        db.session.commit()
    return allocations

def format_payment_allocation_lines(allocations):
    return [f"{a['order'].order_code} +{money_plain(a['amount'])}đ" for a in allocations]

def payment_group_key(payment):
    if payment.batch_id:
        return ('batch', payment.batch_id)
    if payment.receipt_path and payment.receipt_uploaded_at:
        return (
            'receipt',
            payment.receipt_path,
            payment.receipt_uploaded_at.isoformat(),
            (payment.note or '').strip(),
            payment.payment_date.isoformat() if payment.payment_date else '',
            payment.method or '',
        )
    return ('single', payment.id)

def group_payments_for_display(payments):
    """Gộp các dòng payment cùng lần phân bổ thành một dòng hiển thị."""
    if not payments:
        return []
    buckets = {}
    order_keys = []
    for p in payments:
        key = payment_group_key(p)
        if key not in buckets:
            buckets[key] = []
            order_keys.append(key)
        buckets[key].append(p)

    groups = []
    for key in order_keys:
        items = buckets[key]
        head = items[0]
        if key[0] in ('batch', 'receipt') and len(items) > 1:
            items_sorted = sorted(items, key=lambda x: (x.id,))
            groups.append({
                'kind': 'batch',
                'payment_date': head.payment_date,
                'method': head.method,
                'note': head.note,
                'total': sum(i.amount or 0 for i in items),
                'receipt_pid': head.id,
                'receipt_path': head.receipt_path,
                'payment_ids': [i.id for i in items_sorted],
                'allocations': [
                    {'order_code': i.order.order_code if i.order else '—', 'amount': i.amount or 0}
                    for i in items_sorted
                ],
            })
        else:
            for p in items:
                groups.append({
                    'kind': 'single',
                    'payment': p,
                    'payment_date': p.payment_date,
                    'method': p.method,
                    'note': p.note,
                    'total': p.amount or 0,
                    'receipt_pid': p.id,
                    'receipt_path': p.receipt_path,
                    'payment_ids': [p.id],
                })
    return groups

app.jinja_env.globals['group_payments_for_display'] = group_payments_for_display

def build_debt_customer_rows(search_q=''):
    agg = customer_debt_aggregates()
    if not agg:
        return []
    query = Customer.query.filter(Customer.id.in_(agg.keys()))
    if search_q:
        like = f'%{search_q}%'
        query = query.filter(or_(Customer.name.ilike(like), Customer.phone.ilike(like)))
    rows = []
    for c in query.all():
        a = agg.get(c.id, {})
        status = customer_debt_status_from_agg(a)
        rows.append({
            'customer': c,
            'total': a.get('total', 0),
            'paid': a.get('paid', 0),
            'remaining': a.get('remaining', 0),
            'order_count': a.get('order_count', 0),
            'unpaid_count': a.get('unpaid_count', 0),
            'status': status,
            'status_key': customer_debt_status_key(status),
        })
    rows.sort(key=lambda r: (-r['remaining'], (r['customer'].name or '').lower()))
    return rows

class SimplePagination:
    def __init__(self, items, page, per_page):
        self.total = len(items)
        self.per_page = max(per_page, 1)
        self.page = max(page, 1)
        self.pages = max(1, (self.total + self.per_page - 1) // self.per_page) if self.total else 1
        if self.page > self.pages:
            self.page = self.pages
        start = (self.page - 1) * self.per_page
        self.items = items[start:start + self.per_page]
        self.has_prev = self.page > 1
        self.has_next = self.page < self.pages
        self.prev_num = self.page - 1
        self.next_num = self.page + 1

    def iter_pages(self, left_edge=1, right_edge=1, left_current=1, right_current=2):
        last = 0
        for num in range(1, self.pages + 1):
            if (
                num <= left_edge
                or (self.page - left_current - 1 < num < self.page + right_current)
                or num > self.pages - right_edge
            ):
                if last + 1 != num:
                    yield None
                yield num
                last = num


def pagination_neighbor_pages(pagination):
    """Chỉ trang trước, trang hiện tại và trang sau — không dùng ellipsis."""
    if not pagination:
        return []
    page = getattr(pagination, 'page', None)
    if page is None:
        return []
    nums = []
    if getattr(pagination, 'has_prev', False):
        nums.append(pagination.prev_num)
    nums.append(page)
    if getattr(pagination, 'has_next', False):
        nums.append(pagination.next_num)
    return nums


def pagination_page_url(endpoint, page_num, list_args=None, page_param='page'):
    args = {k: v for k, v in (list_args or {}).items() if v and k != page_param}
    args[page_param] = page_num
    return url_for(endpoint, **args)


app.jinja_env.globals['pagination_neighbor_pages'] = pagination_neighbor_pages
app.jinja_env.globals['pagination_page_url'] = pagination_page_url

def debt_redirect(**extra):
    args = {k: v for k, v in request.args.items() if v}
    args.update({k: v for k, v in extra.items() if v is not None and v != ''})
    return redirect(url_for('debt', **args))

def order_stats(query=None):
    base = query if query is not None else Order.query
    total = base.count()
    if total == 0:
        return {'total': 0, 'processing': 0, 'completed': 0, 'cancelled': 0,
                'processing_pct': 0, 'completed_pct': 0, 'cancelled_pct': 0}
    counts = {'processing': 0, 'completed': 0, 'cancelled': 0}
    for o in base.all():
        st = effective_order_status(o)
        if st == 'Hoàn thành':
            counts['completed'] += 1
        elif st == 'Đã hủy':
            counts['cancelled'] += 1
        else:
            counts['processing'] += 1
    def pct(n):
        return round(n * 1000 / total) / 10 if total else 0
    return {
        'total': total,
        **counts,
        'processing_pct': pct(counts['processing']),
        'completed_pct': pct(counts['completed']),
        'cancelled_pct': pct(counts['cancelled']),
    }

def order_filters_from_request():
    return {
        'customer_id': request.args.get('customer_id', ''),
        'date_from': request.args.get('date_from', ''),
        'date_to': request.args.get('date_to', ''),
        'status': request.args.get('status', ''),
        'q': request.args.get('q', '').strip(),
    }

def apply_order_filters(query, filters):
    if filters['customer_id']:
        query = query.filter(Order.customer_id == int(filters['customer_id']))
    if filters['status']:
        st = filters['status']
        status_values = ORDER_STATUS_FILTER_VALUES.get(st, [st])
        query = query.filter(Order.status.in_(status_values))
    date_from = parse_stock_date(filters['date_from'])
    date_to = parse_stock_date(filters['date_to'])
    if date_from:
        query = query.filter(Order.created_at >= datetime.combine(date_from, datetime.min.time()))
    if date_to:
        query = query.filter(Order.created_at <= datetime.combine(date_to, datetime.max.time()))
    if filters['q']:
        like = f"%{filters['q']}%"
        customer_ids = db.session.query(Customer.id).filter(Customer.name.ilike(like))
        query = query.filter(or_(Order.order_code.ilike(like), Order.customer_id.in_(customer_ids)))
    return query

def orders_redirect(**extra):
    f = order_filters_from_request()
    f.update(extra)
    return redirect(url_for('orders', **{k: v for k, v in f.items() if v}))

app.jinja_env.globals['effective_order_status'] = effective_order_status
app.jinja_env.globals['order_payment_status'] = order_payment_status
app.jinja_env.globals['order_status_key'] = order_status_key
app.jinja_env.globals['order_payment_key'] = order_payment_key
app.jinja_env.globals['order_balance'] = order_balance
app.jinja_env.globals['customer_prior_debt'] = customer_prior_debt
app.jinja_env.globals['customer_initials'] = customer_initials
app.jinja_env.globals['customer_debt_status_key'] = customer_debt_status_key
app.jinja_env.globals['order_next_statuses'] = order_next_statuses

def apply_quote_sort(query, sort, order):
    col = getattr(Quote, sort)
    if sort == 'valid_until':
        col = Quote.valid_until
    return query.order_by(col.asc() if order == 'asc' else col.desc())

def recalculate_quote_totals(quote):
    quote.vat_amount = int(max(quote.subtotal - quote.discount, 0) * quote.vat_rate / 100)
    quote.total = max(quote.subtotal - quote.discount, 0) + quote.vat_amount

def apply_quote_filters(query, filters):
    if filters['customer_id']:
        query = query.filter(Quote.customer_id == int(filters['customer_id']))
    if filters['status']:
        if filters['status'] in QUOTE_FILTER_STATUSES:
            legacy = [filters['status']] + [
                old for old, new in LEGACY_QUOTE_STATUS_MAP.items() if new == filters['status']
            ]
            query = query.filter(Quote.status.in_(legacy))
    date_from = parse_stock_date(filters['date_from'])
    date_to = parse_stock_date(filters['date_to'])
    if date_from:
        query = query.filter(Quote.created_at >= datetime.combine(date_from, datetime.min.time()))
    if date_to:
        query = query.filter(Quote.created_at <= datetime.combine(date_to, datetime.max.time()))
    if filters['q']:
        like = f"%{filters['q']}%"
        customer_ids = db.session.query(Customer.id).filter(Customer.name.ilike(like))
        query = query.filter(or_(Quote.quote_code.ilike(like), Quote.customer_id.in_(customer_ids)))
    return query

def quotes_redirect(**extra):
    f = quote_filters_from_request()
    f.update(extra)
    return redirect(url_for('quotes', **{k: v for k, v in f.items() if v}))

def parse_quote_line_items_from_form(quote=None):
    ensure_quote_item_columns()
    line_items = []
    stopped_names = []
    existing_ids = {i.product_id for i in quote.items} if quote else set()
    qty_modes = request.form.getlist('qty_unit_mode')
    variant_labels = request.form.getlist('variant_label')
    for idx, (product_id, qty, price) in enumerate(zip(
        request.form.getlist('product_id'),
        request.form.getlist('qty'),
        request.form.getlist('price'),
    )):
        if not product_id:
            continue
        p = Product.query.get(int(product_id))
        if not p:
            continue
        if not product_is_active(p) and p.id not in existing_ids:
            stopped_names.append(p.name or p.sku)
            continue
        variant_label = (variant_labels[idx] if idx < len(variant_labels) else '').strip()
        mode = (qty_modes[idx] if idx < len(qty_modes) else '').strip().lower()
        if mode not in ('base', 'purchase', 'lot'):
            mode = product_default_sale_unit_mode(p)
        qv = max(parse_qty(qty, 1), 0.0001)
        pv = parse_int(price, p.retail_price)
        line_items.append((p, variant_label, qv, pv, int(round(qv * pv)), mode))
    return line_items, stopped_names

def apply_quote_line_items(quote, line_items):
    ensure_quote_item_columns()
    for old in list(quote.items):
        db.session.delete(old)
    subtotal = 0
    for p, variant_label, qv, pv, amount, mode in line_items:
        subtotal += amount
        display_name = p.name
        if variant_label:
            display_name = f'{p.name} ({variant_label})'
        db.session.add(QuoteItem(
            quote_id=quote.id,
            product_id=p.id,
            variant_label=variant_label or '',
            product_name=display_name,
            sku=p.sku,
            unit=product_unit_label(p, mode),
            qty_unit_mode=mode,
            qty=float(qv),
            price=pv,
            amount=amount,
        ))
    quote.subtotal = subtotal

def build_quote_preview_from_form():
    customer_id = resolve_quote_customer_id(request.form.get('customer_id'))
    customer = Customer.query.get(customer_id) if customer_id else None
    if not customer:
        customer = SimpleNamespace(name=WALKIN_CUSTOMER_NAME, phone='', address='')
    if is_walkin_customer(customer):
        walkin_name = quote_walkin_display_name_from_form()
        if walkin_name:
            customer = SimpleNamespace(
                name=walkin_name,
                phone=getattr(customer, 'phone', '') or '',
                address=getattr(customer, 'address', '') or '',
            )

    vat_rate = parse_int(request.form.get('vat_rate'), 0)
    discount = parse_int(request.form.get('discount'), 0)
    note = (request.form.get('note') or '').strip()
    valid_raw = request.form.get('valid_until', '').strip()
    try:
        valid_until = datetime.strptime(valid_raw, '%Y-%m-%d').date() if valid_raw else date.today() + timedelta(days=30)
    except ValueError:
        valid_until = date.today() + timedelta(days=30)

    line_discounts = request.form.getlist('line_discount')
    variant_labels = request.form.getlist('variant_label')
    items = []
    subtotal = 0
    for idx, (product_id, qty, price) in enumerate(zip(
        request.form.getlist('product_id'),
        request.form.getlist('qty'),
        request.form.getlist('price'),
    )):
        if not product_id:
            continue
        p = Product.query.get(int(product_id))
        if not p:
            continue
        variant_label = (variant_labels[idx] if idx < len(variant_labels) else '').strip()
        line_disc = 0.0
        if idx < len(line_discounts):
            try:
                line_disc = float(line_discounts[idx] or 0)
            except (TypeError, ValueError):
                line_disc = 0.0
        qty_modes = request.form.getlist('qty_unit_mode')
        mode = (qty_modes[idx] if idx < len(qty_modes) else '').strip().lower()
        if mode not in ('base', 'purchase', 'lot'):
            mode = product_default_sale_unit_mode(p)
        qv = max(parse_qty(qty, 1), 0.0001)
        pv = parse_int(price, 0)
        amount = round(qv * pv * (1 - line_disc / 100))
        subtotal += amount
        unit_price = round(pv * (1 - line_disc / 100)) if line_disc else pv
        display_name = f'{p.name} ({variant_label})' if variant_label else p.name
        items.append(SimpleNamespace(
            product_name=display_name,
            sku=p.sku,
            unit=product_unit_label(p, mode),
            qty=qv,
            price=unit_price,
            amount=amount,
            product=p,
        ))

    after_discount = max(subtotal - discount, 0)
    vat_amount = round(after_discount * vat_rate / 100)
    total = after_discount + vat_amount
    quote = SimpleNamespace(
        quote_code='BG-NHÁP',
        created_at=datetime.now(),
        customer_id=customer_id,
        customer=customer,
        vat_rate=vat_rate,
        discount=discount,
        subtotal=subtotal,
        vat_amount=vat_amount,
        total=total,
        note=note,
        items=items,
    )
    return quote, valid_until

def create_quote_from_form():
    ensure_quote_item_columns()
    customer_id = resolve_quote_customer_id(request.form.get('customer_id'))
    vat_rate = parse_int(request.form.get('vat_rate'), 10)
    discount = parse_int(request.form.get('discount'), 0)
    is_draft = request.form.get('save_draft') == '1'
    status = 'Nháp' if is_draft else request.form.get('status', 'Mới tạo') or 'Mới tạo'
    valid_raw = request.form.get('valid_until', '').strip()
    valid_until = datetime.strptime(valid_raw, '%Y-%m-%d').date() if valid_raw else date.today() + timedelta(days=30)
    line_items, stopped_names = parse_quote_line_items_from_form()
    if stopped_names:
        names = ', '.join(stopped_names[:5])
        extra = f' và {len(stopped_names) - 5} sản phẩm khác' if len(stopped_names) > 5 else ''
        raise ValueError(f'Sản phẩm đã ngừng bán không thể thêm vào báo giá: {names}{extra}')
    if not line_items:
        raise ValueError('Vui lòng thêm ít nhất một sản phẩm đang bán vào báo giá')
    walkin = Customer.query.get(customer_id)
    quote = Quote(
        quote_code=next_code('BG', Quote, 'quote_code'),
        customer_id=customer_id,
        status=status,
        vat_rate=vat_rate,
        discount=discount,
        note=request.form.get('note', ''),
        valid_until=valid_until,
        walkin_display_name=quote_walkin_display_name_from_form() if is_walkin_customer(walkin) else '',
    )
    db.session.add(quote)
    db.session.flush()
    apply_quote_line_items(quote, line_items)
    quote.vat_amount = int(max(quote.subtotal - discount, 0) * vat_rate / 100)
    quote.total = max(quote.subtotal - discount, 0) + quote.vat_amount
    generate_quote_documents(quote, commit=False)
    db.session.commit()
    set_quote_session_customer(quote)
    return quote

DEFAULT_WAREHOUSE = 'Kho chính'
STOCK_IN_METHODS = ['Nhập mua hàng', 'Khách trả hàng', 'Kiểm kê (tăng)', 'Chuyển kho đến', 'Khác']
STOCK_OUT_METHODS = ['Bán hàng', 'Xuất hủy/hỏng', 'Kiểm kê (giảm)', 'Chuyển kho đi', 'Khác']
STOCK_FILTER_METHODS = sorted(set(STOCK_IN_METHODS + STOCK_OUT_METHODS + [
    'Nhập tay', 'Trả hàng', 'Kiểm kê', 'Chuyển kho',  # dữ liệu cũ
]))

def stock_method_default(movement_type):
    return STOCK_IN_METHODS[0] if movement_type == 'IN' else STOCK_OUT_METHODS[0]

def normalize_stock_method(movement_type, method):
    allowed = STOCK_IN_METHODS if movement_type == 'IN' else STOCK_OUT_METHODS
    method = (method or '').strip()
    if method in allowed:
        return method
    legacy_map = {
        'IN': {'Nhập tay': 'Nhập mua hàng', 'Trả hàng': 'Khách trả hàng', 'Kiểm kê': 'Kiểm kê (tăng)', 'Chuyển kho': 'Chuyển kho đến'},
        'OUT': {'Nhập tay': 'Khác', 'Bán hàng': 'Bán hàng', 'Kiểm kê': 'Kiểm kê (giảm)', 'Chuyển kho': 'Chuyển kho đi'},
    }
    return legacy_map.get(movement_type, {}).get(method, stock_method_default(movement_type))

def stock_status(product):
    if product.stock <= 0:
        return 'out', 'Hết hàng'
    if product.stock <= (product.low_stock or 5):
        return 'low', 'Sắp hết hàng'
    return 'ok', 'Còn hàng'

def product_inventory_level(product):
    if product.stock <= 0:
        return 'out', 'Hết hàng'
    if product.stock <= (product.low_stock or 5):
        return 'low', 'Sắp hết'
    return 'ok', 'Còn nhiều'

def product_sale_status(product):
    if not getattr(product, 'is_active', True):
        return 'stopped', 'Ngừng bán'
    if product.stock <= 0:
        return 'out', 'Hết hàng'
    return 'selling', 'Đang bán'

def product_list_status(product):
    if getattr(product, 'is_active', True):
        return 'selling', 'Đang bán'
    return 'stopped', 'Ngừng bán'

def product_barcode(product):
    return product.sku or ''

def product_brand_options():
    ensure_category_brand_schema()
    return [b.name for b in Brand.query.order_by(Brand.name).all()]

app.jinja_env.globals['product_inventory_level'] = product_inventory_level
app.jinja_env.globals['product_sale_status'] = product_sale_status
app.jinja_env.globals['product_list_status'] = product_list_status
app.jinja_env.globals['product_barcode'] = product_barcode

app.jinja_env.globals['stock_status'] = stock_status

def parse_stock_date(value):
    if not value:
        return None
    try:
        return datetime.strptime(value, '%Y-%m-%d').date()
    except ValueError:
        return None

def ensure_stock_movement_columns():
    from sqlalchemy import inspect, text
    try:
        cols = {c['name'] for c in inspect(db.engine).get_columns('stock_movement')}
    except Exception:
        return
    with db.engine.begin() as conn:
        if 'method' not in cols:
            conn.execute(text("ALTER TABLE stock_movement ADD COLUMN method VARCHAR(80) DEFAULT 'Nhập tay'"))
        if 'warehouse' not in cols:
            conn.execute(text("ALTER TABLE stock_movement ADD COLUMN warehouse VARCHAR(80) DEFAULT 'Kho chính'"))
        if 'supplier_id' not in cols:
            conn.execute(text("ALTER TABLE stock_movement ADD COLUMN supplier_id INTEGER REFERENCES supplier(id)"))
        if 'purchase_qty' not in cols:
            conn.execute(text('ALTER TABLE stock_movement ADD COLUMN purchase_qty REAL DEFAULT 0'))
        if 'quote_id' not in cols:
            conn.execute(text('ALTER TABLE stock_movement ADD COLUMN quote_id INTEGER REFERENCES quote(id)'))

def ensure_supplier_intake_line_columns():
    from sqlalchemy import inspect, text
    ensure_supplier_intake_schema()
    try:
        cols = {c['name'] for c in inspect(db.engine).get_columns('supplier_intake_line')}
    except Exception:
        return
    with db.engine.begin() as conn:
        if 'purchase_qty' not in cols:
            conn.execute(text('ALTER TABLE supplier_intake_line ADD COLUMN purchase_qty REAL DEFAULT 0'))
            conn.execute(text(
                'UPDATE supplier_intake_line SET purchase_qty = qty '
                'WHERE purchase_qty IS NULL OR purchase_qty = 0'
            ))

def active_suppliers():
    return Supplier.query.filter_by(is_active=True).order_by(Supplier.name).all()

app.jinja_env.globals['active_suppliers'] = active_suppliers

def ensure_product_supplier_schema():
    db.create_all()

def product_suppliers_map(product_ids):
    ensure_product_supplier_schema()
    if not product_ids:
        return {}
    rows = (
        ProductSupplier.query.filter(ProductSupplier.product_id.in_(product_ids))
        .join(Supplier)
        .order_by(ProductSupplier.cost_price.asc(), Supplier.name.asc())
        .all()
    )
    out = {}
    for row in rows:
        out.setdefault(row.product_id, []).append(row)
    return out

def product_supplier_links(product):
    ensure_product_supplier_schema()
    return (
        ProductSupplier.query.filter_by(product_id=product.id)
        .join(Supplier)
        .order_by(ProductSupplier.cost_price.asc(), Supplier.name.asc())
        .all()
    )

def product_primary_supplier_link(product):
    links = product_supplier_links(product)
    if not links:
        return None
    for link in links:
        if link.is_primary:
            return link
    return links[0]

def product_supplier_reference_cost(product):
    """Giá nhập tham chiếu từ NCC (giá thấp nhất trên các link). None nếu chưa có NCC."""
    links = product_supplier_links(product)
    if not links:
        return None
    cheapest = min(links, key=lambda link: (link.cost_price or 0, link.supplier_id))
    return {
        'amount': cheapest.cost_price or 0,
        'supplier_name': cheapest.supplier.name if cheapest.supplier else '',
        'count': len(links),
    }

def sync_product_cost_from_supplier_purchase_price(product, purchase_cost):
    """Lưu Product.cost_price theo đơn vị tồn từ giá nhập / đơn vị nhập của NCC."""
    purchase_cost = int(purchase_cost or 0)
    if purchase_cost <= 0:
        return
    product.cost_price = base_unit_cost_from_purchase_cost(product, purchase_cost)

def product_import_cost_display(product):
    """Giá nhập hiển thị: ưu tiên giá NCC (đơn vị nhập); không có NCC thì từ cost_price SP."""
    ref = product_supplier_reference_cost(product)
    if ref and ref['amount']:
        unit = product_purchase_unit(product) if product_has_unit_conversion(product) else product_base_unit(product)
        base_hint = None
        if product_has_unit_conversion(product):
            base_hint = base_unit_cost_from_purchase_cost(product, ref['amount'])
        return {
            'amount': ref['amount'],
            'unit': unit,
            'source': 'supplier',
            'supplier_name': ref.get('supplier_name') or '',
            'base_hint': base_hint,
        }
    return {
        'amount': product_cost_display_amount(product),
        'unit': product_cost_unit_label(product),
        'source': 'product',
        'supplier_name': '',
        'base_hint': int(product.cost_price or 0) if product_has_unit_conversion(product) else None,
    }

app.jinja_env.globals['product_supplier_reference_cost'] = product_supplier_reference_cost
app.jinja_env.globals['product_import_cost_display'] = product_import_cost_display

def product_supplier_summary(links):
    if not links:
        return {'count': 0, 'min_cost': 0, 'cheapest_name': ''}
    cheapest = min(links, key=lambda link: (link.cost_price or 0, link.supplier_id))
    return {
        'count': len(links),
        'min_cost': cheapest.cost_price or 0,
        'cheapest_name': cheapest.supplier.name if cheapest.supplier else '',
    }

app.jinja_env.globals['product_supplier_summary'] = product_supplier_summary
app.jinja_env.globals['product_supplier_links'] = product_supplier_links
app.jinja_env.globals['product_primary_supplier_link'] = product_primary_supplier_link

def apply_product_suppliers_from_form(product):
    ensure_product_supplier_schema()
    supplier_ids = request.form.getlist('ps_supplier_id')
    cost_prices = request.form.getlist('ps_cost_price')
    supplier_skus = request.form.getlist('ps_supplier_sku')
    notes = request.form.getlist('ps_note')
    primary_id = parse_supplier_id(request.form.get('ps_primary_supplier_id'))

    ProductSupplier.query.filter_by(product_id=product.id).delete(synchronize_session=False)
    seen = set()
    for idx, sid_raw in enumerate(supplier_ids):
        sid = parse_supplier_id(sid_raw)
        if not sid or sid in seen:
            continue
        seen.add(sid)
        db.session.add(ProductSupplier(
            product_id=product.id,
            supplier_id=sid,
            cost_price=parse_int(cost_prices[idx] if idx < len(cost_prices) else 0, 0),
            supplier_sku=(supplier_skus[idx] if idx < len(supplier_skus) else '').strip(),
            note=(notes[idx] if idx < len(notes) else '').strip(),
            is_primary=(sid == primary_id),
        ))

    links = ProductSupplier.query.filter_by(product_id=product.id).all()
    if links and not any(link.is_primary for link in links):
        cheapest = min(links, key=lambda link: (link.cost_price or 0, link.supplier_id))
        cheapest.is_primary = True

def supplier_products_map(supplier_ids):
    ensure_product_supplier_schema()
    if not supplier_ids:
        return {}
    rows = (
        ProductSupplier.query.filter(ProductSupplier.supplier_id.in_(supplier_ids))
        .join(Product)
        .order_by(Product.name.asc())
        .all()
    )
    out = {}
    for row in rows:
        out.setdefault(row.supplier_id, []).append(row)
    return out

def supplier_product_counts(supplier_ids):
    if not supplier_ids:
        return {}
    rows = (
        db.session.query(ProductSupplier.supplier_id, func.count(ProductSupplier.id))
        .filter(ProductSupplier.supplier_id.in_(supplier_ids))
        .group_by(ProductSupplier.supplier_id)
        .all()
    )
    return {sid: cnt for sid, cnt in rows}

def upsert_product_supplier_link(product, supplier, cost_price=0, supplier_sku='', note=''):
    link = ProductSupplier.query.filter_by(product_id=product.id, supplier_id=supplier.id).first()
    if link:
        if cost_price:
            link.cost_price = cost_price
        if supplier_sku:
            link.supplier_sku = supplier_sku
        if note:
            link.note = note
        return link
    link = ProductSupplier(
        product_id=product.id,
        supplier_id=supplier.id,
        cost_price=cost_price or 0,
        supplier_sku=supplier_sku,
        note=note,
        is_primary=False,
    )
    db.session.add(link)
    return link

def supplier_intake_stock_status(intake):
    ensure_supplier_intake_status_column()
    status = (intake.status or '').strip()
    if status in (SUPPLIER_INTAKE_STATUS_PENDING, SUPPLIER_INTAKE_STATUS_RECEIVED):
        return status
    if intake.stock_received_at:
        return SUPPLIER_INTAKE_STATUS_RECEIVED
    has_pending_qty = any(
        (line.qty or 0) > 1e-9 and not line.stock_movement_id
        for line in intake.lines
    )
    if has_pending_qty:
        return SUPPLIER_INTAKE_STATUS_PENDING
    return SUPPLIER_INTAKE_STATUS_RECEIVED

def supplier_intake_stock_status_key(intake):
    st = supplier_intake_stock_status(intake)
    return 'pending' if st == SUPPLIER_INTAKE_STATUS_PENDING else 'received'

def _supplier_intake_line_stock_note(product, purchase_qty, base_qty, unit_mode, intake_note, note, supplier):
    line_note = intake_note
    if note:
        line_note = f'{line_note} — {note}' if line_note else note
    pu = product_purchase_unit(product)
    bu = product_base_unit(product)
    if product_has_unit_conversion(product):
        if unit_mode == 'base':
            conv = f' ({format_qty_display(base_qty)} {bu}'
            if abs(purchase_qty - base_qty) > 1e-6:
                conv += f' = {format_qty_display(purchase_qty)} {pu}'
            conv += ')'
        elif unit_mode == 'lot' and product_has_lot_unit(product):
            lu = product_lot_unit(product)
            lf = product_lot_factor(product)
            lot_qty = (purchase_qty / lf) if lf > 0 else purchase_qty
            conv = (
                f' ({format_qty_display(lot_qty)} {lu} → '
                f'{format_qty_display(purchase_qty)} {pu} → '
                f'{format_qty_display(base_qty)} {bu})'
            )
        else:
            conv = f' ({format_qty_display(purchase_qty)} {pu} → {format_qty_display(base_qty)} {bu})'
        line_note = (line_note or '') + conv
    return line_note or f'Nhập từ NCC {supplier.code}'

def _supplier_intake_apply_line_stock(intake, supplier, line_row, ref_code, intake_note):
    """Ghi nhận tồn kho cho một dòng phiếu NCC (khi duyệt yêu cầu)."""
    ensure_stock_movement_columns()
    product = line_row['product']
    purchase_qty = line_row['purchase_qty']
    base_qty = line_row['base_qty']
    unit_mode = line_row.get('unit_mode', 'purchase')
    note = line_row['note']
    if base_qty <= 1e-9:
        return 0, None
    product.stock = float(product.stock or 0) + base_qty
    movement = StockMovement(
        product_id=product.id,
        supplier_id=supplier.id,
        movement_type='IN',
        qty=base_qty,
        purchase_qty=purchase_qty,
        ref_code=ref_code or intake.ref_code,
        method='Nhập mua hàng',
        warehouse=DEFAULT_WAREHOUSE,
        note=_supplier_intake_line_stock_note(
            product, purchase_qty, base_qty, unit_mode, intake_note, note, supplier,
        ),
    )
    db.session.add(movement)
    db.session.flush()
    return base_qty, movement

def _supplier_intake_process_lines(supplier, pending_lines, ref_code='', intake_note=''):
    """Tạo phiếu nhập NCC — chờ duyệt kho, chưa cộng tồn."""
    ensure_product_supplier_schema()
    ensure_supplier_intake_schema()
    ensure_supplier_intake_status_column()
    if not pending_lines:
        return 0, 0
    intake = SupplierIntake(
        supplier_id=supplier.id,
        intake_code=next_supplier_intake_code(supplier),
        ref_code=ref_code,
        note=intake_note,
        status=SUPPLIER_INTAKE_STATUS_PENDING,
    )
    db.session.add(intake)
    db.session.flush()
    total_amount = 0
    for line in pending_lines:
        product = line['product']
        purchase_qty = line['purchase_qty']
        base_qty = line['base_qty']
        cost = line['cost']
        sku = line['sku']
        note = line['note']
        purchase_cost = cost
        base_cost = base_unit_cost_from_purchase_cost(product, purchase_cost)
        upsert_product_supplier_link(product, supplier, purchase_cost, sku, note)
        if purchase_cost > 0:
            old_cost = product.cost_price or 0
            if old_cost != base_cost:
                pu = product_purchase_unit(product)
                bu = product_base_unit(product)
                hist_note = f'Nhập từ NCC {supplier.code}'
                if product_has_unit_conversion(product) and base_cost != purchase_cost:
                    hist_note += (
                        f' ({money(purchase_cost)}/{pu} → {money(base_cost)}/{bu})'
                    )
                db.session.add(PriceHistory(
                    product_id=product.id,
                    price_type='cost_price',
                    old_price=old_cost,
                    new_price=base_cost,
                    note=hist_note,
                ))
            product.cost_price = base_cost
        bill_qty = purchase_qty if product_has_unit_conversion(product) else base_qty
        line_amount = int(bill_qty * purchase_cost)
        total_amount += line_amount
        db.session.add(SupplierIntakeLine(
            intake_id=intake.id,
            product_id=product.id,
            qty=base_qty,
            purchase_qty=purchase_qty,
            cost_price=purchase_cost,
            line_amount=line_amount,
            supplier_sku=sku,
            note=note,
            stock_movement_id=None,
        ))
    intake.total_amount = total_amount
    return len(pending_lines), 0

def intake_manual_customer_choices():
    """Khách có đơn đang mở — dùng khi phân bổ thủ công."""
    rows = (
        db.session.query(
            Customer.id,
            Customer.name,
            Customer.tax_code,
        )
        .join(Quote, Quote.customer_id == Customer.id)
        .join(Order, Order.quote_id == Quote.id)
        .filter(Quote.status != 'Đã hủy')
        .filter(Quote.status != 'Nháp')
        .filter(Order.status != 'Đã hủy')
        .distinct()
        .order_by(Customer.name.asc())
        .limit(400)
        .all()
    )
    return [
        {
            'customer_id': r.id,
            'customer_name': r.name or '',
            'customer_code': (r.tax_code or '').strip(),
        }
        for r in rows
    ]

def product_customer_demand_batch(product_ids):
    """Nhu cầu SP theo KH từ báo giá/đơn chưa hủy (qty quy về đơn vị tồn)."""
    if not product_ids:
        return {}
    rows = (
        db.session.query(QuoteItem, Quote, Order, Customer)
        .join(Quote, QuoteItem.quote_id == Quote.id)
        .join(Order, Order.quote_id == Quote.id)
        .join(Customer, Quote.customer_id == Customer.id)
        .filter(QuoteItem.product_id.in_(product_ids))
        .filter(Quote.status != 'Đã hủy')
        .filter(Quote.status != 'Nháp')
        .filter(Order.status != 'Đã hủy')
        .all()
    )
    out = {}
    agg = {}
    for item, quote, order, customer in rows:
        pid = item.product_id
        product = item.product if getattr(item, 'product', None) else Product.query.get(pid)
        if not product:
            continue
        base_qty = quote_item_base_qty(product, item)
        if base_qty <= 0:
            continue
        key = (pid, customer.id, order.id, quote.id)
        if key not in agg:
            agg[key] = {
                'customer_id': customer.id,
                'customer_name': customer.name,
                'customer_code': (customer.tax_code or '').strip(),
                'order_id': order.id,
                'order_code': order.order_code or '',
                'quote_id': quote.id,
                'quote_code': quote.quote_code or '',
                'qty_base': 0.0,
                'product': product,
            }
        agg[key]['qty_base'] += base_qty
    for (pid, _cid, _oid, _qid), row in agg.items():
        qty = row['qty_base']
        product = row['product']
        sale_mode = product_default_sale_unit_mode(product)
        entry = out.setdefault(pid, {'total_demand': 0.0, 'customers': []})
        entry['total_demand'] += qty
        entry['customers'].append({
            'customer_id': row['customer_id'],
            'customer_name': row['customer_name'],
            'customer_code': row['customer_code'],
            'order_id': row['order_id'],
            'order_code': row['order_code'],
            'quote_id': row['quote_id'],
            'quote_code': row['quote_code'],
            'qty': qty,
            'qty_base': qty,
            'qty_display': qty_from_base_unit(product, qty, sale_mode),
            'display_unit': product_unit_label(product, sale_mode),
            'unit_mode': sale_mode,
        })
    for pid in product_ids:
        product = Product.query.get(pid)
        if not product:
            continue
        entry = out.setdefault(pid, {'total_demand': 0.0, 'customers': []})
        stock = float(product.stock or 0)
        entry['stock'] = stock
        entry['shortfall'] = max(0.0, entry['total_demand'] - stock)
        entry['customers'].sort(key=lambda c: (-c['qty'], c['customer_name'] or ''))
    return out

def stock_intake_product_gap_map(line_map):
    """{product_id: gap dict} cho các SP trên phiếu nhập kho."""
    product_ids = set()
    for lines in line_map.values():
        for line in lines:
            product_ids.add(line.product_id)
    gap_map = product_customer_demand_batch(product_ids)
    pending_by_product = {}
    for lines in line_map.values():
        for line in lines:
            pending_by_product[line.product_id] = (
                pending_by_product.get(line.product_id, 0.0) + float(line.qty or 0)
            )
    for pid, incoming in pending_by_product.items():
        entry = gap_map.setdefault(pid, {'total_demand': 0.0, 'customers': [], 'stock': 0.0, 'shortfall': 0.0})
        entry['pending_incoming'] = incoming
        product = Product.query.get(pid)
        if product:
            entry['base_unit'] = product_base_unit(product)
    return gap_map

def supplier_intake_saved_allocations_map(line_map):
    """{intake_line_id: [{customer_name, order_code, qty}, ...]}"""
    line_ids = []
    for lines in line_map.values():
        for line in lines:
            line_ids.append(line.id)
    if not line_ids:
        return {}
    rows = SupplierIntakeCustomerAllocation.query.filter(
        SupplierIntakeCustomerAllocation.intake_line_id.in_(line_ids),
    ).all()
    out = {}
    for row in rows:
        customer = row.customer
        order = row.order
        out.setdefault(row.intake_line_id, []).append({
            'customer_name': (customer.name if customer else '') or f'KH#{row.customer_id}',
            'order_code': (order.order_code if order else '') or '',
            'qty': float(row.qty or 0),
        })
    return out

def stock_intake_line_unit_mode(line):
    """Ưu tiên đơn vị đã nhập trên phiếu NCC, không thì đơn vị bán mặc định."""
    product = line.product
    if not product:
        return 'base'
    if product_has_unit_conversion(product) and (line.purchase_qty or 0) > 0:
        factor = product_conversion_factor(product)
        if factor > 0 and abs(float(line.qty or 0) - float(line.purchase_qty or 0) * factor) < 1e-4:
            return 'purchase'
    return product_default_sale_unit_mode(product)

def stock_intake_line_plan(line, gap, unit_mode=None):
    """Kế hoạch nhập/phân bổ — có số liệu theo đơn vị tồn và đơn vị đang chọn."""
    product = line.product
    bu = product_base_unit(product) if product else ''
    mode = unit_mode or stock_intake_line_unit_mode(line)
    stock = float(gap.get('stock', float(product.stock or 0) if product else 0))
    total_demand = float(gap.get('total_demand', 0))
    line_base = float(line.qty or 0)
    shortfall = float(gap.get('shortfall', max(0.0, total_demand - stock)))
    after_stock = stock + line_base
    remain_after = max(0.0, total_demand - after_stock)
    du = product_unit_label(product, mode) if product else bu
    return {
        'base_unit': bu,
        'display_unit': du,
        'unit_mode': mode,
        'unit_options': product_unit_mode_options(product) if product else [],
        'stock': stock,
        'stock_display': qty_from_base_unit(product, stock, mode) if product else stock,
        'total_demand': total_demand,
        'total_demand_display': qty_from_base_unit(product, total_demand, mode) if product else total_demand,
        'shortfall': shortfall,
        'shortfall_display': qty_from_base_unit(product, shortfall, mode) if product else shortfall,
        'line_base': line_base,
        'line_display': qty_from_base_unit(product, line_base, mode) if product else line_base,
        'after_stock': after_stock,
        'remain_after': remain_after,
        'remain_after_display': qty_from_base_unit(product, remain_after, mode) if product else remain_after,
        'pending_incoming': float(gap.get('pending_incoming', line_base)),
        'conversion_factor': product_conversion_factor(product) if product else 1.0,
        'has_conversion': product_has_unit_conversion(product) if product else False,
        'has_lot': product_has_lot_unit(product) if product else False,
    }

app.jinja_env.globals['stock_intake_line_plan'] = stock_intake_line_plan

def parse_intake_customer_allocations_from_form():
    line_ids = request.form.getlist('alloc_line_id')
    customer_ids = request.form.getlist('alloc_customer_id')
    order_ids = request.form.getlist('alloc_order_id')
    quote_ids = request.form.getlist('alloc_quote_id')
    qtys = request.form.getlist('alloc_qty')
    alloc_modes = request.form.getlist('alloc_unit_mode')
    rows = []
    for idx, line_id_raw in enumerate(line_ids):
        qty_entered = parse_qty(qtys[idx] if idx < len(qtys) else 0, 0)
        if qty_entered <= 0:
            continue
        try:
            line_id = int(line_id_raw)
            customer_id = int(customer_ids[idx] if idx < len(customer_ids) else 0)
        except (TypeError, ValueError):
            continue
        if customer_id <= 0:
            continue
        line = SupplierIntakeLine.query.get(line_id)
        product = line.product if line else None
        mode = (alloc_modes[idx] if idx < len(alloc_modes) else '').strip().lower()
        if product:
            if mode not in ('base', 'purchase', 'lot'):
                mode = stock_intake_line_unit_mode(line)
            qty = qty_to_base_unit(product, qty_entered, mode)
        else:
            qty = qty_entered
        order_id = None
        quote_id = None
        if idx < len(order_ids) and order_ids[idx]:
            try:
                order_id = int(order_ids[idx])
            except (TypeError, ValueError):
                pass
        if idx < len(quote_ids) and quote_ids[idx]:
            try:
                quote_id = int(quote_ids[idx])
            except (TypeError, ValueError):
                pass
        rows.append({
            'line_id': line_id,
            'customer_id': customer_id,
            'order_id': order_id,
            'quote_id': quote_id,
            'qty': qty,
        })
    return rows

def allocations_by_intake_line(allocation_rows):
    by_line = {}
    for row in allocation_rows:
        by_line.setdefault(row['line_id'], []).append(row)
    return by_line

def _allocation_note_for_line(alloc_rows):
    parts = []
    for row in alloc_rows:
        customer = Customer.query.get(row['customer_id'])
        label = customer.name if customer else f'KH#{row["customer_id"]}'
        parts.append(f'{label}: {format_qty_display(row["qty"])}')
    return 'Phân bổ KH: ' + ', '.join(parts) if parts else ''

def save_intake_customer_allocations(intake_line_id, alloc_rows):
    for row in alloc_rows:
        db.session.add(SupplierIntakeCustomerAllocation(
            intake_line_id=intake_line_id,
            customer_id=row['customer_id'],
            order_id=row.get('order_id'),
            quote_id=row.get('quote_id'),
            qty=float(row['qty'] or 0),
        ))

def approve_supplier_intake_stock_record(intake, customer_allocations=None):
    """Duyệt phiếu NCC — cộng tồn kho theo các dòng chờ."""
    ensure_supplier_intake_schema()
    ensure_supplier_intake_status_column()
    ensure_stock_movement_columns()
    ensure_supplier_intake_line_columns()
    db.create_all()
    if supplier_intake_stock_status(intake) == SUPPLIER_INTAKE_STATUS_RECEIVED:
        return False, 'Phiếu đã được nhập kho'
    supplier = intake.supplier or Supplier.query.get(intake.supplier_id)
    if not supplier:
        return False, 'Không tìm thấy nhà cung cấp'
    alloc_by_line = allocations_by_intake_line(customer_allocations or [])
    stock_qty = 0
    intake_note = intake.note or ''
    for db_line in intake.lines.order_by(SupplierIntakeLine.id.asc()).all():
        if db_line.stock_movement_id:
            continue
        product = db_line.product or Product.query.get(db_line.product_id)
        if not product:
            continue
        line_allocs = alloc_by_line.get(db_line.id, [])
        line_base = float(db_line.qty or 0)
        if line_allocs and line_base > 0:
            alloc_sum = sum(float(a['qty'] or 0) for a in line_allocs)
            if alloc_sum > line_base + 1e-6:
                return False, (
                    f'Phân bổ vượt SL nhập dòng {product.sku} '
                    f'({format_qty_display(alloc_sum)} > {format_qty_display(line_base)})'
                )
        unit_mode = 'purchase'
        if product_has_unit_conversion(product) and (db_line.purchase_qty or 0) > 0:
            factor = product_conversion_factor(product)
            if factor > 0 and abs((db_line.qty or 0) - (db_line.purchase_qty or 0) * factor) < 1e-6:
                unit_mode = 'purchase'
            elif (db_line.qty or 0) > 0 and abs((db_line.purchase_qty or 0) - (db_line.qty or 0) / factor) < 1e-6:
                unit_mode = 'base'
        line_note = (db_line.note or '').strip()
        alloc_note = _allocation_note_for_line(line_allocs)
        if alloc_note:
            line_note = f'{line_note} | {alloc_note}' if line_note else alloc_note
        row = {
            'product': product,
            'purchase_qty': float(db_line.purchase_qty or 0),
            'base_qty': line_base,
            'unit_mode': unit_mode,
            'note': line_note,
        }
        added, movement = _supplier_intake_apply_line_stock(
            intake, supplier, row, intake.ref_code, intake_note,
        )
        if movement:
            db_line.stock_movement_id = movement.id
            if line_allocs:
                save_intake_customer_allocations(db_line.id, line_allocs)
            stock_qty += added
    intake.status = SUPPLIER_INTAKE_STATUS_RECEIVED
    intake.stock_received_at = datetime.utcnow()
    return True, stock_qty

def product_options_item(p):
    return {
        'id': p.id,
        'sku': p.sku,
        'name': p.name,
        'image_url': product_image_url(p) or '',
        'stock': p.stock or 0,
        'cost_price': p.cost_price or 0,
        'purchase_cost_price': purchase_unit_cost_from_base_cost(p, p.cost_price or 0),
        'base_unit': product_base_unit(p),
        'purchase_unit': product_purchase_unit(p),
        'unit_conversion_enabled': bool(getattr(p, 'unit_conversion_enabled', False))
            or product_has_unit_conversion(p),
        'conversion_factor': product_conversion_factor(p),
        'lot_unit': product_lot_unit(p),
        'lot_factor': product_lot_factor(p),
        'has_lot_unit': product_has_lot_unit(p),
        'sale_unit_mode': product_default_sale_unit_mode(p),
        'unit': product_unit_label(p, product_default_sale_unit_mode(p)),
    }

def supplier_manual_intake_incomplete_rows():
    """Dòng có SL/giá nhưng chưa chọn SP có trong danh mục."""
    product_ids = request.form.getlist('sp_product_id')
    qtys = request.form.getlist('sp_qty')
    cost_prices = request.form.getlist('sp_cost_price')
    bad = []
    for idx, pid_raw in enumerate(product_ids):
        qty = max(parse_qty(qtys[idx] if idx < len(qtys) else 0, 0), 0)
        cost = parse_int(cost_prices[idx] if idx < len(cost_prices) else 0, 0)
        if qty <= 0 and cost <= 0:
            continue
        if not pid_raw:
            bad.append(idx + 1)
            continue
        try:
            pid = int(pid_raw)
        except (TypeError, ValueError):
            bad.append(idx + 1)
            continue
        if not Product.query.get(pid):
            bad.append(idx + 1)
    return bad

def apply_supplier_products_intake(supplier):
    ensure_product_supplier_schema()
    ensure_stock_movement_columns()
    ensure_supplier_intake_schema()
    ensure_supplier_intake_line_columns()
    product_ids = request.form.getlist('sp_product_id')
    row_count = len(product_ids)
    qtys = request.form.getlist('sp_qty')
    qty_modes = request.form.getlist('sp_qty_unit_mode')
    cost_prices = request.form.getlist('sp_cost_price')
    supplier_skus = request.form.getlist('sp_supplier_sku')
    notes = request.form.getlist('sp_note')
    ref_code = request.form.get('ref_code', '').strip()
    intake_note = request.form.get('intake_note', '').strip()
    pending_lines = []
    for idx in range(row_count):
        pid_raw = (product_ids[idx] or '').strip()
        if not pid_raw:
            continue
        try:
            pid = int(pid_raw)
        except (TypeError, ValueError):
            continue
        product = Product.query.get(pid)
        if not product:
            continue
        entered_qty = max(parse_qty(qtys[idx] if idx < len(qtys) else 0, 0), 0)
        unit_mode = (qty_modes[idx] if idx < len(qty_modes) else '').strip() or 'base'
        purchase_qty, base_qty, unit_mode = resolve_supplier_intake_quantities(
            product, entered_qty, unit_mode,
        )
        cost = parse_int(cost_prices[idx] if idx < len(cost_prices) else 0, 0)
        sku = (supplier_skus[idx] if idx < len(supplier_skus) else '').strip()
        note = (notes[idx] if idx < len(notes) else '').strip()
        if base_qty <= 0 and purchase_qty <= 0 and cost <= 0 and not sku:
            continue
        pending_lines.append({
            'product': product,
            'purchase_qty': purchase_qty,
            'base_qty': base_qty,
            'unit_mode': unit_mode,
            'cost': cost,
            'sku': sku,
            'note': note,
        })
    return _supplier_intake_process_lines(supplier, pending_lines, ref_code, intake_note)

def parse_supplier_id(value):
    if value is None or value == '':
        return None
    try:
        sid = int(value)
    except (TypeError, ValueError):
        return None
    if sid <= 0:
        return None
    return sid if Supplier.query.get(sid) else None

def supplier_list_params():
    q = request.args.get('q', '').strip()
    per_page = request.args.get('per_page', 10, type=int)
    per_page = per_page if per_page in (10, 25, 50, 100) else 10
    page = request.args.get('page', 1, type=int)
    sort = request.args.get('sort', 'created_at')
    order = request.args.get('order', 'desc')
    if sort not in ('name', 'code', 'phone', 'created_at'):
        sort = 'created_at'
    if order not in ('asc', 'desc'):
        order = 'desc'
    return q, per_page, page, sort, order

def supplier_sort_url(field, q, per_page, sort, order):
    new_order = 'asc' if sort == field and order == 'desc' else 'desc'
    if sort != field:
        new_order = 'asc'
    return url_for('suppliers', q=q, per_page=per_page, sort=field, order=new_order)

def apply_supplier_form(s, is_new=False):
    s.name = request.form.get('name', '').strip()
    code = request.form.get('code', '').strip()
    if code:
        s.code = code
    elif is_new and not getattr(s, 'code', None):
        s.code = next_code('NCC', Supplier, 'code')
    s.tax_code = request.form.get('tax_code', '').strip()
    s.contact_person = request.form.get('contact_person', '').strip()
    s.phone = request.form.get('phone', '').strip()
    s.email = request.form.get('email', '').strip()
    s.address = request.form.get('address', '').strip()
    s.bank_account = request.form.get('bank_account', '').strip()
    s.bank_name = request.form.get('bank_name', '').strip()
    s.note = request.form.get('note', '').strip()
    if 'is_active' in request.form:
        s.is_active = request.form.get('is_active') == '1'
    elif is_new:
        s.is_active = True

def validate_supplier_form():
    name = request.form.get('name', '').strip()
    if not name:
        return False, 'Vui lòng nhập tên nhà cung cấp'
    code = request.form.get('code', '').strip()
    if code:
        q = Supplier.query.filter(Supplier.code == code)
        sid = (request.view_args or {}).get('sid')
        if sid:
            q = q.filter(Supplier.id != sid)
        if q.first():
            return False, f'Mã NCC "{code}" đã tồn tại'
    return True, ''

def _supplier_redirect_list_params():
    src = request.form if request.method == 'POST' else request.args
    return {
        'q': src.get('_q', ''),
        'page': src.get('_page', 1, type=int) or 1,
        'per_page': src.get('_per_page', 10, type=int) or 10,
        'sort': src.get('_sort', 'created_at'),
        'order': src.get('_order', 'desc'),
    }

def _supplier_redirect_from_detail():
    src = request.form if request.method == 'POST' else request.args
    return src.get('_from_detail') == '1' or request.args.get('_from_detail') == '1'

def suppliers_redirect_after_save(message=None, category='success', **extra):
    if message:
        flash(message, category)
    open_detail_sid = extra.pop('open_detail_sid', None)
    if open_detail_sid:
        tab = extra.pop('detail_tab', None)
        kwargs = {'tab': tab} if tab else {}
        return redirect(url_for('supplier_detail', sid=open_detail_sid, **kwargs))
    open_intake_sid = extra.pop('open_intake_sid', None)
    if open_intake_sid:
        kwargs = {'intake': '1'}
        if extra.pop('open_quick_product', None):
            kwargs['quick_product'] = '1'
        if extra.pop('intake_import_preview', None) == '1':
            kwargs['intake_import_preview'] = '1'
        return redirect(url_for('supplier_detail', sid=open_intake_sid, **kwargs))
    if _supplier_redirect_from_detail():
        sid = parse_int((request.view_args or {}).get('sid'), 0)
        if sid:
            tab = (request.form.get('tab') or request.args.get('tab') or '').strip()
            kwargs = {'tab': tab} if tab else {}
            return redirect(url_for('supplier_detail', sid=sid, **kwargs))
    list_params = _supplier_redirect_list_params()
    return redirect(url_for('suppliers', **list_params, **extra))

def supplier_stats():
    total = Supplier.query.count()
    active = Supplier.query.filter_by(is_active=True).count()
    with_movements = db.session.query(func.count(func.distinct(StockMovement.supplier_id))).filter(
        StockMovement.supplier_id.isnot(None),
    ).scalar() or 0
    return {'total': total, 'active': active, 'with_movements': with_movements}

def supplier_inbound_counts(supplier_ids):
    if not supplier_ids:
        return {}
    rows = db.session.query(
        StockMovement.supplier_id,
        func.count(StockMovement.id),
        func.coalesce(func.sum(StockMovement.qty), 0),
    ).filter(
        StockMovement.supplier_id.in_(supplier_ids),
        StockMovement.movement_type == 'IN',
    ).group_by(StockMovement.supplier_id).all()
    return {sid: {'count': cnt, 'qty': int(qty or 0)} for sid, cnt, qty in rows}

def ensure_supplier_intake_schema():
    db.create_all()
    ensure_supplier_intake_status_column()

def ensure_supplier_intake_status_column():
    """Chỉ ALTER/backfill khi thiếu cột — tránh UPDATE mỗi request (dễ khóa SQLite)."""
    from sqlalchemy import inspect, text
    db.create_all()
    try:
        cols = {c['name'] for c in inspect(db.engine).get_columns('supplier_intake')}
    except Exception:
        return
    needs_backfill = False
    with db.engine.begin() as conn:
        if 'status' not in cols:
            conn.execute(text(
                "ALTER TABLE supplier_intake ADD COLUMN status VARCHAR(40) "
                f"DEFAULT '{SUPPLIER_INTAKE_STATUS_PENDING}'"
            ))
            needs_backfill = True
        if 'stock_received_at' not in cols:
            conn.execute(text(
                'ALTER TABLE supplier_intake ADD COLUMN stock_received_at DATETIME'
            ))
            needs_backfill = True
        if needs_backfill:
            conn.execute(text(f"""
                UPDATE supplier_intake SET status = '{SUPPLIER_INTAKE_STATUS_RECEIVED}'
                WHERE stock_received_at IS NOT NULL
                   OR id IN (
                       SELECT DISTINCT intake_id FROM supplier_intake_line
                       WHERE stock_movement_id IS NOT NULL
                   )
            """))
            conn.execute(text(f"""
                UPDATE supplier_intake SET status = '{SUPPLIER_INTAKE_STATUS_PENDING}'
                WHERE status IS NULL OR status = ''
            """))

def ensure_supplier_payment_columns():
    from sqlalchemy import inspect, text
    db.create_all()
    try:
        cols = {c['name'] for c in inspect(db.engine).get_columns('supplier_payment')}
    except Exception:
        return
    with db.engine.begin() as conn:
        if 'purchase_order_id' not in cols:
            conn.execute(text(
                'ALTER TABLE supplier_payment ADD COLUMN purchase_order_id INTEGER '
                'REFERENCES purchase_order(id)'
            ))

def supplier_intake_balance(intake):
    total = intake.total_amount or 0
    paid = intake.paid_amount or 0
    return max(0, total - paid)

def supplier_intake_payment_status(intake):
    total = intake.total_amount or 0
    paid = intake.paid_amount or 0
    if total <= 0:
        return '—'
    if paid <= 0:
        return 'Chưa thanh toán'
    if paid >= total:
        return 'Đã thanh toán'
    return 'Thanh toán một phần'

def supplier_intake_payment_status_key(intake):
    status = supplier_intake_payment_status(intake)
    if status == 'Đã thanh toán':
        return 'paid'
    if status == 'Thanh toán một phần':
        return 'partial'
    if status == 'Chưa thanh toán':
        return 'unpaid'
    return 'none'

app.jinja_env.globals['supplier_intake_balance'] = supplier_intake_balance
app.jinja_env.globals['supplier_intake_payment_status'] = supplier_intake_payment_status
app.jinja_env.globals['supplier_intake_payment_status_key'] = supplier_intake_payment_status_key
app.jinja_env.globals['supplier_intake_stock_status'] = supplier_intake_stock_status
app.jinja_env.globals['supplier_intake_stock_status_key'] = supplier_intake_stock_status_key

def next_supplier_intake_code(supplier):
    prefix = f'PN{supplier.id}'
    return next_code(prefix, SupplierIntake, 'intake_code')

def recalc_supplier_intake_paid(intake):
    paid = db.session.query(func.coalesce(func.sum(SupplierPayment.amount), 0)).filter(
        SupplierPayment.intake_id == intake.id,
    ).scalar() or 0
    intake.paid_amount = int(paid)
    return intake.paid_amount

def delete_supplier_intake_record(intake):
    """Xóa phiếu nhập NCC: hoàn tồn kho, gỡ phiếu kho và thanh toán liên quan."""
    ensure_supplier_intake_schema()
    ensure_stock_movement_columns()
    code = intake.intake_code or f'#{intake.id}'
    for payment in list(intake.payments):
        unlink_supplier_payment_receipt(payment.receipt_path)
        db.session.delete(payment)
    for line in list(intake.lines):
        if line.stock_movement_id:
            movement = StockMovement.query.get(line.stock_movement_id)
            if movement:
                product = Product.query.get(line.product_id)
                if product and movement.movement_type == 'IN':
                    product.stock = max((product.stock or 0) - float(movement.qty or 0), 0)
                db.session.delete(movement)
    db.session.delete(intake)
    return code

def linked_supplier_stock_movement_ids():
    rows = db.session.query(SupplierIntakeLine.stock_movement_id).filter(
        SupplierIntakeLine.stock_movement_id.isnot(None),
    ).all()
    return {row[0] for row in rows}

def backfill_supplier_intakes_for_supplier(supplier_id):
    ensure_supplier_intake_schema()
    linked = linked_supplier_stock_movement_ids()
    movements = StockMovement.query.filter(
        StockMovement.supplier_id == supplier_id,
        StockMovement.movement_type == 'IN',
    ).order_by(StockMovement.created_at.asc()).all()
    unlinked = [m for m in movements if m.id not in linked]
    if not unlinked:
        return 0
    groups = {}
    for movement in unlinked:
        ts = movement.created_at.replace(microsecond=0) if movement.created_at else datetime.utcnow()
        key = (movement.supplier_id, movement.ref_code or '', ts)
        groups.setdefault(key, []).append(movement)
    created = 0
    supplier = Supplier.query.get(supplier_id)
    if not supplier:
        return 0
    for (_, ref_code, ts), items in groups.items():
        intake = SupplierIntake(
            supplier_id=supplier.id,
            intake_code=next_supplier_intake_code(supplier),
            ref_code=ref_code,
            note=(items[0].note or '') if items else '',
            status=SUPPLIER_INTAKE_STATUS_RECEIVED,
            stock_received_at=ts,
            created_at=ts,
        )
        db.session.add(intake)
        db.session.flush()
        total = 0
        for movement in items:
            product = movement.product or Product.query.get(movement.product_id)
            cost = purchase_unit_cost_from_base_cost(product, product.cost_price or 0) if product else 0
            pq = float(movement.purchase_qty or movement.qty or 0)
            bq = float(movement.qty or 0)
            line_amount = int(pq * cost) if pq else int(bq * cost)
            total += line_amount
            db.session.add(SupplierIntakeLine(
                intake_id=intake.id,
                product_id=movement.product_id,
                qty=bq,
                purchase_qty=pq,
                cost_price=cost,
                line_amount=line_amount,
                note=movement.note or '',
                stock_movement_id=movement.id,
            ))
        intake.total_amount = total
        created += 1
    if created:
        db.session.commit()
    return created

def build_supplier_intake_board(supplier_ids):
    ensure_supplier_intake_schema()
    ensure_supplier_intake_line_columns()
    ensure_stock_movement_columns()
    if not supplier_ids:
        return {'intakes': {}, 'payments': {}, 'summaries': {}, 'line_counts': {}}
    for sid in supplier_ids:
        backfill_supplier_intakes_for_supplier(sid)
    intakes = (
        SupplierIntake.query.filter(SupplierIntake.supplier_id.in_(supplier_ids))
        .order_by(SupplierIntake.created_at.desc())
        .all()
    )
    intake_ids = [i.id for i in intakes]
    line_counts = {}
    if intake_ids:
        rows = (
            db.session.query(SupplierIntakeLine.intake_id, func.count(SupplierIntakeLine.id))
            .filter(SupplierIntakeLine.intake_id.in_(intake_ids))
            .group_by(SupplierIntakeLine.intake_id)
            .all()
        )
        line_counts = {iid: cnt for iid, cnt in rows}
    intakes_by_supplier = {sid: [] for sid in supplier_ids}
    for intake in intakes:
        intakes_by_supplier.setdefault(intake.supplier_id, []).append(intake)
    payments = (
        SupplierPayment.query.filter(SupplierPayment.supplier_id.in_(supplier_ids))
        .order_by(SupplierPayment.payment_date.desc(), SupplierPayment.id.desc())
        .all()
    )
    payments_by_supplier = {sid: [] for sid in supplier_ids}
    for payment in payments:
        payments_by_supplier.setdefault(payment.supplier_id, []).append(payment)
    summaries = {}
    for sid in supplier_ids:
        items = intakes_by_supplier.get(sid, [])
        total = sum(i.total_amount or 0 for i in items)
        paid = sum(i.paid_amount or 0 for i in items)
        summaries[sid] = {
            'total': total,
            'paid': paid,
            'remaining': max(0, total - paid),
            'intake_count': len(items),
            'unpaid_count': sum(
                1 for i in items
                if supplier_intake_balance(i) > 0 and (i.total_amount or 0) > 0
            ),
        }
    return {
        'intakes': intakes_by_supplier,
        'payments': payments_by_supplier,
        'summaries': summaries,
        'line_counts': line_counts,
    }

def supplier_intake_lines(intake_id):
    return (
        SupplierIntakeLine.query.filter_by(intake_id=intake_id)
        .join(Product)
        .order_by(Product.name.asc())
        .all()
    )

def supplier_intake_lines_map(intake_ids):
    if not intake_ids:
        return {}
    rows = (
        SupplierIntakeLine.query.filter(SupplierIntakeLine.intake_id.in_(intake_ids))
        .join(Product)
        .order_by(Product.name.asc())
        .all()
    )
    out = {}
    for row in rows:
        out.setdefault(row.intake_id, []).append(row)
    return out

app.jinja_env.globals['supplier_intake_lines_map'] = supplier_intake_lines_map

def movement_qty_sum(product_id, movement_type, date_from=None, date_to=None, before_date=None):
    q = db.session.query(func.coalesce(func.sum(StockMovement.qty), 0)).filter(
        StockMovement.product_id == product_id,
        StockMovement.movement_type == movement_type,
    )
    if before_date:
        q = q.filter(StockMovement.created_at < datetime.combine(before_date, datetime.min.time()))
    else:
        if date_from:
            q = q.filter(StockMovement.created_at >= datetime.combine(date_from, datetime.min.time()))
        if date_to:
            q = q.filter(StockMovement.created_at <= datetime.combine(date_to, datetime.max.time()))
    return float(q.scalar() or 0)

def product_stock_row(product, date_from=None, date_to=None):
    status_key, status_label = stock_status(product)
    return {
        'product': product,
        'current': product.stock,
        'status_key': status_key,
        'status_label': status_label,
        'warehouse': DEFAULT_WAREHOUSE,
    }

def stock_stats():
    products = Product.query.all()
    return {
        'total_products': len(products),
        'total_qty': sum(p.stock for p in products),
        'available_qty': sum(p.stock for p in products if p.stock > 0),
        'low_stock': Product.query.filter(Product.stock_qty > 0, Product.stock_qty <= Product.low_stock).count(),
        'out_stock': Product.query.filter(Product.stock_qty <= 0).count(),
    }

def product_stats():
    q = Product.query
    return {
        'total': q.count(),
        'selling': q.filter(Product.is_active.is_(True), Product.stock_qty > 0).count(),
        'out_stock': q.filter(Product.stock_qty <= 0).count(),
        'stopped': q.filter(or_(Product.is_active.is_(False), Product.is_active == 0)).count(),
    }

def top_product_category_tabs(limit=7):
    """Danh mục cha có nhiều SP nhất — dùng tab nhanh trên danh sách."""
    tabs = []
    for parent in parent_categories():
        child_ids = [c.id for c in Category.query.filter_by(parent_id=parent.id)]
        ids = [parent.id] + child_ids
        count = Product.query.filter(Product.category_id.in_(ids)).count() if ids else 0
        if count > 0:
            tabs.append({'id': parent.id, 'name': parent.name, 'count': count})
    tabs.sort(key=lambda row: (-row['count'], row['name']))
    return tabs[:limit]

def stock_filters_from_request():
    return {
        'product_id': request.args.get('product_id', ''),
        'movement_type': request.args.get('movement_type', ''),
        'method': request.args.get('method', ''),
        'warehouse': request.args.get('warehouse', ''),
        'supplier_id': request.args.get('supplier_id', ''),
        'date_from': request.args.get('date_from', ''),
        'date_to': request.args.get('date_to', ''),
        'q': request.args.get('q', '').strip(),
        'tab': request.args.get('tab', 'inventory'),
        'status': request.args.get('status', ''),
    }

def stock_redirect_after_save(message=None, category='success'):
    if message:
        flash(message, category)
    src = request.form if request.method == 'POST' else request.args
    return redirect(url_for(
        'stock',
        tab=src.get('_tab') or src.get('tab') or 'inventory',
        product_id=src.get('_product_id') or src.get('product_id') or '',
        movement_type=src.get('_movement_type') or src.get('movement_type') or '',
        method=src.get('_method') or src.get('method') or '',
        warehouse=src.get('_warehouse') or src.get('warehouse') or '',
        supplier_id=src.get('_supplier_id') or src.get('supplier_id') or '',
        date_from=src.get('_date_from') or src.get('date_from') or '',
        date_to=src.get('_date_to') or src.get('date_to') or '',
        q=src.get('_q') or src.get('q') or '',
        status=src.get('_status') or src.get('status') or '',
        page=src.get('_page', 1, type=int) or 1,
        per_page=src.get('_per_page', 10, type=int) or 10,
    ))

def next_code(prefix, model, field):
    """Mã tiếp theo theo năm: lấy số thứ tự lớn nhất + 1 (không dùng count — tránh trùng khi đã xóa bản ghi giữa)."""
    year = datetime.now().year
    pattern = f'{prefix}-{year}-%'
    col = getattr(model, field)
    max_seq = 0
    for (code,) in db.session.query(col).filter(col.like(pattern)):
        if not code:
            continue
        suffix = code.rsplit('-', 1)[-1]
        try:
            max_seq = max(max_seq, int(suffix))
        except ValueError:
            pass
    return f'{prefix}-{year}-{max_seq + 1:04d}'

def parse_qty(value, default=0):
    if value is None or value == '':
        return float(default)
    try:
        cleaned = str(value).strip().replace(',', '.')
        return float(cleaned)
    except (TypeError, ValueError):
        return float(default)

def parse_int(value, default=0):
    """Parse số nguyên VNĐ: số Excel (110000.0), chuỗi 1.100.000, 110.000, v.v."""
    if value is None or value == '':
        return default
    if isinstance(value, bool):
        return default
    if isinstance(value, int):
        return value
    if isinstance(value, float):
        if value != value:
            return default
        return int(round(value))

    s = str(value).strip().replace('\xa0', '').replace(' ', '')
    for ch in ('đ', 'Đ', '₫', 'd', 'D'):
        s = s.replace(ch, '')
    s = s.strip()
    if not s:
        return default

    negative = s.startswith('-')
    if negative:
        s = s[1:].strip()

    dot_count = s.count('.')
    comma_count = s.count(',')

    if dot_count > 1 or comma_count > 1:
        digits = ''.join(c for c in s if c.isdigit())
        n = int(digits) if digits else default
        return -n if negative else n

    for sep in ('.', ','):
        if s.count(sep) != 1:
            continue
        left, _, right = s.partition(sep)
        if not left.isdigit() or not right.isdigit():
            continue
        if len(right) == 3:
            n = int(left + right)
            return -n if negative else n
        if len(right) <= 2:
            try:
                n = int(round(float(s.replace(',', '.'))))
                return -n if negative else n
            except ValueError:
                pass
        break

    digits = ''.join(c for c in s if c.isdigit())
    if not digits:
        return default
    n = int(digits)
    return -n if negative else n

def replace_docx_text(doc, mapping):
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    def replace_in_paragraph(paragraph):
        full = ''.join(run.text for run in paragraph.runs)
        changed = False
        for k, v in mapping.items():
            if k in full:
                full = full.replace(k, str(v))
                changed = True
        if changed:
            for run in paragraph.runs:
                run.text = ''
            if paragraph.runs:
                paragraph.runs[0].text = full
            else:
                paragraph.add_run(full)

    def walk(element):
        if element.tag == qn('w:p'):
            replace_in_paragraph(Paragraph(element, doc))
        elif element.tag == qn('w:tbl'):
            table = Table(element, doc)
            for row in table.rows:
                for cell in row.cells:
                    for child in cell._tc:
                        walk(child)
        elif element.tag == qn('w:sdt'):
            content = element.find(qn('w:sdtContent'))
            if content is not None:
                for child in content:
                    walk(child)
        else:
            for child in element:
                walk(child)

    for child in doc.element.body:
        walk(child)

PARTY_TABLE_COLS = (1350, 4875, 3270)
PARTY_TABLE_WIDTH = sum(PARTY_TABLE_COLS)
PARTY_ROW_SPANS = (1, 2, 1, 2, 2, 2)

def _iter_doc_tables(doc):
    from docx.oxml.ns import qn
    from docx.table import Table

    def walk(element):
        if element.tag == qn('w:tbl'):
            yield Table(element, doc)
        elif element.tag == qn('w:sdt'):
            content = element.find(qn('w:sdtContent'))
            if content is not None:
                for child in content:
                    yield from walk(child)
        else:
            for child in element:
                yield from walk(child)

    for child in doc.element.body:
        yield from walk(child)

def _is_party_info_table(table):
    if len(table.rows) != 6 or len(table.columns) != 3:
        return False
    labels = [row.cells[0].text.strip().lower() for row in table.rows]
    return (
        labels[0].startswith('đại diện')
        and labels[1].startswith('địa chỉ')
        and labels[2].startswith('điện thoại')
        and labels[5].startswith('mst')
        and (labels[3].startswith('số tk') or labels[3].startswith('tài khoản'))
    )

def _ensure_tbl_fixed_layout(tbl):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    layout = tblPr.find(qn('w:tblLayout'))
    if layout is None:
        layout = OxmlElement('w:tblLayout')
        tblPr.append(layout)
    layout.set(qn('w:type'), 'fixed')

    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:w'), str(PARTY_TABLE_WIDTH))
    tblW.set(qn('w:type'), 'dxa')

    tblGrid = tbl.find(qn('w:tblGrid'))
    if tblGrid is None:
        tblGrid = OxmlElement('w:tblGrid')
        tbl.insert(list(tbl).index(tblPr) + 1, tblGrid)
    for child in list(tblGrid):
        if child.tag == qn('w:tblGridChange'):
            tblGrid.remove(child)
    for gc in list(tblGrid.findall(qn('w:gridCol'))):
        tblGrid.remove(gc)
    for width in PARTY_TABLE_COLS:
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(width))
        tblGrid.append(gc)

def _set_cell_width(tc, width, span=1):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)

    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = OxmlElement('w:tcW')
        tcPr.append(tcW)
    tcW.set(qn('w:w'), str(width))
    tcW.set(qn('w:type'), 'dxa')

    gs = tcPr.find(qn('w:gridSpan'))
    if span > 1:
        if gs is None:
            gs = OxmlElement('w:gridSpan')
            tcPr.append(gs)
        gs.set(qn('w:val'), str(span))
    elif gs is not None:
        tcPr.remove(gs)

def _normalize_party_value_cell(cell):
    text = cell.text.strip()
    if not text:
        return
    cell.text = ': ' + text.lstrip(': ').strip()

def _normalize_party_info_table(table):
    w0, w1, w2 = PARTY_TABLE_COLS
    _ensure_tbl_fixed_layout(table._tbl)

    for ri, row in enumerate(table.rows):
        cells = row.cells
        if len(cells) == 3:
            _set_cell_width(cells[0]._tc, w0, 1)
            _set_cell_width(cells[1]._tc, w1, 1)
            _set_cell_width(cells[2]._tc, w2, 1)
            if ri == 0:
                _normalize_party_value_cell(cells[1])
        elif len(cells) == 2:
            _set_cell_width(cells[0]._tc, w0, 1)
            _set_cell_width(cells[1]._tc, w1 + w2, PARTY_ROW_SPANS[ri])
            if ri > 0:
                _normalize_party_value_cell(cells[1])
        cells[0].text = cells[0].text.strip()

def normalize_contract_party_tables(doc):
    changed = False
    for table in _iter_doc_tables(doc):
        if _is_party_info_table(table):
            _normalize_party_info_table(table)
            changed = True
    return changed

def contract_mapping(customer, code, signed, expired, company_id=None):
    is_company = getattr(customer, 'customer_type', 'company') != 'individual'
    company = get_company_profile(company_id)
    return {
        '{{SO_HOP_DONG}}': code,
        '{{NGAY_KY}}': signed.strftime('%d/%m/%Y'),
        '{{TEN_BEN_A}}': customer.name,
        '{{DAI_DIEN_A}}': customer.representative or (customer.name if not is_company else ''),
        '{{CHUC_VU_A}}': customer.position or ('' if not is_company else ''),
        '{{DIA_CHI_A}}': customer.address,
        '{{DIEN_THOAI_A}}': customer.phone,
        '{{SO_TK_A}}': customer.bank_account,
        '{{NGAN_HANG_A}}': customer.bank_name,
        '{{MST_A}}': customer.tax_code if is_company else (customer.id_card or ''),
        '{{TEN_BEN_B}}': company['name'],
        '{{DAI_DIEN_B}}': company.get('director_name') or '',
        '{{CHUC_VU_B}}': company.get('director') or 'Giám đốc',
        '{{DIA_CHI_B}}': company.get('address') or '',
        '{{DIEN_THOAI_B}}': company.get('phone') or '',
        '{{SO_TK_B}}': company.get('bank_account') or '',
        '{{NGAN_HANG_B}}': company.get('bank_name') or '',
        '{{MST_B}}': company.get('tax_code') or '',
        '{{NGAY_HET_HIEU_LUC}}': expired.strftime('%d/%m/%Y') if expired else '',
    }

def ensure_contract_columns():
    from sqlalchemy import inspect, text
    try:
        cols = {c['name'] for c in inspect(db.engine).get_columns('contract')}
    except Exception:
        return
    with db.engine.begin() as conn:
        if 'contract_type' not in cols:
            conn.execute(text("ALTER TABLE contract ADD COLUMN contract_type VARCHAR(20) DEFAULT 'framework'"))
        if 'signed_scan_path' not in cols:
            conn.execute(text("ALTER TABLE contract ADD COLUMN signed_scan_path VARCHAR(255) DEFAULT ''"))
        if 'signed_scan_uploaded_at' not in cols:
            conn.execute(text('ALTER TABLE contract ADD COLUMN signed_scan_uploaded_at DATETIME'))
        if 'company_id' not in cols:
            conn.execute(text('ALTER TABLE contract ADD COLUMN company_id INTEGER'))

def delete_contract_scan_file(scan_path):
    if not scan_path:
        return
    file_path = BASE_DIR / 'static' / scan_path
    if file_path.is_file():
        file_path.unlink(missing_ok=True)

def ensure_order_columns():
    from sqlalchemy import inspect, text
    try:
        cols = {c['name'] for c in inspect(db.engine).get_columns('order')}
    except Exception:
        return
    with db.engine.begin() as conn:
        if 'handover_scan_path' not in cols:
            conn.execute(text("ALTER TABLE \"order\" ADD COLUMN handover_scan_path VARCHAR(512) DEFAULT ''"))
        if 'handover_scan_uploaded_at' not in cols:
            conn.execute(text('ALTER TABLE "order" ADD COLUMN handover_scan_uploaded_at DATETIME'))

def delete_order_handover_scan_file(scan_path):
    delete_contract_scan_file(scan_path)

def order_handover_scan_url(scan_path):
    return contract_signed_scan_url(scan_path)

app.jinja_env.globals['order_handover_scan_url'] = order_handover_scan_url

def ensure_payment_columns():
    from sqlalchemy import inspect, text
    try:
        cols = {c['name'] for c in inspect(db.engine).get_columns('payment')}
    except Exception:
        return
    with db.engine.begin() as conn:
        if 'receipt_path' not in cols:
            conn.execute(text("ALTER TABLE payment ADD COLUMN receipt_path VARCHAR(512) DEFAULT ''"))
        if 'receipt_uploaded_at' not in cols:
            conn.execute(text('ALTER TABLE payment ADD COLUMN receipt_uploaded_at DATETIME'))
        if 'batch_id' not in cols:
            conn.execute(text("ALTER TABLE payment ADD COLUMN batch_id VARCHAR(64) DEFAULT ''"))

def save_payment_receipt(file_storage):
    """Lưu ảnh xác nhận đã nhận tiền. Trả về đường dẫn tương đối hoặc None."""
    if not file_storage or not file_storage.filename:
        flash('Vui lòng upload ảnh xác nhận đã nhận tiền', 'warning')
        return None
    ext = Path(secure_filename(file_storage.filename)).suffix.lower()
    if ext not in ALLOWED_ORDER_HANDOVER_EXTENSIONS:
        flash('Ảnh phải là JPG, PNG hoặc WEBP', 'warning')
        return None
    data = file_storage.read()
    file_storage.seek(0)
    if len(data) > MAX_CONTRACT_SCAN_BYTES:
        flash('File tối đa 10MB', 'warning')
        return None
    PAYMENT_RECEIPT_UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
    filename = f'{uuid.uuid4().hex}{ext}'
    rel_path = f'uploads/payments/receipts/{filename}'
    file_storage.save(PAYMENT_RECEIPT_UPLOAD_DIR / filename)
    return rel_path

def save_supplier_payment_receipt(file_storage):
    if not file_storage or not file_storage.filename:
        flash('Vui lòng chụp/upload ảnh bill thanh toán', 'warning')
        return None
    ext = Path(secure_filename(file_storage.filename)).suffix.lower()
    if ext not in ALLOWED_ORDER_HANDOVER_EXTENSIONS:
        flash('Ảnh bill phải là JPG, PNG hoặc WEBP', 'warning')
        return None
    data = file_storage.read()
    file_storage.seek(0)
    if len(data) > MAX_CONTRACT_SCAN_BYTES:
        flash('File tối đa 10MB', 'warning')
        return None
    SUPPLIER_PAYMENT_RECEIPT_UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
    filename = f'{uuid.uuid4().hex}{ext}'
    rel_path = f'uploads/suppliers/receipts/{filename}'
    file_storage.save(SUPPLIER_PAYMENT_RECEIPT_UPLOAD_DIR / filename)
    return rel_path

def unlink_supplier_payment_receipt(receipt_path):
    if not receipt_path:
        return
    file_path = BASE_DIR / 'static' / receipt_path
    if file_path.is_file():
        try:
            file_path.unlink()
        except OSError:
            pass

def supplier_payment_receipt_url(receipt_path):
    return contract_signed_scan_url(receipt_path)

app.jinja_env.globals['supplier_payment_receipt_url'] = supplier_payment_receipt_url

def payment_receipt_url(receipt_path):
    return contract_signed_scan_url(receipt_path)

app.jinja_env.globals['payment_receipt_url'] = payment_receipt_url

def save_order_handover_scan(order, file_storage):
    if not file_storage or not file_storage.filename:
        flash('Vui lòng upload ảnh biên bản nhận hàng có chữ ký người nhận', 'warning')
        return False
    ext = Path(secure_filename(file_storage.filename)).suffix.lower()
    if ext not in ALLOWED_ORDER_HANDOVER_EXTENSIONS:
        flash('Ảnh phải là JPG, PNG hoặc WEBP', 'warning')
        return False
    data = file_storage.read()
    file_storage.seek(0)
    if len(data) > MAX_CONTRACT_SCAN_BYTES:
        flash('File tối đa 10MB', 'warning')
        return False
    ORDER_HANDOVER_UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
    delete_order_handover_scan_file(order.handover_scan_path)
    filename = f'{order.id}_{uuid.uuid4().hex[:10]}{ext}'
    rel_path = f'uploads/orders/handover/{filename}'
    file_storage.save(ORDER_HANDOVER_UPLOAD_DIR / filename)
    order.handover_scan_path = rel_path
    order.handover_scan_uploaded_at = datetime.utcnow()
    return True

def contract_signed_scan_url(scan_path):
    if not scan_path:
        return ''
    return url_for('static', filename=scan_path)

def signed_scan_is_pdf(scan_path):
    return bool(scan_path) and scan_path.lower().endswith('.pdf')

def signed_scan_mimetype(file_path):
    ext = file_path.suffix.lower()
    return {
        '.pdf': 'application/pdf',
        '.jpg': 'image/jpeg',
        '.jpeg': 'image/jpeg',
        '.png': 'image/png',
        '.webp': 'image/webp',
    }.get(ext, 'application/octet-stream')

def resolve_customer_signed_scan(cid):
    customer = Customer.query.get_or_404(cid)
    contract = get_framework_contract(customer.id)
    if not contract or not contract.signed_scan_path:
        return customer, contract, None
    file_path = BASE_DIR / 'static' / contract.signed_scan_path
    if not file_path.is_file():
        return customer, contract, None
    return customer, contract, file_path

def save_contract_signed_scan(contract, file_storage):
    if not file_storage or not file_storage.filename:
        flash('Vui lòng chọn file hợp đồng đã ký', 'warning')
        return False
    ext = Path(secure_filename(file_storage.filename)).suffix.lower()
    if ext not in ALLOWED_CONTRACT_SCAN_EXTENSIONS:
        flash('File phải là JPG, PNG, WEBP, PDF', 'warning')
        return False
    data = file_storage.read()
    file_storage.seek(0)
    if len(data) > MAX_CONTRACT_SCAN_BYTES:
        flash('File tối đa 10MB', 'warning')
        return False
    CONTRACT_SIGNED_UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
    delete_contract_scan_file(contract.signed_scan_path)
    filename = f'{contract.id}_{uuid.uuid4().hex[:10]}{ext}'
    rel_path = f'uploads/contracts/signed/{filename}'
    file_storage.save(CONTRACT_SIGNED_UPLOAD_DIR / filename)
    contract.signed_scan_path = rel_path
    contract.signed_scan_uploaded_at = datetime.utcnow()
    return True

def default_contract_dates():
    signed = date.today()
    expired = date(date.today().year, 12, 31)
    return signed, expired

def save_contract_docx(doc, code):
    file_name = f"{code.replace('/', '-').replace(' ', '_')}.docx"
    out_path = OUTPUT_DIR / 'contracts' / file_name
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return out_path

def _cleanup_temp_docs(max_age_seconds=3600):
    if not TEMP_DOC_DIR.exists():
        return
    cutoff = time.time() - max_age_seconds
    for pattern in ('*.docx', '*.pdf'):
        for path in TEMP_DOC_DIR.glob(pattern):
            try:
                if path.stat().st_mtime < cutoff:
                    path.unlink()
            except OSError:
                pass

def save_temp_docx(doc, prefix='preview'):
    TEMP_DOC_DIR.mkdir(parents=True, exist_ok=True)
    _cleanup_temp_docs()
    path = TEMP_DOC_DIR / f'{prefix}_{uuid.uuid4().hex[:12]}.docx'
    doc.save(str(path))
    return path

def send_docx_inline(file_path, download_name=None):
    path = Path(file_path)
    if not path.is_file():
        abort(404)
    return send_file(
        path,
        mimetype=DOCX_MIMETYPE,
        as_attachment=False,
        download_name=download_name or path.name,
    )


def send_pdf_inline(file_path, download_name=None):
    path = Path(file_path)
    if not path.is_file():
        abort(404)
    return send_file(
        path,
        mimetype=PDF_MIMETYPE,
        as_attachment=False,
        download_name=download_name or path.name,
    )


def _unlink_preview_file(path):
    if not path:
        return
    p = Path(path)
    if p.is_file():
        try:
            p.unlink()
        except OSError:
            pass
    pdf = p.with_suffix('.pdf')
    if pdf.is_file() and pdf != p:
        try:
            pdf.unlink()
        except OSError:
            pass


def find_libreoffice_executable():
    env = os.environ.get('LIBREOFFICE_PATH', '').strip()
    if env and Path(env).is_file():
        return env
    for name in ('soffice', 'libreoffice'):
        found = shutil.which(name)
        if found:
            return found
    mac_app = Path('/Applications/LibreOffice.app/Contents/MacOS/soffice')
    if mac_app.is_file():
        return str(mac_app)
    return None


def convert_docx_to_pdf(docx_path):
    """Chuyển DOCX → PDF (LibreOffice hoặc docx2pdf trên macOS). Giữ nguyên layout Word."""
    docx_path = Path(docx_path).resolve()
    if not docx_path.is_file():
        raise FileNotFoundError(f'Không tìm thấy file Word: {docx_path}')
    pdf_path = docx_path.with_suffix('.pdf')
    out_dir = docx_path.parent
    out_dir.mkdir(parents=True, exist_ok=True)

    lo = find_libreoffice_executable()
    if lo:
        try:
            subprocess.run(
                [
                    lo,
                    '--headless',
                    '--nologo',
                    '--nofirststartwizard',
                    '--convert-to',
                    'pdf',
                    '--outdir',
                    str(out_dir),
                    str(docx_path),
                ],
                check=True,
                capture_output=True,
                timeout=DOCX_TO_PDF_TIMEOUT,
            )
        except subprocess.TimeoutExpired as exc:
            raise RuntimeError('Chuyển PDF quá thời gian chờ. Thử lại hoặc kiểm tra LibreOffice.') from exc
        except subprocess.CalledProcessError as exc:
            err = (exc.stderr or b'').decode('utf-8', errors='replace')[:500]
            raise RuntimeError(f'LibreOffice không chuyển được sang PDF. {err}'.strip()) from exc
        if pdf_path.is_file():
            return pdf_path

    if os.name == 'darwin':
        try:
            from docx2pdf import convert as docx2pdf_convert
        except ImportError:
            docx2pdf_convert = None
        if docx2pdf_convert is not None:
            docx2pdf_convert(str(docx_path), str(pdf_path))
            if pdf_path.is_file():
                return pdf_path

    raise RuntimeError(
        'Không chuyển được DOCX sang PDF. Cài LibreOffice '
        '(https://www.libreoffice.org) hoặc trên macOS dùng Microsoft Word + pip install docx2pdf.'
    )


def ensure_pdf_for_docx(docx_path, force=False):
    """Tạo/cập nhật PDF cạnh file DOCX nếu DOCX mới hơn hoặc chưa có PDF."""
    docx_path = Path(docx_path).resolve()
    pdf_path = docx_path.with_suffix('.pdf')
    if (
        not force
        and pdf_path.is_file()
        and pdf_path.stat().st_mtime >= docx_path.stat().st_mtime
    ):
        return pdf_path
    return convert_docx_to_pdf(docx_path)


def prepare_document_preview_urls(
    docx_path,
    view_pdf_endpoint,
    view_docx_endpoint,
    download_docx_endpoint,
    download_pdf_endpoint,
    **url_kwargs,
):
    """Chuẩn bị URL preview: ưu tiên PDF; nếu không convert được thì fallback xem Word."""
    docx_path = Path(docx_path)
    if not docx_path.is_file():
        raise FileNotFoundError(f'Không tìm thấy file: {docx_path}')
    urls = {
        'pdf_url': None,
        'docx_view_url': url_for(view_docx_endpoint, **url_kwargs) if view_docx_endpoint else None,
        'docx_download_url': url_for(download_docx_endpoint, **url_kwargs),
        'pdf_download_url': None,
        'pdf_error': None,
    }
    try:
        ensure_pdf_for_docx(docx_path)
        urls['pdf_url'] = url_for(view_pdf_endpoint, **url_kwargs)
        urls['pdf_download_url'] = url_for(download_pdf_endpoint, **url_kwargs)
    except (RuntimeError, FileNotFoundError) as e:
        urls['pdf_error'] = str(e)
    return urls


def render_document_preview_page(doc_title, customer_name, urls, back_url=None, is_draft=False):
    return docx_preview_page(
        doc_title,
        customer_name,
        pdf_url=urls.get('pdf_url'),
        docx_download_url=urls['docx_download_url'],
        pdf_download_url=urls.get('pdf_download_url'),
        docx_view_url=urls.get('docx_view_url'),
        pdf_error=urls.get('pdf_error'),
        back_url=back_url,
        is_draft=is_draft,
    )


def pdf_path_for_docx_or_abort(docx_path):
    try:
        return ensure_pdf_for_docx(docx_path)
    except (RuntimeError, FileNotFoundError) as e:
        abort(503, description=str(e))

def app_back_url(default_endpoint, **default_values):
    """URL quay lại: ưu tiên trang trước (cùng host), không thì list mặc định."""
    ref = request.referrer or ''
    if ref:
        try:
            ref_p = urlparse(ref)
            req_p = urlparse(request.url)
            if ref_p.netloc == req_p.netloc and ref_p.path != req_p.path:
                return ref
        except Exception:
            pass
    return url_for(default_endpoint, **default_values)


def docx_preview_page(
    doc_title,
    customer_name,
    pdf_url=None,
    docx_download_url=None,
    pdf_download_url=None,
    docx_view_url=None,
    pdf_error=None,
    back_url=None,
    is_draft=False,
):
    return render_template(
        'contract_preview.html',
        contract_code=doc_title,
        customer_name=customer_name,
        pdf_url=pdf_url,
        docx_download_url=docx_download_url,
        pdf_download_url=pdf_download_url,
        docx_view_url=docx_view_url,
        pdf_error=pdf_error,
        back_url=back_url or app_back_url('contracts'),
        is_draft=is_draft,
    )


def clear_contract_draft_temp():
    _unlink_preview_file(session.pop(CONTRACT_DRAFT_SESSION_KEY, None))

def _doc_add_center_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def _doc_add_party_block(doc, title, lines):
    doc.add_paragraph(title)
    for line in lines:
        if line:
            doc.add_paragraph(line)

def _doc_add_quote_items_table(doc, items, include_price=True):
    cols = 6 if include_price else 5
    table = doc.add_table(rows=1, cols=cols)
    table.style = 'Table Grid'
    headers = ['STT', 'Tên hàng hóa', 'ĐVT', 'Số lượng']
    if include_price:
        headers += ['Đơn giá (VNĐ)', 'Thành tiền (VNĐ)']
    else:
        headers.append('Ghi chú')
    for i, label in enumerate(headers):
        table.rows[0].cells[i].text = label
    for idx, item in enumerate(items, start=1):
        cells = table.add_row().cells
        cells[0].text = str(idx)
        cells[1].text = item.product_name or ''
        cells[2].text = item.unit or 'cái'
        cells[3].text = str(item.qty or 0)
        if include_price:
            cells[4].text = money_plain(item.price)
            cells[5].text = money_plain(item.amount)
        else:
            cells[4].text = 'Mới, nguyên đai kiện'
    return table

def quote_customer_party_lines(customer):
    is_company = getattr(customer, 'customer_type', 'company') != 'individual'
    rep = customer.representative or (customer.name if not is_company else '')
    pos = customer.position or ('' if is_company else '')
    mst = customer.tax_code if is_company else (customer.id_card or '')
    lines = [
        f'Tên: {customer.name}',
        f'Đại diện: {rep}' + (f'    Chức vụ: {pos}' if pos else ''),
        f'Địa chỉ: {customer.address or "—"}',
        f'Điện thoại: {customer.phone or "—"}',
    ]
    if customer.bank_account or customer.bank_name:
        lines.append(f'Số TK: {customer.bank_account or "—"}    Ngân hàng: {customer.bank_name or "—"}')
    if mst:
        lines.append(f'{"MST" if is_company else "CCCD"}: {mst}')
    return lines

def quote_seller_party_lines(company_id=None):
    company = get_company_profile(company_id)
    lines = [
        f'Tên: {company["name"]}',
        f'Đại diện: {company.get("director_name") or "—"}    Chức vụ: {company.get("director") or "Giám đốc"}',
        f'Địa chỉ: {company.get("address") or "—"}',
        f'Điện thoại: {company.get("phone") or "—"}',
    ]
    if company.get('bank'):
        lines.append(company['bank'])
    if company.get('tax_code'):
        lines.append(f'MST: {company["tax_code"]}')
    return lines

def quote_doc_company_id(quote):
    if quote.doc_company_id:
        return quote.doc_company_id
    row = get_company_row()
    return row.id if row else None

def quote_sale_contract_signed_on(quote):
    if quote.sale_contract_signed_date:
        return quote.sale_contract_signed_date
    return quote.created_at.date() if quote.created_at else date.today()

def quote_handover_on(quote):
    if quote.handover_date:
        return quote.handover_date
    return date.today()

def quote_handover_place_text(quote):
    place = (quote.handover_place or '').strip()
    if place:
        return place
    return (quote.customer.address or '').strip() or 'địa điểm giao hàng'

def parse_optional_date(value):
    value = (value or '').strip()
    if not value:
        return None
    return datetime.strptime(value, '%Y-%m-%d').date()

def apply_quote_sale_contract_form(quote):
    quote.doc_company_id = resolve_company_id_from_form()
    quote.sale_contract_signed_date = parse_optional_date(request.form.get('signed_date'))
    quote.sale_contract_payment_note = request.form.get('payment_note', '').strip()
    quote.sale_contract_delivery_note = request.form.get('delivery_note', '').strip()

def apply_quote_handover_form(quote):
    quote.doc_company_id = resolve_company_id_from_form()
    quote.handover_date = parse_optional_date(request.form.get('handover_date'))
    quote.handover_place = request.form.get('handover_place', '').strip()
    quote.handover_condition_note = request.form.get('condition_note', '').strip()

def quote_doc_redirect(qid):
    if request.form.get('_back') == 'preview':
        return redirect(url_for('quote_preview', qid=qid))
    return quotes_redirect(tab=request.form.get('_tab', 'list'))

def build_quote_sale_contract_doc(quote):
    customer = quote.customer
    items = list(quote.items)
    signed = quote_sale_contract_signed_on(quote)
    valid = quote_valid_until(quote)
    subtotal_net = max(quote.subtotal - quote.discount, 0)
    company_id = quote_doc_company_id(quote)
    payment_note = (quote.sale_contract_payment_note or '').strip() or 'Thanh toán theo thỏa thuận giữa hai bên / theo báo giá đính kèm.'
    delivery_note = (quote.sale_contract_delivery_note or '').strip() or 'Bên B giao hàng cho Bên A trong thời gian thỏa thuận sau khi hợp đồng có hiệu lực.'
    doc = Document()
    _doc_add_center_heading(doc, 'CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM', level=2)
    p = doc.add_paragraph('Độc lập - Tự do - Hạnh phúc')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _doc_add_center_heading(doc, 'HỢP ĐỒNG MUA BÁN HÀNG HÓA', level=1)
    doc.add_paragraph(f'Số hợp đồng: HĐMB/{quote.quote_code}')
    doc.add_paragraph(f'Căn cứ báo giá số {quote.quote_code} ngày {signed.strftime("%d/%m/%Y")}.')
    doc.add_paragraph(f'Hôm nay, ngày {signed.strftime("%d/%m/%Y")}, chúng tôi gồm:')
    _doc_add_party_block(doc, 'BÊN MUA (BÊN A):', quote_customer_party_lines(customer))
    _doc_add_party_block(doc, 'BÊN BÁN (BÊN B):', quote_seller_party_lines(company_id))
    doc.add_paragraph('Hai bên thống nhất ký kết hợp đồng mua bán với nội dung sau:')
    doc.add_paragraph('Điều 1. Hàng hóa, số lượng, chất lượng, giá')
    _doc_add_quote_items_table(doc, items, include_price=True)
    doc.add_paragraph(f'Tổng giá trị hàng hóa (chưa VAT): {money(subtotal_net)}')
    doc.add_paragraph(f'VAT ({quote.vat_rate or 0}%): {money(quote.vat_amount)}')
    doc.add_paragraph(f'Tổng giá trị thanh toán: {money(quote.total)}')
    doc.add_paragraph('Điều 2. Thời hạn giao hàng')
    doc.add_paragraph(delivery_note)
    doc.add_paragraph('Điều 3. Phương thức thanh toán')
    doc.add_paragraph(payment_note)
    doc.add_paragraph('Điều 4. Điều khoản chung')
    doc.add_paragraph('Hợp đồng có hiệu lực kể từ ngày ký đến hết ngày ' + valid.strftime('%d/%m/%Y') + '.')
    doc.add_paragraph('Hợp đồng được lập thành 02 bản có giá trị pháp lý như nhau, mỗi bên giữ 01 bản.')
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = 'ĐẠI DIỆN BÊN MUA\n\n\n\n(Ký, ghi rõ họ tên)'
    table.rows[0].cells[1].text = 'ĐẠI DIỆN BÊN BÁN\n\n\n\n(Ký, đóng dấu)'
    return doc

def quote_handover_context(quote):
    """Dữ liệu cho templates_docx/bien_ban_nhan_hang.docx (docxtpl)."""
    customer = quote.customer
    company = get_company_profile(quote_doc_company_id(quote))
    is_company = getattr(customer, 'customer_type', 'company') != 'individual'
    signed = quote_handover_on(quote)
    items = []
    for item in quote.items:
        items.append({
            'ten_hang': item.product_name or '',
            'dvt': (item.unit or 'cái').strip() or 'cái',
            'so_luong': format_qty_display(item.qty or 0),
        })
    bank_b = company.get('bank') or ''
    if not bank_b and company.get('bank_account'):
        parts = [company.get('bank_account', '')]
        if company.get('bank_name'):
            parts.append(company['bank_name'])
        bank_b = ' - '.join(p for p in parts if p)
    return {
        'TEN_BEN_A': customer.name or '',
        'TEN_BEN_B': company.get('name', ''),
        'DAI_DIEN_A': (customer.representative or '').strip() or (customer.name if not is_company else ''),
        'CHUC_VU_A': (customer.position or '').strip() if is_company else '',
        'DIA_CHI_A': customer.address or '',
        'MST_A': customer.tax_code if is_company else (customer.id_card or ''),
        'DAI_DIEN_B': company.get('director_name', ''),
        'CHUC_VU_B': company.get('director', 'Giám đốc'),
        'DIA_CHI_B': company.get('address', ''),
        'DIEN_THOAI_B': company.get('phone', ''),
        'TAI_KHOAN_B': bank_b,
        'MST_B': company.get('tax_code', ''),
        'NGAY_GIAO': signed.strftime('%d/%m/%Y'),
        'DIA_DIEM': quote_handover_place_text(quote),
        'items': items,
    }


def build_quote_handover_doc(quote):
    from docxtpl import DocxTemplate

    if not TEMPLATE_QUOTE_HANDOVER.is_file():
        raise FileNotFoundError(
            f'Không tìm thấy file mẫu biên bản: {TEMPLATE_QUOTE_HANDOVER}'
        )
    tpl = DocxTemplate(str(TEMPLATE_QUOTE_HANDOVER))
    tpl.render(quote_handover_context(quote))
    return tpl.docx

def regenerate_quote_sale_contract(quote, commit=False):
    old_path = quote.sale_contract_path
    doc = build_quote_sale_contract_doc(quote)
    quote.sale_contract_path = str(save_quote_docx(doc, quote, 'hop-dong-mua-ban'))
    if old_path and old_path != quote.sale_contract_path and Path(old_path).is_file():
        Path(old_path).unlink(missing_ok=True)
    if commit:
        db.session.commit()
    return quote

def regenerate_quote_handover(quote, commit=False):
    old_path = quote.handover_doc_path
    doc = build_quote_handover_doc(quote)
    quote.handover_doc_path = str(save_quote_docx(doc, quote, 'bien-ban-nhan-hang'))
    if old_path and old_path != quote.handover_doc_path:
        if Path(old_path).is_file():
            Path(old_path).unlink(missing_ok=True)
        old_pdf = Path(old_path).with_suffix('.pdf')
        if old_pdf.is_file():
            old_pdf.unlink(missing_ok=True)
    new_pdf = Path(quote.handover_doc_path).with_suffix('.pdf')
    if new_pdf.is_file():
        new_pdf.unlink(missing_ok=True)
    if commit:
        db.session.commit()
    return quote

app.jinja_env.globals['quote_sale_contract_signed_on'] = quote_sale_contract_signed_on
app.jinja_env.globals['quote_handover_on'] = quote_handover_on
app.jinja_env.globals['quote_handover_place_text'] = quote_handover_place_text
app.jinja_env.globals['quote_doc_company_id'] = quote_doc_company_id

def save_quote_docx(doc, quote, suffix):
    QUOTE_DOCS_DIR.mkdir(parents=True, exist_ok=True)
    safe_code = quote.quote_code.replace('/', '-').replace(' ', '_')
    out_path = QUOTE_DOCS_DIR / f'{safe_code}_{suffix}.docx'
    doc.save(out_path)
    return out_path

def generate_quote_documents(quote, commit=True):
    regenerate_quote_sale_contract(quote, commit=False)
    regenerate_quote_handover(quote, commit=False)
    if commit:
        db.session.commit()
    return quote

def docx_template_newer_than(template_path, output_path):
    """File mẫu Word đã sửa sau bản xuất → cần tạo lại (không ghi đè file mẫu)."""
    tpl = Path(template_path)
    out = Path(output_path) if output_path else None
    if not tpl.is_file() or not out or not out.is_file():
        return False
    return tpl.stat().st_mtime > out.stat().st_mtime

def ensure_quote_documents(quote):
    sale_ok = quote.sale_contract_path and Path(quote.sale_contract_path).is_file()
    hand_ok = quote.handover_doc_path and Path(quote.handover_doc_path).is_file()
    if not sale_ok or not hand_ok:
        generate_quote_documents(quote, commit=True)
        return quote
    if hand_ok and docx_template_newer_than(TEMPLATE_QUOTE_HANDOVER, quote.handover_doc_path):
        regenerate_quote_handover(quote, commit=True)
    return quote

def get_framework_contract(customer_id):
    ensure_contract_columns()
    return Contract.query.filter_by(
        customer_id=customer_id,
        contract_type=CONTRACT_TYPE_FRAMEWORK,
    ).order_by(Contract.created_at.desc()).first()

def create_framework_contract_for_customer(customer, commit=True, company_id=None):
    ensure_contract_columns()
    contract = get_framework_contract(customer.id)
    signed, expired = default_contract_dates()
    cid = company_id or (get_company_row().id if get_company_row() else None)
    if contract:
        code = contract.contract_code
        signed = contract.signed_date or signed
        expired = contract.expired_date or expired
    else:
        code = next_code('HD', Contract, 'contract_code')
    out_path = apply_contract_doc(contract, customer, code, signed, expired, cid)
    if contract:
        contract.file_path = out_path
        contract.signed_date = signed
        contract.expired_date = expired
        if cid:
            contract.company_id = cid
    else:
        contract = Contract(
            contract_code=code,
            customer_id=customer.id,
            company_id=cid,
            contract_type=CONTRACT_TYPE_FRAMEWORK,
            signed_date=signed,
            expired_date=expired,
            file_path=out_path,
        )
        db.session.add(contract)
    if commit:
        db.session.commit()
    return contract

def customer_transactions(customer_id, limit=50):
    quotes = Quote.query.filter_by(customer_id=customer_id).order_by(Quote.created_at.desc()).limit(limit).all()
    orders = Order.query.filter_by(customer_id=customer_id).order_by(Order.created_at.desc()).limit(limit).all()
    rows = []
    for q in quotes:
        rows.append({
            'kind': 'quote',
            'code': q.quote_code,
            'date': q.created_at,
            'amount': q.total,
            'status': q.status,
            'url': url_for('quote_preview', qid=q.id),
        })
    for o in orders:
        rows.append({
            'kind': 'order',
            'code': o.order_code,
            'date': o.created_at,
            'amount': o.total,
            'status': o.status,
            'url': url_for('order_preview', oid=o.id),
        })
    rows.sort(key=lambda r: r['date'] or datetime.min, reverse=True)
    return rows[:limit]

def customer_list_extras(customer_ids):
    ensure_contract_columns()
    if not customer_ids:
        return {}, {}, {}
    contracts = {
        c.customer_id: c
        for c in Contract.query.filter(
            Contract.customer_id.in_(customer_ids),
            Contract.contract_type == CONTRACT_TYPE_FRAMEWORK,
        ).all()
    }
    quotes_map = {}
    for q in Quote.query.filter(Quote.customer_id.in_(customer_ids)).order_by(Quote.created_at.desc()):
        quotes_map.setdefault(q.customer_id, []).append(q)
    orders_map = {}
    for o in Order.query.filter(Order.customer_id.in_(customer_ids)).order_by(Order.created_at.desc()):
        orders_map.setdefault(o.customer_id, []).append(o)
    return contracts, quotes_map, orders_map

def ensure_contract_template_layout():
    create_template_docx()
    if not TEMPLATE_DOCX.exists():
        return
    doc = Document(TEMPLATE_DOCX)
    if normalize_contract_party_tables(doc):
        doc.save(TEMPLATE_DOCX)

def build_contract_doc(customer, code, signed, expired, company_id=None):
    ensure_contract_template_layout()
    doc = Document(TEMPLATE_DOCX)
    normalize_contract_party_tables(doc)
    replace_docx_text(doc, contract_mapping(customer, code, signed, expired, company_id=company_id))
    return doc

def parse_contract_form(contract=None):
    customer = Customer.query.get_or_404(int(request.form['customer_id']))
    signed = datetime.strptime(request.form.get('signed_date'), '%Y-%m-%d').date()
    expired_raw = request.form.get('expired_date')
    expired = datetime.strptime(expired_raw, '%Y-%m-%d').date() if expired_raw else None
    code = (request.form.get('contract_code') or '').strip()
    if not code:
        code = contract.contract_code if contract else next_code('HD', Contract, 'contract_code')
    dup_q = Contract.query.filter(Contract.contract_code == code)
    if contract:
        dup_q = dup_q.filter(Contract.id != contract.id)
    if dup_q.first():
        raise ValueError(f'Số hợp đồng "{code}" đã tồn tại')
    quote_id = request.form.get('quote_id') or None
    company_id = resolve_company_id_from_form()
    return customer, code, signed, expired, quote_id, company_id

def apply_contract_doc(contract, customer, code, signed, expired, company_id):
    old_path = contract.file_path if contract else ''
    doc = build_contract_doc(customer, code, signed, expired, company_id=company_id)
    out_path = save_contract_docx(doc, code)
    new_path = str(out_path)
    if old_path and old_path != new_path and Path(old_path).is_file():
        Path(old_path).unlink(missing_ok=True)
        old_pdf = Path(old_path).with_suffix('.pdf')
        if old_pdf.is_file():
            old_pdf.unlink(missing_ok=True)
    new_pdf = Path(new_path).with_suffix('.pdf')
    if new_pdf.is_file():
        new_pdf.unlink(missing_ok=True)
    return new_path

@app.route('/login', methods=['GET', 'POST'])
def login():
    user = get_current_user()
    if user:
        return redirect(default_home_for_role(user.role))
    if request.method == 'POST':
        username = request.form.get('username', '').strip()
        password = request.form.get('password', '')
        user = User.query.filter_by(username=username).first()
        if not user or not user.is_active or not verify_user_password(user, password):
            flash('Tên đăng nhập hoặc mật khẩu không đúng', 'danger')
            return redirect(url_for('login', next=request.form.get('next', '')))
        session.clear()
        session['user_id'] = user.id
        dest = safe_next_url(request.form.get('next', '')) or default_home_for_role(user.role)
        flash(f'Xin chào, {user.full_name or user.username}', 'success')
        return redirect(dest)
    return render_template('login.html', next=request.args.get('next', ''))

@app.route('/logout', methods=['POST'])
def logout():
    session.clear()
    flash('Đã đăng xuất', 'secondary')
    return redirect(url_for('login'))

@app.route('/users', methods=['GET', 'POST'])
def users():
    if request.method == 'POST':
        username = request.form.get('username', '').strip().lower()
        full_name = request.form.get('full_name', '').strip()
        role = request.form.get('role', 'sales').strip()
        password = request.form.get('password', '')
        if role not in USER_ROLES:
            flash('Vai trò không hợp lệ', 'warning')
            return redirect(url_for('users'))
        if not username or len(username) < 3:
            flash('Tên đăng nhập phải có ít nhất 3 ký tự', 'warning')
            return redirect(url_for('users'))
        if len(password) < 6:
            flash('Mật khẩu phải có ít nhất 6 ký tự', 'warning')
            return redirect(url_for('users'))
        if User.query.filter_by(username=username).first():
            flash('Tên đăng nhập đã tồn tại', 'warning')
            return redirect(url_for('users'))
        user = User(username=username, full_name=full_name or username, role=role)
        set_user_password(user, password)
        db.session.add(user)
        db.session.commit()
        flash(f'Đã tạo tài khoản {username}', 'success')
        return redirect(url_for('users'))
    return render_template(
        'users.html',
        users=User.query.order_by(User.created_at.desc()).all(),
        roles=USER_ROLES,
    )

@app.route('/users/<int:uid>/update', methods=['POST'])
def update_user(uid):
    user = User.query.get_or_404(uid)
    current = get_current_user()
    full_name = request.form.get('full_name', '').strip()
    role = request.form.get('role', user.role).strip()
    password = request.form.get('password', '').strip()
    if role not in USER_ROLES:
        flash('Vai trò không hợp lệ', 'warning')
        return redirect(url_for('users'))
    if user.id == current.id and role != 'admin':
        flash('Không thể tự hạ quyền quản trị của chính mình', 'warning')
        return redirect(url_for('users'))
    user.full_name = full_name or user.username
    user.role = role
    if password:
        if len(password) < 6:
            flash('Mật khẩu mới phải có ít nhất 6 ký tự', 'warning')
            return redirect(url_for('users'))
        set_user_password(user, password)
    db.session.commit()
    flash(f'Đã cập nhật tài khoản {user.username}', 'success')
    return redirect(url_for('users'))

@app.route('/users/<int:uid>/toggle', methods=['POST'])
def toggle_user_active(uid):
    user = User.query.get_or_404(uid)
    current = get_current_user()
    if user.id == current.id:
        flash('Không thể khóa tài khoản đang đăng nhập', 'warning')
        return redirect(url_for('users'))
    user.is_active = not user.is_active
    db.session.commit()
    flash(
        f'Đã {"kích hoạt" if user.is_active else "khóa"} tài khoản {user.username}',
        'success' if user.is_active else 'secondary',
    )
    return redirect(url_for('users'))

@app.route('/users/<int:uid>/delete', methods=['POST'])
def delete_user(uid):
    user = User.query.get_or_404(uid)
    current = get_current_user()
    if user.id == current.id:
        flash('Không thể xóa tài khoản đang đăng nhập', 'warning')
        return redirect(url_for('users'))
    name = user.username
    db.session.delete(user)
    db.session.commit()
    flash(f'Đã xóa tài khoản {name}', 'success')
    return redirect(url_for('users'))

@app.route('/company', methods=['GET', 'POST'])
def company_settings():
    companies = ensure_company_profiles()
    edit_id = request.args.get('id', type=int)
    if request.method == 'POST':
        cid = request.form.get('company_id', type=int)
        row = CompanyProfile.query.get_or_404(cid)
        try:
            apply_company_form(row, request.form, request.files)
        except ValueError as e:
            if str(e) == 'logo':
                return redirect(url_for('company_settings', id=cid))
            flash(str(e), 'warning')
            return redirect(url_for('company_settings', id=cid))
        if request.form.get('set_default') == '1':
            CompanyProfile.query.update({CompanyProfile.is_default: False})
            row.is_default = True
        db.session.commit()
        flash('Đã lưu thông tin công ty', 'success')
        return redirect(url_for('company_settings', id=row.id))
    if not edit_id and companies:
        active = session.get(ACTIVE_COMPANY_SESSION_KEY)
        edit_id = active if any(c.id == active for c in companies) else companies[0].id
    row = CompanyProfile.query.get(edit_id) if edit_id else None
    profile = company_profile_to_dict(row) if row else {}
    banks = fetch_vietqr_banks()
    return render_template(
        'company.html',
        companies=companies,
        profile=profile,
        row=row,
        edit_id=edit_id,
        banks=banks,
        bank_value_set=bank_value_set(banks),
    )

@app.route('/company/new', methods=['POST'])
def company_create():
    name = request.form.get('name', '').strip()
    if not name:
        flash('Vui lòng nhập tên công ty mới', 'warning')
        return redirect(url_for('company_settings'))
    row = CompanyProfile(name=name, short_name=request.form.get('short_name', '').strip())
    if CompanyProfile.query.count() == 0:
        row.is_default = True
    db.session.add(row)
    db.session.commit()
    session[ACTIVE_COMPANY_SESSION_KEY] = row.id
    flash(f'Đã thêm công ty {name}', 'success')
    return redirect(url_for('company_settings', id=row.id))

@app.route('/company/<int:cid>/delete', methods=['POST'])
def company_delete(cid):
    row = CompanyProfile.query.get_or_404(cid)
    if CompanyProfile.query.count() <= 1:
        flash('Phải giữ ít nhất một công ty', 'warning')
        return redirect(url_for('company_settings', id=cid))
    delete_company_logo_files(row)
    was_default = row.is_default
    was_active = session.get(ACTIVE_COMPANY_SESSION_KEY) == cid
    name = row.name
    db.session.delete(row)
    db.session.commit()
    remaining = CompanyProfile.query.order_by(CompanyProfile.id).first()
    if remaining and (was_default or was_active):
        remaining.is_default = True
        session[ACTIVE_COMPANY_SESSION_KEY] = remaining.id
        db.session.commit()
    elif was_active:
        session.pop(ACTIVE_COMPANY_SESSION_KEY, None)
    flash(f'Đã xóa công ty {name}', 'success')
    return redirect(url_for('company_settings', id=remaining.id if remaining else None))

@app.route('/company/switch', methods=['POST'])
def switch_company():
    cid = request.form.get('company_id', type=int)
    if cid and CompanyProfile.query.get(cid):
        session[ACTIVE_COMPANY_SESSION_KEY] = cid
        c = CompanyProfile.query.get(cid)
        flash(f'Đang dùng công ty: {c.short_name or c.name}', 'secondary')
    back = request.form.get('_back') or request.referrer or url_for('dashboard')
    return redirect(back)

def dashboard_day_bounds():
    today = date.today()
    start = datetime.combine(today, datetime.min.time())
    end = datetime.combine(today, datetime.max.time())
    return start, end, today

def dashboard_nav_allowed(user, nav_key):
    if not user:
        return False
    if user.role == 'admin':
        return True
    return nav_key in ROLE_NAV_KEYS.get(user.role, set())

def total_supplier_payable():
    ensure_supplier_intake_schema()
    row = db.session.query(
        func.coalesce(func.sum(SupplierIntake.total_amount), 0),
        func.coalesce(func.sum(SupplierIntake.paid_amount), 0),
    ).first()
    if not row:
        return 0
    return max(0, int(row[0] or 0) - int(row[1] or 0))

def open_quotes_count():
    return Quote.query.filter(Quote.status.in_(['Mới tạo', 'Nháp'])).count()

def dashboard_debt_alert_rows(limit=6):
    rows = []
    for cid, agg in customer_debt_aggregates().items():
        remaining = agg.get('remaining', 0)
        if remaining <= 0:
            continue
        customer = db.session.get(Customer, cid)
        if not customer:
            continue
        status = customer_debt_status_from_agg(agg)
        rows.append({
            'customer': customer,
            'remaining': remaining,
            'unpaid_count': agg.get('unpaid_count', 0),
            'status': status,
            'status_key': customer_debt_status_key(status),
        })
    rows.sort(key=lambda r: (-1 if r['status'] == 'Quá hạn' else 0, -r['remaining']))
    return rows[:limit]

def dashboard_low_stock_products(limit=8):
    return (
        Product.query.filter(
            Product.is_active.is_(True),
            Product.stock_qty <= func.coalesce(Product.low_stock, 5),
        )
        .order_by(Product.stock_qty.asc(), Product.name.asc())
        .limit(limit)
        .all()
    )

def build_dashboard_context(user):
    ensure_payment_columns()
    ensure_supplier_intake_schema()
    start, end, today = dashboard_day_bounds()
    can = lambda key: dashboard_nav_allowed(user, key)

    orders_today = Order.query.filter(
        Order.created_at >= start,
        Order.created_at <= end,
    ).all()
    payments_today = Payment.query.filter(Payment.payment_date == today).all()
    supplier_paid_today = SupplierPayment.query.filter(
        SupplierPayment.payment_date == today,
    ).all()

    today_block = {
        'orders_count': len(orders_today),
        'orders_total': sum(o.total or 0 for o in orders_today),
        'collected': sum(p.amount or 0 for p in payments_today),
        'payment_count': len(payments_today),
        'quotes_count': Quote.query.filter(
            Quote.created_at >= start,
            Quote.created_at <= end,
        ).count(),
        'supplier_paid': sum(p.amount or 0 for p in supplier_paid_today),
        'supplier_payment_count': len(supplier_paid_today),
    }

    stk = stock_stats()
    overview = {
        'customers': Customer.query.count(),
        'products': Product.query.count(),
        'quotes': Quote.query.count(),
        'quotes_open': open_quotes_count(),
        'orders': Order.query.count(),
        'debt': total_outstanding_debt(),
        'supplier_debt': total_supplier_payable(),
        'low_stock': stk['low_stock'],
        'out_stock': stk['out_stock'],
        'pending_intake': pending_supplier_intake_count(),
        'suppliers': Supplier.query.filter_by(is_active=True).count(),
    }

    quick_actions = []
    action_defs = [
        ('quotes', 'quotes', 'Báo giá', 'bi-file-earmark-text'),
        ('orders', 'orders', 'Đơn hàng', 'bi-cart3'),
        ('customers', 'customers', 'Khách hàng', 'bi-people'),
        ('debt', 'debt', 'Công nợ KH', 'bi-wallet2'),
        ('products', 'products', 'Sản phẩm', 'bi-box-seam'),
        ('stock', 'stock', 'Tồn kho', 'bi-boxes'),
        ('suppliers', 'suppliers', 'Nhà cung cấp', 'bi-truck'),
        ('stock_intake_requests', 'stock_intake_requests', 'Yêu cầu nhập kho', 'bi-inbox'),
        ('contracts', 'contracts', 'Hợp đồng', 'bi-file-earmark-ruled'),
    ]
    for key, endpoint, label, icon in action_defs:
        if can(key):
            quick_actions.append({
                'key': key,
                'href': url_for(endpoint),
                'label': label,
                'icon': icon,
            })

    kpi_links = []
    if can('customers'):
        kpi_links.append({'key': 'customers', 'href': url_for('customers'), 'label': 'Khách hàng', 'value': overview['customers'], 'money': False})
    if can('products'):
        kpi_links.append({'key': 'products', 'href': url_for('products'), 'label': 'Sản phẩm', 'value': overview['products'], 'money': False})
    if can('quotes'):
        kpi_links.append({'key': 'quotes', 'href': url_for('quotes'), 'label': 'Báo giá mở', 'value': overview['quotes_open'], 'money': False, 'hint': f'{overview["quotes"]} tổng'})
    if can('orders'):
        kpi_links.append({'key': 'orders', 'href': url_for('orders'), 'label': 'Đơn hàng', 'value': overview['orders'], 'money': False})
    if can('debt'):
        kpi_links.append({'key': 'debt', 'href': url_for('debt'), 'label': 'Nợ phải thu', 'value': overview['debt'], 'money': True})
    if can('supplier_debt') or can('suppliers'):
        kpi_links.append({
            'key': 'supplier_debt',
            'href': url_for('suppliers'),
            'label': 'Nợ phải trả NCC',
            'value': overview['supplier_debt'],
            'money': True,
        })
    if can('stock'):
        kpi_links.append({'key': 'stock', 'href': url_for('stock', tab='inventory', status='low'), 'label': 'Tồn thấp', 'value': overview['low_stock'], 'money': False, 'hint': f'{overview["out_stock"]} hết'})
    if can('stock_intake_requests'):
        kpi_links.append({
            'key': 'stock_intake_requests',
            'href': url_for('stock_intake_requests'),
            'label': 'Chờ nhập kho',
            'value': overview['pending_intake'],
            'money': False,
            'warn': overview['pending_intake'] > 0,
        })

    pending_intakes = []
    if can('stock_intake_requests'):
        ensure_supplier_intake_schema()
        for intake in (
            SupplierIntake.query.filter_by(status=SUPPLIER_INTAKE_STATUS_PENDING)
            .order_by(SupplierIntake.created_at.desc())
            .limit(8)
            .all()
        ):
            pending_intakes.append({
                'intake': intake,
                'supplier': intake.supplier,
                'balance': supplier_intake_balance(intake),
            })

    return {
        'today': today,
        'today_iso': today.isoformat(),
        'today_block': today_block,
        'overview': overview,
        'kpi_links': kpi_links,
        'quick_actions': quick_actions,
        'debt_alerts': dashboard_debt_alert_rows() if can('debt') else [],
        'low_stock_products': dashboard_low_stock_products() if can('stock') or can('products') else [],
        'pending_intakes': pending_intakes,
        'latest_quotes': Quote.query.order_by(Quote.created_at.desc()).limit(5).all() if can('quotes') else [],
        'latest_orders': Order.query.order_by(Order.created_at.desc()).limit(5).all() if can('orders') else [],
        'payments_today': sorted(payments_today, key=lambda p: p.id, reverse=True)[:8] if can('debt') else [],
        'show': {
            'today_sales': can('orders') or can('quotes') or can('debt'),
            'today_warehouse': can('stock_intake_requests') or can('suppliers'),
            'debt_alerts': can('debt'),
            'low_stock': can('stock') or can('products'),
            'pending_intake': can('stock_intake_requests'),
            'latest_quotes': can('quotes'),
            'latest_orders': can('orders'),
            'payments_today': can('debt'),
        },
    }

@app.route('/')
def dashboard():
    user = get_current_user()
    ctx = build_dashboard_context(user)
    return render_template(
        'dashboard.html',
        dash=ctx,
        is_admin=user and user.role == 'admin',
        user_role=user.role if user else '',
    )

@app.route('/dashboard/clear-data', methods=['POST'])
def dashboard_clear_data():
    user = get_current_user()
    if not user or user.role != 'admin':
        flash('Chỉ quản trị viên mới được xóa toàn bộ dữ liệu', 'danger')
        return redirect(url_for('dashboard'))
    if request.form.get('confirm_ack') != '1':
        flash('Vui lòng tick xác nhận bạn hiểu hành động này không thể hoàn tác', 'warning')
        return redirect(url_for('dashboard'))
    if (request.form.get('confirm_text') or '').strip().upper() != 'XOA HET':
        flash('Nhập đúng chữ XOA HET (viết hoa, có dấu cách) để xác nhận', 'warning')
        return redirect(url_for('dashboard'))
    password = request.form.get('password') or ''
    if not password or not check_password_hash(user.password_hash, password):
        flash('Mật khẩu đăng nhập không đúng', 'danger')
        return redirect(url_for('dashboard'))
    counts = clear_all_business_data()
    if has_request_context():
        session.pop(ACTIVE_COMPANY_SESSION_KEY, None)
        session.pop(CONTRACT_DRAFT_SESSION_KEY, None)
        session.pop(IMPORT_PREVIEW_SESSION_KEY, None)
        session.pop(CATEGORY_IMPORT_SESSION_KEY, None)
    total_rows = sum(counts.values())
    flash(
        f'Đã xóa toàn bộ dữ liệu ({total_rows} bản ghi nghiệp vụ). '
        f'Tài khoản đăng nhập được giữ nguyên; đã tạo lại khách vãng lai và công ty mặc định.',
        'success',
    )
    return redirect(url_for('dashboard'))

CUSTOMER_SORT_FIELDS = {'name', 'tax_code', 'representative', 'phone', 'address', 'created_at', 'customer_type'}

def ensure_customer_columns():
    from sqlalchemy import inspect, text
    try:
        cols = {c['name'] for c in inspect(db.engine).get_columns('customer')}
    except Exception:
        return
    with db.engine.begin() as conn:
        if 'customer_type' not in cols:
            conn.execute(text("ALTER TABLE customer ADD COLUMN customer_type VARCHAR(20) DEFAULT 'company'"))
        if 'id_card' not in cols:
            conn.execute(text("ALTER TABLE customer ADD COLUMN id_card VARCHAR(32) DEFAULT ''"))

def parse_customer_type():
    t = request.form.get('customer_type', 'company').strip()
    return t if t in ('company', 'individual') else 'company'

WALKIN_CUSTOMER_TYPE = 'walkin'
WALKIN_CUSTOMER_NAME = 'Khách vãng lai'

def ensure_walkin_customer():
    ensure_customer_columns()
    c = Customer.query.filter_by(customer_type=WALKIN_CUSTOMER_TYPE).first()
    if not c:
        c = Customer(
            name=WALKIN_CUSTOMER_NAME,
            customer_type=WALKIN_CUSTOMER_TYPE,
            note='Hệ thống — báo giá không chọn khách hàng',
        )
        db.session.add(c)
        db.session.commit()
    return c

def is_walkin_customer(customer):
    if customer is None:
        return True
    return getattr(customer, 'customer_type', '') == WALKIN_CUSTOMER_TYPE

def quote_walkin_display_name_from_form():
    return (request.form.get('walkin_display_name') or '').strip()

def quote_customer_label(quote_or_customer):
    if hasattr(quote_or_customer, 'customer'):
        customer = quote_or_customer.customer
        quote = quote_or_customer
    else:
        customer = quote_or_customer
        quote = None
    if is_walkin_customer(customer):
        custom = ''
        if quote is not None:
            custom = (getattr(quote, 'walkin_display_name', None) or '').strip()
        return custom or WALKIN_CUSTOMER_NAME
    return customer.name if customer else WALKIN_CUSTOMER_NAME

def customers_for_select():
    ensure_walkin_customer()
    return Customer.query.filter(Customer.customer_type != WALKIN_CUSTOMER_TYPE).order_by(Customer.name).all()

def customers_for_order_filter():
    """Dropdown lọc đơn hàng — gồm cả Khách vãng lai (đơn từ báo giá không chọn KH)."""
    ensure_walkin_customer()
    walkin = Customer.query.filter_by(customer_type=WALKIN_CUSTOMER_TYPE).first()
    others = Customer.query.filter(Customer.customer_type != WALKIN_CUSTOMER_TYPE).order_by(Customer.name).all()
    return ([walkin] if walkin else []) + others

def resolve_quote_customer_id(raw):
    if raw is None or str(raw).strip() == '':
        return ensure_walkin_customer().id
    try:
        cid = int(raw)
    except (TypeError, ValueError):
        return ensure_walkin_customer().id
    c = Customer.query.get(cid)
    if not c:
        return ensure_walkin_customer().id
    return cid

def set_quote_session_customer(quote):
    if quote.customer_id:
        session['debt_customer_id'] = quote.customer_id
        session['orders_customer_id'] = quote.customer_id

def customer_is_company(customer):
    return getattr(customer, 'customer_type', 'company') != 'individual'

def customer_type_label(customer):
    t = getattr(customer, 'customer_type', '')
    if t == 'individual':
        return 'Cá nhân'
    if t == WALKIN_CUSTOMER_TYPE:
        return 'Vãng lai'
    return 'Công ty'

app.jinja_env.globals['customer_is_company'] = customer_is_company
app.jinja_env.globals['customer_type_label'] = customer_type_label
app.jinja_env.globals['is_walkin_customer'] = is_walkin_customer
app.jinja_env.globals['quote_customer_label'] = quote_customer_label
app.jinja_env.globals['contract_signed_scan_url'] = contract_signed_scan_url
app.jinja_env.globals['signed_scan_is_pdf'] = signed_scan_is_pdf
VIETQR_BANKS_URL = 'https://api.vietqr.io/v2/banks'
BANKS_CACHE = {'at': 0, 'data': []}
BANKS_CACHE_TTL = 86400

def fetch_vietqr_banks():
    now = time.time()
    if BANKS_CACHE['data'] and now - BANKS_CACHE['at'] < BANKS_CACHE_TTL:
        return BANKS_CACHE['data']
    try:
        with urlopen(VIETQR_BANKS_URL, timeout=15) as resp:
            payload = json.loads(resp.read().decode('utf-8'))
        if payload.get('code') == '00':
            banks = sorted(payload.get('data', []), key=lambda b: (b.get('shortName') or '').lower())
            BANKS_CACHE['data'] = banks
            BANKS_CACHE['at'] = now
            return banks
    except Exception:
        pass
    return BANKS_CACHE['data']

def bank_value_set(banks):
    values = set()
    for bank in banks:
        if bank.get('name'):
            values.add(bank['name'])
        if bank.get('shortName'):
            values.add(bank['shortName'])
    return values

def customer_list_params():
    q = request.args.get('q', '').strip()
    per_page = request.args.get('per_page', 10, type=int)
    per_page = per_page if per_page in (10, 25, 50, 100) else 10
    page = request.args.get('page', 1, type=int)
    sort = request.args.get('sort', 'created_at')
    sort = sort if sort in CUSTOMER_SORT_FIELDS else 'created_at'
    order = 'asc' if request.args.get('order') == 'asc' else 'desc'
    return q, per_page, page, sort, order

def customer_sort_url(field, q, per_page, sort, order):
    next_order = 'desc' if sort == field and order == 'asc' else 'asc'
    if sort != field:
        next_order = 'asc'
    return url_for('customers', q=q, sort=field, order=next_order, per_page=per_page, page=1)

def validate_customer_form():
    ctype = parse_customer_type()
    name = request.form.get('name', '').strip()
    if not name:
        return False, 'Vui lòng nhập tên khách hàng'
    if ctype == 'company':
        if not request.form.get('tax_code', '').strip():
            return False, 'Vui lòng nhập mã số thuế'
    elif not request.form.get('phone', '').strip():
        return False, 'Vui lòng nhập số điện thoại'
    return True, ''

def apply_customer_form(customer):
    ctype = parse_customer_type()
    customer.customer_type = ctype
    customer.name = request.form.get('name', '').strip()
    customer.address = request.form.get('address', '').strip()
    customer.phone = request.form.get('phone', '').strip()
    customer.email = request.form.get('email', '').strip()
    customer.bank_account = request.form.get('bank_account', '').strip()
    customer.bank_name = request.form.get('bank_name', '').strip()
    customer.note = request.form.get('note', '').strip()
    if ctype == 'company':
        customer.tax_code = request.form.get('tax_code', '').strip()
        customer.representative = request.form.get('representative', '').strip()
        customer.position = request.form.get('position', '').strip()
        customer.id_card = ''
    else:
        customer.id_card = request.form.get('id_card', '').strip()
        customer.tax_code = ''
        customer.representative = ''
        customer.position = ''

def customers_redirect_after_save(message=None, category='success', new_customer_id=None):
    if message:
        flash(message, category)
    if request.form.get('_from') == 'quotes' and new_customer_id:
        return redirect(url_for('quotes', quote_customer_id=new_customer_id))
    return redirect(url_for(
        'customers',
        q=request.form.get('_q', ''),
        page=request.form.get('_page', 1, type=int) or 1,
        per_page=request.form.get('_per_page', 10, type=int) or 10,
        sort=request.form.get('_sort', 'created_at'),
        order=request.form.get('_order', 'desc'),
    ))

@app.route('/customers', methods=['GET', 'POST'])
def customers():
    ensure_customer_columns()
    ensure_contract_columns()
    if request.method == 'POST':
        ok, msg = validate_customer_form()
        if not ok:
            flash(msg, 'warning')
            return redirect(url_for('customers', open_add='customer'))
        c = Customer()
        apply_customer_form(c)
        db.session.add(c)
        db.session.commit()
        return customers_redirect_after_save('Đã thêm khách hàng', new_customer_id=c.id)
    q, per_page, page, sort, order = customer_list_params()
    query = Customer.query.filter(Customer.customer_type != WALKIN_CUSTOMER_TYPE)
    if q:
        like = f'%{q}%'
        query = query.filter(or_(
            Customer.name.ilike(like),
            Customer.tax_code.ilike(like),
            Customer.phone.ilike(like),
            Customer.id_card.ilike(like),
        ))
    sort_col = getattr(Customer, sort)
    query = query.order_by(sort_col.asc() if order == 'asc' else sort_col.desc())
    pagination = query.paginate(page=page, per_page=per_page, error_out=False)
    customer_ids = [c.id for c in pagination.items]
    framework_contracts, customer_quotes, customer_orders = customer_list_extras(customer_ids)
    banks = fetch_vietqr_banks()
    list_args = {'q': q, 'per_page': per_page, 'sort': sort, 'order': order}
    from_quotes = request.args.get('from') == 'quotes'
    open_add_modal = from_quotes or request.args.get('open_add') == 'customer'
    return render_template(
        'customers.html',
        pagination=pagination,
        customers=pagination.items,
        framework_contracts=framework_contracts,
        customer_quotes=customer_quotes,
        customer_orders=customer_orders,
        q=q,
        per_page=per_page,
        sort=sort,
        order=order,
        list_args=list_args,
        banks=banks,
        bank_value_set=bank_value_set(banks),
        sort_url=lambda field: customer_sort_url(field, q, per_page, sort, order),
        from_quotes=from_quotes,
        open_add_modal=open_add_modal,
    )

@app.route('/customers/options.json')
def customer_options_json():
    rows = customers_for_select()
    return {
        'customers': [
            {'id': c.id, 'name': c.name, 'phone': c.phone or ''}
            for c in rows
        ],
    }

@app.route('/customers/<int:cid>/orders.json')
def customer_orders_json(cid):
    Customer.query.get_or_404(cid)
    rows = (
        Order.query.filter_by(customer_id=cid)
        .filter(Order.status != 'Đã hủy')
        .order_by(Order.created_at.desc())
        .limit(200)
        .all()
    )
    return {
        'orders': [
            {
                'id': o.id,
                'order_code': o.order_code,
                'status': o.status,
                'total': o.total or 0,
                'created_at': o.created_at.strftime('%d/%m/%Y') if o.created_at else '',
            }
            for o in rows
        ],
    }

@app.route('/products/<int:pid>/suppliers.json')
def product_suppliers_json(pid):
    ensure_product_supplier_schema()
    product = Product.query.get_or_404(pid)
    links = product_supplier_links(product)
    linked = []
    default_id = None
    for link in links:
        if not link.supplier:
            continue
        linked.append({
            'id': link.supplier_id,
            'code': link.supplier.code,
            'name': link.supplier.name,
            'cost_price': link.cost_price or 0,
            'supplier_sku': link.supplier_sku or '',
            'is_primary': bool(link.is_primary),
        })
    if linked:
        primary = next((x for x in linked if x['is_primary']), None)
        if primary:
            default_id = primary['id']
        elif len(linked) == 1:
            default_id = linked[0]['id']
    active = [
        {'id': s.id, 'code': s.code, 'name': s.name}
        for s in active_suppliers()
    ]
    return {'linked': linked, 'active': active, 'default_supplier_id': default_id}

@app.route('/suppliers', methods=['GET', 'POST'])
def suppliers():
    ensure_stock_movement_columns()
    ensure_supplier_intake_schema()
    ensure_supplier_intake_line_columns()
    ensure_product_columns()
    if request.method == 'POST':
        ok, msg = validate_supplier_form()
        if not ok:
            flash(msg, 'warning')
            return redirect(url_for('suppliers', open_add='supplier'))
        s = Supplier()
        apply_supplier_form(s, is_new=True)
        db.session.add(s)
        db.session.commit()
        return suppliers_redirect_after_save('Đã thêm nhà cung cấp')
    q, per_page, page, sort, order = supplier_list_params()
    query = Supplier.query
    if q:
        like = f'%{q}%'
        query = query.filter(or_(
            Supplier.name.ilike(like),
            Supplier.code.ilike(like),
            Supplier.tax_code.ilike(like),
            Supplier.phone.ilike(like),
            Supplier.contact_person.ilike(like),
        ))
    sort_col = getattr(Supplier, sort)
    query = query.order_by(sort_col.asc() if order == 'asc' else sort_col.desc())
    pagination = query.paginate(page=page, per_page=per_page, error_out=False)
    supplier_ids = [s.id for s in pagination.items]
    inbound_stats = supplier_inbound_counts(supplier_ids)
    supplier_products = supplier_products_map(supplier_ids)
    product_counts = supplier_product_counts(supplier_ids)
    intake_board = build_supplier_intake_board(supplier_ids)
    intake_line_map = supplier_intake_lines_map([
        i.id for sid in supplier_ids for i in intake_board['intakes'].get(sid, [])
    ])
    list_args = {'q': q, 'per_page': per_page, 'sort': sort, 'order': order}
    intake_import_payload = get_supplier_intake_import_session()
    intake_preview_sid = None
    if intake_import_payload and request.args.get('intake_import_preview') == '1':
        intake_preview_sid = intake_import_payload.get('supplier_id')
    open_detail_sid = request.args.get('open_detail_sid', type=int)
    if open_detail_sid:
        tab = request.args.get('detail_tab') or ''
        kwargs = {'tab': tab} if tab in ('intakes', 'catalog', 'payments') else {}
        if request.args.get('intake') == '1':
            kwargs['intake'] = '1'
        return redirect(url_for('supplier_detail', sid=open_detail_sid, **kwargs))
    open_intake_sid = request.args.get('open_intake_sid', type=int) or request.args.get('intake_sid', type=int)
    if open_intake_sid:
        kwargs = {'intake': '1'}
        if request.args.get('intake_import_preview') == '1':
            kwargs['intake_import_preview'] = '1'
        if request.args.get('open_quick_product') == '1':
            kwargs['quick_product'] = '1'
        return redirect(url_for('supplier_detail', sid=open_intake_sid, **kwargs))
    return render_template(
        'suppliers.html',
        pagination=pagination,
        suppliers=pagination.items,
        inbound_stats=inbound_stats,
        supplier_products=supplier_products,
        product_counts=product_counts,
        intake_board=intake_board,
        intake_line_map=intake_line_map,
        stats=supplier_stats(),
        q=q,
        per_page=per_page,
        sort=sort,
        order=order,
        list_args=list_args,
        sort_url=lambda field: supplier_sort_url(field, q, per_page, sort, order),
        today_iso=date.today().isoformat(),
        open_add_modal=request.args.get('open_add') == 'supplier',
        intake_import_payload=intake_import_payload,
        intake_preview_sid=intake_preview_sid,
        pending_stock_intake_count=pending_supplier_intake_count(),
    )

@app.route('/suppliers/<int:sid>')
def supplier_detail(sid):
    ensure_stock_movement_columns()
    ensure_supplier_intake_schema()
    ensure_supplier_intake_line_columns()
    ensure_product_columns()
    ensure_product_supplier_schema()
    supplier = Supplier.query.get_or_404(sid)
    intake_board = build_supplier_intake_board([sid])
    intakes = intake_board['intakes'].get(sid, [])
    intake_line_map = supplier_intake_lines_map([i.id for i in intakes])
    catalog_links = supplier_products_map([sid]).get(sid, [])
    tab = (request.args.get('tab') or 'intakes').strip()
    if tab not in ('intakes', 'catalog', 'payments'):
        tab = 'intakes'
    intake_import_payload = get_supplier_intake_import_session()
    show_intake_preview = (
        intake_import_payload
        and intake_import_payload.get('supplier_id') == sid
        and request.args.get('intake_import_preview') == '1'
    )
    list_ctx = {'from_detail': True}
    return render_template(
        'supplier_detail.html',
        supplier=supplier,
        catalog_links=catalog_links,
        intake_board=intake_board,
        intake_line_map=intake_line_map,
        list_ctx=list_ctx,
        active_tab=tab,
        back_url=app_back_url('suppliers'),
        today_iso=date.today().isoformat(),
        intake_import_payload=intake_import_payload,
        show_intake_preview=show_intake_preview,
    )

@app.route('/suppliers/options.json')
def supplier_options_json():
    rows = active_suppliers()
    return {
        'suppliers': [
            {'id': s.id, 'code': s.code, 'name': s.name, 'phone': s.phone or ''}
            for s in rows
        ],
    }

@app.route('/products/options.json')
def product_options_json():
    rows = Product.query.filter(Product.is_active.is_(True)).order_by(Product.name).all()
    return {'products': [product_options_item(p) for p in rows]}

@app.route('/suppliers/<int:sid>/products/quick-create', methods=['POST'])
def supplier_quick_create_product(sid):
    ensure_category_brand_schema()
    ensure_product_columns()
    ensure_product_supplier_schema()
    supplier = Supplier.query.get_or_404(sid)
    sku = request.form.get('sku', '').strip()
    name = request.form.get('name', '').strip()
    cost = parse_int(request.form.get('cost_price'), 0)
    wants_json = (
        request.headers.get('X-Requested-With') == 'XMLHttpRequest'
        or request.accept_mimetypes.best_match(['application/json', 'text/html']) == 'application/json'
    )
    if not sku or not name:
        msg = 'Vui lòng nhập SKU và tên sản phẩm để thêm vào danh mục cửa hàng.'
        if wants_json:
            return jsonify({'ok': False, 'error': msg}), 400
        flash(msg, 'warning')
        return suppliers_redirect_after_save('', open_intake_sid=sid, open_quick_product=1)
    existing = Product.query.filter_by(sku=sku).first()
    if existing:
        if not existing.is_active:
            existing.is_active = True
        existing.name = name
        apply_product_unit_from_form(existing)
        upsert_product_supplier_link(existing, supplier, cost, '', '')
        if cost > 0:
            existing.cost_price = base_unit_cost_from_purchase_cost(existing, cost)
        db.session.commit()
        if wants_json:
            return jsonify({'ok': True, 'product': product_options_item(existing), 'linked_existing': True})
        flash(f'SKU {sku} đã có — đã liên kết NCC {supplier.code}', 'info')
        return suppliers_redirect_after_save('', open_intake_sid=sid)
    p = Product(sku=sku, name=name, stock=0, is_active=True)
    apply_product_unit_from_form(p)
    if cost > 0:
        p.cost_price = base_unit_cost_from_purchase_cost(p, cost)
    db.session.add(p)
    db.session.flush()
    upsert_product_supplier_link(p, supplier, cost, '', '')
    db.session.commit()
    if wants_json:
        return jsonify({'ok': True, 'product': product_options_item(p)})
    flash(f'Đã thêm sản phẩm {sku} vào danh mục — chọn lại trên dòng nhập', 'success')
    return suppliers_redirect_after_save('', open_intake_sid=sid)

@app.route('/suppliers/<int:sid>/products', methods=['POST'])
def supplier_products_intake(sid):
    ensure_category_brand_schema()
    ensure_product_supplier_schema()
    ensure_supplier_intake_schema()
    ensure_supplier_intake_line_columns()
    supplier = Supplier.query.get_or_404(sid)
    incomplete = supplier_manual_intake_incomplete_rows()
    if incomplete:
        rows = ', '.join(str(n) for n in incomplete[:5])
        extra = f' (+{len(incomplete) - 5} dòng)' if len(incomplete) > 5 else ''
        flash(
            f'Dòng {rows}{extra}: sản phẩm chưa có trong cửa hàng — '
            'bấm 「Thêm sản phẩm mới」 trong ô chọn SP hoặc dùng Import Excel.',
            'warning',
        )
        return suppliers_redirect_after_save('', open_intake_sid=sid)
    try:
        line_count, _ = apply_supplier_products_intake(supplier)
        if line_count <= 0:
            flash(
                'Chưa có dòng nhập hợp lệ. Chọn sản phẩm trong danh mục (hoặc thêm mới) '
                'và nhập số lượng hoặc giá nhập.',
                'warning',
            )
            return suppliers_redirect_after_save('', open_intake_sid=sid)
        db.session.commit()
    except Exception as e:
        db.session.rollback()
        app.logger.exception('supplier_products_intake failed')
        flash(f'Không gửi được yêu cầu nhập kho: {e}', 'danger')
        return suppliers_redirect_after_save('', open_intake_sid=sid)
    return suppliers_redirect_after_save(
        f'Đã gửi yêu cầu nhập kho ({line_count} dòng nhập tay). '
        'Vào menu Kho vận → Yêu cầu nhập kho để phân bổ kg cho khách và duyệt.',
    )

def _supplier_catalog_redirect(sid, message=None, category='success'):
    if _supplier_redirect_from_detail():
        if message:
            flash(message, category)
        return redirect(url_for('supplier_detail', sid=sid, tab='catalog'))
    return suppliers_redirect_after_save(message, category, open_detail_sid=sid, detail_tab='catalog')

def _supplier_catalog_list_ctx_from_request():
    src = request.form if request.method == 'POST' else request.args
    return {
        'q': src.get('_q', ''),
        'page': src.get('_page', 1, type=int) or 1,
        'per_page': src.get('_per_page', 10, type=int) or 10,
        'sort': src.get('_sort', 'created_at'),
        'order': src.get('_order', 'desc'),
    }

def _supplier_catalog_cancel_url(sid, list_ctx):
    if list_ctx.get('from_detail') or request.args.get('_from_detail') == '1':
        return url_for('supplier_detail', sid=sid, tab='catalog')
    return url_for(
        'suppliers',
        open_detail_sid=sid,
        detail_tab='catalog',
        q=list_ctx.get('q', ''),
        page=list_ctx.get('page', 1),
        per_page=list_ctx.get('per_page', 10),
        sort=list_ctx.get('sort', 'created_at'),
        order=list_ctx.get('order', 'desc'),
    )

def supplier_catalog_create_product_from_form(supplier):
    """Tạo SP mới (hoặc liên kết SKU đã có) và gắn vào NCC."""
    ensure_category_brand_schema()
    ensure_product_columns()
    ensure_product_supplier_schema()
    sku = request.form.get('sku', '').strip()
    name = request.form.get('name', '').strip()
    cost = parse_int(request.form.get('cost_price'), 0)
    supplier_sku = (request.form.get('catalog_supplier_sku') or '').strip()
    if not sku or not name:
        return None, 'Vui lòng nhập SKU và tên sản phẩm.'
    if cost <= 0:
        return None, 'Vui lòng nhập giá nhập mặc định.'
    existing = Product.query.filter_by(sku=sku).first()
    if existing:
        if not existing.is_active:
            existing.is_active = True
        existing.name = name
        apply_product_unit_from_form(existing)
        if ProductSupplier.query.filter_by(product_id=existing.id, supplier_id=supplier.id).first():
            return None, f'SKU {sku} đã có trong danh mục NCC này.'
        upsert_product_supplier_link(existing, supplier, cost, supplier_sku, '')
        if cost > 0:
            existing.cost_price = base_unit_cost_from_purchase_cost(existing, cost)
        db.session.commit()
        return existing, f'Đã liên kết SKU {sku} với NCC {supplier.code}.'
    p = Product(sku=sku, name=name, stock=0, is_active=True)
    apply_product_unit_from_form(p)
    if cost > 0:
        p.cost_price = base_unit_cost_from_purchase_cost(p, cost)
    db.session.add(p)
    db.session.flush()
    upsert_product_supplier_link(p, supplier, cost, supplier_sku, '')
    db.session.commit()
    return p, f'Đã thêm {sku} — {name} vào sản phẩm cung cấp của {supplier.name}.'

def _supplier_catalog_new_form_data(source=None):
    src = source or request.form
    return {
        'sku': (src.get('sku') or '').strip(),
        'name': (src.get('name') or '').strip(),
        'cost_price': src.get('cost_price', ''),
        'catalog_supplier_sku': (src.get('catalog_supplier_sku') or '').strip(),
        'unit_conversion_enabled': src.get('unit_conversion_enabled') == '1',
        'base_unit': (src.get('base_unit') or 'kg').strip(),
        'purchase_unit': (src.get('purchase_unit') or 'Thùng').strip(),
        'conversion_factor': src.get('conversion_factor') or '15',
        'lot_enabled': src.get('lot_unit_enabled') == '1' or src.get('lot_enabled') == '1',
        'lot_unit': (src.get('lot_unit') or 'Lô').strip(),
        'lot_factor': src.get('lot_factor') or '10',
        'sale_unit_mode': (src.get('sale_unit_mode') or 'purchase').strip(),
    }

@app.route('/suppliers/<int:sid>/catalog/new', methods=['GET', 'POST'])
def supplier_catalog_new(sid):
    ensure_product_columns()
    supplier = Supplier.query.get_or_404(sid)
    list_ctx = _supplier_catalog_list_ctx_from_request()
    if request.args.get('_from_detail') == '1' or request.form.get('_from_detail') == '1':
        list_ctx['from_detail'] = True
    cancel_url = _supplier_catalog_cancel_url(sid, list_ctx)
    default_form = {
        'sku': '',
        'name': '',
        'cost_price': '',
        'catalog_supplier_sku': '',
        'unit_conversion_enabled': True,
        'base_unit': 'kg',
        'purchase_unit': 'Thùng',
        'conversion_factor': '15',
        'lot_enabled': False,
        'lot_unit': 'Lô',
        'lot_factor': '10',
        'sale_unit_mode': 'purchase',
    }
    if request.method == 'POST':
        product, msg = supplier_catalog_create_product_from_form(supplier)
        if not product:
            flash(msg, 'warning')
            return render_template(
                'supplier_catalog_add.html',
                supplier=supplier,
                cancel_url=cancel_url,
                list_ctx=list_ctx,
                form_data=_supplier_catalog_new_form_data(),
            )
        flash(msg, 'success' if 'Đã thêm' in msg or 'liên kết' in msg else 'info')
        return redirect(cancel_url)
    return render_template(
        'supplier_catalog_add.html',
        supplier=supplier,
        cancel_url=cancel_url,
        list_ctx=list_ctx,
        form_data=default_form,
    )

def supplier_catalog_link_payload(link):
    product = link.product
    pu = product_purchase_unit(product) if product else ''
    bu = product_base_unit(product) if product else ''
    has_conv = product_has_unit_conversion(product) if product else False
    conv_text = ''
    if has_conv and product:
        conv_text = (
            f'1 {pu} = {format_qty_display(product_conversion_factor(product))} {bu}'
        )
    return {
        'product_id': link.product_id,
        'supplier_id': link.supplier_id,
        'cost_price': link.cost_price or 0,
        'cost_display': money_plain(link.cost_price or 0),
        'supplier_sku': link.supplier_sku or '',
        'note': link.note or '',
        'product': {
            'id': product.id,
            'sku': product.sku,
            'name': product.name,
            'image_url': product_image_url(product) or '',
            'detail_url': url_for('product_detail', pid=product.id),
            'edit_url': url_for('product_edit', pid=product.id),
        } if product else None,
        'purchase_unit': pu,
        'base_unit': bu,
        'has_conversion': has_conv,
        'conversion_text': conv_text,
        'base_cost_hint': (
            money_plain(base_unit_cost_from_purchase_cost(product, link.cost_price or 0))
            if has_conv and product else ''
        ),
    }

@app.route('/suppliers/<int:sid>/catalog', methods=['POST'])
def supplier_catalog_add(sid):
    ensure_product_supplier_schema()
    supplier = Supplier.query.get_or_404(sid)
    wants_json = request.headers.get('X-Requested-With') == 'XMLHttpRequest'
    product_id = parse_int(request.form.get('product_id'), 0)
    if not product_id:
        msg = 'Chọn sản phẩm cần gắn với nhà cung cấp.'
        if wants_json:
            return jsonify({'ok': False, 'error': msg}), 400
        return _supplier_catalog_redirect(sid, msg, 'warning')
    product = Product.query.get(product_id)
    if not product:
        msg = 'Sản phẩm không tồn tại.'
        if wants_json:
            return jsonify({'ok': False, 'error': msg}), 400
        return _supplier_catalog_redirect(sid, msg, 'warning')
    cost = parse_int(request.form.get('catalog_cost_price'), 0)
    supplier_sku = (request.form.get('catalog_supplier_sku') or '').strip()
    note = (request.form.get('catalog_note') or '').strip()
    existing_link = ProductSupplier.query.filter_by(product_id=product_id, supplier_id=sid).first()
    try:
        if existing_link:
            link = existing_link
            if cost:
                link.cost_price = cost
                sync_product_cost_from_supplier_purchase_price(product, cost)
            if supplier_sku:
                link.supplier_sku = supplier_sku
            if note:
                link.note = note
            updated = True
        else:
            link = upsert_product_supplier_link(product, supplier, cost, supplier_sku, note)
            db.session.flush()
            if cost:
                sync_product_cost_from_supplier_purchase_price(product, cost)
            updated = False
        db.session.commit()
    except Exception as e:
        db.session.rollback()
        app.logger.exception('supplier_catalog_add failed')
        msg = f'Không thêm được sản phẩm: {e}'
        if wants_json:
            return jsonify({'ok': False, 'error': msg}), 500
        return _supplier_catalog_redirect(sid, msg, 'danger')
    if updated:
        msg = f'Đã cập nhật {product.sku} trong danh mục NCC {supplier.name}.'
    else:
        msg = f'Đã thêm {product.sku} — {product.name} vào sản phẩm cung cấp của {supplier.name}.'
    if wants_json:
        link = ProductSupplier.query.filter_by(product_id=product_id, supplier_id=sid).first()
        return jsonify({
            'ok': True,
            'message': msg,
            'link': supplier_catalog_link_payload(link),
            'updated': updated,
        })
    return _supplier_catalog_redirect(sid, msg)

@app.route('/suppliers/<int:sid>/catalog/<int:pid>/remove', methods=['POST'])
def supplier_catalog_remove(sid, pid):
    ensure_product_supplier_schema()
    Supplier.query.get_or_404(sid)
    link = ProductSupplier.query.filter_by(supplier_id=sid, product_id=pid).first()
    if not link:
        return _supplier_catalog_redirect(sid, 'Liên kết sản phẩm không tồn tại.', 'warning')
    product = link.product
    db.session.delete(link)
    db.session.commit()
    label = f'{product.sku}' if product else 'Sản phẩm'
    return _supplier_catalog_redirect(sid, f'Đã gỡ {label} khỏi danh mục NCC.')

@app.route('/suppliers/<int:sid>/products/import/template')
def supplier_intake_import_template(sid):
    Supplier.query.get_or_404(sid)
    rows = [supplier_intake_sample_row()]
    fmt = (request.args.get('format') or 'xlsx').lower()
    if fmt == 'csv':
        return Response(
            build_supplier_intake_csv_text(rows),
            mimetype='text/csv; charset=utf-8',
            headers={'Content-Disposition': 'attachment; filename=mau-nhap-ncc.csv'},
        )
    try:
        buf = build_supplier_intake_xlsx_bytes(rows)
        return send_file(
            buf,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name='mau-nhap-ncc.xlsx',
        )
    except ImportError:
        return Response(
            build_supplier_intake_csv_text(rows),
            mimetype='text/csv; charset=utf-8',
            headers={'Content-Disposition': 'attachment; filename=mau-nhap-ncc.csv'},
        )

@app.route('/suppliers/<int:sid>/products/import/preview', methods=['POST'])
def supplier_intake_import_preview(sid):
    ensure_category_brand_schema()
    supplier = Supplier.query.get_or_404(sid)
    file_storage = request.files.get('import_file')
    if not file_storage or not file_storage.filename:
        flash('Vui lòng chọn file Excel hoặc CSV', 'warning')
        return suppliers_redirect_after_save('')
    update_existing = request.form.get('update_existing') == '1'
    try:
        rows = read_product_import_sheet(file_storage)
        preview = build_supplier_intake_import_preview(rows, update_existing=update_existing)
        save_supplier_intake_import_session(
            supplier.id, file_storage.filename, update_existing, preview,
        )
    except ValueError as e:
        flash(str(e), 'danger')
        return suppliers_redirect_after_save('')
    except Exception as e:
        flash(f'Không đọc được file: {e}', 'danger')
        return suppliers_redirect_after_save('')
    return suppliers_redirect_after_save(
        '',
        intake_import_preview=1,
        intake_sid=supplier.id,
    )

@app.route('/suppliers/<int:sid>/products/import/confirm', methods=['POST'])
def supplier_intake_import_confirm(sid):
    ensure_category_brand_schema()
    supplier = Supplier.query.get_or_404(sid)
    payload = get_supplier_intake_import_session()
    if not payload or payload.get('supplier_id') != supplier.id:
        flash('Phiên import đã hết hạn. Vui lòng tải file lại.', 'warning')
        return suppliers_redirect_after_save('')
    selected_ids = set(request.form.getlist('import_row_id'))
    update_existing = payload.get('update_existing', False)
    to_import = [
        r for r in payload.get('rows', [])
        if r.get('row_id') in selected_ids and r.get('importable')
    ]
    if not to_import:
        flash('Chưa chọn dòng nào để gửi yêu cầu', 'warning')
        return suppliers_redirect_after_save(
            '', intake_import_preview=1, intake_sid=supplier.id,
        )
    ref_code = request.form.get('ref_code', '').strip()
    intake_note = request.form.get('intake_note', '').strip()
    try:
        updated, stock_qty = apply_supplier_intake_from_import_rows(
            supplier, to_import, update_existing, ref_code, intake_note,
        )
    except ValueError as e:
        flash(str(e), 'danger')
        return suppliers_redirect_after_save(
            '', intake_import_preview=1, intake_sid=supplier.id,
        )
    clear_supplier_intake_import_session()
    db.session.commit()
    msg = (
        f'Đã gửi yêu cầu nhập kho ({updated} dòng từ Excel). '
        'Vào menu Kho vận → Yêu cầu nhập kho để phân bổ kg cho khách và duyệt.'
    )
    return suppliers_redirect_after_save(msg)

@app.route('/suppliers/<int:sid>/products/import/cancel', methods=['POST'])
def supplier_intake_import_cancel(sid):
    Supplier.query.get_or_404(sid)
    clear_supplier_intake_import_session()
    return suppliers_redirect_after_save('')

@app.route('/suppliers/<int:sid>/intakes/<int:iid>/payment', methods=['POST'])
def supplier_intake_payment(sid, iid):
    ensure_supplier_intake_schema()
    supplier = Supplier.query.get_or_404(sid)
    intake = SupplierIntake.query.filter_by(id=iid, supplier_id=supplier.id).first_or_404()
    balance = supplier_intake_balance(intake)
    if balance <= 0:
        flash('Phiếu nhập đã thanh toán đủ', 'warning')
        return suppliers_redirect_after_save('')
    amount = parse_int(request.form.get('amount'))
    if amount <= 0:
        flash('Số tiền không hợp lệ', 'danger')
        return suppliers_redirect_after_save('')
    receipt_path = save_supplier_payment_receipt(request.files.get('payment_receipt'))
    if not receipt_path:
        return suppliers_redirect_after_save('')
    pay_amount = min(amount, balance)
    if amount > balance:
        flash(
            f'Số tiền vượt còn nợ ({money(balance)}). Chỉ ghi nhận {money(pay_amount)}.',
            'warning',
        )
    method = request.form.get('method', 'Chuyển khoản').strip() or 'Chuyển khoản'
    note = request.form.get('note', '').strip()
    payment_date_raw = request.form.get('payment_date', '').strip()
    payment_date = date.today()
    if payment_date_raw:
        try:
            payment_date = datetime.strptime(payment_date_raw, '%Y-%m-%d').date()
        except ValueError:
            pass
    db.session.add(SupplierPayment(
        supplier_id=supplier.id,
        intake_id=intake.id,
        amount=pay_amount,
        method=method,
        note=note,
        payment_date=payment_date,
        receipt_path=receipt_path,
        receipt_uploaded_at=datetime.utcnow(),
    ))
    recalc_supplier_intake_paid(intake)
    db.session.commit()
    status = supplier_intake_payment_status(intake)
    flash(f'Đã ghi nhận thanh toán {money(pay_amount)} — {status}', 'success')
    return suppliers_redirect_after_save('')

@app.route('/suppliers/<int:sid>/intakes/<int:iid>/delete', methods=['POST'])
def delete_supplier_intake(sid, iid):
    ensure_supplier_intake_schema()
    ensure_stock_movement_columns()
    supplier = Supplier.query.get_or_404(sid)
    intake = SupplierIntake.query.filter_by(id=iid, supplier_id=supplier.id).first_or_404()
    code = delete_supplier_intake_record(intake)
    db.session.commit()
    return suppliers_redirect_after_save(f'Đã xóa phiếu nhập {code}')

@app.route('/suppliers/payments/<int:pid>/receipt')
def view_supplier_payment_receipt(pid):
    ensure_supplier_intake_schema()
    payment = SupplierPayment.query.get_or_404(pid)
    if not payment.receipt_path:
        flash('Chưa có ảnh bill', 'warning')
        return redirect(url_for('suppliers'))
    file_path = BASE_DIR / 'static' / payment.receipt_path
    if not file_path.is_file():
        flash('Không tìm thấy file bill', 'warning')
        return redirect(url_for('suppliers'))
    return send_file(file_path, mimetype=signed_scan_mimetype(file_path))

@app.route('/suppliers/<int:sid>/update', methods=['POST'])
def update_supplier(sid):
    ensure_stock_movement_columns()
    s = Supplier.query.get_or_404(sid)
    from_detail = request.form.get('_from_detail') == '1'
    ok, msg = validate_supplier_form()
    if not ok:
        flash(msg, 'warning')
        if from_detail:
            return redirect(url_for('supplier_detail', sid=sid))
        return suppliers_redirect_after_save('')
    apply_supplier_form(s)
    db.session.commit()
    if from_detail:
        flash('Đã cập nhật nhà cung cấp', 'success')
        return redirect(url_for('supplier_detail', sid=sid))
    return suppliers_redirect_after_save('Đã cập nhật nhà cung cấp')

@app.route('/suppliers/<int:sid>/delete', methods=['POST'])
def delete_supplier(sid):
    ensure_stock_movement_columns()
    s = Supplier.query.get_or_404(sid)
    linked = StockMovement.query.filter_by(supplier_id=s.id).count()
    product_links = ProductSupplier.query.filter_by(supplier_id=s.id).count()
    if linked or product_links:
        parts = []
        if product_links:
            parts.append(f'{product_links} sản phẩm')
        if linked:
            parts.append(f'{linked} phiếu nhập kho')
        flash('Không thể xóa — NCC đang liên kết: ' + ', '.join(parts), 'danger')
        return suppliers_redirect_after_save('')
    db.session.delete(s)
    db.session.commit()
    return suppliers_redirect_after_save('Đã xóa nhà cung cấp')

@app.route('/customers/<int:cid>/update', methods=['POST'])
def update_customer(cid):
    ensure_customer_columns()
    c = Customer.query.get_or_404(cid)
    if is_walkin_customer(c):
        flash('Không thể chỉnh sửa khách hàng hệ thống', 'warning')
        return customers_redirect_after_save('')
    ok, msg = validate_customer_form()
    if not ok:
        flash(msg, 'warning')
        return customers_redirect_after_save('')
    apply_customer_form(c)
    db.session.commit()
    return customers_redirect_after_save('Đã cập nhật khách hàng')

def contracts_create_url(customer_id=None, open_add=True):
    args = {}
    if open_add:
        args['open_add'] = 'contract'
    if customer_id:
        args['customer_id'] = customer_id
    return url_for('contracts', **args)


@app.route('/customers/<int:cid>/contract/preview')
def preview_customer_framework_contract(cid):
    customer = Customer.query.get_or_404(cid)
    contract = get_framework_contract(customer.id)
    if not contract:
        flash('Chưa có hợp đồng. Vui lòng tạo tại màn Hợp đồng.', 'info')
        return redirect(contracts_create_url(customer.id))
    try:
        urls = prepare_document_preview_urls(
            contract.file_path,
            'view_customer_framework_contract_pdf',
            'view_customer_framework_contract',
            'download_customer_framework_contract',
            'download_customer_framework_contract_pdf',
            cid=customer.id,
        )
    except FileNotFoundError as e:
        flash(str(e), 'danger')
        return redirect(url_for('customers'))
    return render_document_preview_page(
        contract.contract_code,
        customer.name,
        urls,
        url_for('customers'),
    )

@app.route('/customers/<int:cid>/contract/view')
def view_customer_framework_contract(cid):
    customer = Customer.query.get_or_404(cid)
    contract = get_framework_contract(customer.id)
    if not contract or not contract.file_path or not Path(contract.file_path).exists():
        flash('Chưa có hợp đồng. Vui lòng tạo tại màn Hợp đồng.', 'info')
        return redirect(contracts_create_url(customer.id))
    return send_docx_inline(contract.file_path, f'{contract.contract_code}.docx')

@app.route('/customers/<int:cid>/contract/view.pdf')
def view_customer_framework_contract_pdf(cid):
    customer = Customer.query.get_or_404(cid)
    contract = get_framework_contract(customer.id)
    if not contract or not contract.file_path or not Path(contract.file_path).is_file():
        abort(404)
    return send_pdf_inline(
        pdf_path_for_docx_or_abort(contract.file_path),
        f'{contract.contract_code}.pdf',
    )

@app.route('/customers/<int:cid>/contract/download')
def download_customer_framework_contract(cid):
    customer = Customer.query.get_or_404(cid)
    contract = get_framework_contract(customer.id)
    if not contract or not contract.file_path or not Path(contract.file_path).exists():
        flash('Chưa có hợp đồng. Vui lòng tạo tại màn Hợp đồng.', 'info')
        return redirect(contracts_create_url(customer.id))
    return send_file(contract.file_path, as_attachment=True, download_name=Path(contract.file_path).name)

@app.route('/customers/<int:cid>/contract/download.pdf')
def download_customer_framework_contract_pdf(cid):
    customer = Customer.query.get_or_404(cid)
    contract = get_framework_contract(customer.id)
    if not contract or not contract.file_path or not Path(contract.file_path).is_file():
        flash('Chưa có hợp đồng. Vui lòng tạo tại màn Hợp đồng.', 'info')
        return redirect(contracts_create_url(customer.id))
    pdf_path = ensure_pdf_for_docx(contract.file_path)
    return send_file(
        pdf_path,
        mimetype=PDF_MIMETYPE,
        as_attachment=True,
        download_name=Path(contract.file_path).stem + '.pdf',
    )

def customers_list_redirect(message=None, category='success', hub_id=None, hub_tab=None):
    if message:
        flash(message, category)
    src = request.form if request.method == 'POST' else request.args
    extra = {}
    hid = hub_id if hub_id is not None else src.get('_hub', type=int) or None
    htab = hub_tab if hub_tab is not None else src.get('_hub_tab', '')
    if hid:
        extra['hub'] = hid
    if htab:
        extra['hub_tab'] = htab
    return redirect(url_for(
        'customers',
        q=src.get('_q', ''),
        page=src.get('_page', 1, type=int) or 1,
        per_page=src.get('_per_page', 10, type=int) or 10,
        sort=src.get('_sort', 'created_at'),
        order=src.get('_order', 'desc'),
        **extra,
    ))

@app.route('/customers/<int:cid>/contract/signed-upload', methods=['POST'])
def upload_customer_signed_contract(cid):
    ensure_contract_columns()
    customer = Customer.query.get_or_404(cid)
    contract = get_framework_contract(customer.id)
    if not contract:
        flash('Chưa có hợp đồng. Tạo hợp đồng trước khi tải bản scan đã ký.', 'warning')
        return redirect(contracts_create_url(customer.id))
    if save_contract_signed_scan(contract, request.files.get('signed_scan')):
        db.session.commit()
        return customers_list_redirect('Đã tải lên hợp đồng đã ký', hub_id=cid, hub_tab='contract')
    db.session.rollback()
    return customers_list_redirect('', 'warning')

@app.route('/customers/<int:cid>/contract/signed')
def view_customer_signed_contract(cid):
    ensure_contract_columns()
    customer, contract, file_path = resolve_customer_signed_scan(cid)
    if not file_path:
        flash('Chưa có file hợp đồng đã ký hoặc file không tồn tại', 'warning')
        return redirect(url_for('customers'))
    return send_file(
        file_path,
        mimetype=signed_scan_mimetype(file_path),
        as_attachment=False,
        download_name=file_path.name,
    )

@app.route('/customers/<int:cid>/contract/signed/preview')
def preview_customer_signed_contract(cid):
    ensure_contract_columns()
    customer, contract, file_path = resolve_customer_signed_scan(cid)
    if not file_path:
        flash('Chưa có file hợp đồng đã ký hoặc file không tồn tại', 'warning')
        return redirect(url_for('customers'))
    return render_template(
        'signed_scan_preview.html',
        customer=customer,
        contract=contract,
        is_pdf=signed_scan_is_pdf(contract.signed_scan_path),
        view_url=url_for('view_customer_signed_contract', cid=cid),
        download_url=url_for('download_customer_signed_contract', cid=cid),
        back_url=url_for('customers', hub=cid, hub_tab='contract'),
    )

@app.route('/customers/<int:cid>/contract/signed-download')
def download_customer_signed_contract(cid):
    ensure_contract_columns()
    customer, contract, file_path = resolve_customer_signed_scan(cid)
    if not file_path:
        flash('Chưa có file hợp đồng đã ký hoặc file không tồn tại', 'warning')
        return redirect(url_for('customers'))
    return send_file(file_path, as_attachment=True, download_name=file_path.name)

@app.route('/customers/<int:cid>/contract/signed-delete', methods=['POST'])
def delete_customer_signed_contract(cid):
    ensure_contract_columns()
    customer = Customer.query.get_or_404(cid)
    contract = get_framework_contract(customer.id)
    if contract and contract.signed_scan_path:
        delete_contract_scan_file(contract.signed_scan_path)
        contract.signed_scan_path = ''
        contract.signed_scan_uploaded_at = None
        db.session.commit()
        return customers_list_redirect('Đã xóa file hợp đồng đã ký', 'warning')
    return customers_list_redirect()

def product_filters_from_request():
    return {
        'category': request.args.get('category', '').strip(),
        'brand': request.args.get('brand', '').strip(),
        'category_parent_id': request.args.get('category_parent_id', type=int),
        'category_id': request.args.get('category_id', type=int),
        'brand_id': request.args.get('brand_id', type=int),
        'supplier_id': request.args.get('supplier_id', type=int),
        'status': request.args.get('status', '').strip(),
        'sort': request.args.get('sort', 'newest').strip() or 'newest',
        'q': request.args.get('q', '').strip(),
    }


def product_list_url_args(per_page, filters, page=1, exclude=()):
    """Tham số query cho url_for('products', **args)."""
    skip = set(exclude or ())
    args = {'per_page': per_page}
    if page and page > 1 and 'page' not in skip:
        args['page'] = page
    for key in ('q', 'category_parent_id', 'category_id', 'brand_id', 'status', 'sort'):
        if key in skip:
            continue
        val = (filters or {}).get(key)
        if not val:
            continue
        if key == 'sort' and val == 'newest':
            continue
        args[key] = val
    return args


def products_list_url(per_page, filters, page=1, exclude=()):
    return url_for('products', **product_list_url_args(per_page, filters, page, exclude))


def products_filter_url(per_page, filters, exclude=()):
    """URL danh sách sản phẩm, bỏ một hoặc vài tham số lọc (dùng cho chip xóa từng bộ lọc)."""
    return products_list_url(per_page, filters, page=1, exclude=exclude)


def product_list_ctx_from_request():
    """Ngữ cảnh danh sách từ query string (khi mở chi tiết / sửa từ list)."""
    filters = product_filters_from_request()
    per_page = request.args.get('per_page', 10, type=int)
    per_page = per_page if per_page in (10, 25, 50, 100) else 10
    page = request.args.get('page', 1, type=int) or 1
    return {'filters': filters, 'per_page': per_page, 'page': page}


def product_list_ctx_from_form():
    """Ngữ cảnh danh sách từ form POST (sau thao tác trên list / detail)."""
    filters = {
        'q': request.form.get('_q', '').strip(),
        'category_parent_id': request.form.get('_category_parent_id', type=int),
        'category_id': request.form.get('_category_id', type=int),
        'brand_id': request.form.get('_brand_id', type=int),
        'status': request.form.get('_status', '').strip(),
        'sort': request.form.get('_sort', 'newest').strip() or 'newest',
    }
    per_page = request.form.get('_per_page', 10, type=int) or 10
    if per_page not in (10, 25, 50, 100):
        per_page = 10
    page = request.form.get('_page', 1, type=int) or 1
    return {'filters': filters, 'per_page': per_page, 'page': page}


def product_detail_url_kwargs(list_ctx):
    """Query params giữ trang / bộ lọc khi link sang chi tiết hoặc sửa."""
    if not list_ctx:
        return {}
    return product_list_url_args(
        list_ctx.get('per_page', 10),
        list_ctx.get('filters') or {},
        list_ctx.get('page', 1),
    )


app.jinja_env.globals['products_filter_url'] = products_filter_url
app.jinja_env.globals['products_list_url'] = products_list_url
app.jinja_env.globals['product_detail_url_kwargs'] = product_detail_url_kwargs


def _product_search_field_match(column, token):
    raw = (token or '').strip()
    if not raw:
        return None
    folded = vietnamese_search_fold(raw)
    parts = [column.ilike(f'%{raw}%')]
    if folded:
        parts.append(func.vn_fold(column).like(f'%{folded}%'))
    return or_(*parts)


def apply_product_text_search(query, q):
    """Tìm gần đúng: bỏ dấu, mỗi từ trong ô tìm phải khớp ít nhất một trường."""
    tokens = [t for t in re.split(r'\s+', (q or '').strip()) if t]
    if not tokens:
        return query
    columns = (
        Product.name,
        Product.sku,
        Product.model,
        Product.variant,
        Product.brand,
        Product.category,
    )
    for token in tokens:
        field_clauses = [_product_search_field_match(col, token) for col in columns]
        field_clauses = [c for c in field_clauses if c is not None]
        if field_clauses:
            query = query.filter(or_(*field_clauses))
    return query


def apply_product_filters(query, filters):
    q = filters.get('q', '')
    if q:
        query = apply_product_text_search(query, q)
    if filters.get('category_id'):
        query = query.filter(Product.category_id == filters['category_id'])
    elif filters.get('category_parent_id'):
        child_ids = [c.id for c in Category.query.filter_by(parent_id=filters['category_parent_id'])]
        child_ids.append(filters['category_parent_id'])
        query = query.filter(Product.category_id.in_(child_ids))
    elif filters.get('category'):
        query = query.filter(Product.category == filters['category'])
    if filters.get('brand_id'):
        bid = filters['brand_id']
        bid_s = str(bid)
        query = query.filter(or_(
            Product.brand_id == bid,
            Product.variant_brand_ids == bid_s,
            Product.variant_brand_ids.like(f'{bid_s},%'),
            Product.variant_brand_ids.like(f'%, {bid_s},%'),
            Product.variant_brand_ids.like(f'%, {bid_s}'),
        ))
    elif filters.get('brand'):
        query = query.filter(Product.brand == filters['brand'])
    if filters.get('supplier_id'):
        pid_rows = db.session.query(ProductSupplier.product_id).filter(
            ProductSupplier.supplier_id == filters['supplier_id'],
        )
        query = query.filter(Product.id.in_(pid_rows))
    st = filters.get('status', '')
    if st == 'selling':
        query = query.filter(Product.is_active.is_(True))
    elif st == 'in_stock':
        query = query.filter(Product.is_active.is_(True), Product.stock_qty > 0)
    elif st == 'low':
        query = query.filter(
            Product.stock_qty > 0,
            Product.stock_qty <= func.coalesce(Product.low_stock, 5),
        )
    elif st == 'out':
        query = query.filter(Product.stock_qty <= 0)
    elif st == 'stopped':
        query = query.filter(or_(Product.is_active.is_(False), Product.is_active == 0))
    return query

def product_list_order(query, sort_key):
    sort_key = (sort_key or 'newest').strip()
    if sort_key == 'name_asc':
        return query.order_by(Product.name.asc(), Product.id.asc())
    if sort_key == 'name_desc':
        return query.order_by(Product.name.desc(), Product.id.desc())
    if sort_key == 'stock_asc':
        return query.order_by(Product.stock_qty.asc(), Product.name.asc())
    if sort_key == 'stock_desc':
        return query.order_by(Product.stock_qty.desc(), Product.name.asc())
    return query.order_by(Product.created_at.desc(), Product.id.desc())

def products_redirect_after_save(message=None, category='success'):
    if message:
        flash(message, category)
    list_ctx = product_list_ctx_from_form()
    nav = product_detail_url_kwargs(list_ctx)
    if request.form.get('_from_edit') == '1':
        pid = request.form.get('_pid', type=int)
        if pid:
            return redirect(url_for('product_edit', pid=pid, **nav))
    if request.form.get('_from_detail') == '1':
        pid = request.form.get('_pid', type=int)
        if pid:
            return redirect(url_for('product_detail', pid=pid, **nav))
    return redirect(products_list_url(
        list_ctx['per_page'],
        list_ctx['filters'],
        list_ctx['page'],
    ))

PRICE_TYPE_LABELS = {
    'retail_price': 'Giá bán lẻ',
    'project_price': 'Giá dự án',
    'dealer_price': 'Giá đại lý',
    'cost_price': 'Giá nhập',
}

PRODUCT_PRICE_FIELDS = ('cost_price', 'project_price', 'dealer_price', 'retail_price')

def price_type_label(key):
    return PRICE_TYPE_LABELS.get(key, key)

app.jinja_env.globals['price_type_label'] = price_type_label

def apply_product_price_from_form(product, note='', fields=None):
    """Cập nhật giá từ form; ghi PriceHistory khi có thay đổi. Trả về list field đã đổi."""
    changed = []
    for field in fields or PRODUCT_PRICE_FIELDS:
        if field not in request.form:
            continue
        old = getattr(product, field) or 0
        if field == 'cost_price':
            old_vc = (getattr(product, 'variant_cost_prices', None) or '').strip()
            ok, err = apply_product_variant_cost_prices(product, request.form.get('cost_price'))
            if not ok:
                flash(err, 'warning')
                continue
            new = product.cost_price or 0
            new_vc = (getattr(product, 'variant_cost_prices', None) or '').strip()
            if new != old or new_vc != old_vc:
                db.session.add(PriceHistory(
                    product_id=product.id,
                    price_type=field,
                    old_price=old,
                    new_price=new,
                    note=note or (new_vc if new_vc else ''),
                ))
                changed.append(field)
            continue
        elif field == 'retail_price':
            old_vp = (getattr(product, 'variant_prices', None) or '').strip()
            ok, err = apply_product_variant_retail_prices(product, request.form.get('retail_price'))
            if not ok:
                flash(err, 'warning')
                continue
            new = product.retail_price or 0
            new_vp = (getattr(product, 'variant_prices', None) or '').strip()
            if new != old or new_vp != old_vp:
                db.session.add(PriceHistory(
                    product_id=product.id,
                    price_type=field,
                    old_price=old,
                    new_price=new,
                    note=note or (new_vp if new_vp else ''),
                ))
                changed.append(field)
            continue
        elif field == 'dealer_price':
            old_vd = (getattr(product, 'variant_dealer_prices', None) or '').strip()
            ok, err = apply_product_variant_dealer_prices(product, request.form.get('dealer_price'))
            if not ok:
                flash(err, 'warning')
                continue
            new = product.dealer_price or 0
            new_vd = (getattr(product, 'variant_dealer_prices', None) or '').strip()
            if new != old or new_vd != old_vd:
                db.session.add(PriceHistory(
                    product_id=product.id,
                    price_type=field,
                    old_price=old,
                    new_price=new,
                    note=note or (new_vd if new_vd else ''),
                ))
                changed.append(field)
            continue
        elif field == 'project_price':
            old_vp = (getattr(product, 'variant_project_prices', None) or '').strip()
            ok, err = apply_product_variant_project_prices(product, request.form.get('project_price'))
            if not ok:
                flash(err, 'warning')
                continue
            new = product.project_price or 0
            new_vp = (getattr(product, 'variant_project_prices', None) or '').strip()
            if new != old or new_vp != old_vp:
                db.session.add(PriceHistory(
                    product_id=product.id,
                    price_type=field,
                    old_price=old,
                    new_price=new,
                    note=note or (new_vp if new_vp else ''),
                ))
                changed.append(field)
            continue
        else:
            new = parse_int(request.form.get(field), old)
        if new != old:
            db.session.add(PriceHistory(
                product_id=product.id,
                price_type=field,
                old_price=old,
                new_price=new,
                note=note,
            ))
            setattr(product, field, new)
            changed.append(field)
    return changed

def apply_product_image_url_from_form(product):
    raw = (request.form.get('image_url') or '').strip()
    product.image_url = raw if is_external_image_url(raw) else ''


def apply_product_form(product):
    product.name = request.form.get('name', '').strip()
    assign_product_taxonomy_from_form(product)
    product.model = request.form.get('model', '')
    product.variant = request.form.get('variant', '')
    apply_product_unit_from_form(product)
    product.warranty = request.form.get('warranty', '')
    variant_labels = split_variant_labels(request.form.get('variant', '') or product.variant or '')
    if len(variant_labels) > 1:
        ok_img, err_img = apply_product_variant_image_urls(
            product, request.form.get('variant_image_urls', '')
        )
        if not ok_img:
            flash(err_img, 'warning')
        ok_brand, err_brand = apply_product_variant_brand_ids(
            product, request.form.get('variant_brand_ids', '')
        )
        if not ok_brand:
            flash(err_brand, 'warning')
    else:
        apply_product_image_url_from_form(product)
        product.variant_image_urls = ''
        product.variant_brand_ids = ''
    if 'cost_price' in request.form:
        ok_cost, err_cost = apply_product_variant_cost_prices(product, request.form.get('cost_price'))
        if not ok_cost:
            flash(err_cost, 'warning')
    ok, err = apply_product_variant_retail_prices(product, request.form.get('retail_price'))
    if not ok:
        flash(err, 'warning')
    if 'dealer_price' in request.form:
        ok_dealer, err_dealer = apply_product_variant_dealer_prices(
            product, request.form.get('dealer_price')
        )
        if not ok_dealer:
            flash(err_dealer, 'warning')
    if 'project_price' in request.form:
        ok_project, err_project = apply_product_variant_project_prices(
            product, request.form.get('project_price')
        )
        if not ok_project:
            flash(err_project, 'warning')
    ok_stock, err_stock = apply_product_variant_stocks(product, request.form.get('stock'))
    if not ok_stock:
        flash(err_stock, 'warning')
    product.low_stock = parse_qty(request.form.get('low_stock'), product.low_stock or 5)

def category_product_counts():
    counts = {}
    for cat in Category.query.all():
        path = category_display_path(cat.id)
        counts[cat.id] = Product.query.filter(
            or_(Product.category_id == cat.id, Product.category == path, Product.category == cat.name)
        ).count()
    return counts

def brand_product_counts():
    counts = {}
    for brand in Brand.query.all():
        counts[brand.id] = Product.query.filter(
            or_(Product.brand_id == brand.id, Product.brand == brand.name)
        ).count()
    return counts

def parent_category_product_count(parent_id):
    child_ids = [c.id for c in Category.query.filter_by(parent_id=parent_id)]
    ids = [parent_id] + child_ids
    if not ids:
        return 0
    return Product.query.filter(Product.category_id.in_(ids)).count()

app.jinja_env.globals['parent_category_product_count'] = parent_category_product_count

PRODUCT_IMPORT_FIELDS = [
    ('sku', 'SKU'),
    ('name', 'Tên sản phẩm'),
    ('category_parent', 'Danh mục cha'),
    ('category_child', 'Danh mục con'),
    ('brand', 'Thương hiệu'),
    ('model', 'Model'),
    ('variant', 'Biến thể'),
    ('image_url', 'URL hình ảnh'),
    ('unit', 'Đơn vị tồn kho'),
    ('purchase_unit', 'Đơn vị nhập'),
    ('conversion_factor', 'Hệ số quy đổi'),
    ('warranty', 'Bảo hành'),
    ('cost_price', 'Giá nhập'),
    ('retail_price', 'Giá bán lẻ'),
    ('dealer_price', 'Giá đại lý'),
    ('project_price', 'Giá dự án'),
    ('stock', 'Tồn kho'),
    ('low_stock', 'Tồn tối thiểu'),
    ('status', 'Trạng thái'),
]

PRODUCT_IMPORT_SAMPLE_ROW = [
    'SP-CLE-17', 'Cờ lê 17', 'Dụng cụ', 'Cờ lê', 'Stanley',
    '', '17mm', 'https://example.com/images/co-le-17.jpg', 'Chiếc', 'Hộp', 20,
    '12 tháng', 85000, 120000, 110000, 115000, 0, 10, 'Đang bán',
]

PRODUCT_IMPORT_COL_WIDTHS = {
    'sku': 14,
    'name': 30,
    'category_parent': 16,
    'category_child': 18,
    'brand': 14,
    'model': 12,
    'variant': 14,
    'image_url': 44,
    'unit': 12,
    'purchase_unit': 14,
    'conversion_factor': 14,
    'warranty': 12,
    'cost_price': 12,
    'retail_price': 12,
    'dealer_price': 12,
    'project_price': 12,
    'stock': 10,
    'low_stock': 12,
    'status': 12,
    'intake_qty': 14,
}

PRODUCT_IMPORT_HEADER_ALIASES = {
    'sku': {'sku', 'mã sku', 'ma sku', 'mã vạch', 'ma vach'},
    'name': {'tên sản phẩm', 'ten san pham', 'tên', 'ten', 'name', 'sản phẩm'},
    'category_parent': {'danh mục cha', 'danh muc cha', 'category parent', 'nhóm danh mục', 'ngành hàng'},
    'category_child': {'danh mục con', 'danh muc con', 'category child', 'loại sản phẩm'},
    'category': {'danh mục', 'danh muc', 'category', 'loại'},
    'brand': {'thương hiệu', 'thuong hieu', 'brand', 'hãng'},
    'model': {'model', 'mẫu'},
    'variant': {'biến thể', 'bien the', 'variant'},
    'image_url': {
        'url hình ảnh', 'url hinh anh', 'image url', 'image_url', 'ảnh url', 'anh url',
        'link ảnh', 'link anh', 'hình ảnh url', 'hinh anh url', 'url ảnh', 'url anh',
    },
    'unit': {
        'đơn vị', 'don vi', 'unit', 'dvt',
        'đơn vị tồn kho', 'don vi ton kho', 'base unit', 'base_unit',
    },
    'purchase_unit': {
        'đơn vị nhập', 'don vi nhap', 'purchase unit', 'purchase_unit',
        'đơn vị đóng gói', 'don vi dong goi', 'đóng gói', 'dong goi',
        'đơn vị thùng', 'don vi thung',
    },
    'conversion_factor': {
        'hệ số quy đổi', 'he so quy doi', 'conversion factor', 'conversion_factor',
        'số chiếc/hộp', 'so chiec/hop', 'số cái/hộp', 'so cai/hop',
        'số cái trong hộp', 'so cai trong hop', 'số chiếc trong hộp',
        'sl trong hộp', 'sl trong hop', 'qty per box', 'pcs per box',
        'số mét/cuộn', 'so met/cuon', 'met moi cuon',
    },
    'warranty': {'bảo hành', 'bao hanh', 'warranty'},
    'cost_price': {'giá nhập', 'gia nhap', 'cost', 'cost_price'},
    'retail_price': {'giá bán lẻ', 'gia ban le', 'giá bán', 'retail', 'retail_price'},
    'dealer_price': {'giá đại lý', 'gia dai ly', 'dealer', 'dealer_price'},
    'project_price': {'giá dự án', 'gia du an', 'project', 'project_price'},
    'stock': {'tồn kho', 'ton kho', 'stock', 'sl tồn'},
    'intake_qty': {
        'số lượng nhập', 'so luong nhap', 'sl nhập', 'sl nhap', 'qty nhập', 'qty nhap',
        'số lượng', 'so luong', 'quantity',
    },
    'low_stock': {'tồn tối thiểu', 'ton toi thieu', 'low_stock', 'tồn min'},
    'status': {'trạng thái', 'trang thai', 'status'},
}

def normalize_import_header(value):
    return (str(value or '').strip().lower().replace('_', ' '))

def build_product_import_header_map(headers):
    col_map = {}
    for idx, raw in enumerate(headers):
        key = normalize_import_header(raw)
        if not key:
            continue
        for field, aliases in PRODUCT_IMPORT_HEADER_ALIASES.items():
            if key in aliases or key == field:
                col_map[field] = idx
                break
    return col_map

def parse_import_active_status(value):
    s = normalize_import_header(value)
    if s in ('ngừng bán', 'ngung ban', 'stopped', '0', 'false', 'no', 'n'):
        return False
    if s in ('đang bán', 'dang ban', 'selling', '1', 'true', 'yes', 'y'):
        return True
    return True

def ensure_category_name(name):
    ensure_category_from_path(name)

def _import_row_cells(row):
    return ['' if c is None else c for c in row]


def _row_looks_like_import_header(row):
    """Bỏ qua dòng tiêu đề lặp lại giữa các sheet."""
    col_map = build_product_import_header_map(row)
    if 'sku' not in col_map or 'name' not in col_map:
        return False
    sku_idx = col_map['sku']
    name_idx = col_map['name']
    sku_cell = normalize_import_header(row[sku_idx] if sku_idx < len(row) else '')
    name_cell = normalize_import_header(row[name_idx] if name_idx < len(row) else '')
    if sku_cell in PRODUCT_IMPORT_HEADER_ALIASES['sku']:
        return True
    if name_cell in PRODUCT_IMPORT_HEADER_ALIASES['name']:
        return True
    return False


def _find_import_header_in_sheet(sheet_rows):
    for idx, row in enumerate(sheet_rows):
        col_map = build_product_import_header_map(row)
        if 'sku' in col_map and 'name' in col_map:
            return idx, row
    return None, None


def merge_product_import_workbook_sheets(wb):
    """Gộp mọi sheet Excel có cùng định dạng cột (SKU + Tên sản phẩm)."""
    merged = []
    sheets_meta = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        sheet_rows = [_import_row_cells(r) for r in ws.iter_rows(values_only=True)]
        if not sheet_rows:
            continue
        header_idx, header_row = _find_import_header_in_sheet(sheet_rows)
        if header_idx is None:
            continue
        if not merged:
            merged.append(header_row)
        data_count = 0
        for row in sheet_rows[header_idx + 1:]:
            if not any(str(c).strip() for c in row):
                continue
            if _row_looks_like_import_header(row):
                continue
            merged.append(row)
            data_count += 1
        if data_count > 0 or header_idx is not None:
            sheets_meta.append({'name': sheet_name, 'rows': data_count})
    return merged, sheets_meta


def read_product_import_sheet(file_storage, with_meta=False):
    filename = secure_filename(file_storage.filename or '')
    ext = Path(filename).suffix.lower()
    sheets_meta = []
    if ext == '.csv':
        raw = file_storage.read()
        if raw[:3] == b'\xef\xbb\xbf':
            raw = raw[3:]
        text = raw.decode('utf-8-sig', errors='replace')
        reader = csv.reader(io.StringIO(text))
        rows = [list(row) for row in reader]
    elif ext in ('.xlsx', '.xlsm'):
        try:
            from openpyxl import load_workbook
        except ImportError:
            raise ValueError('Thiếu thư viện openpyxl. Dùng file .csv hoặc cài: pip install openpyxl')
        file_storage.seek(0)
        wb = load_workbook(file_storage, read_only=True, data_only=True)
        rows, sheets_meta = merge_product_import_workbook_sheets(wb)
        wb.close()
        if not rows:
            raise ValueError(
                'Không tìm thấy sheet nào có cột SKU và Tên sản phẩm. '
                'Mỗi tab cần cùng định dạng như file mẫu.'
            )
    else:
        raise ValueError('Chỉ hỗ trợ file Excel (.xlsx) hoặc CSV (.csv)')
    if with_meta:
        return rows, {'sheets': sheets_meta, 'sheet_count': len(sheets_meta)}
    return rows

def product_category_parts(product):
    """Tách danh mục cha / con để export Excel."""
    if getattr(product, 'category_id', None):
        cat = Category.query.get(product.category_id)
        if cat and cat.parent_id:
            parent = Category.query.get(cat.parent_id)
            return (parent.name if parent else '', cat.name)
        if cat:
            return (cat.name, '')
    path = (product.category or '').strip()
    if '>' in path:
        parts = [p.strip() for p in path.split('>') if p.strip()]
        if len(parts) >= 2:
            return (parts[0], parts[-1])
    return (path, '')

def resolve_import_category_path(data):
    """Gộp danh mục từ cột cha/con hoặc cột danh mục gộp (Cha > Con)."""
    parent = (data.get('category_parent') or '').strip()
    child = (data.get('category_child') or '').strip()
    combined = (data.get('category') or '').strip()
    if parent and child:
        return f'{parent} > {child}'
    if child and not parent:
        return child
    if parent and not child:
        return parent
    return combined

def product_row_from_import(row, col_map):
    def cell(field, default=''):
        idx = col_map.get(field)
        if idx is None or idx >= len(row):
            return default
        val = row[idx]
        if val is None:
            return default
        return str(val).strip()

    raw = {
        'sku': cell('sku'),
        'name': cell('name'),
        'category_parent': cell('category_parent'),
        'category_child': cell('category_child'),
        'category': cell('category'),
        'brand': cell('brand'),
        'model': cell('model'),
        'variant': cell('variant'),
        'image_url_raw': cell('image_url'),
        'image_url': cell('image_url'),
        'unit': cell('unit', 'cái') or 'cái',
        'purchase_unit': cell('purchase_unit') or cell('unit', 'cái') or 'cái',
        'conversion_factor': parse_qty(cell('conversion_factor'), 1.0),
        'warranty': cell('warranty'),
        'cost_price_raw': cell('cost_price'),
        'cost_price': parse_int(cell('cost_price'), 0),
        'retail_price_raw': cell('retail_price'),
        'retail_price': parse_int(cell('retail_price'), 0),
        'dealer_price': parse_int(cell('dealer_price'), 0),
        'project_price': parse_int(cell('project_price'), 0),
        'stock_raw': cell('stock'),
        'stock': parse_qty(cell('stock'), 0),
        'low_stock': parse_qty(cell('low_stock'), 5),
        'is_active': parse_import_active_status(cell('status', 'Đang bán')),
    }
    path = resolve_import_category_path(raw)
    raw['category'] = path
    parts = path.split('>') if path and '>' in path else []
    if len(parts) >= 2 and not raw['category_parent']:
        raw['category_parent'] = parts[0].strip()
        raw['category_child'] = parts[-1].strip()
    elif path and not raw['category_parent'] and not raw['category_child']:
        raw['category_parent'] = path
    retail_prices = parse_price_csv(raw['retail_price_raw'])
    if retail_prices:
        raw['retail_price'] = retail_prices[0]
    raw['variant_prices'] = _import_variant_prices_from_data(raw)
    cost_prices = parse_price_csv(raw['cost_price_raw'])
    if cost_prices:
        raw['cost_price'] = cost_prices[0]
    raw['variant_cost_prices'] = _import_variant_cost_prices_from_data(raw)
    stock_qtys = parse_qty_csv(raw['stock_raw'])
    if stock_qtys:
        raw['stock'] = sum(stock_qtys) if len(stock_qtys) > 1 else stock_qtys[0]
    raw['variant_stocks'] = _import_variant_stocks_from_data(raw)
    raw['variant_image_urls'] = _import_variant_image_urls_from_data(raw)
    raw['variant_brand_ids'] = _import_variant_brand_ids_from_data(raw)
    return raw

def _apply_import_image_url(product, data):
    vi = (data.get('variant_image_urls') or '').strip()
    if vi:
        product.variant_image_urls = vi
        urls = product_stored_variant_image_urls(product)
        product.image_url = urls[0] if urls else ''
        return
    product.variant_image_urls = ''
    raw = (data.get('image_url') or '').strip()
    product.image_url = raw if is_external_image_url(raw) else ''

def _apply_import_unit_fields(product, data):
    base = (data.get('unit') or data.get('base_unit') or 'cái').strip() or 'cái'
    purchase = (data.get('purchase_unit') or base).strip() or base
    factor = parse_qty(data.get('conversion_factor'), 1.0)
    if factor <= 0:
        factor = 1.0
    enabled = purchase.lower() != base.lower() or factor != 1.0
    product.unit = base
    product.base_unit = base
    product.purchase_unit = purchase if enabled else base
    product.unit_conversion_enabled = enabled
    product.conversion_factor = factor if enabled else 1.0

def product_to_export_row(product):
    parent_name, child_name = product_category_parts(product)
    _, sale_label = product_sale_status(product)
    variant_imgs = product_stored_variant_image_urls(product)
    image_export = '||'.join(variant_imgs) if variant_imgs else (getattr(product, 'image_url', None) or '').strip()
    return [
        product.sku,
        product.name,
        parent_name,
        child_name,
        product_variant_brands_export_value(product),
        product.model or '',
        product.variant or '',
        image_export,
        product_base_unit(product),
        product_purchase_unit(product) if product_has_unit_conversion(product) else product_base_unit(product),
        product_conversion_factor(product) if product_has_unit_conversion(product) else '',
        product.warranty or '',
        (getattr(product, 'variant_cost_prices', None) or '').strip() or (product.cost_price or 0),
        (getattr(product, 'variant_prices', None) or '').strip() or (product.retail_price or 0),
        product.dealer_price or 0,
        product.project_price or 0,
        (getattr(product, 'variant_stocks', None) or '').strip() or (product.stock or 0),
        product.low_stock or 5,
        sale_label,
    ]

def product_import_field_labels():
    return [label for _, label in PRODUCT_IMPORT_FIELDS]

def product_import_column_widths(field_items=None):
    items = field_items or PRODUCT_IMPORT_FIELDS
    return [PRODUCT_IMPORT_COL_WIDTHS.get(key, 12) for key, _ in items]


def _style_product_import_worksheet(ws, field_items):
    from openpyxl.styles import Font, PatternFill, Alignment
    from openpyxl.utils import get_column_letter
    header_fill = PatternFill('solid', fgColor='E8F0FE')
    header_font = Font(bold=True)
    for cell in ws[1]:
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='center')
    ws.freeze_panes = 'A2'
    for idx, width in enumerate(product_import_column_widths(field_items), start=1):
        ws.column_dimensions[get_column_letter(idx)].width = width


def build_products_xlsx_bytes(rows):
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = 'San pham'
    headers = product_import_field_labels()
    ws.append(headers)
    for row in rows:
        ws.append(row)
    _style_product_import_worksheet(ws, PRODUCT_IMPORT_FIELDS)
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf

def build_products_csv_text(rows):
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow(product_import_field_labels())
    for row in rows:
        writer.writerow(row)
    return '\ufeff' + buf.getvalue()


def sync_product_import_template_files(rows=None):
    """Ghi bản sao file mẫu vào static/templates/ (csv + xlsx nếu có openpyxl)."""
    rows = rows or [PRODUCT_IMPORT_SAMPLE_ROW]
    template_dir = BASE_DIR / 'static' / 'templates'
    try:
        template_dir.mkdir(parents=True, exist_ok=True)
        csv_text = build_products_csv_text(rows)
        (template_dir / 'mau-import-san-pham.csv').write_text(csv_text, encoding='utf-8')
        try:
            xlsx_buf = build_products_xlsx_bytes(rows)
            (template_dir / 'mau-import-san-pham.xlsx').write_bytes(xlsx_buf.getvalue())
        except ImportError:
            pass
    except OSError:
        pass

def build_product_import_preview(rows, update_existing=False):
    if not rows:
        raise ValueError('File trống')
    header_row = rows[0]
    col_map = build_product_import_header_map(header_row)
    if 'sku' not in col_map or 'name' not in col_map:
        raise ValueError('File phải có cột SKU và Tên sản phẩm (xem file mẫu)')
    preview = []
    seen_skus = set()
    for line_no, raw_row in enumerate(rows[1:], start=2):
        if not any(str(c).strip() for c in raw_row):
            continue
        data = product_row_from_import(raw_row, col_map)
        sku = data['sku']
        name = data['name']
        row_id = uuid.uuid4().hex[:12]
        if not sku and not name:
            continue
        if not sku:
            preview.append({
                'row_id': row_id, 'line_no': line_no, 'data': data,
                'action': 'error', 'action_label': 'Lỗi', 'note': 'Thiếu SKU',
                'importable': False,
            })
            continue
        if not name:
            preview.append({
                'row_id': row_id, 'line_no': line_no, 'data': data,
                'action': 'error', 'action_label': 'Lỗi', 'note': 'Thiếu tên sản phẩm',
                'importable': False,
            })
            continue
        sku_key = sku.lower()
        if sku_key in seen_skus:
            preview.append({
                'row_id': row_id, 'line_no': line_no, 'data': data,
                'action': 'error', 'action_label': 'Lỗi', 'note': 'SKU trùng trong file',
                'importable': False,
            })
            continue
        seen_skus.add(sku_key)
        existing = Product.query.filter_by(sku=sku).first()
        if existing:
            if update_existing:
                preview.append({
                    'row_id': row_id, 'line_no': line_no, 'data': data,
                    'action': 'update', 'action_label': 'Cập nhật', 'note': 'SKU đã có trong hệ thống',
                    'importable': True,
                })
            else:
                preview.append({
                    'row_id': row_id, 'line_no': line_no, 'data': data,
                    'action': 'skip', 'action_label': 'Bỏ qua', 'note': 'SKU đã tồn tại',
                    'importable': False,
                })
            continue
        preview.append({
            'row_id': row_id, 'line_no': line_no, 'data': data,
            'action': 'create', 'action_label': 'Thêm mới', 'note': '',
            'importable': True,
        })
    if not preview:
        raise ValueError('Không có dòng dữ liệu hợp lệ trong file')
    return preview

def supplier_intake_import_fields():
    """File mẫu nhập NCC: cột sản phẩm + Số lượng nhập (sau Giá nhập)."""
    fields = []
    for key, label in PRODUCT_IMPORT_FIELDS:
        fields.append((key, label))
        if key == 'cost_price':
            fields.append(('intake_qty', 'Số lượng nhập'))
    return fields

def supplier_intake_import_field_labels():
    return [label for _, label in supplier_intake_import_fields()]

def supplier_intake_sample_row():
    row = list(PRODUCT_IMPORT_SAMPLE_ROW)
    cost_idx = next(i for i, (k, _) in enumerate(PRODUCT_IMPORT_FIELDS) if k == 'cost_price')
    row.insert(cost_idx + 1, 2)
    return row

def build_supplier_intake_xlsx_bytes(rows):
    from openpyxl import Workbook
    field_items = supplier_intake_import_fields()
    wb = Workbook()
    ws = wb.active
    ws.title = 'Nhap NCC'
    headers = supplier_intake_import_field_labels()
    ws.append(headers)
    for row in rows:
        ws.append(row)
    _style_product_import_worksheet(ws, field_items)
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf

def build_supplier_intake_csv_text(rows):
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow(supplier_intake_import_field_labels())
    for row in rows:
        writer.writerow(row)
    return '\ufeff' + buf.getvalue()

def build_supplier_intake_import_col_map(headers):
    return build_product_import_header_map(headers)

def supplier_intake_row_from_import(row, col_map):
    data = product_row_from_import(row, col_map)
    if 'intake_qty' in col_map:
        idx = col_map['intake_qty']
        val = row[idx] if idx is not None and idx < len(row) else 0
        data['intake_qty'] = parse_qty(val, 0)
        if data['intake_qty'] <= 0:
            data['intake_qty'] = parse_qty(data.get('stock', 0), 0)
    else:
        data['intake_qty'] = parse_qty(data.get('stock', 0), 0)
    return data

def build_supplier_intake_import_preview(rows, update_existing=False):
    if not rows:
        raise ValueError('File trống')
    header_row = rows[0]
    col_map = build_supplier_intake_import_col_map(header_row)
    if 'sku' not in col_map or 'name' not in col_map:
        raise ValueError('File phải có cột SKU và Tên sản phẩm (dùng file mẫu import sản phẩm)')
    preview = []
    seen_skus = set()
    for line_no, raw_row in enumerate(rows[1:], start=2):
        if not any(str(c).strip() for c in raw_row):
            continue
        data = supplier_intake_row_from_import(raw_row, col_map)
        sku = data['sku']
        name = data['name']
        intake_qty = float(data.get('intake_qty') or 0)
        cost = int(data.get('cost_price') or 0)
        row_id = uuid.uuid4().hex[:12]
        if not sku and not name:
            continue
        if not sku:
            preview.append({
                'row_id': row_id, 'line_no': line_no, 'data': data,
                'action': 'error', 'action_label': 'Lỗi', 'note': 'Thiếu SKU',
                'importable': False,
            })
            continue
        if not name:
            preview.append({
                'row_id': row_id, 'line_no': line_no, 'data': data,
                'action': 'error', 'action_label': 'Lỗi', 'note': 'Thiếu tên sản phẩm',
                'importable': False,
            })
            continue
        sku_key = sku.lower()
        if sku_key in seen_skus:
            preview.append({
                'row_id': row_id, 'line_no': line_no, 'data': data,
                'action': 'error', 'action_label': 'Lỗi', 'note': 'SKU trùng trong file',
                'importable': False,
            })
            continue
        seen_skus.add(sku_key)
        if intake_qty <= 0 and cost <= 0:
            preview.append({
                'row_id': row_id, 'line_no': line_no, 'data': data,
                'action': 'error', 'action_label': 'Lỗi',
                'note': 'Cần cột Số lượng nhập (hoặc Tồn kho) hoặc Giá nhập > 0',
                'importable': False,
            })
            continue
        existing = Product.query.filter_by(sku=sku).first()
        if existing:
            note = 'SKU đã có — nhập kho theo file'
            action = 'update'
            action_label = 'Nhập kho'
            if update_existing:
                note = 'Cập nhật thông tin SP + nhập kho'
                action_label = 'Cập nhật & nhập'
            preview.append({
                'row_id': row_id, 'line_no': line_no, 'data': data,
                'action': action, 'action_label': action_label, 'note': note,
                'importable': True,
            })
            continue
        preview.append({
            'row_id': row_id, 'line_no': line_no, 'data': data,
            'action': 'create', 'action_label': 'Thêm SP & nhập', 'note': '',
            'importable': True,
        })
    if not preview:
        raise ValueError('Không có dòng dữ liệu hợp lệ trong file')
    return preview

def ensure_product_for_supplier_intake_row(data, update_existing=False):
    ensure_category_brand_schema()
    ensure_product_columns()
    sku = data['sku']
    existing = Product.query.filter_by(sku=sku).first()
    if existing:
        if update_existing:
            existing.name = data['name']
            existing.model = data['model']
            existing.variant = data['variant']
            _apply_import_unit_fields(existing, data)
            existing.warranty = data['warranty']
            existing.cost_price = data['cost_price']
            existing.retail_price = data['retail_price']
            existing.variant_prices = data.get('variant_prices') or ''
            existing.variant_stocks = data.get('variant_stocks') or ''
            existing.variant_cost_prices = data.get('variant_cost_prices') or ''
            existing.variant_brand_ids = data.get('variant_brand_ids') or ''
            existing.dealer_price = data['dealer_price']
            existing.project_price = data['project_price']
            existing.stock = data['stock']
            existing.low_stock = data['low_stock']
            existing.is_active = data['is_active']
            _apply_import_taxonomy(existing, data)
            _apply_import_image_url(existing, data)
        return existing
    p = Product(
        sku=sku,
        name=data['name'],
        model=data['model'],
        variant=data['variant'],
        variant_prices=data.get('variant_prices') or '',
        variant_stocks=data.get('variant_stocks') or '',
        variant_cost_prices=data.get('variant_cost_prices') or '',
        variant_brand_ids=data.get('variant_brand_ids') or '',
        warranty=data['warranty'],
        cost_price=data['cost_price'],
        retail_price=data['retail_price'],
        dealer_price=data['dealer_price'],
        project_price=data['project_price'],
        stock=data['stock'],
        low_stock=data['low_stock'],
        is_active=data['is_active'],
    )
    _apply_import_unit_fields(p, data)
    _apply_import_taxonomy(p, data)
    _apply_import_image_url(p, data)
    db.session.add(p)
    db.session.flush()
    return p

def apply_supplier_intake_from_import_rows(supplier, items, update_existing, ref_code='', intake_note=''):
    pending_lines = []
    for item in items:
        data = item.get('data') or item
        product = ensure_product_for_supplier_intake_row(data, update_existing=update_existing)
        entered_qty = float(data.get('intake_qty') or 0)
        purchase_qty, base_qty, unit_mode = resolve_supplier_intake_quantities(
            product, entered_qty, 'purchase',
        )
        cost = int(data.get('cost_price') or 0)
        if base_qty <= 0 and purchase_qty <= 0 and cost <= 0:
            continue
        pending_lines.append({
            'product': product,
            'purchase_qty': purchase_qty,
            'base_qty': base_qty,
            'unit_mode': unit_mode,
            'cost': cost,
            'sku': data.get('sku') or product.sku or '',
            'note': 'Import Excel',
        })
    if not pending_lines:
        raise ValueError('Không có dòng hợp lệ để nhập kho')
    return _supplier_intake_process_lines(supplier, pending_lines, ref_code, intake_note)

def _apply_import_taxonomy(product, data):
    category_path = resolve_import_category_path(data)
    if category_path:
        product.category_id = ensure_category_from_path(category_path)
    else:
        product.category_id = None
    data['category_id'] = product.category_id
    variant_brand_ids = (data.get('variant_brand_ids') or '').strip()
    labels = split_csv_field(data.get('variant') or '')
    if len(labels) > 1 and variant_brand_ids:
        product.variant_brand_ids = variant_brand_ids
        ids = product_stored_variant_brand_ids(product)
        product.brand_id = ids[0] if ids else None
    else:
        product.variant_brand_ids = ''
        brand_name = (data.get('brand') or '').strip()
        if brand_name and ',' not in brand_name:
            product.brand_id = ensure_brand_name(brand_name, product.category_id)
        elif brand_name:
            first_brand = brand_name.split(',')[0].strip()
            product.brand_id = ensure_brand_name(first_brand, product.category_id) if first_brand else None
        else:
            product.brand_id = None
    sync_product_taxonomy_strings(product)

def import_product_data_list(items, update_existing=False):
    created = updated = skipped = errors = 0
    error_lines = []
    for item in items:
        data = item.get('data') or item
        sku = data['sku']
        name = data['name']
        existing = Product.query.filter_by(sku=sku).first()
        if existing:
            if not update_existing:
                skipped += 1
                continue
            existing.name = name
            existing.model = data['model']
            existing.variant = data['variant']
            _apply_import_unit_fields(existing, data)
            existing.warranty = data['warranty']
            existing.cost_price = data['cost_price']
            existing.retail_price = data['retail_price']
            existing.variant_prices = data.get('variant_prices') or ''
            existing.variant_stocks = data.get('variant_stocks') or ''
            existing.variant_cost_prices = data.get('variant_cost_prices') or ''
            existing.variant_brand_ids = data.get('variant_brand_ids') or ''
            existing.dealer_price = data['dealer_price']
            existing.project_price = data['project_price']
            existing.stock = data['stock']
            existing.low_stock = data['low_stock']
            existing.is_active = data['is_active']
            _apply_import_taxonomy(existing, data)
            _apply_import_image_url(existing, data)
            updated += 1
            continue
        p = Product(
            sku=sku, name=name,
            model=data['model'], variant=data['variant'],
            variant_prices=data.get('variant_prices') or '',
            variant_stocks=data.get('variant_stocks') or '',
            variant_cost_prices=data.get('variant_cost_prices') or '',
            variant_brand_ids=data.get('variant_brand_ids') or '',
            warranty=data['warranty'], cost_price=data['cost_price'],
            retail_price=data['retail_price'], dealer_price=data['dealer_price'],
            project_price=data['project_price'], stock=data['stock'],
            low_stock=data['low_stock'], is_active=data['is_active'],
        )
        _apply_import_unit_fields(p, data)
        _apply_import_taxonomy(p, data)
        _apply_import_image_url(p, data)
        db.session.add(p)
        created += 1
    if created == 0 and updated == 0:
        raise ValueError('Không có sản phẩm nào được import')
    db.session.commit()
    return {
        'created': created, 'updated': updated, 'skipped': skipped,
        'errors': errors, 'error_lines': error_lines[:8],
    }

def refresh_quote_subtotals(quote):
    quote.subtotal = sum((item.amount or 0) for item in quote.items)
    recalculate_quote_totals(quote)

def delete_quote_doc_files(quote):
    for path_attr in ('sale_contract_path', 'handover_doc_path'):
        path = getattr(quote, path_attr, '') or ''
        if path and Path(path).is_file():
            Path(path).unlink(missing_ok=True)

def delete_contract_files(contract):
    if contract.file_path and Path(contract.file_path).is_file():
        Path(contract.file_path).unlink(missing_ok=True)
    delete_contract_scan_file(contract.signed_scan_path)

def delete_payment_record_files(payment):
    if payment.receipt_path:
        fp = BASE_DIR / 'static' / payment.receipt_path
        if fp.is_file():
            fp.unlink(missing_ok=True)

def delete_order_record(order, reset_quote=True):
    quote = db.session.get(Quote, order.quote_id) if order.quote_id else None
    for payment in Payment.query.filter_by(order_id=order.id).all():
        delete_payment_record_files(payment)
    Payment.query.filter_by(order_id=order.id).delete(synchronize_session=False)
    delete_order_handover_scan_file(order.handover_scan_path)
    db.session.delete(order)
    if reset_quote and quote and quote_display_status(quote) == 'Đã chốt':
        quote.status = 'Mới tạo'

def delete_quote_record(quote):
    order = get_quote_order(quote)
    if order and (order.paid_amount or 0) > 0:
        raise ValueError(
            f'Báo giá đang liên kết đơn {order.order_code} đã có thanh toán công nợ. '
            'Xóa thanh toán công nợ trước.'
        )
    if order:
        delete_order_record(order, reset_quote=False)
    delete_quote_doc_files(quote)
    db.session.delete(quote)

def reverse_payment_record(payment):
    order = payment.order
    if order:
        order.paid_amount = max(int(order.paid_amount or 0) - int(payment.amount or 0), 0)
    delete_payment_record_files(payment)
    db.session.delete(payment)

def delete_payment_group(payment_id):
    payment = Payment.query.get_or_404(payment_id)
    customer_id = payment.order.customer_id if payment.order else None
    if payment.batch_id:
        payments = Payment.query.filter_by(batch_id=payment.batch_id).all()
    else:
        payments = [payment]
    total = sum(p.amount or 0 for p in payments)
    for p in payments:
        reverse_payment_record(p)
    return customer_id, len(payments), total

def delete_product_record(product):
    affected_ids = [
        row[0] for row in db.session.query(QuoteItem.quote_id).filter_by(product_id=product.id).distinct().all()
    ]
    delete_all_product_image_files(product)
    QuoteItem.query.filter_by(product_id=product.id).delete(synchronize_session=False)
    StockMovement.query.filter_by(product_id=product.id).delete(synchronize_session=False)
    ProductSupplier.query.filter_by(product_id=product.id).delete(synchronize_session=False)
    PriceHistory.query.filter_by(product_id=product.id).delete(synchronize_session=False)
    db.session.delete(product)
    for qid in affected_ids:
        q = db.session.get(Quote, qid)
        if q:
            refresh_quote_subtotals(q)
            generate_quote_documents(q, commit=False)

def delete_customer_record(customer):
    for order in Order.query.filter_by(customer_id=customer.id).all():
        delete_order_record(order, reset_quote=True)
    for quote in Quote.query.filter_by(customer_id=customer.id).all():
        delete_quote_doc_files(quote)
        db.session.delete(quote)
    for contract in Contract.query.filter_by(customer_id=customer.id).all():
        delete_contract_files(contract)
        db.session.delete(contract)
    db.session.delete(customer)

def delete_all_products_data():
    """Xóa toàn bộ sản phẩm và dữ liệu liên quan (danh mục, kho, báo giá chi tiết)."""
    product_count = Product.query.count()
    QuoteItem.query.delete()
    StockMovement.query.delete()
    ProductSupplier.query.delete()
    PriceHistory.query.delete()
    Product.query.delete()
    Brand.query.delete()
    Category.query.delete()
    db.session.commit()
    clear_product_import_session()
    if PRODUCT_UPLOAD_DIR.exists():
        for path in PRODUCT_UPLOAD_DIR.iterdir():
            if path.is_file():
                path.unlink()
    return product_count

def empty_directory_tree(directory):
    """Xóa mọi file/thư mục con; giữ nguyên thư mục gốc."""
    if not directory.exists():
        return
    for path in directory.iterdir():
        try:
            if path.is_file():
                path.unlink()
            elif path.is_dir():
                shutil.rmtree(path)
        except OSError:
            pass

def clear_all_business_data():
    """Xóa toàn bộ dữ liệu nghiệp vụ và file đính kèm; giữ bảng User (tài khoản đăng nhập)."""
    counts = {
        'supplier_payments': SupplierPayment.query.count(),
        'purchase_orders': PurchaseOrder.query.count(),
        'supplier_intakes': SupplierIntake.query.count(),
        'payments': Payment.query.count(),
        'stock_movements': StockMovement.query.count(),
        'quotes': Quote.query.count(),
        'orders': Order.query.count(),
        'contracts': Contract.query.count(),
        'products': Product.query.count(),
        'customers': Customer.query.count(),
        'suppliers': Supplier.query.count(),
    }
    SupplierPayment.query.delete(synchronize_session=False)
    PurchaseOrderLine.query.delete(synchronize_session=False)
    PurchaseOrder.query.delete(synchronize_session=False)
    SupplierIntakeLine.query.delete(synchronize_session=False)
    SupplierIntake.query.delete(synchronize_session=False)
    Payment.query.delete(synchronize_session=False)
    StockMovement.query.delete(synchronize_session=False)
    QuoteItem.query.delete(synchronize_session=False)
    Order.query.delete(synchronize_session=False)
    Quote.query.delete(synchronize_session=False)
    Contract.query.delete(synchronize_session=False)
    ProductSupplier.query.delete(synchronize_session=False)
    PriceHistory.query.delete(synchronize_session=False)
    Product.query.delete(synchronize_session=False)
    Supplier.query.delete(synchronize_session=False)
    Customer.query.delete(synchronize_session=False)
    Brand.query.delete(synchronize_session=False)
    Category.query.delete(synchronize_session=False)
    CompanyProfile.query.delete(synchronize_session=False)
    db.session.commit()

    for upload_dir in (
        PRODUCT_UPLOAD_DIR,
        COMPANY_LOGO_DIR,
        CONTRACT_SIGNED_UPLOAD_DIR,
        ORDER_HANDOVER_UPLOAD_DIR,
        PAYMENT_RECEIPT_UPLOAD_DIR,
        SUPPLIER_PAYMENT_RECEIPT_UPLOAD_DIR,
        IMPORT_PREVIEW_DIR,
        TEMP_DOC_DIR,
    ):
        empty_directory_tree(upload_dir)
    empty_directory_tree(QUOTE_DOCS_DIR)
    empty_directory_tree(OUTPUT_DIR / 'contracts')

    ensure_walkin_customer()
    ensure_company_profiles()
    return counts

def _import_preview_path(preview_id):
    return IMPORT_PREVIEW_DIR / f'{preview_id}.json'

def _cleanup_old_import_previews(max_age_seconds=3600):
    if not IMPORT_PREVIEW_DIR.exists():
        return
    cutoff = time.time() - max_age_seconds
    for path in IMPORT_PREVIEW_DIR.glob('*.json'):
        try:
            if path.stat().st_mtime < cutoff:
                path.unlink()
        except OSError:
            pass

def clear_product_import_session():
    preview_id = session.pop(IMPORT_PREVIEW_SESSION_KEY, None)
    if preview_id:
        path = _import_preview_path(preview_id)
        if path.exists():
            try:
                path.unlink()
            except OSError:
                pass

def clear_supplier_intake_import_session():
    preview_id = session.pop(SUPPLIER_INTAKE_IMPORT_SESSION_KEY, None)
    if preview_id:
        path = _import_preview_path(preview_id)
        if path.exists():
            try:
                path.unlink()
            except OSError:
                pass

def get_supplier_intake_import_session():
    preview_id = session.get(SUPPLIER_INTAKE_IMPORT_SESSION_KEY)
    if not preview_id:
        return None
    path = _import_preview_path(preview_id)
    if not path.exists():
        session.pop(SUPPLIER_INTAKE_IMPORT_SESSION_KEY, None)
        return None
    try:
        return json.loads(path.read_text(encoding='utf-8'))
    except (json.JSONDecodeError, OSError):
        clear_supplier_intake_import_session()
        return None

def save_supplier_intake_import_session(supplier_id, filename, update_existing, preview_rows):
    IMPORT_PREVIEW_DIR.mkdir(parents=True, exist_ok=True)
    clear_supplier_intake_import_session()
    preview_id = uuid.uuid4().hex
    payload = {
        'supplier_id': supplier_id,
        'filename': filename,
        'update_existing': update_existing,
        'rows': preview_rows,
    }
    _import_preview_path(preview_id).write_text(
        json.dumps(payload, ensure_ascii=False),
        encoding='utf-8',
    )
    session[SUPPLIER_INTAKE_IMPORT_SESSION_KEY] = preview_id
    _cleanup_old_import_previews()

def get_product_import_session():
    preview_id = session.get(IMPORT_PREVIEW_SESSION_KEY)
    if not preview_id:
        return None
    path = _import_preview_path(preview_id)
    if not path.exists():
        session.pop(IMPORT_PREVIEW_SESSION_KEY, None)
        return None
    try:
        return json.loads(path.read_text(encoding='utf-8'))
    except (json.JSONDecodeError, OSError):
        clear_product_import_session()
        return None

def save_product_import_session(filename, update_existing, preview_rows, import_meta=None):
    IMPORT_PREVIEW_DIR.mkdir(parents=True, exist_ok=True)
    clear_product_import_session()
    preview_id = uuid.uuid4().hex
    payload = {
        'filename': filename,
        'update_existing': update_existing,
        'rows': preview_rows,
    }
    if import_meta:
        payload['import_meta'] = import_meta
    _import_preview_path(preview_id).write_text(
        json.dumps(payload, ensure_ascii=False),
        encoding='utf-8',
    )
    session[IMPORT_PREVIEW_SESSION_KEY] = preview_id
    _cleanup_old_import_previews()

def import_products_from_rows(rows, update_existing=False):
    preview = build_product_import_preview(rows, update_existing)
    importable = [r for r in preview if r.get('importable')]
    if not importable:
        raise ValueError('Không có dòng hợp lệ để import')
    return import_product_data_list(importable, update_existing)

def redirect_after_category_manage(category_import_preview=False, modal_tab=None):
    """Quay lại trang sản phẩm và mở lại modal quản lý danh mục."""
    ref = request.referrer or url_for('products')
    parts = urlparse(ref)
    qs = dict(parse_qsl(parts.query, keep_blank_values=True))
    qs['category_modal'] = '1'
    if category_import_preview:
        qs['category_import_preview'] = '1'
    if modal_tab in ('categories', 'brands'):
        qs['category_modal_tab'] = modal_tab
    return redirect(urlunparse(parts._replace(query=urlencode(qs))))

CATEGORY_IMPORT_FIELDS = [
    ('category_parent', 'Danh mục cha'),
    ('category_child', 'Danh mục con'),
    ('brand', 'Thương hiệu'),
]

CATEGORY_IMPORT_SAMPLE_ROWS = [
    ['Điện thoại', 'Smartphone', 'Samsung'],
    ['Điện thoại', 'Tablet', 'Apple'],
    ['Điện thoại', 'Phụ kiện điện thoại', ''],
    ['Camera', 'Action camera', 'DJI'],
    ['Camera', 'Máy ảnh', 'Sony'],
    ['Laptop', '', 'Asus'],
]

CATEGORY_IMPORT_HEADER_ALIASES = {
    'category_parent': {'danh mục cha', 'danh muc cha', 'nhóm', 'ngành hàng', 'parent'},
    'category_child': {'danh mục con', 'danh muc con', 'loại', 'child'},
    'brand': {'thương hiệu', 'thuong hieu', 'brand', 'hãng'},
    'category': {'danh mục', 'danh muc', 'category'},
}

def build_category_import_header_map(headers):
    col_map = {}
    for idx, raw in enumerate(headers):
        key = normalize_import_header(raw)
        if not key:
            continue
        for field, aliases in CATEGORY_IMPORT_HEADER_ALIASES.items():
            if key in aliases or key == field:
                col_map[field] = idx
                break
    return col_map

def category_row_from_import(row, col_map):
    def cell(field, default=''):
        idx = col_map.get(field)
        if idx is None or idx >= len(row):
            return default
        val = row[idx]
        if val is None:
            return default
        return str(val).strip()

    data = {
        'category_parent': cell('category_parent'),
        'category_child': cell('category_child'),
        'brand': cell('brand'),
    }
    combined = cell('category')
    if combined and not data['category_parent'] and not data['category_child']:
        if '>' in combined:
            parts = [p.strip() for p in combined.split('>') if p.strip()]
            data['category_parent'] = parts[0]
            data['category_child'] = parts[-1] if len(parts) > 1 else ''
        else:
            data['category_parent'] = combined
    return data

def _category_import_notes(parent_name, child_name, brand_name, parent_exists, child_exists, brand_exists):
    parts = []
    if parent_name and not parent_exists:
        parts.append(f'Thêm danh mục cha: {parent_name}')
    if child_name and not child_exists:
        parts.append(f'Thêm danh mục con: {child_name}')
    if brand_name and not brand_exists:
        parts.append(f'Thêm thương hiệu: {brand_name}')
    if not parts:
        return 'Đã có đủ trong hệ thống'
    return ' · '.join(parts)

def build_category_import_preview(rows):
    if not rows:
        raise ValueError('File trống')
    col_map = build_category_import_header_map(rows[0])
    if 'category_parent' not in col_map and 'category' not in col_map:
        raise ValueError('File phải có cột Danh mục cha (xem file mẫu)')
    preview = []
    for line_no, raw_row in enumerate(rows[1:], start=2):
        if not any(str(c).strip() for c in raw_row):
            continue
        data = category_row_from_import(raw_row, col_map)
        parent_name = data['category_parent']
        child_name = data['category_child']
        brand_name = data['brand']
        row_id = uuid.uuid4().hex[:12]
        if not parent_name and not child_name and not brand_name:
            continue
        if child_name and not parent_name:
            preview.append({
                'row_id': row_id, 'line_no': line_no, 'data': data,
                'action': 'error', 'action_label': 'Lỗi',
                'note': 'Có danh mục con nhưng thiếu danh mục cha',
                'importable': False,
            })
            continue
        if brand_name and not parent_name and not child_name:
            preview.append({
                'row_id': row_id, 'line_no': line_no, 'data': data,
                'action': 'error', 'action_label': 'Lỗi',
                'note': 'Thương hiệu cần gắn với danh mục cha hoặc con',
                'importable': False,
            })
            continue
        parent = None
        if parent_name:
            parent = Category.query.filter(
                Category.parent_id.is_(None), Category.name == parent_name
            ).first()
        child = None
        if parent and child_name:
            child = Category.query.filter_by(parent_id=parent.id, name=child_name).first()
        scope_id = child.id if child else (parent.id if parent else None)
        brand = Brand.query.filter_by(name=brand_name).first() if brand_name else None
        parent_exists = parent is not None
        child_exists = child is not None if child_name else True
        brand_exists = brand is not None if brand_name else True
        will_create = (
            (parent_name and not parent_exists)
            or (child_name and not child_exists)
            or (brand_name and not brand_exists)
        )
        if not will_create:
            preview.append({
                'row_id': row_id, 'line_no': line_no, 'data': data,
                'action': 'skip', 'action_label': 'Bỏ qua',
                'note': _category_import_notes(
                    parent_name, child_name, brand_name,
                    parent_exists, child_exists, brand_exists,
                ),
                'importable': False,
            })
            continue
        preview.append({
            'row_id': row_id, 'line_no': line_no, 'data': data,
            'action': 'create', 'action_label': 'Thêm mới',
            'note': _category_import_notes(
                parent_name, child_name, brand_name,
                parent_exists, child_exists, brand_exists,
            ),
            'importable': True,
        })
    if not preview:
        raise ValueError('Không có dòng dữ liệu hợp lệ trong file')
    return preview

def apply_category_import_row(data):
    created = {'parent': 0, 'child': 0, 'brand': 0}
    parent_name = (data.get('category_parent') or '').strip()
    child_name = (data.get('category_child') or '').strip()
    brand_name = (data.get('brand') or '').strip()
    parent_id = child_id = None
    if parent_name:
        parent = Category.query.filter(
            Category.parent_id.is_(None), Category.name == parent_name
        ).first()
        if not parent:
            parent = Category(name=parent_name)
            db.session.add(parent)
            db.session.flush()
            created['parent'] += 1
        parent_id = parent.id
    if child_name:
        child = Category.query.filter_by(parent_id=parent_id, name=child_name).first()
        if not child:
            child = Category(name=child_name, parent_id=parent_id)
            db.session.add(child)
            db.session.flush()
            created['child'] += 1
        child_id = child.id
    scope_id = child_id or parent_id
    if brand_name:
        brand = Brand.query.filter_by(name=brand_name).first()
        if not brand:
            db.session.add(Brand(name=brand_name, category_id=scope_id))
            created['brand'] += 1
        elif scope_id and not brand.category_id:
            brand.category_id = scope_id
    return created

def import_category_data_list(items):
    totals = {'parent': 0, 'child': 0, 'brand': 0}
    for item in items:
        data = item.get('data') or item
        row_created = apply_category_import_row(data)
        for k in totals:
            totals[k] += row_created[k]
    if sum(totals.values()) == 0:
        raise ValueError('Không có danh mục/thương hiệu mới nào được import')
    db.session.commit()
    return totals

def clear_category_import_session():
    preview_id = session.pop(CATEGORY_IMPORT_SESSION_KEY, None)
    if preview_id:
        path = _import_preview_path(f'cat_{preview_id}')
        if path.exists():
            try:
                path.unlink()
            except OSError:
                pass

def get_category_import_session():
    preview_id = session.get(CATEGORY_IMPORT_SESSION_KEY)
    if not preview_id:
        return None
    path = _import_preview_path(f'cat_{preview_id}')
    if not path.exists():
        session.pop(CATEGORY_IMPORT_SESSION_KEY, None)
        return None
    try:
        return json.loads(path.read_text(encoding='utf-8'))
    except (json.JSONDecodeError, OSError):
        clear_category_import_session()
        return None

def save_category_import_session(filename, preview_rows):
    IMPORT_PREVIEW_DIR.mkdir(parents=True, exist_ok=True)
    clear_category_import_session()
    preview_id = uuid.uuid4().hex
    payload = {'filename': filename, 'rows': preview_rows}
    _import_preview_path(f'cat_{preview_id}').write_text(
        json.dumps(payload, ensure_ascii=False),
        encoding='utf-8',
    )
    session[CATEGORY_IMPORT_SESSION_KEY] = preview_id
    _cleanup_old_import_previews()

def category_import_field_labels():
    return [label for _, label in CATEGORY_IMPORT_FIELDS]

def build_category_import_xlsx_bytes(rows):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    from openpyxl.utils import get_column_letter
    wb = Workbook()
    ws = wb.active
    ws.title = 'Danh muc'
    headers = category_import_field_labels()
    ws.append(headers)
    for row in rows:
        ws.append(row)
    for cell in ws[1]:
        cell.font = Font(bold=True)
        cell.fill = PatternFill('solid', fgColor='E8F0FE')
        cell.alignment = Alignment(horizontal='center')
    ws.freeze_panes = 'A2'
    for idx, width in enumerate([18, 20, 16], start=1):
        ws.column_dimensions[get_column_letter(idx)].width = width
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf

def build_category_import_csv_text(rows):
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow(category_import_field_labels())
    for row in rows:
        writer.writerow(row)
    return '\ufeff' + buf.getvalue()

def taxonomy_export_rows():
    rows = []
    seen_brands = set()
    for parent in parent_categories():
        rows.append([parent.name, '', ''])
        for child in child_categories(parent.id):
            brand_names = [
                b.name for b in Brand.query.filter_by(category_id=child.id).order_by(Brand.name)
            ]
            if brand_names:
                for bn in brand_names:
                    rows.append([parent.name, child.name, bn])
                    seen_brands.add(bn)
            else:
                rows.append([parent.name, child.name, ''])
        for brand in Brand.query.filter_by(category_id=parent.id).order_by(Brand.name):
            if brand.name not in seen_brands:
                rows.append([parent.name, '', brand.name])
                seen_brands.add(brand.name)
    for brand in Brand.query.filter(Brand.category_id.is_(None)).order_by(Brand.name):
        if brand.name not in seen_brands:
            rows.append(['', '', brand.name])
    return rows

@app.route('/categories', methods=['POST'])
def create_category():
    ensure_category_brand_schema()
    keep_modal = request.form.get('_keep_modal') == '1'
    name = request.form.get('name', '').strip()
    parent_id = request.form.get('parent_id', type=int)
    if not name:
        flash('Vui lòng nhập tên danh mục', 'warning')
    elif category_name_exists(name, parent_id):
        flash('Tên danh mục đã tồn tại trong nhóm này', 'danger')
    elif parent_id and not Category.query.get(parent_id):
        flash('Danh mục cha không hợp lệ', 'danger')
    else:
        db.session.add(Category(name=name, parent_id=parent_id))
        db.session.commit()
        label = 'danh mục con' if parent_id else 'danh mục cha'
        flash(f'Đã tạo {label}', 'success')
    if keep_modal:
        return redirect_after_category_manage(modal_tab='categories')
    return redirect(request.referrer or url_for('products'))

@app.route('/categories/<int:cid>/delete', methods=['POST'])
def delete_category(cid):
    ensure_category_brand_schema()
    cat = Category.query.get_or_404(cid)
    name = cat.name
    children = Category.query.filter_by(parent_id=cat.id).count()
    if children:
        flash(f'Không thể xóa "{name}" vì còn {children} danh mục con', 'danger')
        if request.form.get('_keep_modal') == '1':
            return redirect_after_category_manage(modal_tab='categories')
        return redirect(request.referrer or url_for('products'))
    path = category_display_path(cat.id)
    affected = Product.query.filter(
        or_(Product.category_id == cat.id, Product.category == path, Product.category == cat.name)
    ).update({Product.category_id: None, Product.category: ''})
    Brand.query.filter_by(category_id=cat.id).update({Brand.category_id: None})
    db.session.delete(cat)
    db.session.commit()
    if affected:
        flash(f'Đã xóa danh mục "{name}". {affected} sản phẩm đã được bỏ danh mục.', 'warning')
    else:
        flash(f'Đã xóa danh mục "{name}"', 'success')
    if request.form.get('_keep_modal') == '1':
        return redirect_after_category_manage(modal_tab='categories')
    return redirect(request.referrer or url_for('products'))

@app.route('/brands', methods=['POST'])
def create_brand():
    ensure_category_brand_schema()
    keep_modal = request.form.get('_keep_modal') == '1'
    name = request.form.get('name', '').strip()
    category_id = request.form.get('category_id', type=int)
    if not name:
        flash('Vui lòng nhập tên thương hiệu', 'warning')
    elif Brand.query.filter_by(name=name).first():
        flash('Thương hiệu đã tồn tại', 'danger')
    else:
        db.session.add(Brand(name=name, category_id=category_id or None))
        db.session.commit()
        flash('Đã tạo thương hiệu', 'success')
    if keep_modal:
        return redirect_after_category_manage(modal_tab='brands')
    return redirect(request.referrer or url_for('products'))

@app.route('/brands/<int:bid>/delete', methods=['POST'])
def delete_brand(bid):
    ensure_category_brand_schema()
    brand = Brand.query.get_or_404(bid)
    name = brand.name
    affected = Product.query.filter(
        or_(Product.brand_id == brand.id, Product.brand == name)
    ).update({Product.brand_id: None, Product.brand: ''})
    db.session.delete(brand)
    db.session.commit()
    if affected:
        flash(f'Đã xóa thương hiệu "{name}". {affected} sản phẩm đã được bỏ thương hiệu.', 'warning')
    else:
        flash(f'Đã xóa thương hiệu "{name}"', 'success')
    if request.form.get('_keep_modal') == '1':
        return redirect_after_category_manage(modal_tab='brands')
    return redirect(request.referrer or url_for('products'))

@app.route('/categories/export')
def categories_export():
    ensure_category_brand_schema()
    rows = taxonomy_export_rows()
    fmt = (request.args.get('format') or 'xlsx').lower()
    if fmt == 'csv':
        return Response(
            build_category_import_csv_text(rows),
            mimetype='text/csv; charset=utf-8',
            headers={'Content-Disposition': 'attachment; filename=danh-muc.csv'},
        )
    try:
        buf = build_category_import_xlsx_bytes(rows)
        return send_file(
            buf,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name='danh-muc.xlsx',
        )
    except ImportError:
        return Response(
            build_category_import_csv_text(rows),
            mimetype='text/csv; charset=utf-8',
            headers={'Content-Disposition': 'attachment; filename=danh-muc.csv'},
        )

@app.route('/categories/import/template')
def categories_import_template():
    ensure_category_brand_schema()
    rows = list(CATEGORY_IMPORT_SAMPLE_ROWS)
    fmt = (request.args.get('format') or 'xlsx').lower()
    if fmt == 'csv':
        return Response(
            build_category_import_csv_text(rows),
            mimetype='text/csv; charset=utf-8',
            headers={'Content-Disposition': 'attachment; filename=mau-import-danh-muc.csv'},
        )
    try:
        buf = build_category_import_xlsx_bytes(rows)
        return send_file(
            buf,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name='mau-import-danh-muc.xlsx',
        )
    except ImportError:
        return Response(
            build_category_import_csv_text(rows),
            mimetype='text/csv; charset=utf-8',
            headers={'Content-Disposition': 'attachment; filename=mau-import-danh-muc.csv'},
        )

@app.route('/categories/import/preview', methods=['POST'])
def categories_import_preview():
    ensure_category_brand_schema()
    file_storage = request.files.get('import_file')
    if not file_storage or not file_storage.filename:
        flash('Vui lòng chọn file Excel hoặc CSV', 'warning')
        return redirect_after_category_manage()
    try:
        rows = read_product_import_sheet(file_storage)
        preview = build_category_import_preview(rows)
        save_category_import_session(file_storage.filename, preview)
    except ValueError as e:
        flash(str(e), 'danger')
        return redirect_after_category_manage()
    except Exception as e:
        flash(f'Không đọc được file: {e}', 'danger')
        return redirect_after_category_manage()
    return redirect_after_category_manage(category_import_preview=True)

@app.route('/categories/import/confirm', methods=['POST'])
def categories_import_confirm():
    ensure_category_brand_schema()
    payload = get_category_import_session()
    if not payload:
        flash('Phiên import đã hết hạn. Vui lòng tải file lại.', 'warning')
        return redirect_after_category_manage()
    selected_ids = set(request.form.getlist('import_row_id'))
    to_import = [
        r for r in payload.get('rows', [])
        if r.get('row_id') in selected_ids and r.get('importable')
    ]
    if not to_import:
        flash('Chưa chọn dòng nào để import', 'warning')
        return redirect_after_category_manage(category_import_preview=True)
    try:
        totals = import_category_data_list(to_import)
    except ValueError as e:
        flash(str(e), 'danger')
        return redirect_after_category_manage(category_import_preview=True)
    clear_category_import_session()
    parts = []
    if totals['parent']:
        parts.append(f"{totals['parent']} danh mục cha")
    if totals['child']:
        parts.append(f"{totals['child']} danh mục con")
    if totals['brand']:
        parts.append(f"{totals['brand']} thương hiệu")
    flash('Đã import: ' + ', '.join(parts), 'success')
    return redirect_after_category_manage()

@app.route('/categories/import/cancel', methods=['POST'])
def categories_import_cancel():
    clear_category_import_session()
    return redirect_after_category_manage()

@app.route('/products', methods=['GET', 'POST'])
def products():
    ensure_category_brand_schema()
    ensure_product_image_column()
    if request.method == 'POST':
        sku = request.form.get('sku', '').strip()
        name = request.form.get('name', '').strip()
        if not sku or not name:
            flash('Vui lòng nhập SKU và tên sản phẩm', 'warning')
            return redirect(url_for('products'))
        if Product.query.filter_by(sku=sku).first():
            flash('SKU đã tồn tại', 'danger')
            return redirect(url_for('products'))
        p = Product(
            sku=sku, name=name,
            warranty=request.form.get('warranty', ''),
            cost_price=parse_int(request.form.get('cost_price')),
            retail_price=parse_int(request.form.get('retail_price')),
            project_price=parse_int(request.form.get('project_price')),
            stock=parse_qty(request.form.get('stock')),
        )
        assign_product_taxonomy_from_form(p)
        apply_product_unit_from_form(p)
        p.model = request.form.get('model', '')
        p.variant = request.form.get('variant', '')
        p.low_stock = parse_qty(request.form.get('low_stock'), 5)
        apply_product_image_url_from_form(p)
        db.session.add(p)
        db.session.commit()
        flash('Đã thêm sản phẩm', 'success')
        return redirect(url_for('products'))
    filters = product_filters_from_request()
    per_page = request.args.get('per_page', 10, type=int)
    per_page = per_page if per_page in (10, 25, 50, 100) else 10
    page = request.args.get('page', 1, type=int)
    query = apply_product_filters(Product.query, filters)
    query = product_list_order(query, filters.get('sort'))
    pagination = query.paginate(page=page, per_page=per_page, error_out=False)
    categories = Category.query.order_by(Category.name).all()
    list_args = {k: v for k, v in filters.items() if v}
    list_args['per_page'] = per_page
    if pagination.page > 1:
        list_args['page'] = pagination.page
    product_list_ctx = {'filters': filters, 'per_page': per_page, 'page': pagination.page}
    product_list_nav = product_detail_url_kwargs(product_list_ctx)
    export_args = {k: v for k, v in list_args.items() if k != 'page'}
    price_histories = {}
    if pagination.items:
        ids = [p.id for p in pagination.items]
        for h in PriceHistory.query.filter(PriceHistory.product_id.in_(ids)).order_by(PriceHistory.created_at.desc()).all():
            price_histories.setdefault(h.product_id, []).append(h)
    return render_template(
        'products.html',
        pagination=pagination,
        products=pagination.items,
        categories=categories,
        category_counts=category_product_counts(),
        brand_counts=brand_product_counts(),
        category_tree=build_category_tree(),
        parent_categories=parent_categories(),
        brands_list=Brand.query.order_by(Brand.name).all(),
        filters=filters,
        per_page=per_page,
        list_args=list_args,
        product_list_nav=product_list_nav,
        export_args=export_args,
        brands=product_brand_options(),
        taxonomy_catalog_json=json.dumps(categories_catalog_json(), ensure_ascii=False),
        brands_catalog_json=json.dumps(brands_catalog_json(), ensure_ascii=False),
        price_histories=price_histories,
        stock_in_methods=STOCK_IN_METHODS,
        stock_out_methods=STOCK_OUT_METHODS,
        all_suppliers=active_suppliers(),
        import_preview_payload=get_product_import_session(),
        show_import_preview=bool(request.args.get('import_preview')),
        category_import_preview_payload=get_category_import_session(),
        show_category_import_preview=bool(request.args.get('category_import_preview')),
        stats=product_stats(),
        quick_category_tabs=top_product_category_tabs(),
        can_quotes=role_can_access(get_current_user().role, 'quotes') if get_current_user() else False,
        quote_catalog=quote_product_catalog(
            Product.query.filter(Product.is_active.is_(True)).order_by(Product.name).all()
        ) if get_current_user() and role_can_access(get_current_user().role, 'quotes') else [],
        customers=customers_for_select() if get_current_user() and role_can_access(get_current_user().role, 'quotes') else [],
        walkin_customer_id=ensure_walkin_customer().id if get_current_user() and role_can_access(get_current_user().role, 'quotes') else None,
        default_valid=(date.today() + timedelta(days=30)).isoformat(),
    )

@app.route('/products/<int:pid>/edit')
def product_edit(pid):
    ensure_category_brand_schema()
    ensure_product_image_column()
    ensure_product_supplier_schema()
    ensure_stock_movement_columns()
    p = Product.query.get_or_404(pid)
    price_histories = (
        PriceHistory.query.filter_by(product_id=p.id)
        .order_by(PriceHistory.created_at.desc())
        .limit(50)
        .all()
    )
    stock_movements = (
        StockMovement.query.filter_by(product_id=p.id)
        .order_by(StockMovement.created_at.desc())
        .limit(50)
        .all()
    )
    list_ctx = product_list_ctx_from_request()
    product_list_nav = product_detail_url_kwargs(list_ctx)
    return render_template(
        'product_edit.html',
        p=p,
        parent_categories=parent_categories(),
        brands_list=Brand.query.order_by(Brand.name).all(),
        price_histories=price_histories,
        stock_movements=stock_movements,
        stock_in_methods=STOCK_IN_METHODS,
        stock_out_methods=STOCK_OUT_METHODS,
        all_suppliers=active_suppliers(),
        product_list_nav=product_list_nav,
        products_back_url=products_list_url(
            list_ctx['per_page'], list_ctx['filters'], list_ctx['page'],
        ),
        taxonomy_catalog_json=json.dumps(categories_catalog_json(), ensure_ascii=False),
        brands_catalog_json=json.dumps(brands_catalog_json(), ensure_ascii=False),
        product_edit_snapshot_json=json.dumps(
            product_edit_form_snapshot(p), ensure_ascii=False,
        ),
    )

@app.route('/products/<int:pid>')
def product_detail(pid):
    ensure_category_brand_schema()
    ensure_product_image_column()
    ensure_product_supplier_schema()
    p = Product.query.get_or_404(pid)
    list_ctx = product_list_ctx_from_request()
    product_list_nav = product_detail_url_kwargs(list_ctx)
    price_histories = {
        p.id: PriceHistory.query.filter_by(product_id=p.id)
        .order_by(PriceHistory.created_at.desc())
        .limit(50)
        .all(),
    }
    return render_template(
        'product_detail.html',
        p=p,
        supplier_links=product_supplier_links(p),
        primary_supplier=product_primary_supplier_link(p),
        parent_categories=parent_categories(),
        brands_list=Brand.query.order_by(Brand.name).all(),
        price_histories=price_histories,
        stock_in_methods=STOCK_IN_METHODS,
        stock_out_methods=STOCK_OUT_METHODS,
        all_suppliers=active_suppliers(),
        filters=list_ctx['filters'],
        pagination=type('Pg', (), {'page': list_ctx['page']})(),
        per_page=list_ctx['per_page'],
        product_list_nav=product_list_nav,
        products_back_url=products_list_url(
            list_ctx['per_page'], list_ctx['filters'], list_ctx['page'],
        ),
        from_detail=True,
        taxonomy_catalog_json=json.dumps(categories_catalog_json(), ensure_ascii=False),
    )

@app.route('/products/<int:pid>/update', methods=['POST'])
def update_product(pid):
    ensure_category_brand_schema()
    ensure_product_image_column()
    p = Product.query.get_or_404(pid)
    sku = request.form.get('sku', '').strip()
    name = request.form.get('name', '').strip()
    if not sku or not name:
        flash('Vui lòng nhập SKU và tên sản phẩm', 'warning')
        return products_redirect_after_save()
    if Product.query.filter(Product.sku == sku, Product.id != pid).first():
        flash('SKU đã tồn tại', 'danger')
        return products_redirect_after_save()
    p.sku = sku
    apply_product_form(p)
    if request.form.getlist('ps_supplier_id'):
        apply_product_suppliers_from_form(p)
    db.session.commit()
    return products_redirect_after_save('Đã cập nhật sản phẩm')

@app.route('/products/<int:pid>/price', methods=['POST'])
def update_price(pid):
    p = Product.query.get_or_404(pid)
    changed = apply_product_price_from_form(p)
    db.session.commit()
    if changed:
        return products_redirect_after_save('Đã cập nhật giá')
    return products_redirect_after_save('Không có thay đổi giá', 'info')

@app.route('/products/export')
def products_export():
    ensure_category_brand_schema()
    filters = product_filters_from_request()
    query = apply_product_filters(Product.query, filters).order_by(Product.created_at.desc())
    rows = [product_to_export_row(p) for p in query.all()]
    fmt = (request.args.get('format') or 'xlsx').lower()
    if fmt == 'csv':
        return Response(
            build_products_csv_text(rows),
            mimetype='text/csv; charset=utf-8',
            headers={'Content-Disposition': 'attachment; filename=san-pham.csv'},
        )
    try:
        buf = build_products_xlsx_bytes(rows)
        return send_file(
            buf,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name='san-pham.xlsx',
        )
    except ImportError:
        return Response(
            build_products_csv_text(rows),
            mimetype='text/csv; charset=utf-8',
            headers={'Content-Disposition': 'attachment; filename=san-pham.csv'},
        )

@app.route('/products/import/template')
def products_import_template():
    ensure_category_brand_schema()
    rows = [PRODUCT_IMPORT_SAMPLE_ROW]
    sync_product_import_template_files(rows)
    fmt = (request.args.get('format') or 'xlsx').lower()
    if fmt == 'csv':
        return Response(
            build_products_csv_text(rows),
            mimetype='text/csv; charset=utf-8',
            headers={'Content-Disposition': 'attachment; filename=mau-import-san-pham.csv'},
        )
    try:
        buf = build_products_xlsx_bytes(rows)
        return send_file(
            buf,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name='mau-import-san-pham.xlsx',
        )
    except ImportError:
        return Response(
            build_products_csv_text(rows),
            mimetype='text/csv; charset=utf-8',
            headers={'Content-Disposition': 'attachment; filename=mau-import-san-pham.csv'},
        )

@app.route('/products/import/preview', methods=['POST'])
def products_import_preview():
    ensure_category_brand_schema()
    file_storage = request.files.get('import_file')
    if not file_storage or not file_storage.filename:
        flash('Vui lòng chọn file Excel hoặc CSV', 'warning')
        return redirect(url_for('products'))
    update_existing = request.form.get('update_existing') == '1'
    try:
        rows, import_meta = read_product_import_sheet(file_storage, with_meta=True)
        preview = build_product_import_preview(rows, update_existing=update_existing)
        save_product_import_session(
            file_storage.filename, update_existing, preview, import_meta=import_meta,
        )
        sheet_count = import_meta.get('sheet_count') or 0
        if sheet_count > 1:
            total = sum(s.get('rows', 0) for s in import_meta.get('sheets', []))
            flash(
                f'Đã gộp {sheet_count} sheet Excel ({total} dòng sản phẩm). '
                'Kiểm tra xem trước rồi xác nhận import.',
                'info',
            )
    except ValueError as e:
        flash(str(e), 'danger')
        return redirect(url_for('products'))
    except Exception as e:
        flash(f'Không đọc được file: {e}', 'danger')
        return redirect(url_for('products'))
    return redirect(url_for('products', import_preview=1))

@app.route('/products/import/confirm', methods=['POST'])
def products_import_confirm():
    ensure_category_brand_schema()
    payload = get_product_import_session()
    if not payload:
        flash('Phiên import đã hết hạn. Vui lòng tải file lại.', 'warning')
        return redirect(url_for('products'))
    selected_ids = set(request.form.getlist('import_row_id'))
    update_existing = payload.get('update_existing', False)
    to_import = [
        r for r in payload.get('rows', [])
        if r.get('row_id') in selected_ids and r.get('importable')
    ]
    if not to_import:
        flash('Chưa chọn sản phẩm nào để import', 'warning')
        return redirect(url_for('products', import_preview=1))
    try:
        result = import_product_data_list(to_import, update_existing=update_existing)
    except ValueError as e:
        flash(str(e), 'danger')
        return redirect(url_for('products', import_preview=1))
    clear_product_import_session()
    parts = []
    if result['created']:
        parts.append(f"thêm mới {result['created']}")
    if result['updated']:
        parts.append(f"cập nhật {result['updated']}")
    msg = 'Đã import: ' + ', '.join(parts) + '.'
    flash(msg, 'success')
    return redirect(url_for('products'))

@app.route('/products/import/cancel', methods=['POST'])
def products_import_cancel():
    clear_product_import_session()
    return redirect(url_for('products'))

@app.route('/products/delete-all', methods=['POST'])
def products_delete_all():
    ensure_product_columns()
    count = Product.query.count()
    if count == 0:
        flash('Không có sản phẩm nào để xóa', 'info')
        return redirect(url_for('products'))
    deleted = delete_all_products_data()
    flash(f'Đã xóa toàn bộ {deleted} sản phẩm và dữ liệu liên quan', 'success')
    return redirect(url_for('products'))

@app.route('/products/<int:pid>/toggle-active', methods=['POST'])
def toggle_product_active(pid):
    ensure_product_columns()
    p = Product.query.get_or_404(pid)
    p.is_active = not bool(p.is_active)
    db.session.commit()
    label = 'Đang bán' if p.is_active else 'Ngừng bán'
    return products_redirect_after_save(f'Đã chuyển sang trạng thái: {label}')

@app.route('/products/<int:pid>/delete', methods=['POST'])
def delete_product(pid):
    ensure_product_image_column()
    p = Product.query.get_or_404(pid)
    name = p.name or p.sku
    delete_product_record(p)
    db.session.commit()
    flash(f'Đã xóa sản phẩm "{name}"', 'success')
    if request.form.get('_from_detail') == '1':
        list_ctx = product_list_ctx_from_form()
        return redirect(products_list_url(
            list_ctx['per_page'], list_ctx['filters'], list_ctx['page'],
        ))
    return products_redirect_after_save()

@app.route('/customers/<int:cid>/delete', methods=['POST'])
def delete_customer(cid):
    ensure_customer_columns()
    ensure_contract_columns()
    c = Customer.query.get_or_404(cid)
    if is_walkin_customer(c):
        flash('Không thể xóa khách hàng hệ thống', 'warning')
        return customers_redirect_after_save()
    name = c.name
    delete_customer_record(c)
    db.session.commit()
    flash(f'Đã xóa khách hàng "{name}" và dữ liệu liên quan', 'success')
    return customers_redirect_after_save()

@app.route('/orders/<int:oid>/delete', methods=['POST'])
def delete_order(oid):
    ensure_order_columns()
    ensure_payment_columns()
    o = Order.query.get_or_404(oid)
    code = o.order_code
    delete_order_record(o, reset_quote=True)
    db.session.commit()
    flash(f'Đã xóa đơn hàng / hóa đơn {code}', 'success')
    if request.form.get('_from') == 'debt':
        return debt_redirect(
            customer_id=request.form.get('_customer_id', type=int),
            page=request.form.get('_page', 1, type=int) or 1,
            order_page=request.form.get('_order_page', 1, type=int) or 1,
        )
    if request.form.get('_from') == 'preview':
        return redirect(url_for('orders'))
    return orders_redirect(page=request.form.get('_page', 1, type=int) or 1)

@app.route('/products/<int:pid>/stock', methods=['POST'])
def adjust_product_stock(pid):
    ensure_stock_movement_columns()
    p = Product.query.get_or_404(pid)
    qty = max(parse_qty(request.form.get('qty')), 0)
    mtype = request.form.get('movement_type', 'IN')
    if qty <= 0:
        return products_redirect_after_save('Số lượng phải lớn hơn 0', 'warning')
    if mtype == 'OUT' and p.stock < qty - 1e-9:
        return products_redirect_after_save('Tồn kho không đủ để xuất', 'danger')
    if mtype == 'IN':
        p.stock += qty
    else:
        p.stock -= qty
    supplier_id = parse_supplier_id(request.form.get('supplier_id')) if mtype == 'IN' else None
    db.session.add(StockMovement(
        product_id=p.id,
        supplier_id=supplier_id,
        movement_type=mtype,
        qty=qty,
        purchase_qty=0,
        ref_code=request.form.get('ref_code', '').strip(),
        method=normalize_stock_method(mtype, request.form.get('method')),
        warehouse=DEFAULT_WAREHOUSE,
        note=request.form.get('note', '').strip(),
    ))
    stock_note = request.form.get('note', '').strip()
    price_note = f'Cùng lúc {("nhập" if mtype == "IN" else "xuất")} kho' + (f': {stock_note}' if stock_note else '')
    price_changed = apply_product_price_from_form(p, note=price_note)
    db.session.commit()
    label = 'nhập' if mtype == 'IN' else 'xuất'
    if price_changed:
        return products_redirect_after_save(f'Đã {label} kho và cập nhật giá')
    return products_redirect_after_save(f'Đã {label} kho thành công')

@app.route('/quotes/<int:qid>/delete', methods=['POST'])
def delete_quote(qid):
    ensure_quote_columns()
    quote = Quote.query.get_or_404(qid)
    code = quote.quote_code
    tab = request.form.get('_tab') or ('draft' if quote.status == 'Nháp' else 'list')
    try:
        delete_quote_record(quote)
    except ValueError as e:
        flash(str(e), 'danger')
        return quotes_redirect(tab=tab)
    db.session.commit()
    flash(f'Đã xóa báo giá {code}', 'success')
    if request.form.get('_from') == 'preview':
        return redirect(url_for('quotes', tab=tab))
    return quotes_redirect(tab=tab)

@app.route('/quotes/preview', methods=['POST'])
def quote_create_preview():
    quote, valid_until = build_quote_preview_from_form()
    return render_template(
        'partials/quote_create_preview_doc.html',
        quote=quote,
        company=get_company_profile(),
        valid_until=valid_until,
    )

@app.route('/quotes', methods=['GET', 'POST'])
def quotes():
    ensure_quote_columns()
    normalize_quote_statuses()
    if request.method == 'POST':
        try:
            quote = create_quote_from_form()
        except ValueError as e:
            flash(str(e) or 'Vui lòng kiểm tra lại thông tin báo giá', 'warning')
            return redirect(url_for('quotes', tab='list'))
        except KeyError:
            flash('Vui lòng kiểm tra lại thông tin báo giá', 'warning')
            return redirect(url_for('quotes', tab='list'))
        flash('Đã lưu nháp báo giá' if quote.status == 'Nháp' else 'Đã tạo báo giá',
              'secondary' if quote.status == 'Nháp' else 'success')
        return redirect(url_for('quote_preview', qid=quote.id))

    filters = quote_filters_from_request()
    per_page = request.args.get('per_page', 10, type=int)
    per_page = per_page if per_page in (10, 25, 50) else 10
    page = request.args.get('page', 1, type=int)

    query = Quote.query
    if filters['tab'] == 'draft':
        query = query.filter(Quote.status == 'Nháp')
    else:
        query = query.filter(Quote.status != 'Nháp')
    query = apply_quote_filters(query, filters)

    stats_query = Quote.query.filter(Quote.status != 'Nháp')
    query = apply_quote_sort(query, filters['sort'], filters['order'])
    pagination = query.paginate(page=page, per_page=per_page, error_out=False)
    orders_by_quote = orders_by_quote_ids([q.id for q in pagination.items])

    list_args = {k: v for k, v in filters.items() if v and k != 'tab'}
    list_args['per_page'] = per_page
    list_args_list = {**list_args, 'tab': 'list'}
    list_args_draft = {**list_args, 'tab': 'draft'}
    export_args = {k: v for k, v in list_args.items()}
    list_args_tab = {**list_args, 'tab': filters['tab']}

    return render_template(
        'quotes.html',
        quotes=pagination.items,
        orders_by_quote=orders_by_quote,
        pagination=pagination,
        stats=quote_stats(stats_query),
        filters=filters,
        list_args=list_args_tab,
        list_args_list=list_args_list,
        list_args_draft=list_args_draft,
        export_args=export_args,
        per_page=per_page,
        sort=filters['sort'],
        order=filters['order'],
        customers=customers_for_select(),
        products=Product.query.order_by(Product.name).all(),
        quote_catalog=quote_product_catalog(
            Product.query.filter(Product.is_active.is_(True)).order_by(Product.name).all()
        ),
        quote_filter_statuses=QUOTE_FILTER_STATUSES,
        quote_edit_statuses=QUOTE_FILTER_STATUSES,
        default_valid=(date.today() + timedelta(days=30)).isoformat(),
        quote_customer_id=request.args.get('quote_customer_id', type=int),
        walkin_customer_id=ensure_walkin_customer().id,
        companies=list_company_profiles(),
        active_company_id=get_company_row().id if get_company_row() else None,
        doc_tab=filters['tab'],
    )

@app.route('/quotes/export')
def quotes_export():
    ensure_quote_columns()
    normalize_quote_statuses()
    filters = quote_filters_from_request()
    query = Quote.query.filter(Quote.status != 'Nháp')
    query = apply_quote_filters(query, filters)
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow(['Mã báo giá', 'Khách hàng', 'Ngày tạo', 'Hiệu lực đến', 'Tổng tiền', 'Trạng thái'])
    for q in query.order_by(Quote.created_at.desc()).all():
        writer.writerow([
            q.quote_code,
            quote_customer_label(q),
            q.created_at.strftime('%d/%m/%Y') if q.created_at else '',
            quote_valid_until(q).strftime('%d/%m/%Y'),
            q.total,
            quote_display_status(q),
        ])
    return Response(
        '\ufeff' + buf.getvalue(),
        mimetype='text/csv; charset=utf-8',
        headers={'Content-Disposition': 'attachment; filename=bao-gia.csv'},
    )

@app.route('/quotes/<int:qid>/status', methods=['POST'])
def update_quote_status(qid):
    ensure_quote_columns()
    normalize_quote_statuses()
    quote = Quote.query.get_or_404(qid)
    new_status = request.form.get('status', '').strip()
    current = quote_display_status(quote)
    allowed = {
        'Nháp': {'Mới tạo'},
        'Mới tạo': {'Đã chốt', 'Đã hủy'},
    }
    if new_status in QUOTE_STATUSES and new_status in allowed.get(current, set()):
        if new_status == 'Đã chốt':
            order, created, stock_deducted, had_shortage = chot_quote(quote)
            set_quote_session_customer(quote)
            db.session.commit()
            if created:
                msg = f'Đã chốt báo giá và tạo đơn hàng {order.order_code}'
            else:
                msg = f'Đã chốt báo giá (đơn hàng {order.order_code})'
            detail = quote_chot_flash_detail(stock_deducted, had_shortage)
            if detail:
                msg += f' — {detail}'
            flash(msg, 'success')
            shortage_msg = quote_stock_shortage_flash_message(quote) if had_shortage else None
            if shortage_msg:
                flash(shortage_msg, 'warning')
        else:
            quote.status = new_status
            db.session.commit()
            flash_status_updated(new_status)
    elif new_status in QUOTE_STATUSES:
        flash('Không thể chuyển trạng thái báo giá theo hướng này', 'warning')
    return quotes_redirect(tab=request.form.get('_tab', 'list'))

@app.route('/quotes/<int:qid>/update', methods=['POST'])
def update_quote(qid):
    ensure_quote_columns()
    quote = Quote.query.get_or_404(qid)
    if quote_display_status(quote) == 'Đã hủy':
        flash('Không thể chỉnh sửa báo giá đã hủy', 'warning')
        return quotes_redirect(tab=request.form.get('_tab', 'list'))
    quote.customer_id = resolve_quote_customer_id(request.form.get('customer_id'))
    quote.vat_rate = parse_int(request.form.get('vat_rate'), 10)
    quote.discount = parse_int(request.form.get('discount'), 0)
    quote.note = request.form.get('note', '').strip()
    valid_raw = request.form.get('valid_until', '').strip()
    if valid_raw:
        quote.valid_until = datetime.strptime(valid_raw, '%Y-%m-%d').date()
    line_items, stopped_names = parse_quote_line_items_from_form(quote)
    if stopped_names:
        names = ', '.join(stopped_names[:5])
        extra = f' và {len(stopped_names) - 5} sản phẩm khác' if len(stopped_names) > 5 else ''
        flash(f'Sản phẩm đã ngừng bán không thể thêm vào báo giá: {names}{extra}', 'warning')
        return quotes_redirect(tab=request.form.get('_tab', 'list'))
    if not line_items:
        flash('Vui lòng thêm ít nhất một sản phẩm vào báo giá', 'warning')
        return quotes_redirect(tab=request.form.get('_tab', 'list'))
    apply_quote_line_items(quote, line_items)
    new_status = request.form.get('status', '').strip()
    old_status = quote_display_status(quote)
    if new_status in QUOTE_STATUSES:
        if new_status == 'Đã chốt' and old_status != 'Đã chốt':
            _, _, stock_deducted, had_shortage = chot_quote(quote)
            detail = quote_chot_flash_detail(stock_deducted, had_shortage)
            if detail:
                flash(f'Đã chốt báo giá — {detail}', 'success')
            shortage_msg = quote_stock_shortage_flash_message(quote) if had_shortage else None
            if shortage_msg:
                flash(shortage_msg, 'warning')
        elif new_status != 'Đã chốt':
            quote.status = new_status
    recalculate_quote_totals(quote)
    order = get_quote_order(quote)
    if order:
        order.total = quote.total
        order.customer_id = quote.customer_id
    generate_quote_documents(quote, commit=False)
    db.session.commit()
    flash('Đã cập nhật báo giá', 'success')
    return quotes_redirect(tab=request.form.get('_tab', 'list'))

def _quote_doc_preview_page(
    quote,
    file_path,
    doc_title,
    download_docx_endpoint,
    download_pdf_endpoint,
    view_pdf_endpoint,
    view_docx_endpoint,
):
    if not file_path or not Path(file_path).is_file():
        flash('Không tìm thấy file tài liệu', 'danger')
        return redirect(url_for('quote_preview', qid=quote.id))
    try:
        urls = prepare_document_preview_urls(
            file_path,
            view_pdf_endpoint,
            view_docx_endpoint,
            download_docx_endpoint,
            download_pdf_endpoint,
            qid=quote.id,
        )
    except FileNotFoundError as e:
        flash(str(e), 'danger')
        return redirect(url_for('quote_preview', qid=quote.id))
    return render_document_preview_page(
        doc_title,
        quote_customer_label(quote),
        urls,
        url_for('quote_preview', qid=quote.id),
    )

@app.route('/quotes/<int:qid>/sale-contract/preview')
def preview_quote_sale_contract(qid):
    ensure_quote_columns()
    quote = ensure_quote_documents(Quote.query.get_or_404(qid))
    return _quote_doc_preview_page(
        quote, quote.sale_contract_path,
        f'HĐMB/{quote.quote_code}',
        'download_quote_sale_contract',
        'download_quote_sale_contract_pdf',
        'view_quote_sale_contract_pdf',
        'view_quote_sale_contract',
    )

@app.route('/quotes/<int:qid>/sale-contract/view')
def view_quote_sale_contract(qid):
    ensure_quote_columns()
    quote = ensure_quote_documents(Quote.query.get_or_404(qid))
    if not quote.sale_contract_path or not Path(quote.sale_contract_path).is_file():
        from flask import abort
        abort(404)
    return send_docx_inline(quote.sale_contract_path, f'HĐMB-{quote.quote_code}.docx')

@app.route('/quotes/<int:qid>/sale-contract/view.pdf')
def view_quote_sale_contract_pdf(qid):
    ensure_quote_columns()
    quote = ensure_quote_documents(Quote.query.get_or_404(qid))
    if not quote.sale_contract_path or not Path(quote.sale_contract_path).is_file():
        abort(404)
    return send_pdf_inline(
        pdf_path_for_docx_or_abort(quote.sale_contract_path),
        f'HĐMB-{quote.quote_code}.pdf',
    )

@app.route('/quotes/<int:qid>/sale-contract/download')
def download_quote_sale_contract(qid):
    ensure_quote_columns()
    quote = ensure_quote_documents(Quote.query.get_or_404(qid))
    if not quote.sale_contract_path or not Path(quote.sale_contract_path).is_file():
        flash('Không tìm thấy hợp đồng mua bán', 'danger')
        return redirect(url_for('quote_preview', qid=qid))
    return send_file(quote.sale_contract_path, as_attachment=True, download_name=Path(quote.sale_contract_path).name)

@app.route('/quotes/<int:qid>/sale-contract/download.pdf')
def download_quote_sale_contract_pdf(qid):
    ensure_quote_columns()
    quote = ensure_quote_documents(Quote.query.get_or_404(qid))
    if not quote.sale_contract_path or not Path(quote.sale_contract_path).is_file():
        flash('Không tìm thấy hợp đồng mua bán', 'danger')
        return redirect(url_for('quote_preview', qid=qid))
    pdf_path = ensure_pdf_for_docx(quote.sale_contract_path)
    return send_file(
        pdf_path,
        mimetype=PDF_MIMETYPE,
        as_attachment=True,
        download_name=Path(quote.sale_contract_path).stem + '.pdf',
    )

@app.route('/quotes/<int:qid>/handover/preview')
def preview_quote_handover(qid):
    ensure_quote_columns()
    quote = Quote.query.get_or_404(qid)
    try:
        ensure_quote_documents(quote)
    except Exception as e:
        app.logger.exception('ensure_quote_documents handover')
        flash(f'Không tạo được biên bản giao nhận: {e}', 'danger')
        return redirect(url_for('quote_preview', qid=qid))
    return _quote_doc_preview_page(
        quote, quote.handover_doc_path,
        f'BBGH/{quote.quote_code}',
        'download_quote_handover',
        'download_quote_handover_pdf',
        'view_quote_handover_pdf',
        'view_quote_handover',
    )

@app.route('/quotes/<int:qid>/handover/view')
def view_quote_handover(qid):
    ensure_quote_columns()
    quote = ensure_quote_documents(Quote.query.get_or_404(qid))
    if not quote.handover_doc_path or not Path(quote.handover_doc_path).is_file():
        from flask import abort
        abort(404)
    return send_docx_inline(quote.handover_doc_path, f'BBGH-{quote.quote_code}.docx')

@app.route('/quotes/<int:qid>/handover/view.pdf')
def view_quote_handover_pdf(qid):
    ensure_quote_columns()
    quote = ensure_quote_documents(Quote.query.get_or_404(qid))
    if not quote.handover_doc_path or not Path(quote.handover_doc_path).is_file():
        abort(404)
    return send_pdf_inline(
        pdf_path_for_docx_or_abort(quote.handover_doc_path),
        f'BBGH-{quote.quote_code}.pdf',
    )

@app.route('/quotes/<int:qid>/handover/download')
def download_quote_handover(qid):
    ensure_quote_columns()
    quote = ensure_quote_documents(Quote.query.get_or_404(qid))
    if not quote.handover_doc_path or not Path(quote.handover_doc_path).is_file():
        flash('Không tìm thấy biên bản nhận hàng', 'danger')
        return redirect(url_for('quote_preview', qid=qid))
    return send_file(quote.handover_doc_path, as_attachment=True, download_name=Path(quote.handover_doc_path).name)

@app.route('/quotes/<int:qid>/handover/download.pdf')
def download_quote_handover_pdf(qid):
    ensure_quote_columns()
    quote = ensure_quote_documents(Quote.query.get_or_404(qid))
    if not quote.handover_doc_path or not Path(quote.handover_doc_path).is_file():
        flash('Không tìm thấy biên bản nhận hàng', 'danger')
        return redirect(url_for('quote_preview', qid=qid))
    pdf_path = ensure_pdf_for_docx(quote.handover_doc_path)
    return send_file(
        pdf_path,
        mimetype=PDF_MIMETYPE,
        as_attachment=True,
        download_name=Path(quote.handover_doc_path).stem + '.pdf',
    )

QUOTE_SALE_DRAFT_KEY = 'quote_sale_draft_path'
QUOTE_HANDOVER_DRAFT_KEY = 'quote_handover_draft_path'

def _quote_doc_draft_preview(
    quote,
    doc,
    title,
    view_pdf_endpoint,
    view_docx_endpoint,
    download_docx_endpoint,
    download_pdf_endpoint,
    back_url,
    session_key,
):
    _unlink_preview_file(session.pop(session_key, None))
    try:
        temp_path = save_temp_docx(doc, 'quote_doc_draft')
    except Exception as e:
        flash(f'Không tạo được file Word: {e}', 'danger')
        return redirect(back_url)
    session[session_key] = str(temp_path)
    try:
        urls = prepare_document_preview_urls(
            temp_path,
            view_pdf_endpoint,
            view_docx_endpoint,
            download_docx_endpoint,
            download_pdf_endpoint,
        )
    except FileNotFoundError as e:
        _unlink_preview_file(str(temp_path))
        flash(str(e), 'danger')
        return redirect(back_url)
    return render_document_preview_page(
        title,
        quote_customer_label(quote),
        urls,
        back_url,
        is_draft=True,
    )

@app.route('/quotes/<int:qid>/sale-contract/update', methods=['POST'])
def update_quote_sale_contract(qid):
    ensure_quote_columns()
    quote = Quote.query.get_or_404(qid)
    apply_quote_sale_contract_form(quote)
    regenerate_quote_sale_contract(quote, commit=True)
    flash('Đã cập nhật hợp đồng mua bán', 'success')
    return quote_doc_redirect(qid)

@app.route('/quotes/<int:qid>/handover/update', methods=['POST'])
def update_quote_handover(qid):
    ensure_quote_columns()
    quote = Quote.query.get_or_404(qid)
    apply_quote_handover_form(quote)
    regenerate_quote_handover(quote, commit=True)
    flash('Đã cập nhật biên bản giao nhận hàng', 'success')
    return quote_doc_redirect(qid)

@app.route('/quotes/<int:qid>/sale-contract/preview-draft', methods=['POST'])
def preview_quote_sale_contract_draft(qid):
    ensure_quote_columns()
    quote = Quote.query.get_or_404(qid)
    saved = {
        'doc_company_id': quote.doc_company_id,
        'sale_contract_signed_date': quote.sale_contract_signed_date,
        'sale_contract_payment_note': quote.sale_contract_payment_note,
        'sale_contract_delivery_note': quote.sale_contract_delivery_note,
    }
    apply_quote_sale_contract_form(quote)
    doc = build_quote_sale_contract_doc(quote)
    for k, v in saved.items():
        setattr(quote, k, v)
    back = url_for('quote_preview', qid=qid) if request.form.get('_back') == 'preview' else url_for('quotes')
    return _quote_doc_draft_preview(
        quote, doc, f'HĐMB/{quote.quote_code}',
        'view_quote_sale_contract_draft_pdf',
        'view_quote_sale_contract_draft',
        'download_quote_sale_contract_draft',
        'download_quote_sale_contract_draft_pdf',
        back, QUOTE_SALE_DRAFT_KEY,
    )

@app.route('/quotes/<int:qid>/handover/preview-draft', methods=['POST'])
def preview_quote_handover_draft(qid):
    ensure_quote_columns()
    quote = Quote.query.get_or_404(qid)
    saved = {
        'doc_company_id': quote.doc_company_id,
        'handover_date': quote.handover_date,
        'handover_place': quote.handover_place,
        'handover_condition_note': quote.handover_condition_note,
    }
    apply_quote_handover_form(quote)
    doc = build_quote_handover_doc(quote)
    for k, v in saved.items():
        setattr(quote, k, v)
    back = url_for('quote_preview', qid=qid) if request.form.get('_back') == 'preview' else url_for('quotes')
    return _quote_doc_draft_preview(
        quote, doc, f'BBGH/{quote.quote_code}',
        'view_quote_handover_draft_pdf',
        'view_quote_handover_draft',
        'download_quote_handover_draft',
        'download_quote_handover_draft_pdf',
        back, QUOTE_HANDOVER_DRAFT_KEY,
    )

def _draft_docx_path(session_key):
    path = session.get(session_key)
    if not path or not Path(path).is_file():
        return None
    return Path(path)

@app.route('/quotes/sale-contract/draft/view')
def view_quote_sale_contract_draft():
    path = _draft_docx_path(QUOTE_SALE_DRAFT_KEY)
    if not path:
        abort(404)
    return send_docx_inline(path, 'hop-dong-mua-ban-nhap.docx')

@app.route('/quotes/sale-contract/draft/view.pdf')
def view_quote_sale_contract_draft_pdf():
    path = _draft_docx_path(QUOTE_SALE_DRAFT_KEY)
    if not path:
        abort(404)
    return send_pdf_inline(ensure_pdf_for_docx(path), 'hop-dong-mua-ban-nhap.pdf')

@app.route('/quotes/sale-contract/draft/download')
def download_quote_sale_contract_draft():
    path = _draft_docx_path(QUOTE_SALE_DRAFT_KEY)
    if not path:
        flash('Không tìm thấy bản nháp', 'danger')
        return redirect(url_for('quotes'))
    return send_file(path, as_attachment=True, download_name=path.name)

@app.route('/quotes/sale-contract/draft/download.pdf')
def download_quote_sale_contract_draft_pdf():
    path = _draft_docx_path(QUOTE_SALE_DRAFT_KEY)
    if not path:
        flash('Không tìm thấy bản nháp', 'danger')
        return redirect(url_for('quotes'))
    pdf_path = ensure_pdf_for_docx(path)
    return send_file(pdf_path, mimetype=PDF_MIMETYPE, as_attachment=True, download_name=path.stem + '.pdf')

@app.route('/quotes/handover/draft/view')
def view_quote_handover_draft():
    path = _draft_docx_path(QUOTE_HANDOVER_DRAFT_KEY)
    if not path:
        abort(404)
    return send_docx_inline(path, 'bien-ban-nhan-hang-nhap.docx')

@app.route('/quotes/handover/draft/view.pdf')
def view_quote_handover_draft_pdf():
    path = _draft_docx_path(QUOTE_HANDOVER_DRAFT_KEY)
    if not path:
        abort(404)
    return send_pdf_inline(pdf_path_for_docx_or_abort(path), 'bien-ban-nhan-hang-nhap.pdf')

@app.route('/quotes/handover/draft/download')
def download_quote_handover_draft():
    path = _draft_docx_path(QUOTE_HANDOVER_DRAFT_KEY)
    if not path:
        flash('Không tìm thấy bản nháp', 'danger')
        return redirect(url_for('quotes'))
    return send_file(path, as_attachment=True, download_name=path.name)

@app.route('/quotes/handover/draft/download.pdf')
def download_quote_handover_draft_pdf():
    path = _draft_docx_path(QUOTE_HANDOVER_DRAFT_KEY)
    if not path:
        flash('Không tìm thấy bản nháp', 'danger')
        return redirect(url_for('quotes'))
    pdf_path = ensure_pdf_for_docx(path)
    return send_file(pdf_path, mimetype=PDF_MIMETYPE, as_attachment=True, download_name=path.stem + '.pdf')

@app.route('/quotes/<int:qid>/preview')
def quote_preview(qid):
    ensure_quote_columns()
    quote = ensure_quote_documents(Quote.query.get_or_404(qid))
    st = quote_display_status(quote)
    return render_template(
        'quote_preview_page.html',
        quote=quote,
        valid_until=quote_valid_until(quote),
        status=st,
        expired=quote_is_expired(quote),
        linked_order=get_quote_order(quote),
        companies=list_company_profiles(),
        active_company_id=get_company_row().id if get_company_row() else None,
        doc_tab='list',
        back_url=app_back_url('quotes'),
    )

@app.route('/quotes/<int:qid>/print')
def quote_print(qid):
    ensure_quote_columns()
    quote = Quote.query.get_or_404(qid)
    return render_template(
        'quote_print.html',
        quote=quote,
        valid_until=quote_valid_until(quote),
    )

@app.route('/quotes/<int:qid>/order')
def create_order_from_quote(qid):
    ensure_quote_columns()
    normalize_quote_statuses()
    quote = Quote.query.get_or_404(qid)
    existing = get_quote_order(quote)
    if existing:
        flash(f'Báo giá đã có đơn hàng {existing.order_code}', 'info')
        return redirect(url_for('order_preview', oid=existing.id))
    if quote_display_status(quote) != 'Đã chốt':
        flash('Vui lòng chốt báo giá trước khi tạo đơn hàng', 'warning')
        return redirect(url_for('quotes'))
    had_shortage = bool(quote_shortage_lines(quote))
    order, created = create_order_from_quote_record(quote)
    stock_deducted = apply_quote_stock_deduction(quote, order)
    db.session.commit()
    if created:
        msg = f'Đã tạo đơn hàng {order.order_code}'
    else:
        msg = f'Đơn hàng {order.order_code} đã tồn tại'
    detail = quote_chot_flash_detail(stock_deducted, had_shortage)
    if detail:
        msg += f' — {detail}'
    flash(msg, 'success' if created or stock_deducted else 'info')
    shortage_msg = quote_stock_shortage_flash_message(quote) if had_shortage else None
    if shortage_msg:
        flash(shortage_msg, 'warning')
    return redirect(url_for('order_preview', oid=order.id))

@app.route('/contracts', methods=['GET', 'POST'])
def contracts():
    ensure_contract_columns()
    if request.method == 'POST':
        try:
            customer, code, signed, expired, quote_id, company_id = parse_contract_form()
        except ValueError as e:
            flash(str(e), 'warning')
            return redirect(url_for('contracts', open_add='contract'))
        out_path = apply_contract_doc(None, customer, code, signed, expired, company_id)
        contract = Contract(
            contract_code=code,
            customer_id=customer.id,
            company_id=company_id,
            quote_id=int(quote_id) if quote_id else None,
            signed_date=signed,
            expired_date=expired,
            file_path=out_path,
        )
        db.session.add(contract)
        db.session.commit()
        flash(f'Đã tạo hợp đồng {code}', 'success')
        return redirect(url_for('contracts'))
    today = date.today().isoformat()
    year_end = date(date.today().year, 12, 31).isoformat()
    return render_template(
        'contracts.html',
        contracts=Contract.query.order_by(Contract.created_at.desc()).limit(100).all(),
        customers=customers_for_select(),
        quotes=Quote.query.order_by(Quote.created_at.desc()).limit(50).all(),
        companies=list_company_profiles(),
        active_company_id=get_company_row().id if get_company_row() else None,
        today=today,
        year_end=year_end,
        open_add_modal=request.args.get('open_add') == 'contract',
        preselect_customer_id=request.args.get('customer_id', type=int),
    )

@app.route('/contracts/preview', methods=['POST'])
def preview_contract_draft():
    try:
        cid = request.form.get('contract_id', type=int)
        contract = Contract.query.get(cid) if cid else None
        customer, code, signed, expired, _, company_id = parse_contract_form(contract)
    except ValueError as e:
        flash(str(e), 'warning')
        return redirect(url_for('contracts', open_add='contract'))
    doc = build_contract_doc(customer, code, signed, expired, company_id=company_id)
    clear_contract_draft_temp()
    temp_path = save_temp_docx(doc, 'contract_draft')
    session[CONTRACT_DRAFT_SESSION_KEY] = str(temp_path)
    try:
        urls = prepare_document_preview_urls(
            temp_path,
            'view_contract_draft_pdf',
            'view_contract_draft',
            'download_contract_draft',
            'download_contract_draft_pdf',
        )
    except FileNotFoundError as e:
        _unlink_preview_file(str(temp_path))
        flash(str(e), 'danger')
        return redirect(url_for('contracts', open_add='contract'))
    return render_document_preview_page(
        code,
        customer.name,
        urls,
        url_for('contracts', open_add='contract'),
        is_draft=True,
    )

@app.route('/contracts/<int:cid>/update', methods=['POST'])
def update_contract(cid):
    ensure_contract_columns()
    contract = Contract.query.get_or_404(cid)
    try:
        customer, code, signed, expired, quote_id, company_id = parse_contract_form(contract)
    except ValueError as e:
        flash(str(e), 'warning')
        return redirect(url_for('contracts'))
    if (
        contract.contract_type == CONTRACT_TYPE_FRAMEWORK
        and customer.id != contract.customer_id
        and get_framework_contract(customer.id)
    ):
        flash('Khách hàng này đã có hợp đồng nguyên tắc', 'warning')
        return redirect(url_for('contracts'))
    contract.file_path = apply_contract_doc(contract, customer, code, signed, expired, company_id)
    contract.contract_code = code
    contract.customer_id = customer.id
    contract.company_id = company_id
    contract.quote_id = int(quote_id) if quote_id else None
    contract.signed_date = signed
    contract.expired_date = expired
    db.session.commit()
    flash(f'Đã cập nhật hợp đồng {code}', 'success')
    return redirect(url_for('contracts'))

@app.route('/contracts/preview/view')
def view_contract_draft():
    path = session.get(CONTRACT_DRAFT_SESSION_KEY)
    if not path or not Path(path).is_file():
        abort(404)
    return send_docx_inline(path, 'hop-dong-nhap.docx')

@app.route('/contracts/preview/view.pdf')
def view_contract_draft_pdf():
    path = session.get(CONTRACT_DRAFT_SESSION_KEY)
    if not path or not Path(path).is_file():
        abort(404)
    return send_pdf_inline(pdf_path_for_docx_or_abort(path), 'hop-dong-nhap.pdf')

@app.route('/contracts/preview/download')
def download_contract_draft():
    path = session.get(CONTRACT_DRAFT_SESSION_KEY)
    if not path or not Path(path).is_file():
        flash('Không tìm thấy bản nháp hợp đồng', 'danger')
        return redirect(url_for('contracts'))
    return send_file(path, as_attachment=True, download_name=Path(path).name)

@app.route('/contracts/preview/download.pdf')
def download_contract_draft_pdf():
    path = session.get(CONTRACT_DRAFT_SESSION_KEY)
    if not path or not Path(path).is_file():
        flash('Không tìm thấy bản nháp hợp đồng', 'danger')
        return redirect(url_for('contracts'))
    pdf_path = ensure_pdf_for_docx(path)
    return send_file(
        pdf_path,
        mimetype=PDF_MIMETYPE,
        as_attachment=True,
        download_name=Path(path).stem + '.pdf',
    )

@app.route('/contracts/<int:cid>/preview')
def preview_contract(cid):
    c = Contract.query.get_or_404(cid)
    if not c.file_path or not Path(c.file_path).is_file():
        flash('Không tìm thấy file hợp đồng', 'danger')
        return redirect(url_for('contracts'))
    try:
        urls = prepare_document_preview_urls(
            c.file_path,
            'view_contract_pdf',
            'view_contract',
            'download_contract',
            'download_contract_pdf',
            cid=c.id,
        )
    except FileNotFoundError as e:
        flash(str(e), 'danger')
        return redirect(url_for('contracts'))
    return render_document_preview_page(
        c.contract_code,
        c.customer.name,
        urls,
        app_back_url('contracts'),
    )

@app.route('/contracts/<int:cid>/view')
def view_contract(cid):
    c = Contract.query.get_or_404(cid)
    if not c.file_path or not Path(c.file_path).is_file():
        abort(404)
    return send_docx_inline(c.file_path, f'{c.contract_code}.docx')

@app.route('/contracts/<int:cid>/view.pdf')
def view_contract_pdf(cid):
    c = Contract.query.get_or_404(cid)
    if not c.file_path or not Path(c.file_path).is_file():
        abort(404)
    return send_pdf_inline(
        pdf_path_for_docx_or_abort(c.file_path),
        f'{c.contract_code}.pdf',
    )

@app.route('/download/contract/<int:cid>')
def download_contract(cid):
    c = Contract.query.get_or_404(cid)
    if not c.file_path or not Path(c.file_path).is_file():
        abort(404)
    return send_file(c.file_path, as_attachment=True, download_name=Path(c.file_path).name)

@app.route('/download/contract/<int:cid>.pdf')
def download_contract_pdf(cid):
    c = Contract.query.get_or_404(cid)
    if not c.file_path or not Path(c.file_path).is_file():
        abort(404)
    pdf_path = ensure_pdf_for_docx(c.file_path)
    return send_file(
        pdf_path,
        mimetype=PDF_MIMETYPE,
        as_attachment=True,
        download_name=Path(c.file_path).stem + '.pdf',
    )

@app.route('/orders', methods=['GET', 'POST'])
def orders():
    ensure_order_columns()
    normalize_order_statuses()
    if request.method == 'POST':
        customer_id = int(request.form['customer_id'])
        total = parse_int(request.form.get('total'))
        status = 'Đang xử lý'
        order = Order(
            order_code=next_code('DH', Order, 'order_code'),
            customer_id=customer_id,
            total=total,
            status=status,
        )
        db.session.add(order)
        db.session.commit()
        flash('Đã tạo đơn hàng', 'success')
        return orders_redirect()

    filters = order_filters_from_request()
    if request.args.get('clear'):
        session.pop('orders_customer_id', None)
        filters['customer_id'] = ''
    elif 'customer_id' in request.args:
        if filters['customer_id']:
            session['orders_customer_id'] = filters['customer_id']
        else:
            session.pop('orders_customer_id', None)
    elif session.get('orders_customer_id'):
        try:
            remembered = int(session.get('orders_customer_id'))
        except (TypeError, ValueError):
            remembered = None
        if remembered and Customer.query.get(remembered):
            filters['customer_id'] = str(remembered)
        else:
            session.pop('orders_customer_id', None)
    per_page = request.args.get('per_page', 10, type=int)
    per_page = per_page if per_page in (10, 25, 50) else 10
    page = request.args.get('page', 1, type=int)

    query = apply_order_filters(Order.query, filters)
    pagination = query.order_by(Order.created_at.desc()).paginate(page=page, per_page=per_page, error_out=False)

    list_args = {k: v for k, v in filters.items() if v}
    list_args['per_page'] = per_page
    export_args = {k: v for k, v in filters.items() if v}

    return render_template(
        'orders.html',
        orders=pagination.items,
        pagination=pagination,
        stats=order_stats(query),
        filters=filters,
        list_args=list_args,
        export_args=export_args,
        per_page=per_page,
        customers=customers_for_order_filter(),
        order_filter_statuses=ORDER_FILTER_STATUSES,
    )

@app.route('/orders/export')
def orders_export():
    filters = order_filters_from_request()
    query = apply_order_filters(Order.query, filters)
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow(['Mã đơn', 'Khách hàng', 'Ngày tạo', 'Tổng tiền', 'Đã trả', 'Còn nợ', 'Trạng thái', 'Thanh toán'])
    for o in query.order_by(Order.created_at.desc()).all():
        writer.writerow([
            o.order_code,
            o.customer.name,
            o.created_at.strftime('%d/%m/%Y') if o.created_at else '',
            o.total,
            o.paid_amount,
            order_balance(o),
            effective_order_status(o),
            order_payment_status(o),
        ])
    return Response(
        '\ufeff' + buf.getvalue(),
        mimetype='text/csv; charset=utf-8',
        headers={'Content-Disposition': 'attachment; filename=don-hang.csv'},
    )

@app.route('/orders/<int:oid>/status', methods=['POST'])
def update_order_status(oid):
    ensure_order_columns()
    normalize_order_statuses()
    o = Order.query.get_or_404(oid)
    new_status = request.form.get('status', '').strip()
    current = effective_order_status(o)
    allowed = ORDER_ALLOWED_TRANSITIONS.get(current, set())
    if new_status == 'Hoàn thành':
        flash('Vui lòng upload biên bản nhận hàng có chữ ký để hoàn thành đơn', 'warning')
        return orders_redirect(page=request.form.get('_page', 1, type=int))
    if new_status in allowed:
        o.status = new_status
        db.session.commit()
        flash_status_updated(new_status)
    elif new_status in ORDER_STATUSES:
        flash('Không thể chuyển trạng thái đơn hàng theo hướng này', 'warning')
    return orders_redirect(page=request.form.get('_page', 1, type=int))

@app.route('/orders/<int:oid>/complete', methods=['POST'])
def complete_order(oid):
    ensure_order_columns()
    normalize_order_statuses()
    o = Order.query.get_or_404(oid)
    current = effective_order_status(o)
    if 'Hoàn thành' not in ORDER_ALLOWED_TRANSITIONS.get(current, set()):
        flash('Không thể hoàn thành đơn hàng ở trạng thái hiện tại', 'warning')
        return orders_redirect(page=request.form.get('_page', 1, type=int))
    if not save_order_handover_scan(o, request.files.get('handover_scan')):
        return orders_redirect(page=request.form.get('_page', 1, type=int))
    o.status = 'Hoàn thành'
    db.session.commit()
    flash('Đã hoàn thành đơn hàng và lưu biên bản nhận hàng', 'success')
    return orders_redirect(page=request.form.get('_page', 1, type=int))

@app.route('/orders/<int:oid>/handover-scan')
def view_order_handover_scan(oid):
    ensure_order_columns()
    o = Order.query.get_or_404(oid)
    if not o.handover_scan_path:
        flash('Chưa có biên bản nhận hàng', 'warning')
        return redirect(url_for('order_preview', oid=oid))
    file_path = BASE_DIR / 'static' / o.handover_scan_path
    if not file_path.is_file():
        flash('Không tìm thấy file biên bản nhận hàng', 'danger')
        return redirect(url_for('order_preview', oid=oid))
    return send_file(file_path, mimetype=signed_scan_mimetype(file_path))

@app.route('/payments/<int:pid>/receipt')
def view_payment_receipt(pid):
    ensure_payment_columns()
    p = Payment.query.get_or_404(pid)
    if not p.receipt_path:
        flash('Chưa có ảnh xác nhận thanh toán', 'warning')
        if p.order_id:
            return redirect(url_for('order_preview', oid=p.order_id))
        return redirect(url_for('debt'))
    file_path = BASE_DIR / 'static' / p.receipt_path
    if not file_path.is_file():
        flash('Không tìm thấy file ảnh xác nhận', 'danger')
        return redirect(url_for('debt'))
    return send_file(file_path, mimetype=signed_scan_mimetype(file_path))

@app.route('/orders/<int:oid>/payment', methods=['POST'])
def add_payment(oid):
    ensure_payment_columns()
    o = Order.query.get_or_404(oid)
    page = request.form.get('_page', 1, type=int)
    from_debt = request.form.get('_from') == 'debt'
    if not order_counts_toward_debt(o):
        flash('Đơn hàng đã hủy, không ghi nhận thanh toán / công nợ', 'warning')
        if from_debt:
            return debt_redirect(customer_id=request.form.get('customer_id') or o.customer_id, page=page)
        return orders_redirect(page=page)
    amount = parse_int(request.form.get('amount'))
    if amount <= 0:
        flash('Số tiền không hợp lệ', 'danger')
        if from_debt:
            return debt_redirect(customer_id=request.form.get('customer_id') or o.customer_id, page=page)
        return orders_redirect(page=page)
    receipt_path = save_payment_receipt(request.files.get('payment_receipt'))
    if not receipt_path:
        if from_debt:
            return debt_redirect(customer_id=request.form.get('customer_id') or o.customer_id, page=page)
        return orders_redirect(page=page)
    receipt_at = datetime.utcnow()
    result = apply_order_payment(
        o, amount, request.form.get('method', 'Chuyển khoản'), request.form.get('note', ''),
        receipt_path=receipt_path, receipt_uploaded_at=receipt_at,
    )
    if not result:
        flash('Số tiền không hợp lệ hoặc đơn đã thanh toán đủ', 'danger')
        if from_debt:
            return debt_redirect(customer_id=request.form.get('customer_id') or o.customer_id, page=page)
        return orders_redirect(page=page)
    db.session.commit()
    flash(f"Đã ghi nhận thanh toán {money(result['amount'])} cho {o.order_code}", 'success')
    if from_debt:
        return debt_redirect(
            customer_id=request.form.get('customer_id') or o.customer_id,
            page=request.form.get('_page', 1, type=int),
        )
    return orders_redirect(page=request.form.get('_page', 1, type=int))

@app.route('/orders/<int:oid>/preview')
def order_preview(oid):
    ensure_order_columns()
    normalize_order_statuses()
    o = Order.query.get_or_404(oid)
    quote = None
    valid_until = None
    if o.quote_id:
        ensure_quote_columns()
        quote = Quote.query.get(o.quote_id)
        if quote:
            valid_until = quote_valid_until(quote)
    return render_template(
        'order_preview_page.html',
        order=o,
        quote=quote,
        valid_until=valid_until,
        status=effective_order_status(o),
        payment=order_payment_status(o),
        balance=order_balance(o),
        back_url=app_back_url('orders'),
    )

@app.route('/orders/<int:oid>/print')
def order_print(oid):
    o = Order.query.get_or_404(oid)
    if not o.quote_id:
        flash('Đơn hàng không có báo giá liên kết', 'warning')
        return redirect(url_for('order_preview', oid=oid))
    ensure_quote_columns()
    quote = Quote.query.get_or_404(o.quote_id)
    return render_template(
        'order_print.html',
        order=o,
        quote=quote,
        valid_until=quote_valid_until(quote),
    )

def pending_supplier_intake_count():
    ensure_supplier_intake_status_column()
    return SupplierIntake.query.filter(
        SupplierIntake.status == SUPPLIER_INTAKE_STATUS_PENDING,
    ).count()

def stock_intake_list_filters_from_request():
    status = (request.args.get('status') or request.form.get('list_status') or 'pending').strip()
    q = (request.args.get('q') or request.form.get('list_q') or '').strip()
    page = request.args.get('page', type=int) or request.form.get('list_page', type=int) or 1
    if status not in ('pending', 'received', 'all'):
        status = 'pending'
    if page < 1:
        page = 1
    return {'status': status, 'q': q, 'page': page}

def redirect_stock_intake_list(extra=None):
    flt = stock_intake_list_filters_from_request()
    args = {k: v for k, v in flt.items() if v}
    if extra:
        args.update(extra)
    return redirect(url_for('stock_intake_requests', **args))

def redirect_stock_intake_detail(iid):
    flt = stock_intake_list_filters_from_request()
    args = {k: v for k, v in flt.items() if k != 'page' or v > 1}
    return redirect(url_for('stock_intake_request_detail', iid=iid, **args))

@app.route('/stock/intake-requests')
def stock_intake_requests():
    ensure_supplier_intake_schema()
    ensure_supplier_intake_line_columns()
    status_filter = request.args.get('status', 'pending').strip()
    q = request.args.get('q', '').strip()
    per_page = request.args.get('per_page', 15, type=int)
    page = request.args.get('page', 1, type=int)
    if per_page not in (10, 15, 25, 50):
        per_page = 15
    query = (
        SupplierIntake.query.join(Supplier)
        .order_by(SupplierIntake.created_at.desc())
    )
    if status_filter == 'pending':
        query = query.filter(SupplierIntake.status == SUPPLIER_INTAKE_STATUS_PENDING)
    elif status_filter == 'received':
        query = query.filter(SupplierIntake.status == SUPPLIER_INTAKE_STATUS_RECEIVED)
    elif status_filter != 'all':
        query = query.filter(SupplierIntake.status == SUPPLIER_INTAKE_STATUS_PENDING)
        status_filter = 'pending'
    if q:
        like = f'%{q}%'
        query = query.filter(or_(
            SupplierIntake.intake_code.ilike(like),
            SupplierIntake.ref_code.ilike(like),
            Supplier.name.ilike(like),
            Supplier.code.ilike(like),
        ))
    pagination = query.paginate(page=page, per_page=per_page, error_out=False)
    intake_ids = [i.id for i in pagination.items]
    line_map = supplier_intake_lines_map(intake_ids)
    product_gap_map = stock_intake_product_gap_map(line_map)
    pending_count = pending_supplier_intake_count()
    shortage_products = []
    for pid, gap in product_gap_map.items():
        if gap.get('shortfall', 0) <= 1e-9:
            continue
        product = Product.query.get(pid)
        shortage_products.append({
            'product_id': pid,
            'product': product,
            **gap,
        })
    shortage_products.sort(key=lambda x: -x['shortfall'])
    shortage_products = shortage_products[:30]
    return render_template(
        'stock_intake_requests.html',
        pagination=pagination,
        intakes=pagination.items,
        line_map=line_map,
        product_gap_map=product_gap_map,
        shortage_products=shortage_products,
        filters={'status': status_filter, 'q': q},
        list_args={'status': status_filter, 'q': q, 'per_page': per_page},
        per_page=per_page,
        pending_count=pending_count,
        back_url=app_back_url('stock'),
    )

@app.route('/stock/intake-requests/<int:iid>')
def stock_intake_request_detail(iid):
    ensure_supplier_intake_schema()
    ensure_supplier_intake_line_columns()
    ensure_product_columns()
    intake = SupplierIntake.query.get_or_404(iid)
    list_filters = stock_intake_list_filters_from_request()
    line_map = supplier_intake_lines_map([intake.id])
    lines = line_map.get(intake.id, [])
    product_gap_map = stock_intake_product_gap_map({intake.id: lines})
    saved_alloc_map = supplier_intake_saved_allocations_map({intake.id: lines})
    stock_st = supplier_intake_stock_status(intake)
    back_args = {k: v for k, v in list_filters.items() if v and (k != 'page' or v > 1)}
    return render_template(
        'stock_intake_request_detail.html',
        intake=intake,
        lines=lines,
        stock_st=stock_st,
        product_gap_map=product_gap_map,
        saved_alloc_map=saved_alloc_map,
        intake_customer_choices=intake_manual_customer_choices(),
        list_filters=list_filters,
        back_url=url_for('stock_intake_requests', **back_args),
    )

@app.route('/stock/intake-requests/<int:iid>/approve', methods=['POST'])
def approve_supplier_intake_stock(iid):
    ensure_supplier_intake_schema()
    intake = SupplierIntake.query.get_or_404(iid)
    allocations = parse_intake_customer_allocations_from_form()
    ok, result = approve_supplier_intake_stock_record(intake, customer_allocations=allocations)
    if not ok:
        flash(result, 'warning')
        return redirect_stock_intake_detail(iid)
    db.session.commit()
    code = intake.intake_code or intake.id
    msg = f'Đã duyệt nhập kho phiếu {code}'
    if result:
        msg += f' — +{format_qty_display(result)} (đơn vị tồn)'
    flash(msg, 'success')
    return redirect_stock_intake_list()

@app.route('/stock/intake-requests/<int:iid>/reject', methods=['POST'])
def reject_supplier_intake_stock(iid):
    ensure_supplier_intake_schema()
    intake = SupplierIntake.query.get_or_404(iid)
    if supplier_intake_stock_status(intake) == SUPPLIER_INTAKE_STATUS_RECEIVED:
        flash('Phiếu đã nhập kho — không thể hủy', 'warning')
        return redirect_stock_intake_detail(iid)
    code = delete_supplier_intake_record(intake)
    db.session.commit()
    flash(f'Đã hủy yêu cầu nhập kho {code}', 'success')
    return redirect_stock_intake_list()

@app.route('/stock', methods=['GET', 'POST'])
def stock():
    ensure_stock_movement_columns()
    if request.method == 'POST':
        p = Product.query.get_or_404(int(request.form['product_id']))
        qty = max(parse_qty(request.form.get('qty')), 0)
        mtype = request.form.get('movement_type', 'IN')
        if qty <= 0:
            return stock_redirect_after_save('Số lượng phải lớn hơn 0', 'danger')
        if mtype == 'OUT' and p.stock < qty - 1e-9:
            return stock_redirect_after_save('Tồn kho không đủ để xuất', 'danger')
        if mtype == 'IN':
            p.stock += qty
        else:
            p.stock -= qty
        supplier_id = parse_supplier_id(request.form.get('supplier_id')) if mtype == 'IN' else None
        db.session.add(StockMovement(
            product_id=p.id,
            supplier_id=supplier_id,
            movement_type=mtype,
            qty=qty,
            purchase_qty=0,
            ref_code=request.form.get('ref_code', '').strip(),
            method=normalize_stock_method(mtype, request.form.get('method')),
            warehouse=request.form.get('warehouse', DEFAULT_WAREHOUSE).strip() or DEFAULT_WAREHOUSE,
            note=request.form.get('note', '').strip(),
        ))
        db.session.commit()
        label = 'nhập' if mtype == 'IN' else 'xuất'
        return stock_redirect_after_save(f'Đã {label} kho thành công')

    filters = stock_filters_from_request()
    date_from = parse_stock_date(filters['date_from'])
    date_to = parse_stock_date(filters['date_to'])
    per_page = request.args.get('per_page', 10, type=int)
    per_page = per_page if per_page in (10, 25, 50, 100) else 10
    page = request.args.get('page', 1, type=int)

    query = Product.query
    if filters['product_id']:
        query = query.filter(Product.id == int(filters['product_id']))
    if filters['q']:
        like = f"%{filters['q']}%"
        query = query.filter(or_(Product.sku.ilike(like), Product.name.ilike(like)))
    if filters['status'] == 'ok':
        query = query.filter(Product.stock_qty > Product.low_stock)
    elif filters['status'] == 'low':
        query = query.filter(Product.stock_qty > 0, Product.stock_qty <= Product.low_stock)
    elif filters['status'] == 'out':
        query = query.filter(Product.stock_qty <= 0)

    pagination = query.order_by(Product.name).paginate(page=page, per_page=per_page, error_out=False)
    inventory_rows = [product_stock_row(p, date_from, date_to) for p in pagination.items]

    mov_query = StockMovement.query.join(Product)
    if filters['product_id']:
        mov_query = mov_query.filter(StockMovement.product_id == int(filters['product_id']))
    if filters['movement_type']:
        mov_query = mov_query.filter(StockMovement.movement_type == filters['movement_type'])
    if filters['method']:
        mov_query = mov_query.filter(StockMovement.method == filters['method'])
    if filters['warehouse']:
        mov_query = mov_query.filter(StockMovement.warehouse == filters['warehouse'])
    if filters['supplier_id']:
        mov_query = mov_query.filter(StockMovement.supplier_id == int(filters['supplier_id']))
    if date_from:
        mov_query = mov_query.filter(StockMovement.created_at >= datetime.combine(date_from, datetime.min.time()))
    if date_to:
        mov_query = mov_query.filter(StockMovement.created_at <= datetime.combine(date_to, datetime.max.time()))
    if filters['q']:
        like = f"%{filters['q']}%"
        mov_query = mov_query.filter(or_(Product.sku.ilike(like), Product.name.ilike(like)))
    history_pagination = mov_query.order_by(StockMovement.created_at.desc()).paginate(
        page=page if filters['tab'] == 'history' else 1,
        per_page=per_page,
        error_out=False,
    )

    list_args = {k: v for k, v in filters.items() if v}
    list_args['per_page'] = per_page
    base_args = {k: v for k, v in list_args.items() if k not in ('tab', 'status')}
    inventory_args = {**base_args, 'tab': 'inventory'}
    history_args = {**base_args, 'tab': 'history'}
    low_stock_args = {**base_args, 'tab': 'inventory', 'status': 'low'}
    out_stock_args = {**base_args, 'tab': 'inventory', 'status': 'out'}
    export_args = {k: v for k, v in list_args.items() if k != 'tab'}
    all_products = Product.query.order_by(Product.name).all()
    all_suppliers = active_suppliers()
    ensure_supplier_intake_status_column()
    return render_template(
        'stock.html',
        stats=stock_stats(),
        pending_intake_count=pending_supplier_intake_count(),
        inventory_rows=inventory_rows,
        pagination=pagination,
        history_pagination=history_pagination,
        movements=history_pagination.items,
        all_products=all_products,
        all_suppliers=all_suppliers,
        filters=filters,
        list_args=list_args,
        inventory_args=inventory_args,
        history_args=history_args,
        low_stock_args=low_stock_args,
        out_stock_args=out_stock_args,
        export_args=export_args,
        per_page=per_page,
        stock_in_methods=STOCK_IN_METHODS,
        stock_out_methods=STOCK_OUT_METHODS,
        stock_filter_methods=STOCK_FILTER_METHODS,
        warehouses=[DEFAULT_WAREHOUSE],
    )

@app.route('/stock/export')
def stock_export():
    ensure_stock_movement_columns()
    filters = stock_filters_from_request()
    date_from = parse_stock_date(filters['date_from'])
    date_to = parse_stock_date(filters['date_to'])
    query = Product.query.order_by(Product.name)
    if filters['q']:
        like = f"%{filters['q']}%"
        query = query.filter(or_(Product.sku.ilike(like), Product.name.ilike(like)))
    rows = [product_stock_row(p, date_from, date_to) for p in query.all()]
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow(['SKU', 'Sản phẩm', 'Kho', 'Tồn hiện tại', 'Đơn vị', 'Trạng thái'])
    for r in rows:
        p = r['product']
        writer.writerow([
            p.sku, p.name, r['warehouse'], r['current'], product_base_unit(p), r['status_label'],
        ])
    return Response(
        '\ufeff' + buf.getvalue(),
        mimetype='text/csv; charset=utf-8',
        headers={'Content-Disposition': 'attachment; filename=ton-kho.csv'},
    )

@app.route('/supplier-debt')
def supplier_debt():
    return render_template('supplier_debt.html')

@app.route('/debt')
def debt():
    ensure_payment_columns()
    normalize_order_statuses()
    search_q = request.args.get('q', '').strip()
    page = request.args.get('page', 1, type=int)
    per_page = request.args.get('per_page', 10, type=int)
    order_page = request.args.get('order_page', 1, type=int)
    per_page = min(max(per_page, 5), 50)
    customer_id = request.args.get('customer_id', type=int)
    if request.args.get('clear'):
        session.pop('debt_customer_id', None)
        customer_id = None
    elif customer_id:
        session['debt_customer_id'] = customer_id
    elif session.get('debt_customer_id'):
        try:
            customer_id = int(session.get('debt_customer_id'))
        except (TypeError, ValueError):
            customer_id = None
        if customer_id and not Customer.query.get(customer_id):
            session.pop('debt_customer_id', None)
            customer_id = None

    all_rows = build_debt_customer_rows(search_q)
    pagination = SimplePagination(all_rows, page, per_page)
    list_args = {'q': search_q, 'per_page': per_page}
    if customer_id:
        list_args['customer_id'] = customer_id
    if order_page > 1:
        list_args['order_page'] = order_page

    selected = None
    selected_summary = None
    customer_orders = []
    order_pagination = None
    payments = []

    if customer_id:
        selected = Customer.query.get(customer_id)
    elif pagination.items:
        selected = pagination.items[0]['customer']
        customer_id = selected.id
        list_args['customer_id'] = customer_id

    if selected:
        agg = customer_debt_aggregates().get(selected.id, {})
        selected_summary = {
            'total': agg.get('total', 0),
            'paid': agg.get('paid', 0),
            'remaining': agg.get('remaining', 0),
            'unpaid_count': agg.get('unpaid_count', 0),
            'order_count': agg.get('order_count', 0),
            'status': customer_debt_status_from_agg(agg),
            'status_key': customer_debt_status_key(customer_debt_status_from_agg(agg)),
        }
        orders_all = [
            o for o in Order.query.filter_by(customer_id=selected.id)
            .order_by(Order.created_at.desc()).all()
            if order_counts_toward_debt(o)
        ]
        order_pagination = SimplePagination(orders_all, order_page, 8)
        customer_orders = order_pagination.items
        payments = (
            Payment.query.join(Order)
            .filter(Order.customer_id == selected.id)
            .order_by(Payment.payment_date.desc(), Payment.id.desc())
            .limit(20)
            .all()
        )

    unpaid_orders = customer_unpaid_orders_oldest_first(selected.id) if selected else []
    customer_outstanding = customer_outstanding_balance(selected.id) if selected else 0

    return render_template(
        'debt.html',
        debt_rows=pagination.items,
        pagination=pagination,
        list_args=list_args,
        per_page=per_page,
        search_q=search_q,
        selected=selected,
        selected_summary=selected_summary,
        customer_orders=customer_orders,
        order_pagination=order_pagination,
        payments=payments,
        payment_groups=group_payments_for_display(payments) if payments else [],
        unpaid_orders=unpaid_orders,
        customer_outstanding=customer_outstanding,
        customer_id=customer_id,
    )

@app.route('/debt/payment', methods=['POST'])
def debt_allocate_payment():
    ensure_payment_columns()
    normalize_order_statuses()
    customer_id = request.form.get('customer_id', type=int)
    page = request.form.get('_page', 1, type=int)
    if not customer_id:
        flash('Không xác định được khách hàng', 'danger')
        return debt_redirect(page=page)
    Customer.query.get_or_404(customer_id)
    amount = parse_int(request.form.get('amount'))
    method = request.form.get('method', 'Chuyển khoản')
    note = request.form.get('note', '')
    outstanding = customer_outstanding_balance(customer_id)
    if outstanding <= 0:
        flash('Khách hàng không còn công nợ', 'warning')
        return debt_redirect(customer_id=customer_id, page=page)
    if amount <= 0:
        flash('Số tiền không hợp lệ', 'danger')
        return debt_redirect(customer_id=customer_id, page=page)
    receipt_path = save_payment_receipt(request.files.get('payment_receipt'))
    if not receipt_path:
        return debt_redirect(customer_id=customer_id, page=page)
    receipt_at = datetime.utcnow()
    pay_amount = min(amount, outstanding)
    if amount > outstanding:
        flash(
            f'Số tiền vượt tổng còn nợ ({money(outstanding)}). Chỉ phân bổ {money(pay_amount)}.',
            'warning',
        )
    allocations = allocate_customer_payment(
        customer_id, pay_amount, method, note,
        receipt_path=receipt_path, receipt_uploaded_at=receipt_at,
    )
    if not allocations:
        flash('Không thể phân bổ thanh toán', 'danger')
        return debt_redirect(customer_id=customer_id, page=page)
    lines = format_payment_allocation_lines(allocations)
    total_paid = sum(a['amount'] for a in allocations)
    flash(
        f'Đã ghi nhận thanh toán {money(total_paid)}.\nPhân bổ thanh toán:\n' + '\n'.join(lines),
        'success',
    )
    return debt_redirect(customer_id=customer_id, page=page)

@app.route('/debt/payments/<int:pid>/delete', methods=['POST'])
def delete_debt_payment(pid):
    ensure_payment_columns()
    page = request.form.get('_page', 1, type=int)
    order_page = request.form.get('_order_page', 1, type=int)
    customer_id, count, total = delete_payment_group(pid)
    db.session.commit()
    flash(
        f'Đã xóa {count} dòng thanh toán ({money(total)}). Công nợ đơn hàng đã được cập nhật lại.',
        'success',
    )
    return debt_redirect(customer_id=customer_id, page=page, order_page=order_page)

@app.route('/debt/export')
def debt_export():
    normalize_order_statuses()
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow(['Khách hàng', 'SĐT', 'Tổng công nợ', 'Đã thanh toán', 'Còn nợ', 'Số đơn', 'Trạng thái'])
    for row in build_debt_customer_rows(request.args.get('q', '').strip()):
        c = row['customer']
        writer.writerow([
            c.name,
            c.phone or '',
            row['total'],
            row['paid'],
            row['remaining'],
            row['order_count'],
            row['status'],
        ])
    return Response(
        '\ufeff' + buf.getvalue(),
        mimetype='text/csv; charset=utf-8',
        headers={'Content-Disposition': 'attachment; filename=cong-no.csv'},
    )

@app.cli.command('init-db')
def init_db_command():
    init_db()
    print('Initialized database')

def create_quote_doc_templates():
    TEMPLATE_QUOTE_SALE_CONTRACT.parent.mkdir(parents=True, exist_ok=True)
    if not TEMPLATE_QUOTE_SALE_CONTRACT.exists():
        doc = Document()
        _doc_add_center_heading(doc, 'HỢP ĐỒNG MUA BÁN HÀNG HÓA', level=1)
        doc.add_paragraph('Số hợp đồng: {{SO_HOP_DONG}}')
        doc.add_paragraph('Căn cứ báo giá số {{SO_BAO_GIA}} ngày {{NGAY_KY}}.')
        doc.add_paragraph('BÊN MUA: {{TEN_BEN_A}}')
        doc.add_paragraph('BÊN BÁN: {{TEN_BEN_B}}')
        doc.add_paragraph('Tổng thanh toán: {{TONG_THANH_TOAN}}')
        doc.save(TEMPLATE_QUOTE_SALE_CONTRACT)
    if not TEMPLATE_QUOTE_HANDOVER.exists():
        doc = Document()
        _doc_add_center_heading(doc, 'BIÊN BẢN GIAO NHẬN HÀNG HÓA', level=1)
        doc.add_paragraph('Căn cứ vào hợp đồng mua bán giữa {{TEN_BEN_B}} và {{TEN_BEN_A}}')
        doc.add_paragraph('Hôm nay, ngày {{NGAY_GIAO}}, tại {{DIA_DIEM}}, chúng tôi đại diện cho các bên ký hợp đồng gồm có:')
        doc.add_paragraph('BÊN MUA (BÊN A): {{TEN_BEN_A}}')
        doc.add_paragraph('BÊN BÁN (BÊN B): {{TEN_BEN_B}}')
        doc.add_paragraph('Đại diện A: {{DAI_DIEN_A}} — Chức vụ: {{CHUC_VU_A}}')
        doc.add_paragraph('Đại diện B: {{DAI_DIEN_B}} — Chức vụ: {{CHUC_VU_B}}')
        doc.add_paragraph('(Bảng hàng: dùng {%tr for item in items %} trong file Word thật)')
        doc.save(TEMPLATE_QUOTE_HANDOVER)

def create_template_docx():
    if TEMPLATE_DOCX.exists(): return
    doc = Document()
    doc.add_paragraph('CỘNG HOÀ XÃ HỘI CHỦ NGHĨA VIỆT NAM').alignment = 1
    doc.add_paragraph('Độc lập - Tự do - Hạnh phúc').alignment = 1
    doc.add_heading('HỢP ĐỒNG NGUYÊN TẮC', level=1).alignment = 1
    doc.add_paragraph('Số: {{SO_HOP_DONG}}')
    doc.add_paragraph('Hôm nay, ngày {{NGAY_KY}}, chúng tôi gồm có:')
    doc.add_paragraph('BÊN MUA (BÊN A): {{TEN_BEN_A}}')
    doc.add_paragraph('Đại diện: {{DAI_DIEN_A}}    Chức vụ: {{CHUC_VU_A}}')
    doc.add_paragraph('Địa chỉ: {{DIA_CHI_A}}')
    doc.add_paragraph('Điện thoại: {{DIEN_THOAI_A}}')
    doc.add_paragraph('Số TK: {{SO_TK_A}}')
    doc.add_paragraph('Mở tại: {{NGAN_HANG_A}}')
    doc.add_paragraph('MST: {{MST_A}}')
    doc.add_paragraph('BÊN BÁN (BÊN B): CÔNG TY TNHH THƯƠNG MẠI SEOUL – HÀ NỘI')
    doc.add_paragraph('Đại diện: Bà Trần Thị Hương Giang    Chức vụ: Giám đốc')
    doc.add_paragraph('Địa chỉ: Số 270 Ngô Gia Tự, phường Đức Giang, Quận Long Biên, TP Hà Nội')
    doc.add_paragraph('Điều 1: Điều khoản chung')
    doc.add_paragraph('Hai bên thống nhất ký kết hợp đồng nguyên tắc mua bán hàng hóa.')
    doc.add_paragraph('Điều 2: Hàng hóa và bảo hành')
    doc.add_paragraph('Hàng hóa đảm bảo đúng chủng loại, chất lượng và thông số kỹ thuật.')
    doc.add_paragraph('Điều 3: Thanh toán')
    doc.add_paragraph('Thanh toán theo thỏa thuận của từng đơn hàng/báo giá.')
    doc.add_paragraph('Điều 4: Hiệu lực')
    doc.add_paragraph('Hợp đồng này có hiệu lực đến hết ngày {{NGAY_HET_HIEU_LUC}}.')
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = 'ĐẠI DIỆN BÊN A'
    table.rows[0].cells[1].text = 'ĐẠI DIỆN BÊN B'
    doc.save(TEMPLATE_DOCX)

def init_db():
    DB_PATH.parent.mkdir(exist_ok=True)
    PRODUCT_UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
    COMPANY_LOGO_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_DIR.joinpath('contracts').mkdir(parents=True, exist_ok=True)
    QUOTE_DOCS_DIR.mkdir(parents=True, exist_ok=True)
    create_template_docx()
    create_quote_doc_templates()
    ensure_contract_template_layout()
    db.create_all()
    ensure_contract_columns()
    ensure_order_columns()
    ensure_payment_columns()
    CONTRACT_SIGNED_UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
    ORDER_HANDOVER_UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
    PAYMENT_RECEIPT_UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
    SUPPLIER_PAYMENT_RECEIPT_UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
    ensure_category_brand_schema()
    ensure_product_supplier_schema()
    ensure_supplier_intake_schema()
    ensure_supplier_intake_line_columns()
    ensure_supplier_payment_columns()
    ensure_company_profiles()
    ensure_walkin_customer()
    if User.query.count() == 0:
        admin = User(
            username='admin',
            full_name='Quản trị viên',
            role='admin',
        )
        set_user_password(admin, 'admin123')
        db.session.add(admin)
    db.session.commit()

if __name__ == '__main__':
    with app.app_context():
        init_db()
    app.run(host='0.0.0.0', port=5050, debug=True)
